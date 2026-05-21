"""
Build script for Paper 38: Admissibility Gate.
Generates both PUBLIC and FULL (confidential) .docx versions.
Framework F29: Binary PERMIT/DENY enforcement for deployment decisions.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_document_formatting(doc):
    """Set page size, margins, and default font."""
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.25
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    """Add a heading with Arial font."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text):
    """Add a justified paragraph with correct formatting."""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para_format = para.paragraph_format
    para_format.line_spacing = 1.25
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return para


def add_footer(doc, footer_text):
    """Add footer text to all sections."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = footer_text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_header(doc, header_text):
    """Add header text to all sections."""
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = header_text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_table(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)
                run.bold = True

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)

    return table


# ============================================================
# SECTION CONTENT
# ============================================================

ABSTRACT = [
    "This paper defines the Admissibility Gate for MTCP evaluation outcomes. "
    "The Gate is a binary decision engine. "
    "It takes four inputs and returns PERMIT or DENY.",

    "Inputs are BIS (behavioural integrity), CSAS (cross-system admissibility), "
    "JRS (jurisdiction resolution), and TDS (temporal drift status). "
    "All four must meet the threshold for the target deployment context.",

    "Five deployment contexts are defined with distinct thresholds. "
    "Critical infrastructure requires the highest scores. "
    "General enterprise requires the lowest. "
    "The Gate applies uniformly within each context.",

    "The Admissibility Lemma states that partial compliance is not admissibility. "
    "Meeting 3 of 4 thresholds is DENY. "
    "There is no weighted average or override mechanism. "
    "The decision is binary and final until expiry.",

    "Gate decisions are hash-chained using BEC architecture. "
    "Each decision record is immutable and auditable. "
    "Decisions expire when TDS validity closes or scores drift. "
    "Re-evaluation produces a new decision record.",
]

INTRODUCTION = [
    "MTCP evaluates 32 models from 13 providers across multiple metrics. "
    "These metrics produce grades and scores. "
    "Grades inform decisions but do not enforce them.",

    "Evaluation without enforcement is incomplete. "
    "A model with Grade F can still be deployed. "
    "No mechanism prevents it. "
    "The Admissibility Gate provides that mechanism.",

    "The Gate sits on top of all evaluation constructs. "
    "It combines BIS, CSAS, JRS, and TDS into a single binary outcome. "
    "PERMIT means the model meets all requirements. "
    "DENY means it does not.",

    "This paper specifies the Gate architecture. "
    "It defines threshold tables for five deployment contexts. "
    "It establishes the Admissibility Lemma. "
    "It specifies the hash-chained decision record format.",

    "The Gate does not replace evaluation. "
    "It depends on evaluation. "
    "BIS must be computed before the Gate can reference it. "
    "The Gate is the enforcement layer above the measurement layer.",
]

ENFORCEMENT_PROBLEM = [
    "Grades exist but do not prevent deployment. "
    "A model receiving BIS Grade D can still be deployed in healthcare. "
    "A model with Critical temporal drift can still serve financial transactions. "
    "No binary gate exists to prevent inappropriate deployment.",

    "The gap between measurement and enforcement creates liability. "
    "Organizations have evaluation data showing a model is unsuitable. "
    "They deploy it anyway because no enforcement mechanism says no. "
    "When failure occurs the evaluation data becomes evidence of negligence.",

    "Regulators require enforcement not just measurement. "
    "EU AI Act Article 9 requires risk management systems. "
    "Risk management requires the ability to prevent deployment. "
    "Measurement without prevention is monitoring without management.",

    "The Admissibility Gate fills this gap. "
    "It converts evaluation scores into a binary deployment decision. "
    "It provides the enforcement that measurement alone cannot. "
    "It creates a verifiable record that the decision was made.",
]

FORMAL_FRAMEWORK = [
    "The Admissibility Gate is a binary decision function G. "
    "It takes four inputs and a deployment context. "
    "It returns PERMIT or DENY. "
    "There are no other possible outcomes.",

    "G(BIS, CSAS, JRS, TDS, C) produces PERMIT when all thresholds are met. "
    "G produces DENY when any threshold is not met. "
    "The function is conjunctive. "
    "All conditions must hold simultaneously.",

    "BIS provides behavioural integrity. "
    "It measures constraint persistence across multi-turn interactions. "
    "It is always required regardless of deployment pattern.",

    "CSAS provides cross-system admissibility. "
    "It measures whether constraints transfer between systems. "
    "It is required when the model operates in a coordination pair.",

    "JRS provides jurisdiction resolution. "
    "It measures whether constraint authority is explicitly established. "
    "It is required when the model operates at a system boundary.",

    "TDS provides temporal stability status. "
    "It measures whether model behaviour is stable over time. "
    "It is always required. "
    "Its validity window of 90 days also controls decision expiry.",
]

GATE_THRESHOLDS = [
    "Five deployment contexts define distinct threshold requirements. "
    "Higher-risk contexts require higher scores. "
    "The thresholds reflect the consequence of failure in each context.",

    "Critical infrastructure requires Grade A (90%) for BIS and CSAS. "
    "JRS must be 0.9 or above. "
    "TDS must be Stable. "
    "This is the highest threshold. "
    "Failure in critical infrastructure has catastrophic consequences.",

    "Financial services requires Grade B (80%) for BIS and CSAS. "
    "JRS must be at the standard 0.75 threshold. "
    "TDS must be Stable or Marginal. "
    "Financial systems handle sensitive transactions and personal data.",

    "Healthcare requires Grade B (80%) for BIS and CSAS. "
    "JRS must be 0.75 or above. "
    "TDS must be Stable or Marginal. "
    "Clinical systems affect patient safety and treatment decisions.",

    "Government services requires Grade B (80%) for BIS. "
    "CSAS requires Grade C (70%). "
    "JRS threshold is 0.5. "
    "TDS must not be Critical. "
    "Government systems serve citizens and require accountability.",

    "General enterprise requires Grade C (70%) for BIS and CSAS. "
    "JRS threshold is 0.5. "
    "TDS must not be Critical. "
    "Enterprise tools have lower consequence of failure.",
]

ADMISSIBILITY_LEMMA = [
    "The Admissibility Lemma: a model is deployment-admissible ONLY when "
    "all four gate inputs meet the threshold for the target context. "
    "Partial compliance is not admissibility.",

    "Meeting 3 of 4 thresholds is DENY not PERMIT. "
    "A model with BIS Grade A and CSAS Grade A and JRS 0.95 "
    "but TDS Critical is DENIED for all contexts. "
    "One failure is sufficient for denial.",

    "There is no weighted average mechanism. "
    "There is no override mechanism. "
    "There is no exception process. "
    "The only path from DENY to PERMIT is re-evaluation with improved scores.",

    "The Lemma is strict by design. "
    "Deployment decisions require certainty. "
    "Partial compliance creates ambiguity. "
    "The Gate eliminates ambiguity through binary enforcement.",

    "BEC integrity is a prerequisite for PERMIT. "
    "If input scores come from a compromised chain then "
    "those scores are inadmissible under the Integrity Lemma. "
    "Inadmissible inputs force DENY regardless of their values.",
]

GATE_DECISION_RECORD = [
    "Each gate evaluation produces an immutable decision record. "
    "Records are hash-chained using BEC architecture from Framework F28. "
    "Each record includes the hash of the previous record.",

    "Record fields are: gate_id, model_id, provider, evaluation_timestamp, "
    "bis_score, csas_score, jrs_score, tds_status, deployment_context, "
    "decision, decision_timestamp, previous_gate_hash, and gate_hash.",

    "The gate_hash is computed as SHA-256 of all fields concatenated "
    "with pipe delimiters. "
    "This follows the same pattern as BEC record hashes. "
    "Null fields are represented as the string 'null'.",

    "The genesis decision record has previous_gate_hash equal to 64 zeros. "
    "Each subsequent record references its predecessor. "
    "Modification of any record breaks the chain from that point forward.",

    "DENY decisions include deny_reasons listing which thresholds failed. "
    "This provides auditability. "
    "Reviewers can see exactly why a model was denied. "
    "The reasons are stored as part of the immutable record.",
]

GATE_EXPIRY = [
    "Gate decisions are not permanent. "
    "They expire under defined conditions. "
    "An expired PERMIT requires re-evaluation before continued deployment.",

    "TDS validity window expiry: TDS evaluations are valid for 90 days. "
    "When the window closes the TDS input is no longer current. "
    "The gate decision expires automatically.",

    "Score drift expiry: if any input score changes by more than "
    "the drift tolerance for the context then the decision expires. "
    "Critical infrastructure has 2 percentage point tolerance. "
    "General enterprise has 10 percentage point tolerance.",

    "Expired PERMIT does not automatically become DENY. "
    "It becomes 'expired -- re-evaluation required.' "
    "However continued deployment without re-evaluation is a compliance violation.",

    "Expiry is recorded in the gate expiry log with the reason. "
    "Re-evaluation produces a new gate decision record. "
    "The new record is appended to the hash chain. "
    "The old record remains in the chain as historical evidence.",
]

API_CONTRACT = [
    "The Gate exposes two API endpoints for programmatic access. "
    "These are specification-level definitions. "
    "Implementation details are private.",

    "GET /gate/evaluate/{model}/{context} evaluates the Gate. "
    "It retrieves current scores for the specified model. "
    "It applies thresholds for the specified context. "
    "It returns PERMIT or DENY with full reasoning.",

    "GET /gate/history/{model} returns all gate decisions for a model. "
    "It includes decisions across all deployment contexts. "
    "It includes both current and expired decisions. "
    "It provides the full audit trail for the model.",

    "Response format is JSON. "
    "The evaluate endpoint returns: decision, scores, thresholds applied, "
    "reasons (if DENY), gate_hash, and expiry timestamp (if PERMIT). "
    "The history endpoint returns an array of decision records.",
]

RELATIONSHIP_EXISTING = [
    "BIS provides the primary input to the Gate. "
    "Every gate evaluation requires a current BIS score. "
    "BIS measures constraint persistence. "
    "The Gate enforces minimum constraint persistence for each context.",

    "CSAS provides the coordination input. "
    "When models work in pairs the CSAS score matters. "
    "The Gate enforces minimum cross-system admissibility. "
    "Models that cannot maintain constraints across system boundaries "
    "are denied for contexts requiring coordination.",

    "JRS provides the governance input. "
    "Jurisdiction must be established before coordination. "
    "The Gate enforces minimum jurisdiction resolution. "
    "Models without clear governance structures are denied for "
    "high-stakes contexts.",

    "TDS provides both stability input and temporal control. "
    "The drift classification feeds into the gate evaluation. "
    "The validity window controls decision expiry. "
    "TDS is the only input that directly affects decision lifetime.",

    "BEC provides the integrity infrastructure. "
    "Gate decision records are hash-chained using BEC architecture. "
    "Input score integrity depends on BEC verification. "
    "The Gate is built on top of BEC not alongside it.",
]

IMPLICATIONS = [
    "Procurement decisions become verifiable. "
    "Organizations can require PERMIT status for a specific context "
    "before purchasing or deploying a model. "
    "The gate record proves the model met all requirements at that time.",

    "SLAs become enforceable. "
    "Service level agreements can reference gate status. "
    "A provider guarantees PERMIT for financial services. "
    "If the gate later returns DENY the SLA is breached.",

    "Regulatory enforcement becomes operationalized. "
    "Regulators can verify that appropriate gate checks were performed. "
    "The hash-chained records prove decisions were made before deployment. "
    "This satisfies EU AI Act requirements for risk management.",

    "Liability boundaries become clear. "
    "An organization that deploys after DENY has evidence of negligence. "
    "An organization that deploys after PERMIT has evidence of due diligence. "
    "The gate record is evidence in both directions.",
]

LIMITATIONS = [
    "The Gate enforces thresholds but does not improve scores. "
    "A denied model must be improved by its provider. "
    "The Gate provides no remediation guidance.",

    "Threshold selection is a policy decision not a scientific one. "
    "The specific thresholds represent MTCP assessment of risk. "
    "Different organizations may require different values. "
    "The architecture supports configurable thresholds.",

    "The Gate evaluates individual models not systems. "
    "System-level admissibility requires evaluating each model. "
    "A system is admissible only when all its models are admissible.",

    "The 90-day TDS validity window is conservative. "
    "Some contexts may need shorter windows. "
    "The current specification does not support per-context windows.",

    "Input score accuracy is assumed. "
    "If evaluation produces incorrect scores the Gate "
    "enforces incorrect thresholds faithfully. "
    "The Gate does not validate input accuracy.",
]

CONCLUSION = [
    "The Admissibility Gate provides the enforcement layer above "
    "MTCP evaluation metrics. "
    "It converts scores into a binary deployment decision. "
    "PERMIT or DENY. Nothing else.",

    "The Admissibility Lemma ensures no partial compliance. "
    "All four inputs must meet thresholds for the target context. "
    "One failure is sufficient for DENY. "
    "This eliminates ambiguity in deployment decisions.",

    "Gate decisions are hash-chained and immutable. "
    "They expire when inputs change or validity windows close. "
    "Re-evaluation produces new decisions on the chain. "
    "The full history is preserved for audit.",

    "Evaluation without enforcement is incomplete. "
    "The Gate completes the MTCP architecture by providing "
    "the binary decision that deployment contexts require. "
    "Models are either admissible or they are not.",
]

REFERENCES = [
    "Abby, A. (2026). Multi-Turn Constraint Persistence (MTCP) V1.0. "
    "DOI: 10.17605/OSF.IO/DXGK5.",

    "Abby, A. (2026, Framework F4). Behavioural Integrity Score Definition. "
    "MTCP Research Programme.",

    "Abby, A. (2026, Framework F21). Cross-System Admissibility Score. "
    "MTCP Research Programme.",

    "Abby, A. (2026, Framework F22). Jurisdiction Resolution Score. "
    "MTCP Research Programme.",

    "Abby, A. (2026, Framework F23). Temporal Drift Status. "
    "MTCP Research Programme.",

    "Abby, A. (2026, Framework F28). Blockchain Evidence Chain Definition. "
    "MTCP Research Programme.",

    "Abby, A. (2026, Framework F29). Admissibility Gate Definition. "
    "MTCP Research Programme.",

    "EU AI Act. Regulation (EU) 2024/1689 of the European Parliament "
    "and of the Council. Official Journal of the European Union.",

    "ISO/IEC 42001:2023. Information technology -- Artificial intelligence "
    "-- Management system. International Organization for Standardization.",

    "NIST AI 100-1. Artificial Intelligence Risk Management Framework. "
    "National Institute of Standards and Technology.",
]

IMPLEMENTATION_NOTES = [
    "The Gate is implemented as a PostgreSQL-backed service. "
    "Four tables support the architecture: gate_decisions, gate_thresholds, "
    "gate_expiry_log, and gate_registry.",

    "gate_decisions stores every decision record with full inputs and hash chain. "
    "gate_thresholds stores configurable thresholds per deployment context. "
    "gate_expiry_log tracks when decisions expire and why. "
    "gate_registry tracks current state per model per context.",

    "Hash computation follows the BEC pattern. "
    "SHA-256 of pipe-delimited fields produces a 64-character hex string. "
    "The genesis record uses 64 zeros as previous hash. "
    "Each subsequent record references its predecessor.",

    "The gate_registry table uses a unique constraint on model_id "
    "plus deployment_context. "
    "This ensures one current decision per model per context. "
    "Historical decisions remain in gate_decisions for audit.",

    "Expiry monitoring runs as a scheduled check. "
    "It queries gate_decisions for PERMIT decisions past expires_at. "
    "Expired decisions are logged to gate_expiry_log. "
    "Notifications trigger re-evaluation workflows.",

    "The API server exposes two endpoints. "
    "GET /api/gate/evaluate/{model}/{context} runs the gate live. "
    "GET /api/gate/history/{model} returns all historical decisions. "
    "Both endpoints return JSON responses.",
]


# ============================================================
# THRESHOLD TABLE DATA
# ============================================================

THRESHOLD_HEADERS = ["Context", "Min BIS", "Grade", "Min CSAS", "Min JRS", "Max TDS"]
THRESHOLD_ROWS = [
    ["Critical Infrastructure", "90%", "A", "90%", "0.9", "Stable"],
    ["Financial Services", "80%", "B", "80%", "0.75", "Marginal"],
    ["Healthcare", "80%", "B", "80%", "0.75", "Marginal"],
    ["Government Services", "80%", "B", "70%", "0.5", "Significant"],
    ["General Enterprise", "70%", "C", "70%", "0.5", "Significant"],
]


# ============================================================
# BUILD FUNCTIONS
# ============================================================

def build_public():
    """Build the PUBLIC version of Paper 38."""
    doc = Document()
    set_document_formatting(doc)
    add_header(doc, "MTCP V1.0 -- Admissibility Gate")

    # Title page
    add_heading_styled(doc, "MTCP Paper 38", level=1)
    add_heading_styled(doc, "Admissibility Gate", level=2)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "Status: NOT for OSF")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(
        doc,
        "This paper defines the Admissibility Gate for MTCP. "
        "The Gate provides binary PERMIT/DENY enforcement for deployment decisions. "
        "It evaluates 32 models from 13 providers against context-specific thresholds. "
        "This is a specification document. "
        "No implementation code is included."
    )

    doc.add_page_break()

    # 1. Abstract
    add_heading_styled(doc, "1. Abstract", level=1)
    for para_text in ABSTRACT:
        add_paragraph_justified(doc, para_text)

    # 2. Introduction
    add_heading_styled(doc, "2. Introduction", level=1)
    for para_text in INTRODUCTION:
        add_paragraph_justified(doc, para_text)

    # 3. The Enforcement Problem
    add_heading_styled(doc, "3. The Enforcement Problem", level=1)
    for para_text in ENFORCEMENT_PROBLEM:
        add_paragraph_justified(doc, para_text)

    # 4. Formal Framework
    add_heading_styled(doc, "4. Formal Framework", level=1)
    for para_text in FORMAL_FRAMEWORK:
        add_paragraph_justified(doc, para_text)

    # 5. Gate Threshold Specification
    add_heading_styled(doc, "5. Gate Threshold Specification", level=1)
    for para_text in GATE_THRESHOLDS:
        add_paragraph_justified(doc, para_text)
    add_paragraph_justified(doc, "")
    add_table(doc, THRESHOLD_HEADERS, THRESHOLD_ROWS)
    add_paragraph_justified(doc, "")

    # 6. Admissibility Lemma
    add_heading_styled(doc, "6. The Admissibility Lemma", level=1)
    for para_text in ADMISSIBILITY_LEMMA:
        add_paragraph_justified(doc, para_text)

    # 7. Gate Decision Record
    add_heading_styled(doc, "7. Gate Decision Record", level=1)
    for para_text in GATE_DECISION_RECORD:
        add_paragraph_justified(doc, para_text)

    # 8. Gate Expiry and Re-evaluation
    add_heading_styled(doc, "8. Gate Expiry and Re-evaluation", level=1)
    for para_text in GATE_EXPIRY:
        add_paragraph_justified(doc, para_text)

    # 9. API Contract
    add_heading_styled(doc, "9. API Contract", level=1)
    for para_text in API_CONTRACT:
        add_paragraph_justified(doc, para_text)

    # 10. Relationship to Existing Constructs
    add_heading_styled(doc, "10. Relationship to Existing Constructs", level=1)
    for para_text in RELATIONSHIP_EXISTING:
        add_paragraph_justified(doc, para_text)

    # 11. Implications
    add_heading_styled(doc, "11. Implications", level=1)
    for para_text in IMPLICATIONS:
        add_paragraph_justified(doc, para_text)

    # 12. Limitations
    add_heading_styled(doc, "12. Limitations", level=1)
    for para_text in LIMITATIONS:
        add_paragraph_justified(doc, para_text)

    # 13. Conclusion
    add_heading_styled(doc, "13. Conclusion", level=1)
    for para_text in CONCLUSION:
        add_paragraph_justified(doc, para_text)

    # 14. References
    add_heading_styled(doc, "14. References", level=1)
    for ref in REFERENCES:
        add_paragraph_justified(doc, ref)

    output_path = os.path.expanduser(
        "~/Desktop/Paper38_Admissibility_Gate_PUBLIC.docx"
    )
    doc.save(output_path)
    print(f"PUBLIC version saved: {output_path}")
    return output_path


def build_full():
    """Build the FULL (confidential) version of Paper 38."""
    doc = Document()
    set_document_formatting(doc)
    add_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
    add_footer(
        doc,
        "MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved."
    )

    # Title page
    add_heading_styled(doc, "MTCP Paper 38", level=1)
    add_heading_styled(doc, "Admissibility Gate", level=2)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "Status: CONFIDENTIAL -- NDA REQUIRED")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(
        doc,
        "This paper defines the Admissibility Gate for MTCP. "
        "The Gate provides binary PERMIT/DENY enforcement for deployment decisions. "
        "It evaluates 32 models from 13 providers against context-specific thresholds. "
        "This version includes implementation architecture notes. "
        "Distribution restricted to NDA holders."
    )

    doc.add_page_break()

    # 1. Abstract
    add_heading_styled(doc, "1. Abstract", level=1)
    for para_text in ABSTRACT:
        add_paragraph_justified(doc, para_text)
    add_paragraph_justified(
        doc,
        "The Gate enforces thresholds across the full MTCP model roster. "
        "This includes GPT-4o, GPT-4o-mini, GPT-3.5, Claude Sonnet 4, "
        "Claude Haiku 4.5, Gemini 2.0 Flash, Gemini 2.5, Gemini 2.5 Flash, "
        "Gemma 2 27B, Gemma 3 27B, Llama 3.1 8B, Llama 3.1 70B, "
        "Llama 3.3 70B, Llama 4 Scout, Llama 4 Maverick, Mistral Large, "
        "Mistral Small, Magistral Medium, Command R+, DeepSeek R1, "
        "Grok 3 Mini, Nova Micro, Nova Lite, Nova Pro, NVIDIA Llama 3.1, "
        "Phi-4, Kimi, Qwen 2.5, Qwen 3 8B, Qwen 3 32B, and Granite."
    )

    # 2. Introduction
    add_heading_styled(doc, "2. Introduction", level=1)
    for para_text in INTRODUCTION:
        add_paragraph_justified(doc, para_text)
    add_paragraph_justified(
        doc,
        "Providers evaluated include OpenAI, Anthropic, Google, Meta, "
        "Mistral, Cohere, Amazon, xAI, Cerebras, Fireworks, NVIDIA, "
        "Moonshot, and IBM. "
        "Each provider's models are subject to gate evaluation "
        "before deployment in regulated contexts."
    )

    # 3. The Enforcement Problem
    add_heading_styled(doc, "3. The Enforcement Problem", level=1)
    for para_text in ENFORCEMENT_PROBLEM:
        add_paragraph_justified(doc, para_text)
    add_paragraph_justified(
        doc,
        "Concrete example: Llama 3.1 8B received BIS Grade C (72.1%). "
        "Without the Gate this model could be deployed in healthcare. "
        "The Gate prevents this. "
        "Healthcare requires Grade B minimum. "
        "The Gate returns DENY with explicit reasoning."
    )

    # 4. Formal Framework
    add_heading_styled(doc, "4. Formal Framework", level=1)
    for para_text in FORMAL_FRAMEWORK:
        add_paragraph_justified(doc, para_text)

    # 5. Gate Threshold Specification
    add_heading_styled(doc, "5. Gate Threshold Specification", level=1)
    for para_text in GATE_THRESHOLDS:
        add_paragraph_justified(doc, para_text)
    add_paragraph_justified(doc, "")
    add_table(doc, THRESHOLD_HEADERS, THRESHOLD_ROWS)
    add_paragraph_justified(doc, "")

    # 6. Admissibility Lemma
    add_heading_styled(doc, "6. The Admissibility Lemma", level=1)
    for para_text in ADMISSIBILITY_LEMMA:
        add_paragraph_justified(doc, para_text)

    # 7. Gate Decision Record
    add_heading_styled(doc, "7. Gate Decision Record", level=1)
    for para_text in GATE_DECISION_RECORD:
        add_paragraph_justified(doc, para_text)

    # 8. Gate Expiry and Re-evaluation
    add_heading_styled(doc, "8. Gate Expiry and Re-evaluation", level=1)
    for para_text in GATE_EXPIRY:
        add_paragraph_justified(doc, para_text)

    # 9. API Contract
    add_heading_styled(doc, "9. API Contract", level=1)
    for para_text in API_CONTRACT:
        add_paragraph_justified(doc, para_text)

    # 10. Relationship to Existing Constructs
    add_heading_styled(doc, "10. Relationship to Existing Constructs", level=1)
    for para_text in RELATIONSHIP_EXISTING:
        add_paragraph_justified(doc, para_text)

    # 11. Implications
    add_heading_styled(doc, "11. Implications", level=1)
    for para_text in IMPLICATIONS:
        add_paragraph_justified(doc, para_text)

    # 12. Limitations
    add_heading_styled(doc, "12. Limitations", level=1)
    for para_text in LIMITATIONS:
        add_paragraph_justified(doc, para_text)

    # 13. Conclusion
    add_heading_styled(doc, "13. Conclusion", level=1)
    for para_text in CONCLUSION:
        add_paragraph_justified(doc, para_text)

    # 14. Implementation Notes (FULL version only)
    add_heading_styled(doc, "14. Implementation Notes", level=1)
    add_paragraph_justified(
        doc,
        "The following architectural notes describe the implementation "
        "of the Admissibility Gate within the MTCP platform. "
        "These notes are restricted to NDA holders."
    )
    for para_text in IMPLEMENTATION_NOTES:
        add_paragraph_justified(doc, para_text)

    # 15. References
    add_heading_styled(doc, "15. References", level=1)
    for ref in REFERENCES:
        add_paragraph_justified(doc, ref)

    output_path = os.path.expanduser(
        "~/Desktop/Paper38_Admissibility_Gate_FULL.docx"
    )
    doc.save(output_path)
    print(f"FULL version saved: {output_path}")
    return output_path


if __name__ == '__main__':
    print("Building Paper 38: Admissibility Gate")
    print("=" * 50)
    pub = build_public()
    full = build_full()
    print("=" * 50)
    print("Done.")
    print(f"  PUBLIC: {pub}")
    print(f"  FULL:   {full}")
