"""
Build script for Paper 32: Remediation Effectiveness Score.
Generates both PUBLIC and FULL (confidential) .docx versions.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os


def set_document_formatting(doc):
    """Set page size, margins, and default font."""
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.25
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    """Add a heading with Arial font."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text):
    """Add a justified paragraph with correct formatting."""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para_format = para.paragraph_format
    para_format.line_spacing = 1.25
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return para


def add_footer(doc, footer_text):
    """Add footer text to all sections."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = footer_text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_header(doc, header_text):
    """Add header text to all sections."""
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = header_text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


# ============================================================
# PAPER CONTENT
# ============================================================

ABSTRACT = (
    "MTCP identifies constraint persistence failures across 35 models and 14 providers. "
    "Identification alone is insufficient for deployers. "
    "This paper introduces the Remediation Effectiveness Score (RES). "
    "RES measures percentage point improvement in BIS achieved by a specific intervention. "
    "Four intervention types are defined. "
    "Prompt engineering, system prompt injection, temperature reduction, and instruction reordering each produce different expected RES ranges. "
    "The Remediation Lemma establishes that RES is model-specific and intervention-specific. "
    "Fixes do not generalise across models. "
    "A remediation ceiling exists for each model. "
    "No operational intervention exceeds this ceiling. "
    "For architectural failures, provider-level training intervention is the only path beyond the ceiling. "
    "RES is theoretical. "
    "No systematic remediation experiments have been conducted. "
    "Expected values derive from MTCP temperature sensitivity analysis and PRP design principles."
)

INTRODUCTION = [
    "MTCP evaluates constraint persistence across 183,924 evaluations. It assigns grades from A to F. It identifies which models fail and how severely they fail. It distinguishes stochastic from architectural failure patterns. This is necessary work. It is not sufficient work.",
    "A deployer receives a Grade D evaluation. The model violates constraints 31 to 40 percent of the time. The deployer now knows the problem exists. The deployer does not know what to do about it. MTCP provides diagnosis without treatment.",
    "This gap is not accidental. Rigorous evaluation must precede remediation guidance. Measuring the problem is prerequisite to solving it. Papers 1 through 31 establish the measurement framework. Paper 32 begins the remediation framework.",
    "The question this paper addresses is direct. Given a model with a known constraint persistence failure, what interventions are available and how effective are they? The answer requires a formal metric for intervention effectiveness. That metric is RES.",
]

REMEDIATION_PROBLEM = [
    "Deployers who receive sub-Grade-B evaluations face a practical decision. They can accept the failure rate. They can switch models. They can attempt to fix the problem. They can abandon the use case.",
    "Accepting a 30 percent constraint violation rate is negligent in regulated domains. Switching models is expensive and disruptive. Abandoning the use case wastes prior investment. Attempting remediation is the rational choice.",
    "Without formal guidance, remediation attempts are ad hoc. A deployer rewrites the prompt. Performance improves slightly. Is the improvement statistically significant? Is it the maximum achievable? Would a different approach work better?",
    "The remediation problem has a deeper structural issue. Not all failures are remediable through operational controls. Paper 9 established that architectural failures are temperature-invariant. They persist regardless of sampling configuration. Prompt engineering cannot address training-level deficits.",
    "Deployers need three things. First, a metric that measures intervention effectiveness. Second, a taxonomy that classifies available interventions. Third, a ceiling concept that identifies when operational intervention is futile and provider-level change is required. RES provides all three.",
]

FORMAL_FRAMEWORK = [
    "RES is defined relative to baseline BIS. Let M be a model with baseline BIS score B_before. Let I be an intervention applied to M. Let B_after be the BIS score after intervention. RES is computed as follows.",
    "RES(M, I) = (B_after -- B_before) / (100 -- B_before).",
    "The denominator normalises by the remaining improvement space. A model at BIS 60 has 40 percentage points of possible improvement. A model at BIS 80 has 20 percentage points. RES measures what fraction of available improvement was captured.",
    "RES = 1.0 means full remediation. The failure is completely resolved. RES = 0.0 means no improvement. The intervention had no effect. Negative RES means the intervention made performance worse.",
    "This normalisation is deliberate. A 10 percentage point improvement from 60 to 70 represents greater relative effectiveness than a 10 point improvement from 80 to 90. The first captures 25 percent of available space. The second captures 50 percent. RES reflects this difference.",
    "The remediation ceiling C(M) is defined as the maximum BIS achievable through any operational intervention. When C(M) falls below the Grade B threshold of 80 percent, no operational intervention achieves deployment readiness. Provider-level training change is required.",
]

REMEDIATION_LEMMA = [
    "The Remediation Lemma states that RES is model-specific and intervention-specific. An intervention producing RES = 0.3 on model M provides no information about its effectiveness on model N.",
    "This follows directly from the stochastic versus architectural distinction. Paper 9 established two failure patterns. Pattern 1 is architectural. The failure is embedded in model weights. It persists at all temperatures. Pattern 2 is stochastic. The failure emerges from sampling variance. It responds to temperature control.",
    "Temperature reduction produces positive RES for stochastic failures. Expected range is 0.0 to 0.5. Temperature reduction produces RES near zero for architectural failures. Expected range is 0.0 to 0.05. The same intervention has fundamentally different effectiveness depending on failure type.",
    "This non-transferability extends beyond temperature. Prompt engineering that resolves a constraint violation in one model may have no effect on another model. The linguistic representation of constraints interacts with model-specific training in ways that do not generalise.",
    "The practical consequence is clear. Every remediation claim must be validated per-model and per-constraint. Vendor statements that a fix works generally are not credible without per-model evidence. RES provides the metric for that evidence.",
]

INTERVENTION_TAXONOMY = [
    "Four operational intervention types are defined within RES. Each has a distinct mechanism and expected effectiveness range.",
    "Type 1 is prompt engineering. The constraint instruction is rephrased without changing its semantic content. Directness, specificity, and imperative framing are typical modifications. Expected RES ranges from 0.0 to 0.3. The effect is highly model-specific. Some constraint phrasings outperform others but the optimal phrasing differs across models.",
    "Type 2 is system prompt injection. The constraint is added to the system prompt so it is present at every inference call. This reinforces the constraint continuously rather than relying on single-turn instruction following. Expected RES ranges from 0.1 to 0.4. This is the most effective single operational intervention because it addresses the attention decay mechanism.",
    "Type 3 is temperature reduction. The inference temperature is lowered from the deployment default. This reduces sampling variance and makes outputs more deterministic. Expected RES ranges from 0.0 to 0.5 for stochastic failures. Expected RES is 0.0 to 0.05 for architectural failures. Temperature reduction serves both remediation and diagnostic functions.",
    "Type 4 is instruction reordering. The constraint position relative to the task is changed. Constraint-before-task versus task-before-constraint. This exploits positional attention biases. Expected RES ranges from 0.05 to 0.15. The effect is small but consistent across most models.",
    "PRP (Posture Reauthorization Protocol) is a special intervention outside the standard taxonomy. Defined in Paper 20, PRP is a wrapper-layer mitigation designed specifically for architectural failures. Expected RES ranges from 0.2 to 0.5 on architectural models. PRP is the only intervention with positive expected RES for Pattern 1 failures.",
]

CEILING_IDENTIFICATION = [
    "The remediation ceiling is the most important practical concept in RES. It answers the question that matters most to deployers. How good can this model get?",
    "Some models have architectural ceilings well below Grade A. No prompt engineering exceeds this ceiling. No system prompt injection exceeds it. No temperature configuration exceeds it. The ceiling is structural.",
    "Identifying the ceiling requires exhaustive intervention testing. All four types must be attempted individually and in combination. The highest observed BIS_after across all interventions approximates the ceiling.",
    "When the ceiling falls below Grade B, the deployer faces a binary choice. Accept a permanently sub-B model with operational mitigations. Or require the provider to conduct training-level intervention. There is no third option.",
    "The ceiling concept reframes the conversation between deployers and providers. Instead of asking whether a model can be fixed, the question becomes where the ceiling sits. A Grade D model with a ceiling at Grade B is remediable through operations. A Grade D model with a ceiling at Grade C is not.",
    "Provider-level training intervention is the only path beyond an architectural ceiling. Fine-tuning, RLHF adjustment, or constitutional AI revision may raise the ceiling. These require provider cooperation and resources. RES cannot measure their effectiveness until after implementation.",
]

COMMERCIAL_FRAMEWORK = [
    "RES enables grade-specific remediation guidance. Each BIS grade below B has a defined minimum RES required to reach deployment readiness.",
    "For Grade D models (BIS 60 to 69 percent), the minimum RES to reach Grade B is 0.52 to 0.63. This means more than half the failure gap must be closed. The recommended sequence begins with temperature reduction as a diagnostic. If effective, the failure is stochastic and the combination of temperature reduction with system prompt injection is likely sufficient.",
    "For Grade F models (BIS below 60 percent), the minimum RES to reach Grade B is 0.50 to 1.00 depending on baseline. For a model at BIS 40 percent, RES must exceed 0.67 to reach Grade B. This is achievable for stochastic failures but not for architectural ones.",
    "Procurement contracts should incorporate RES requirements. Baseline BIS evaluation before deployment. Documented RES evaluation if grade is below B. Vendor acknowledgment of architectural limitation if RES grade is F. Timeline commitment for provider-level intervention.",
    "Vendor SLAs can reference RES directly. A provider commits to achieving RES greater than 0.5 within 60 days of Grade D notification. If RES remains below 0.2 after operational intervention, the provider acknowledges architectural failure and commits to training-level remediation.",
]

RELATIONSHIP_CONSTRUCTS = [
    "RES builds on and relates to multiple existing MTCP constructs. BIS is the dependent variable. RES measures change in BIS. Without BIS there is no RES. The relationship is foundational.",
    "Ve (Framework F2) measures within-conversation failure escalation. High Ve means the model cannot self-correct during interaction. RES addresses the question of what happens outside the conversation. Operational interventions modify conditions before conversation begins. RES measures whether those modifications reduce the failures that Ve detects.",
    "The stochastic versus architectural distinction (Paper 9) is the primary predictor of RES outcomes. Stochastic failures respond to temperature-based interventions. Architectural failures do not. RES empirically validates the Paper 9 classification when temperature reduction is tested.",
    "PRP (Paper 20) is the only intervention designed for architectural failures. All other intervention types (prompt engineering, system prompt injection, temperature reduction, instruction reordering) target the operational layer. PRP intervenes at the wrapper layer. Its expected RES of 0.2 to 0.5 on architectural models makes it the primary tool for Pattern 1 remediation.",
    "TDS (Framework F23) monitors BIS drift over time without intervention. RES measures BIS change because of intervention. The two must be distinguished. If TDS is high, measured RES may include drift effects unrelated to the intervention. The RES measurement protocol requires temporal controls to account for this.",
]

IMPLICATIONS = [
    "RES has direct implications for AI procurement. Buyers can require RES documentation alongside BIS grades. A model with Grade D BIS and Grade A RES is remediable. A model with Grade D BIS and Grade F RES is not. Procurement decisions should weight both.",
    "Vendor SLAs gain objective remediation targets. Rather than vague commitments to improvement, providers can commit to specific RES thresholds within defined timeframes. Failure to meet RES targets triggers escalation to provider-level intervention.",
    "PRP (Paper 20) is reframed as an intervention with measurable effectiveness. Its expected RES of 0.2 to 0.5 on architectural models gives deployers a quantitative expectation. If PRP achieves less than expected, the architectural failure may be more severe than initially classified.",
    "The remediation ceiling concept changes how deployers think about model selection. A model with lower baseline BIS but higher ceiling may be preferable to a model with higher baseline BIS but lower ceiling. The ceiling determines long-term viability.",
    "Regulatory frameworks can incorporate RES. If a jurisdiction requires minimum BIS for deployment, RES provides the formal mechanism for demonstrating that sub-threshold models can be remediated to compliance. Without RES, the only compliant response to a failed evaluation is model replacement.",
]

LIMITATIONS = [
    "RES is theoretical. This is the primary limitation. No systematic remediation experiments have been conducted. Expected RES values are derived from MTCP temperature sensitivity analysis and PRP design principles. They are not from controlled intervention trials.",
    "The expected RES ranges stated in this paper are estimates. Actual RES values may differ significantly when empirical measurement is conducted. The ranges are informed by temperature sensitivity patterns observed across 183,924 evaluations but they are not direct measurements of intervention effectiveness.",
    "RES assumes BIS measurement is stable. If BIS itself has high variance (measured by TDS), then RES computation inherits that variance. Small RES values may be indistinguishable from measurement noise.",
    "The intervention taxonomy is not exhaustive. Other intervention types exist. Few-shot examples in system prompts. Chain-of-thought prompting. Output format constraints. These are not currently classified within RES. Future versions may expand the taxonomy.",
    "Interaction effects between interventions are not modelled. Combining temperature reduction with system prompt injection may produce RES different from either alone. The current framework evaluates interventions independently. Combination effects require separate empirical study.",
    "RES does not address the cost of intervention. Temperature reduction is free. PRP adds latency and token cost. System prompt injection consumes context window. A complete remediation framework would include cost-effectiveness analysis. This is deferred to future work.",
]

CONCLUSION = [
    "RES extends MTCP from diagnosis to treatment guidance. It provides a formal metric for intervention effectiveness. It defines an intervention taxonomy. It establishes the remediation ceiling concept. It introduces the Remediation Lemma that constrains generalisation claims.",
    "The key findings are structural. Temperature reduction is diagnostic. System prompt injection is the most effective single operational intervention. PRP is the only intervention designed for architectural failures. The ceiling exists and cannot be exceeded through operational means alone.",
    "For deployers, RES provides actionable guidance. Test temperature reduction first. If effective, the failure is stochastic and remediable. If ineffective, the failure is architectural and requires PRP or provider-level intervention.",
    "For providers, RES creates accountability. Remediation claims require per-model evidence. General statements about improvement are insufficient. The RES metric makes intervention effectiveness falsifiable.",
    "Empirical validation is the immediate next step. Controlled intervention experiments across the MTCP model set will calibrate expected RES ranges, identify actual remediation ceilings, and test the Remediation Lemma across model families.",
]

REFERENCES = [
    "Abby, A. (2026a). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 9). Stochastic versus architectural constraint failure in language models. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 20). Posture Reauthorization Protocol: Wrapper-layer mitigation for architectural constraint failures. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 31). Constraint Conflict Detection: Multi-constraint prioritisation analysis for AI governance. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Three-Layer). Three-layer constraint failure in AI systems. DOI: 10.17605/OSF.IO/DXGK5.",
    "Framework F2 -- Ve Metric. MTCP Research Programme.",
    "Framework F4 -- BIS Definition. MTCP Research Programme.",
    "Framework F23 -- TDS Definition. MTCP Research Programme.",
    "Framework F24 -- CCS Definition. MTCP Research Programme.",
    "Framework F25 -- RES Definition. MTCP Research Programme.",
]


def build_paper(doc, is_public=True):
    """Build the paper content into the document."""
    set_document_formatting(doc)

    if is_public:
        add_header(doc, "MTCP V1.0 -- Remediation Effectiveness Score")
    else:
        add_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
        add_footer(
            doc,
            "MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved.",
        )

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Remediation Effectiveness Score")
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Measuring Intervention Impact on Constraint Persistence Failures"
    )
    run.font.name = "Arial"
    run.font.size = Pt(13)

    # Author info
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("A. Abby")
    run.font.name = "Arial"
    run.font.size = Pt(12)

    if is_public:
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = contact.add_run("admin@mtcp.live")
        run.font.name = "Arial"
        run.font.size = Pt(11)

    doi = doc.add_paragraph()
    doi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doi.add_run("DOI: 10.17605/OSF.IO/DXGK5")
    run.font.name = "Arial"
    run.font.size = Pt(11)

    series = doc.add_paragraph()
    series.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = series.add_run("MTCP Paper Series -- Paper 32")
    run.font.name = "Arial"
    run.font.size = Pt(11)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run("May 2026")
    run.font.name = "Arial"
    run.font.size = Pt(11)

    # Abstract
    add_heading_styled(doc, "1. Abstract", level=1)
    add_paragraph_justified(doc, ABSTRACT)

    # Introduction
    add_heading_styled(doc, "2. Introduction", level=1)
    for para_text in INTRODUCTION:
        add_paragraph_justified(doc, para_text)

    # The Remediation Problem
    add_heading_styled(doc, "3. The Remediation Problem", level=1)
    for para_text in REMEDIATION_PROBLEM:
        add_paragraph_justified(doc, para_text)

    # Formal Framework
    add_heading_styled(doc, "4. Formal Framework", level=1)
    for para_text in FORMAL_FRAMEWORK:
        add_paragraph_justified(doc, para_text)

    # Remediation Lemma
    add_heading_styled(doc, "5. Remediation Lemma", level=1)
    for para_text in REMEDIATION_LEMMA:
        add_paragraph_justified(doc, para_text)

    # Intervention Taxonomy
    add_heading_styled(doc, "6. Intervention Taxonomy", level=1)
    for para_text in INTERVENTION_TAXONOMY:
        add_paragraph_justified(doc, para_text)

    # Remediation Ceiling Identification
    add_heading_styled(doc, "7. Remediation Ceiling Identification", level=1)
    for para_text in CEILING_IDENTIFICATION:
        add_paragraph_justified(doc, para_text)

    # Commercial Framework
    add_heading_styled(doc, "8. Commercial Framework", level=1)
    for para_text in COMMERCIAL_FRAMEWORK:
        add_paragraph_justified(doc, para_text)

    # Relationship to Existing Constructs
    add_heading_styled(doc, "9. Relationship to Existing Constructs", level=1)
    for para_text in RELATIONSHIP_CONSTRUCTS:
        add_paragraph_justified(doc, para_text)

    # Implications
    add_heading_styled(doc, "10. Implications", level=1)
    for para_text in IMPLICATIONS:
        add_paragraph_justified(doc, para_text)

    # Limitations
    add_heading_styled(doc, "11. Limitations", level=1)
    for para_text in LIMITATIONS:
        add_paragraph_justified(doc, para_text)

    # Conclusion
    add_heading_styled(doc, "12. Conclusion", level=1)
    for para_text in CONCLUSION:
        add_paragraph_justified(doc, para_text)

    # References
    add_heading_styled(doc, "13. References", level=1)
    for ref in REFERENCES:
        add_paragraph_justified(doc, ref)


def main():
    desktop = os.path.expanduser("~/Desktop")

    # Build PUBLIC version
    doc_public = Document()
    build_paper(doc_public, is_public=True)
    public_path = os.path.join(
        desktop, "Paper32_Remediation_Effectiveness_PUBLIC.docx"
    )
    doc_public.save(public_path)
    print(f"PUBLIC version saved: {public_path}")

    # Build FULL version
    doc_full = Document()
    build_paper(doc_full, is_public=False)
    full_path = os.path.join(
        desktop, "Paper32_Remediation_Effectiveness_FULL.docx"
    )
    doc_full.save(full_path)
    print(f"FULL version saved: {full_path}")


if __name__ == "__main__":
    main()
