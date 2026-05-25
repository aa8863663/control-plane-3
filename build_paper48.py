"""Build Paper 48: Agent Substrate Measurement. Framework F38."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
DESKTOP = os.path.expanduser("~/Desktop")

def setup(doc):
    for s in doc.sections:
        s.page_width=Inches(8.5);s.page_height=Inches(11)
        s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Inches(1)
    st=doc.styles["Normal"];st.font.name="Arial";st.font.size=Pt(12)
    st.paragraph_format.line_spacing=1.25;st.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

def h(doc,t,l=1):
    hd=doc.add_heading(t,level=l)
    for r in hd.runs: r.font.name="Arial"

def p(doc,t):
    pa=doc.add_paragraph(t);pa.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    pa.paragraph_format.line_spacing=1.25
    for r in pa.runs: r.font.name="Arial";r.font.size=Pt(12)

def hdr(doc,t):
    for s in doc.sections:
        pa=s.header.paragraphs[0] if s.header.paragraphs else s.header.add_paragraph()
        pa.text=t;pa.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in pa.runs: r.font.name="Arial";r.font.size=Pt(9)

def ftr(doc,t):
    for s in doc.sections:
        pa=s.footer.paragraphs[0] if s.footer.paragraphs else s.footer.add_paragraph()
        pa.text=t;pa.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in pa.runs: r.font.name="Arial";r.font.size=Pt(9)

C = [
 ("1. Abstract",[
  "This paper defines Agent Substrate Measurement. "
  "A knowledge graph holon has four graphs: Knowledge, Context, Schema, Projection. "
  "These define what an agent operates against. "
  "The fifth element measures whether the agent followed its constraints.",
  "Four substrate measurement vectors map onto the four graphs. "
  "Knowledge Integrity measures consistency across domains. "
  "Context Provenance measures consistency across data sources. "
  "Schema Persistence is the core BIS measurement. "
  "Projection Consistency measures consistency across output formats.",
  "The Substrate Measurement Lemma establishes the requirement. "
  "A holon without measurement is architecturally specified but empirically unverified. "
  "Both are required for the holon to be load-bearing.",
 ]),
 ("2. Introduction",[
  "Knowledge graph holons define agent operating environments. "
  "They specify what data the agent has access to. "
  "They specify what context governs the interaction. "
  "They specify what schema the agent must follow. "
  "They specify what projection formats are expected.",
  "None of this specifies whether the agent actually follows the schema. "
  "Specification is not verification. "
  "A perfectly specified holon with an unreliable agent is still unreliable. "
  "The specification does not make the agent better.",
  "Agent Substrate Measurement closes this gap. "
  "It measures constraint adherence empirically. "
  "It maps findings onto the four-graph structure. "
  "It produces a composite Holon Completeness Score. "
  "The score determines deployment admissibility.",
 ]),
 ("3. Four Substrate Vectors",[
  "Vector 1: Knowledge Integrity Score.",
  "Measures consistency across knowledge domains within the graph. "
  "A model that holds constraints on general topics but fails on technical domains scores low. "
  "Computed from per-vector pass rate variance across evaluation probe categories. "
  "High variance means domain-dependent constraint persistence.",
  "Vector 2: Context Provenance Score.",
  "Measures consistency when context provenance changes. "
  "Different upstream sources may produce different constraint persistence. "
  "Maps directly to CSAS coordination findings. "
  "Computed from per-provider pass rate variance.",
  "Vector 3: Schema Constraint Persistence Score.",
  "This is the core MTCP BIS measurement. "
  "Whether the agent holds schema constraints under correction pressure. "
  "Computed as the normalised mean pass rate across all temperatures. "
  "This is the primary substrate measurement vector.",
  "Vector 4: Projection Consistency Score.",
  "Measures consistency across output contexts. "
  "Different evaluation datasets test different projection requirements. "
  "A model inconsistent across projections has format-dependent constraints. "
  "Computed from per-dataset pass rate variance.",
 ]),
 ("4. The Substrate Measurement Lemma",[
  "A holon without agent substrate measurement is incomplete.",
  "The four graphs specify what the agent should do. "
  "Agent Substrate Measurement proves whether it did. "
  "Specification without verification is assumption. "
  "Verification without specification is noise.",
  "Both are required for the holon to be load-bearing end to end. "
  "A load-bearing holon can support deployment decisions. "
  "A non-load-bearing holon is documentation only. "
  "The Substrate Measurement Lemma distinguishes these cases.",
 ]),
 ("5. Holon Completeness Score",[
  "The composite score combines all four vectors equally.",
  "Holon Completeness equals the mean of all four substrate scores. "
  "Each vector contributes 25 percent. "
  "All four must be measured for the score to be valid.",
  "Grading follows the standard MTCP scale. "
  "Grade A: 0.90 or above. "
  "Grade B: 0.80 to 0.89. "
  "Grade C: 0.70 to 0.79. "
  "Grade D: 0.60 to 0.69. "
  "Grade F: below 0.60.",
  "Deployment admissibility requires Grade C or above. "
  "A holon scoring below 0.70 is not deployment-admissible. "
  "The four graphs may be perfectly specified. "
  "The agent simply does not follow them reliably enough.",
 ]),
 ("6. First Run Results",[
  "Substrate measurement was computed for all 32 public models.",
  "12 models achieve Grade C or above (deployment-admissible). "
  "13 models achieve Grade D (marginal). "
  "10 models achieve Grade F (not admissible). "
  "No model achieves Grade A or B.",
  "Schema Persistence (Vector 3) is the strongest differentiator. "
  "Models with high BIS tend to have high Holon Completeness. "
  "But Knowledge Integrity and Projection Consistency reveal additional weakness. "
  "A model can have Grade B BIS and Grade D Holon Completeness.",
  "This demonstrates the value of multi-vector measurement. "
  "BIS alone provides an incomplete picture. "
  "The substrate vectors reveal domain-specific and format-specific weaknesses. "
  "These are invisible to single-score evaluation.",
 ]),
 ("7. Implications",[
  "Knowledge graph architects should require substrate measurement. "
  "A holon without it is unverified. "
  "No deployment decision should rely on specification alone.",
  "Agent selection for holon deployment should use Holon Completeness Score. "
  "Not all models are suitable for all holons. "
  "Domain-specific weakness (low Knowledge Integrity) matters. "
  "Format-specific weakness (low Projection Consistency) matters.",
  "The score enables principled agent-holon matching. "
  "Match agent strengths to holon requirements. "
  "Avoid deploying agents with known substrate weaknesses into matching contexts.",
 ]),
 ("8. Limitations",[
  "Vector computation uses proxy metrics from existing evaluation data. "
  "Direct measurement of each vector would require purpose-built probes. "
  "Current scores are derived, not directly measured.",
  "Equal weighting across four vectors is proposed without calibration. "
  "Some holons may weight Schema Persistence higher than Knowledge Integrity. "
  "Context-specific weighting is deferred to future work.",
 ]),
 ("9. Conclusion",[
  "Agent Substrate Measurement is the fifth element of the holon. "
  "Without it the holon is specified but unverified. "
  "With it the holon is load-bearing.",
  "Four vectors measure four dimensions of constraint reliability. "
  "The Holon Completeness Score provides a composite admissibility determination. "
  "32 models have been measured. "
  "The substrate layer is now empirically populated.",
 ]),
 ("10. References",[
  "Abby, A. (2026). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.",
  "Abby, A. (2026, Paper 28). Cross-System Admissibility Score. DOI: 10.17605/OSF.IO/DXGK5.",
  "Abby, A. (2026, Paper 47). Permanence Architecture. DOI: 10.17605/OSF.IO/DXGK5.",
  "Framework F4 -- BIS Definition. MTCP Research Programme.",
  "Framework F21 -- CSAS Definition. MTCP Research Programme.",
  "Framework F38 -- Substrate Measurement Definition. MTCP Research Programme.",
 ]),
]

def build(is_public=True):
    doc=Document();setup(doc)
    if is_public: hdr(doc,"MTCP V1.0 -- Agent Substrate Measurement")
    else: hdr(doc,"CONFIDENTIAL -- NDA REQUIRED");ftr(doc,"MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby.")
    h(doc,"MTCP Paper 48",1);h(doc,"Agent Substrate Measurement",2)
    p(doc,"")
    if not is_public: p(doc,"CONFIDENTIAL -- NDA REQUIRED")
    p(doc,"Author: A. Abby");p(doc,"Contact: admin@mtcp.live")
    p(doc,"DOI: 10.17605/OSF.IO/DXGK5");p(doc,"Date: May 2026");p(doc,"Version: 1.0")
    p(doc,"");p(doc,"MTCP Research Programme -- Paper 48");p(doc,"Framework F38")
    doc.add_page_break()
    for title,paras in C:
        h(doc,title,1)
        for t in paras: p(doc,t)
    return doc

if __name__=="__main__":
    print("Building Paper 48: Agent Substrate Measurement")
    build(True).save(os.path.join(DESKTOP,"Paper48_Agent_Substrate_Measurement_PUBLIC.docx"))
    build(False).save(os.path.join(DESKTOP,"Paper48_Agent_Substrate_Measurement_FULL.docx"))
    print("Done.")
