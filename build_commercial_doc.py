from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

DARK = RGBColor(0x1A, 0x1A, 0x2E)
PURPLE = RGBColor(0x6C, 0x5C, 0xE7)
BODY = RGBColor(0x2C, 0x2C, 0x3A)
GREY = RGBColor(0x66, 0x66, 0x77)
AMBER = RGBColor(0xC0, 0x7B, 0x00)

def setup_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Calibri'; s.font.size = Pt(10.5); s.font.color.rgb = BODY
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.15
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
    return doc

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = DARK if level == 1 else PURPLE; r.font.name = 'Calibri'

def para(doc, text, bold=False, italic=False, color=BODY, size=Pt(10.5), after=Pt(6)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = size; r.font.color.rgb = color; r.bold = bold; r.italic = italic; r.font.name = 'Calibri'
    p.paragraph_format.space_after = after

def bullet(doc, text, bold_prefix="", color=BODY):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = color; r.font.name = 'Calibri'
        r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = BODY; r.font.name = 'Calibri'
    else:
        r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = color; r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(3)

def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9); r.font.name = 'Calibri'; r.font.color.rgb = DARK
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ""
            r = c.paragraphs[0].add_run(str(val)); r.font.size = Pt(9); r.font.name = 'Calibri'; r.font.color.rgb = BODY
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width = Cm(w)

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '1', qn('w:color'): '6C5CE7'})
    pBdr.append(bottom); pPr.append(pBdr)

def sign(doc):
    doc.add_paragraph(); divider(doc)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MTCP Research Programme  |  A. Abby"); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK; r.font.name = 'Calibri'
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("mtcp.live  |  DOI: 10.17605/OSF.IO/DXGK5\n© 2026 A. Abby. All Rights Reserved."); r.font.size = Pt(9); r.font.color.rgb = GREY; r.font.name = 'Calibri'


doc = setup_doc()

# TITLE
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MTCP"); r.font.size = Pt(42); r.bold = True; r.font.color.rgb = DARK; r.font.name = 'Calibri'
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Commercial Value and Engagement Pricing"); r.font.size = Pt(16); r.font.color.rgb = PURPLE; r.font.name = 'Calibri'
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Confidential — Prepared for sirar by stc"); r.font.size = Pt(12); r.font.color.rgb = GREY; r.font.name = 'Calibri'
for _ in range(2): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("April 2026"); r.font.size = Pt(11); r.font.color.rgb = GREY; r.font.name = 'Calibri'
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A. Abby  |  MTCP Research Programme\nmtcp.live  |  DOI: 10.17605/OSF.IO/DXGK5"); r.font.size = Pt(9); r.font.color.rgb = GREY; r.font.name = 'Calibri'
doc.add_page_break()


# --- 1. TWO TRACKS ---
heading(doc, "1. Engagement Structure")
divider(doc)
para(doc,
    "The engagement has two distinct components. They are scoped, priced, and agreed separately.",
    bold=True
)

heading(doc, "Track A: Subject Matter Expert Role", level=2)
para(doc,
    "Ongoing advisory, interpretation, and service design work. Compensated via retainer or salary. "
    "Covers the expertise required to operate MTCP within sirar's service offering."
)
bullet(doc, "Define scope, model selection, and success criteria for each client engagement", bold_prefix="Evaluation design — ")
bullet(doc, "Explain grades, failure patterns, CPD, and Ve findings to clients", bold_prefix="Results interpretation — ")
bullet(doc, "Executive summaries, risk assessments, and deployment recommendations", bold_prefix="Report production — ")
bullet(doc, "Package MTCP evidence for sirar's client-facing service delivery", bold_prefix="Client framing — ")
bullet(doc, "EU AI Act, Saudi NCA/SDAIA, ISO 42001 compliance guidance", bold_prefix="Regulatory mapping — ")
bullet(doc, "Shape the AI Assurance service workflow and offering structure", bold_prefix="Service design — ")
bullet(doc, "Train sirar staff on interpreting and presenting MTCP outputs", bold_prefix="Team training — ")

heading(doc, "Track B: Platform Licensing", level=2)
para(doc,
    "Access to the MTCP platform — evaluation engine, proprietary probe sets, constraint detector, "
    "Evidence Pack generator, grading system, and research estate. This is a pre-existing commercial "
    "asset with independent value. It is not included in SME compensation.",
    bold=True
)


# --- 2. COMPARABLE RATES ---
heading(doc, "2. Comparable Market Rates")
divider(doc)

heading(doc, "SME Role Comparables", level=2)

table(doc,
    ["Comparable Role", "Gulf / Saudi (tax-free)", "Global Equivalent"],
    [
        ["AI Assurance Director — Big 4 (Riyadh)", "SAR 900K–1.2M/yr ($240K–$320K)", "$270K–$380K (US, pre-tax)"],
        ["Principal Cybersecurity Consultant (KSA)", "SAR 480K–720K/yr ($128K–$192K)", "$150K–$250K (US)"],
        ["AI Safety Researcher (Anthropic / OpenAI)", "No Gulf market exists", "$300K–$700K total comp"],
        ["Independent AI Consultant — Gulf day rate", "$2,000–$5,000/day", "$1,500–$4,000/day (US/UK)"],
        ["Big 4 client billing rate (Director level)", "$750–$1,200/day", "$500–$1,000/day (Manager)"],
    ],
    col_widths=[5.5, 4.5, 4.5]
)

doc.add_paragraph()
para(doc,
    "There is no established market for this specific role. Post-correction constraint persistence "
    "evaluation is a new field defined by the MTCP research programme. These are the closest comparables.",
    italic=True, color=GREY, size=Pt(9.5)
)

heading(doc, "Platform Licensing Comparables", level=2)

table(doc,
    ["Platform", "Category", "Annual Licence"],
    [
        ["Credo AI", "AI governance", "$100K–$300K/year"],
        ["Holistic AI", "AI assurance", "$75K–$250K/year"],
        ["Arthur AI", "Model monitoring", "$50K–$200K/year"],
        ["Robust Intelligence (Cisco)", "AI security / testing", "$100K–$400K/year"],
        ["ServiceNow GRC", "Compliance platform", "$150K–$500K/year"],
        ["Splunk (security tier)", "Security analytics", "$150K–$1M+/year"],
    ],
    col_widths=[4, 3, 3.5]
)

doc.add_paragraph()
para(doc,
    "None of these platforms provide post-correction constraint persistence evaluation. "
    "MTCP occupies a distinct category with no direct competitor.",
    italic=True, color=GREY, size=Pt(9.5)
)


# --- 3. PRICING ---
heading(doc, "3. Pricing")
divider(doc)

heading(doc, "SME Role", level=2)

table(doc,
    ["Engagement Shape", "Rate", "Basis"],
    [
        ["Retainer (2–3 days/week)", "$8,000–$15,000/month", "Below Big 4 Director rate. Above senior consultant. Niche specialism."],
        ["Full-time equivalent", "$180,000–$280,000/year", "Benchmarked to AI Assurance Director (Big 4, Riyadh). Tax-free."],
        ["Project-based day rate", "$2,500–$4,000/day", "Independent AI/cyber specialist (Gulf). Niche premium."],
    ],
    col_widths=[3.5, 3.5, 7.5]
)

doc.add_paragraph()
para(doc,
    "The appropriate rate depends on time commitment, exclusivity, and whether sirar requires "
    "dedicated or shared availability. The recommended starting point is a retainer covering a "
    "defined number of days per month, with overages billed at day rate."
)

heading(doc, "Platform Licensing", level=2)

table(doc,
    ["Model", "Description", "Indicative Range"],
    [
        ["Pilot", "3–5 models, 2–4 weeks, Evidence Packs + reports", "$15,000–$30,000 (one-off)"],
        ["Evaluation-as-a-Service", "Per-model evaluations delivered as Evidence Packs", "$5,000–$15,000 per model"],
        ["Annual Platform Access", "Controlled access to run evaluations directly", "$100,000–$250,000/year"],
        ["White-Label Licence", "sirar offers MTCP under own brand to clients", "$200,000–$500,000/year + revenue share"],
        ["Exclusivity (territory)", "Exclusive deployment rights for a defined region", "Scoped separately"],
    ],
    col_widths=[3.5, 5.5, 5.5]
)

doc.add_paragraph()
para(doc,
    "Final pricing depends on scope, territory, duration, and exclusivity terms."
)


# --- 4. BUILD VS ACCESS ---
heading(doc, "4. Build vs Access")
divider(doc)
para(doc,
    "The alternative to engaging MTCP is to build equivalent capability from scratch:"
)

table(doc,
    ["Component", "Estimated Cost", "Timeline"],
    [
        ["AI safety research team (2–3 researchers)", "$450K–$750K/year ongoing", "12+ months to recruit"],
        ["Platform development", "$200K–$500K", "6–12 months"],
        ["API costs (184,387+ evaluations, 13 providers)", "$50K–$100K+", "3–6 months"],
        ["Probe development (532 scenarios)", "Included in research team", "12+ months"],
        ["Academic publishing (5+ papers, DOI)", "Included in research team", "6–18 months per paper"],
        ["Independent validation", "Cannot be purchased", "Organic credibility over time"],
        ["Regulatory mapping (EU AI Act, Saudi)", "$100K–$200K", "3–6 months"],
    ],
    col_widths=[6, 3.5, 4]
)

doc.add_paragraph()
para(doc,
    "Estimated total: $1–2M+ and 18–24 months to reach the current state of the MTCP platform, "
    "dataset, and research estate.",
    bold=True, color=PURPLE
)
para(doc,
    "The EU AI Act compliance deadline is August 2026. Saudi AI governance frameworks are being "
    "written now. An 18–24 month build timeline means missing this market window."
)


# --- 5. WHAT MTCP BRINGS ---
heading(doc, "5. What MTCP Brings")
divider(doc)

table(doc,
    ["Asset", "Detail"],
    [
        ["184,387 evaluations", "Across 32 production models from 13 providers"],
        ["532 proprietary probes", "Across 5 evaluation vectors — core IP"],
        ["Live platform", "mtcp.live — in production"],
        ["Published research", "5 papers on OSF with DOI. 22 more completed."],
        ["Grading system", "A+ to F — no competitor has this"],
        ["Evidence Packs", "SHA-256 signed, EU AI Act Annex IV ready"],
        ["Independent validation", "6 practitioners in 48 hours. 1 formal citation. 1 formal spec."],
        ["Sovereign AI integration", "Already integrated into Gulf governance architectures"],
        ["Failure classification", "4-pattern taxonomy, independently validated"],
        ["Regulatory mapping", "EU AI Act article-level. Saudi transferable. ISO 42001 compatible."],
        ["Public dataset", "HuggingFace — independently verifiable"],
    ],
    col_widths=[3.5, 11]
)


# --- 6. RECOMMENDED STRUCTURE ---
heading(doc, "6. Recommended Engagement Structure")
divider(doc)

para(doc, "Phase 1: Pilot (immediate)", bold=True, color=PURPLE)
bullet(doc, "3–5 models, 2–4 weeks, full MTCP evaluation")
bullet(doc, "Deliverables: Evidence Packs, grades, executive summary, failure classification, deployment recommendation")
bullet(doc, "$15,000–$30,000")

para(doc, "Phase 2: SME Retainer (post-pilot)", bold=True, color=PURPLE)
bullet(doc, "Ongoing advisory: evaluation design, interpretation, client framing, regulatory mapping")
bullet(doc, "Retainer at $8,000–$15,000/month or equivalent full-time package")

para(doc, "Phase 3: Platform Licensing (post-pilot)", bold=True, color=PURPLE)
bullet(doc, "Licensing model selected based on pilot experience and service requirements")
bullet(doc, "Evaluation-as-a-Service, Platform Access, White-Label, or Strategic Integration")
bullet(doc, "Priced as a separate commercial agreement")

para(doc,
    "Each phase is a distinct agreement. Proceeding to Phase 2 or 3 depends on mutual alignment.",
    italic=True, color=GREY, size=Pt(10)
)

sign(doc)

out = os.path.expanduser("~/Desktop/MTCP_Commercial_Value_for_sirar.docx")
doc.save(out)
print(f"Saved: {out}")
