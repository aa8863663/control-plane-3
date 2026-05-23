"""
Build Paper 46: Runtime Behavioural Monitoring.
Framework F36: PRP Live Wrapper and session coherence.
Generates PUBLIC and FULL .docx versions.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DESKTOP = os.path.expanduser("~/Desktop")


def set_document_formatting(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.25
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return para


def add_header(doc, text):
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_footer(doc, text):
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


ABSTRACT = [
    "This paper defines Runtime Behavioural Monitoring for MTCP. "
    "Application-layer monitoring without a baseline is governance theatre. "
    "Any system can log responses and sign them cryptographically. "
    "That proves a record was kept, not that the model was constrained.",

    "MTCP closes this gap through two mechanisms. "
    "First: 183,924 evaluations establishing empirical baseline BIS scores. "
    "Second: the PRP wrapper computing rolling compliance during deployment. "
    "The baseline makes the runtime signal meaningful.",

    "Five constructs are defined. "
    "Rolling BIS measures live session compliance continuously. "
    "Drift Delta compares production to evaluation baseline. "
    "The Session Coherence Record proves session integrity cryptographically. "
    "The Runtime Admissibility Condition defines context-specific tolerances. "
    "The Runtime Monitoring Lemma establishes why both are required.",
]

INTRODUCTION = [
    "Production AI systems operate continuously after deployment. "
    "Evaluation occurs before deployment. "
    "Nothing in the pre-deployment framework measures post-deployment behaviour. "
    "This gap has remained open since Paper 1.",

    "Multiple vendors offer runtime monitoring products. "
    "They log API responses. They compute statistics. "
    "They sign records cryptographically. "
    "They prove that a record was created.",

    "They do not prove the model was behaving within governance boundaries. "
    "The distinction is fundamental. "
    "Recording what a model said is not the same as knowing it was constrained. "
    "The former is logging. The latter is governance.",

    "The difference is the baseline. "
    "Without a pre-deployment evaluation baseline, runtime signals are noise. "
    "A model producing 85 percent compliance in production means nothing. "
    "Is 85 percent above or below its evaluated capability. "
    "Without the baseline, the question has no answer.",

    "MTCP provides the baseline. "
    "183,924 evaluations across 32 models at four temperatures. "
    "Empirical BIS scores for each model under each constraint type. "
    "The PRP wrapper computes rolling compliance against this baseline. "
    "Drift becomes measurable. Degradation becomes detectable.",
]

GOVERNANCE_THEATRE = [
    "Application-layer monitoring without baseline comparison is widespread. "
    "It creates an appearance of governance without substance. "
    "Three properties make it governance theatre.",

    "First: no reference point for normal behaviour. "
    "The monitor logs what the model does. "
    "It cannot determine whether that behaviour is acceptable. "
    "Acceptability requires a reference. The reference is the baseline.",

    "Second: no distinction between constraint failure and normal variance. "
    "Models produce different outputs at different temperatures. "
    "A monitor without a baseline cannot distinguish normal variance from drift. "
    "Everything looks like signal. Nothing is calibrated.",

    "Third: cryptographic integrity without semantic integrity. "
    "A signed log proves the log was not tampered with. "
    "It does not prove the logged behaviour was governed. "
    "Tamper-evidence is necessary but not sufficient.",

    "MTCP resolves all three. "
    "The baseline provides the reference point. "
    "Temperature-specific evaluation data calibrates normal variance. "
    "The Session Coherence Record provides tamper-evidence. "
    "Together they constitute governance. Separately they constitute logging.",
]

ROLLING_BIS = [
    "Rolling BIS is the runtime compliance metric.",

    "Definition: Rolling BIS equals compliant turns divided by total turns. "
    "It is computed continuously within a live session. "
    "Each turn updates the rolling score.",

    "A turn is compliant if the response satisfies the registered constraint. "
    "Compliance checking uses the same logic as evaluation probing. "
    "The constraint is identified by constraint_id from the COS registry. "
    "The same constraint governs evaluation and production.",

    "Rolling BIS starts at 1.0 (first turn assumed compliant until checked). "
    "It decreases as non-compliant turns accumulate. "
    "It can recover if subsequent turns are compliant. "
    "The metric reflects current session state, not historical state.",
]

DRIFT_DELTA = [
    "Drift Delta measures production degradation against baseline.",

    "Definition: Drift Delta equals Rolling BIS minus Baseline BIS. "
    "Negative delta means production is worse than evaluation predicted. "
    "Positive delta means production is better (unexpected, investigate).",

    "The baseline is the mean BIS from pre-deployment evaluation. "
    "It represents the model's established constraint persistence capability. "
    "It was measured under controlled conditions across four temperatures. "
    "It is the empirical reference for production comparison.",

    "Drift thresholds vary by deployment context. "
    "Critical infrastructure: 2 percentage point tolerance. "
    "Financial services: 5 percentage points. "
    "Healthcare: 5 percentage points. "
    "Government services: 5 percentage points. "
    "General enterprise: 10 percentage points.",
]

SESSION_COHERENCE = [
    "The Session Coherence Record proves session integrity.",

    "Definition: SHA-256 hash of all response hashes in sequence. "
    "Each response is individually hashed. "
    "The concatenation of all hashes in order is then hashed. "
    "This produces one hash representing the entire session.",

    "The record proves three things. "
    "Completeness: no responses were deleted from the record. "
    "Order: responses were not reordered. "
    "Integrity: no response content was modified after capture.",

    "It does not prove constraint compliance. "
    "It proves the record on which compliance was computed is intact. "
    "Compliance determination requires the rolling BIS calculation. "
    "The coherence record proves the inputs to that calculation are valid.",

    "This is the application-layer answer to session coherence proof. "
    "It requires no model internals. No weight access. No hidden state. "
    "It uses only observable API responses. "
    "It is legitimate, verifiable, and independently reproducible.",
]

RUNTIME_ADMISSIBILITY = [
    "The Runtime Admissibility Condition defines when a session is acceptable.",

    "A session is runtime-admissible when rolling BIS stays within tolerance. "
    "Tolerance is set per deployment context. "
    "Exceeding tolerance triggers governance alerts.",

    "Warning alert: drift exceeds 10 percentage points below baseline. "
    "The session continues but stakeholders are notified. "
    "Increased monitoring frequency is recommended.",

    "Critical alert: drift exceeds 20 percentage points below baseline. "
    "The session is flagged for immediate review. "
    "A gate expiry event is logged. "
    "The model may require re-evaluation before continued deployment.",

    "Runtime admissibility is distinct from deployment admissibility. "
    "Deployment admissibility (DRA/Gate) is pre-deployment. "
    "Runtime admissibility is during deployment. "
    "Both must hold for full governance coverage.",
]

MONITORING_LEMMA = [
    "The Runtime Monitoring Lemma establishes formal requirements.",

    "Session coherence records alone cannot establish governance compliance. "
    "They establish record integrity only. "
    "Record integrity is necessary but not sufficient.",

    "Governance compliance requires two conditions simultaneously. "
    "Condition 1: record integrity (coherence record proves untampered data). "
    "Condition 2: baseline comparison (drift delta proves within tolerance).",

    "MTCP provides both conditions. "
    "The PRP wrapper generates coherence records (Condition 1). "
    "Baseline BIS from 183,924 evaluations enables comparison (Condition 2). "
    "No other runtime monitoring system provides both.",

    "Systems providing only Condition 1 produce governance theatre. "
    "They prove records exist. They cannot prove records indicate compliance. "
    "The distinction is between a tamper-evident log and a governance signal. "
    "MTCP provides the governance signal.",
]

INTEGRATION = [
    "Runtime monitoring integrates with the full MTCP stack.",

    "Gate Integration: the wrapper checks gate status before each session. "
    "A DENY decision blocks the session from starting. "
    "Gate enforcement extends from deployment decision to runtime.",

    "BEC Integration: session coherence hashes feed into the BEC chain. "
    "Runtime evidence is hash-chained alongside evaluation evidence. "
    "The entire evidence history is a single verifiable chain.",

    "Manifest Integration: two new fields added. "
    "runtime_monitor_enabled indicates active monitoring. "
    "last_session_coherence_hash links to most recent session proof.",

    "MCP Integration: get_runtime_status tool provides live compliance. "
    "Decision-makers can query current production state in one call. "
    "The tool returns rolling BIS, drift delta, and coherence status.",
]

LIMITATIONS = [
    "Compliance checking uses simplified constraint detection. "
    "The full output classification layer is NDA-protected. "
    "The wrapper uses scope-based heuristics. "
    "False positives and negatives are possible.",

    "Baseline BIS assumes model behaviour is stationary within TDS windows. "
    "If the model changes between evaluation and deployment, drift appears. "
    "The drift may reflect model change, not governance failure. "
    "TDS monitoring addresses this at the evaluation level.",

    "The wrapper monitors API responses, not model internals. "
    "It cannot detect internal state changes that do not manifest in output. "
    "This is a fundamental limitation of application-layer monitoring. "
    "It is accepted by design.",

    "Session coherence requires continuous capture. "
    "Any gap in response logging breaks the coherence hash. "
    "The wrapper must capture every response without exception. "
    "Network failures or timeouts create coherence gaps.",
]

CONCLUSION = [
    "Runtime monitoring without baseline is governance theatre. "
    "MTCP closes the gap with empirical baselines and live comparison. "
    "Rolling BIS measures production compliance continuously. "
    "Drift Delta detects degradation against evaluated capability.",

    "The Session Coherence Record proves session integrity. "
    "The Runtime Monitoring Lemma proves why both are needed. "
    "Record integrity alone is logging, not governance. "
    "Baseline comparison makes it governance.",

    "The PRP wrapper is the first runtime monitor with an evaluation baseline. "
    "183,924 evaluations provide the reference that makes monitoring meaningful. "
    "Without the evaluations, the wrapper would be another logging tool. "
    "With them, it is a governance instrument.",
]

REFERENCES = [
    "Abby, A. (2026a). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 30). Temporal Drift Monitoring. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 37). Blockchain Evidence Chain Integrity. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 38). Admissibility Gate. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 40). Constraint Object Specification. DOI: 10.17605/OSF.IO/DXGK5.",
    "Framework F2 -- Ve Metric. MTCP Research Programme.",
    "Framework F4 -- BIS Definition. MTCP Research Programme.",
    "Framework F15 -- PRP Design. MTCP Research Programme.",
    "Framework F28 -- BEC Definition. MTCP Research Programme.",
    "Framework F29 -- Gate Definition. MTCP Research Programme.",
    "Framework F36 -- Runtime Monitoring Definition. MTCP Research Programme.",
]


def build_paper(is_public=True):
    doc = Document()
    set_document_formatting(doc)

    if is_public:
        add_header(doc, "MTCP V1.0 -- Runtime Behavioural Monitoring")
    else:
        add_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
        add_footer(doc, "MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved.")

    add_heading_styled(doc, "MTCP Paper 46", level=1)
    add_heading_styled(doc, "Runtime Behavioural Monitoring", level=2)
    add_paragraph_justified(doc, "")
    if not is_public:
        add_paragraph_justified(doc, "CONFIDENTIAL -- NDA REQUIRED")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "MTCP Research Programme -- Paper 46")
    add_paragraph_justified(doc, "Framework F36")

    doc.add_page_break()

    add_heading_styled(doc, "1. Abstract", level=1)
    for para in ABSTRACT:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "2. Introduction", level=1)
    for para in INTRODUCTION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "3. Why Monitoring Without Baseline is Theatre", level=1)
    for para in GOVERNANCE_THEATRE:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "4. Rolling BIS", level=1)
    for para in ROLLING_BIS:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "5. Drift Delta", level=1)
    for para in DRIFT_DELTA:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "6. Session Coherence Record", level=1)
    for para in SESSION_COHERENCE:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "7. Runtime Admissibility Condition", level=1)
    for para in RUNTIME_ADMISSIBILITY:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "8. Runtime Monitoring Lemma", level=1)
    for para in MONITORING_LEMMA:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "9. Integration", level=1)
    for para in INTEGRATION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "10. Limitations", level=1)
    for para in LIMITATIONS:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "11. Conclusion", level=1)
    for para in CONCLUSION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "12. References", level=1)
    for ref in REFERENCES:
        add_paragraph_justified(doc, ref)

    return doc


if __name__ == "__main__":
    print("Building Paper 46: Runtime Behavioural Monitoring")
    print("=" * 60)

    doc_pub = build_paper(is_public=True)
    pub_path = os.path.join(DESKTOP, "Paper46_Runtime_Behavioural_Monitoring_PUBLIC.docx")
    doc_pub.save(pub_path)
    print(f"PUBLIC version saved: {pub_path}")

    doc_full = build_paper(is_public=False)
    full_path = os.path.join(DESKTOP, "Paper46_Runtime_Behavioural_Monitoring_FULL.docx")
    doc_full.save(full_path)
    print(f"FULL version saved: {full_path}")

    print("=" * 60)
    print("Done.")
