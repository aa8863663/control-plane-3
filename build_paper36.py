"""
Build script for Paper 36: Sovereign AI Deployment Standard.
Generates both PUBLIC and FULL (confidential) .docx versions.
This is a formal specification document with numbered requirements.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_document_formatting(doc):
    """Set page size, margins, and default font."""
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.25
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    """Add a heading with Arial font."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text):
    """Add a justified paragraph with correct formatting."""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para_format = para.paragraph_format
    para_format.line_spacing = 1.25
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return para


def add_requirement(doc, req_id, text):
    """Add a numbered requirement line."""
    full_text = f"{req_id}: {text}"
    para = doc.add_paragraph(full_text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para_format = para.paragraph_format
    para_format.line_spacing = 1.25
    para_format.left_indent = Inches(0.5)
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return para


def add_footer(doc, footer_text):
    """Add footer text to all sections."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = footer_text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_header(doc, header_text):
    """Add header text to all sections."""
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = header_text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


# ============================================================
# SECTION CONTENT
# ============================================================

SCOPE_AND_PURPOSE = [
    "This standard defines minimum requirements for AI model deployment "
    "in sovereign infrastructure. "
    "It unifies findings from MTCP Papers 26 through 35 into a single "
    "deployment specification. "
    "It provides measurable thresholds for deployment decisions.",

    "Target audience: national AI regulators, sovereign wealth funds, "
    "government CIOs, and procurement authorities. "
    "This standard is intended for decision-makers responsible for "
    "AI deployment in national infrastructure.",

    "This standard applies to any AI system deployed in government services, "
    "critical infrastructure, financial services, healthcare, or "
    "citizen-facing applications under national jurisdiction. "
    "It applies regardless of model provider or deployment architecture.",

    "The standard uses MTCP V1.0 as the measurement framework. "
    "MTCP provides reproducible, evidence-based evaluation of AI model "
    "constraint persistence. "
    "The dataset comprises 183,924 evaluations across 35 models and 14 providers. "
    "All requirements reference metrics defined within MTCP V1.0.",

    "Compliance with this standard is binary. "
    "A deployment either meets all applicable requirements or it does not. "
    "Partial compliance is not recognized. "
    "Requirements use SHALL for mandatory items and MAY for optional items.",
]

NORMATIVE_REFERENCES = [
    "MTCP V1.0 -- Multi-Turn Constraint Persistence Protocol "
    "(DOI: 10.17605/OSF.IO/DXGK5). "
    "Primary measurement framework for all requirements in this standard.",

    "EU AI Act (Regulation 2024/1689). "
    "European Union regulation on artificial intelligence. "
    "Referenced for high-risk AI system requirements.",

    "NDMO AI Governance Framework. "
    "National Data Management Office (Saudi Arabia) governance requirements. "
    "Referenced for data sovereignty and AI governance requirements.",

    "NCA Essential Cybersecurity Controls (ECC-1:2018). "
    "National Cybersecurity Authority (Saudi Arabia) security requirements. "
    "Referenced for adversarial resistance and security evaluation.",

    "SDAIA AI Ethics Principles. "
    "Saudi Data and Artificial Intelligence Authority ethical AI framework. "
    "Referenced for responsible AI deployment requirements.",

    "ISO/IEC 42001:2023. "
    "Information technology -- Artificial intelligence -- Management system. "
    "Referenced for AI management system certification.",

    "MAS Technology Risk Management Guidelines. "
    "Monetary Authority of Singapore technology governance framework. "
    "Referenced for financial services AI deployment.",

    "UK DSIT AI Regulation Framework. "
    "Department for Science, Innovation and Technology AI governance approach. "
    "Referenced for cross-border deployment compliance.",
]

DEFINITIONS = [
    "BIS (Behavioural Integrity Score): Mean constraint persistence rate across "
    "four temperature settings, expressed as a percentage and graded A through F.",

    "Ve (Violation Extinction): Number of consecutive corrections required before "
    "a model recovers a violated constraint. "
    "Ve >= 2 indicates persistent failure.",

    "CPD (Contamination-Persistence Delta): Difference in percentage points between "
    "benchmark contamination indicators and constraint persistence scores. "
    "Used to identify inflated evaluation results.",

    "CSAS (Cross-System Admissibility Score): Measure of whether constraint "
    "persistence holds when models operate in coordinated multi-model systems. "
    "Graded A through F.",

    "JRS (Jurisdiction Resolution Score): Proportion of coordination boundaries "
    "where jurisdiction over constraint enforcement is clearly assigned. "
    "Range 0 to 1.",

    "TDS (Temporal Drift Status): Classification of constraint persistence stability "
    "over time. "
    "Categories: Stable, Marginal, Significant, Critical.",

    "CCS (Constraint Conflict Score): Proportion of conflicting constraints "
    "that are resolved without constraint violation. "
    "Range 0 to 1.",

    "RES (Remediation Effectiveness Score): Proportion of identified constraint "
    "failures that can be addressed through operational controls. "
    "Range 0 to 1.",

    "ACPS (Adversarial Constraint Persistence Score): Constraint persistence rate "
    "under hostile adversarial attack conditions. "
    "Graded A through F. Always bounded above by BIS.",

    "CAF (Cascade Amplification Factor): Ratio of downstream constraint violations "
    "to upstream violations in a coordinated system. "
    "CAF > 1.0 indicates amplification risk.",

    "PRP (Programmatic Remediation Protocol): Operational control layer deployed "
    "around a model to enforce constraints that the model cannot maintain natively.",

    "Sigma-Forensics: Five-stage audit framework ensuring cryptographic integrity "
    "and chain-of-custody for all MTCP evaluation records.",
]

# Requirements grouped by section
GENERAL_REQUIREMENTS = [
    ("REQ-G-001",
     "All AI models deployed in sovereign infrastructure SHALL be evaluated "
     "under MTCP V1.0 before deployment authorization is granted."),
    ("REQ-G-002",
     "Evaluation SHALL cover all 5 MTCP vectors (BIS, Ve, CPD, TDS, CSAS) "
     "at minimum 2 temperature settings (T=0.0 and T=0.8)."),
    ("REQ-G-003",
     "Results SHALL be recorded with SHA-256 hash integrity verification "
     "under the Sigma-Forensics audit framework."),
    ("REQ-G-004",
     "MTCP grades SHALL be re-evaluated at intervals not exceeding 90 days "
     "to monitor for temporal drift (TDS)."),
    ("REQ-G-005",
     "All evaluation records SHALL be retained for a minimum of 5 years "
     "from the date of evaluation."),
    ("REQ-G-006",
     "Evaluation results SHALL identify the specific model version tested "
     "using the provider's API version string or model identifier."),
    ("REQ-G-007",
     "Evaluation SHALL be conducted by personnel with documented training "
     "in MTCP V1.0 evaluation methodology."),
]

DEPLOYMENT_CONTEXT_REQUIREMENTS = [
    ("REQ-D-001",
     "Critical national infrastructure: minimum Grade A (BIS >= 90%), "
     "CSAS Grade A, and ACPS Grade B or higher."),
    ("REQ-D-002",
     "Financial services: minimum Grade B (BIS >= 80%), "
     "CSAS Grade B, and CCS >= 0.8."),
    ("REQ-D-003",
     "Healthcare: minimum Grade B (BIS >= 80%), "
     "ACPS Grade B, and CPD < 30 percentage points."),
    ("REQ-D-004",
     "Government services (Gulf jurisdictions): minimum Grade B (BIS >= 80%), "
     "and Arabic LANG pass rate >= 90%."),
    ("REQ-D-005",
     "Public-facing applications: minimum Grade C (BIS >= 70%), "
     "TDS status Stable or Marginal, and valid evaluation within 90 days."),
    ("REQ-D-006",
     "Defence and intelligence: minimum Grade A (BIS >= 90%), "
     "ACPS Grade A, CSAS Grade A, and TDS Stable."),
    ("REQ-D-007",
     "Education: minimum Grade C (BIS >= 70%) "
     "with CPD < 20 percentage points to ensure no benchmark contamination."),
]

CROSS_SYSTEM_REQUIREMENTS = [
    ("REQ-CS-001",
     "Multi-model systems SHALL be evaluated under CSAS before any "
     "coordinated deployment is authorized."),
    ("REQ-CS-002",
     "CSAS Grade B is the minimum acceptable score for any coordinated "
     "deployment involving two or more AI models."),
    ("REQ-CS-003",
     "JRS SHALL be explicitly assigned for every coordination boundary. "
     "JRS >= 0.75 is the minimum acceptable threshold."),
    ("REQ-CS-004",
     "Coordination boundaries SHALL be re-evaluated when any component "
     "model in the system is updated by its provider."),
    ("REQ-CS-005",
     "Cascade amplification (CAF > 1.0) SHALL trigger immediate system "
     "review and suspension of the affected coordination path."),
    ("REQ-CS-006",
     "Each model in a coordinated system SHALL meet the individual "
     "deployment requirements for its deployment context independently."),
    ("REQ-CS-007",
     "System-level CSAS evaluation SHALL test constraint persistence "
     "across at least 3 coordination boundaries per model pair."),
]

TEMPORAL_REQUIREMENTS = [
    ("REQ-T-001",
     "MTCP grades expire after 90 days without re-evaluation. "
     "Expired grades SHALL NOT be used for deployment authorization."),
    ("REQ-T-002",
     "TDS Critical drift (greater than 10 percentage points) SHALL trigger "
     "immediate suspension of deployment pending re-evaluation."),
    ("REQ-T-003",
     "Model updates by providers SHALL trigger mandatory re-evaluation "
     "within 30 days of the update notification."),
    ("REQ-T-004",
     "TDS monitoring SHALL run continuously for critical infrastructure "
     "deployments with automated alerting on status change."),
    ("REQ-T-005",
     "Marginal TDS status SHALL trigger increased monitoring frequency "
     "to every 30 days instead of the standard 90-day cycle."),
    ("REQ-T-006",
     "Significant TDS status SHALL trigger formal review within 14 days "
     "and deployment restriction pending that review."),
]

ADVERSARIAL_REQUIREMENTS = [
    ("REQ-A-001",
     "High-consequence deployments (REQ-D-001, REQ-D-006) SHALL undergo "
     "full ACPS evaluation covering all four attack types."),
    ("REQ-A-002",
     "ACPS Grade B is the minimum for government-facing systems. "
     "ACPS Grade A is required for critical infrastructure and defence."),
    ("REQ-A-003",
     "Context flooding resistance SHALL be specifically tested for any "
     "system that processes external user input or documents."),
    ("REQ-A-004",
     "Models failing ACPS evaluation MAY still be deployed with approved "
     "PRP mitigation if the operational risk assessment justifies deployment."),
    ("REQ-A-005",
     "ACPS evaluation SHALL use pre-scripted attack sequences for "
     "reproducibility under Sigma-Forensics audit requirements."),
    ("REQ-A-006",
     "Adversarial failure mode distribution SHALL be documented. "
     "Injection Capture, Authority Acceptance, and Context Displacement "
     "rates SHALL be reported separately."),
]

LANGUAGE_REQUIREMENTS = [
    ("REQ-L-001",
     "Arabic constraint persistence SHALL be evaluated using LANG_Arabic "
     "probes with a minimum of 50 probes per evaluation."),
    ("REQ-L-002",
     "Arabic LANG pass rate SHALL be >= 90% for citizen-facing systems "
     "in Gulf jurisdictions."),
    ("REQ-L-003",
     "Arabic LANG pass rate SHALL be >= 95% for government-internal systems "
     "in Gulf jurisdictions."),
    ("REQ-L-004",
     "Models SHALL be evaluated for first-turn compliance rate separately "
     "from overall multi-turn pass rate. "
     "Both rates SHALL be reported."),
    ("REQ-L-005",
     "Technical vocabulary retention SHALL be tested for domain-specific "
     "terminology relevant to the deployment jurisdiction."),
    ("REQ-L-006",
     "Models with Arabic LANG below 85% SHALL NOT be deployed in any "
     "Arabic-language government context without PRP mitigation."),
]

EVIDENCE_REQUIREMENTS = [
    ("REQ-E-001",
     "Board-level evidence pack SHALL contain: model grade, regulatory "
     "compliance matrix, CSAS scores, TDS validity, ACPS scores, "
     "and deployment recommendation with risk assessment."),
    ("REQ-E-002",
     "All evidence SHALL be cryptographically hashed using SHA-256. "
     "Hash values SHALL be recorded in the Sigma-Forensics audit trail."),
    ("REQ-E-003",
     "Evidence packs SHALL identify the specific model version evaluated "
     "using the provider API version string or checkpoint identifier."),
    ("REQ-E-004",
     "Evidence packs SHALL state a valid-until date based on the TDS "
     "monitoring window. Default expiry is 90 days from evaluation."),
    ("REQ-E-005",
     "Evidence packs for multi-jurisdiction deployment SHALL map compliance "
     "against all relevant jurisdictions simultaneously."),
    ("REQ-E-006",
     "Evidence packs SHALL include a clear DEPLOY or DO NOT DEPLOY "
     "recommendation with rationale referencing specific requirements."),
    ("REQ-E-007",
     "Evidence packs SHALL be version-controlled. "
     "Any update to the evidence pack SHALL generate a new version "
     "with a new SHA-256 hash."),
]

REMEDIATION_REQUIREMENTS = [
    ("REQ-R-001",
     "Models scoring below deployment threshold SHALL have a documented "
     "remediation path identifying specific gaps and corrective actions."),
    ("REQ-R-002",
     "RES assessment SHALL identify whether the gap is addressable through "
     "operational controls or requires provider-level intervention."),
    ("REQ-R-003",
     "Architectural failures (temperature-invariant constraint violations) "
     "SHALL be documented as requiring provider-level remediation. "
     "These cannot be addressed through operational controls alone."),
    ("REQ-R-004",
     "PRP deployment SHALL be considered for architectural failures where "
     "model substitution is not feasible due to capability requirements."),
    ("REQ-R-005",
     "Remediation timelines SHALL be documented. "
     "Re-evaluation SHALL occur within 30 days of remediation completion."),
    ("REQ-R-006",
     "Models that fail re-evaluation after remediation SHALL be escalated "
     "to the deployment authority for substitution review."),
]

COMPLIANCE_DECLARATION_TEXT = [
    "This section defines the compliance declaration that deploying "
    "organisations must complete before deployment authorization.",

    "The deploying organisation SHALL attest to the following statements. "
    "All statements must be true at the time of signing. "
    "False attestation invalidates the deployment authorization.",

    "Statement 1: The AI model identified in this declaration has been "
    "evaluated under MTCP V1.0 by qualified personnel.",

    "Statement 2: The evaluation results meet or exceed all requirements "
    "applicable to the intended deployment context as defined in this standard.",

    "Statement 3: The evaluation evidence pack is complete, cryptographically "
    "hashed, and available for audit upon request.",

    "Statement 4: A re-evaluation schedule has been established "
    "in accordance with REQ-T-001 through REQ-T-006.",

    "Statement 5: The deploying organisation accepts responsibility for "
    "maintaining compliance throughout the deployment lifecycle.",

    "Statement 6: Any material change to the deployment context, model version, "
    "or system architecture will trigger re-evaluation before continued operation.",

    "The compliance declaration SHALL be signed by the Chief Information Officer "
    "or equivalent authority. "
    "The declaration SHALL be countersigned by the information security officer. "
    "Both signatures are required for deployment authorization.",

    "The compliance declaration is valid for the duration stated in the "
    "evidence pack valid-until date. "
    "Renewal requires fresh evaluation and a new declaration.",
]

# Annex content
ANNEX_A_GRADING = [
    "Grade A: BIS 90% to 100%. Constraint persistence is robust across all "
    "temperature settings. Suitable for critical infrastructure deployment.",

    "Grade B: BIS 80% to 89%. Constraint persistence is strong with minor "
    "degradation at high temperatures. Suitable for government and financial services.",

    "Grade C: BIS 70% to 79%. Constraint persistence is adequate with moderate "
    "degradation under pressure. Suitable for public-facing applications with monitoring.",

    "Grade D: BIS 60% to 69%. Constraint persistence is marginal. "
    "Not suitable for sovereign deployment without PRP mitigation.",

    "Grade F: BIS below 60%. Constraint persistence is insufficient. "
    "Not suitable for sovereign deployment. Requires provider-level remediation.",

    "The same grading scale applies to CSAS and ACPS independently. "
    "Each metric is graded separately. "
    "A model may have different grades across different metrics.",
]

ANNEX_B_REGULATORY = [
    "EU AI Act (2024/1689): High-risk systems require BIS Grade B minimum. "
    "Conformity assessment maps to MTCP evidence pack (REQ-E-001). "
    "Post-market monitoring maps to TDS (REQ-T-001 through REQ-T-006).",

    "NDMO AI Governance Framework: Data sovereignty requirements map to "
    "evaluation record retention (REQ-G-005). "
    "Transparency requirements map to evidence pack disclosure (REQ-E-001).",

    "NCA Essential Cybersecurity Controls: Security evaluation maps to "
    "ACPS (REQ-A-001 through REQ-A-006). "
    "ACPS Grade A required for government-facing systems.",

    "SDAIA AI Ethics Principles: Responsible AI deployment maps to "
    "remediation requirements (REQ-R-001 through REQ-R-006). "
    "Accountability maps to compliance declaration (Section 12).",

    "MAS Technology Risk Management: Model risk management maps to "
    "BIS and TDS evaluation. "
    "Third-party risk maps to evidence pack requirements (REQ-E-003).",

    "UK DSIT AI Regulation: Proportionality principle maps to deployment "
    "context tiers (REQ-D-001 through REQ-D-007). "
    "Cross-border deployment maps to multi-jurisdiction evidence (REQ-E-005).",
]

ANNEX_C_TEMPLATE = [
    "Evidence Pack Template Structure:",

    "Section 1 -- Executive Summary: Deployment recommendation (DEPLOY or "
    "DO NOT DEPLOY), model identification, evaluation date, valid-until date.",

    "Section 2 -- Model Identification: Provider name, model name, API version "
    "string, deployment architecture, intended use case.",

    "Section 3 -- MTCP Results: BIS grade and percentage, Ve summary, "
    "CPD assessment, temperature-specific scores.",

    "Section 4 -- Cross-System Assessment: CSAS grade (if applicable), "
    "JRS scores per coordination boundary, CAF assessment.",

    "Section 5 -- Temporal Validity: TDS status, last evaluation date, "
    "drift trajectory, next scheduled evaluation.",

    "Section 6 -- Adversarial Assessment: ACPS grade (if applicable), "
    "attack type resistance breakdown, failure mode distribution.",

    "Section 7 -- Language Assessment: LANG pass rates by language "
    "(if applicable), first-turn compliance, vocabulary retention.",

    "Section 8 -- Regulatory Compliance Matrix: Requirement-by-requirement "
    "compliance mapping against all applicable jurisdictions.",

    "Section 9 -- Risk Assessment: Identified gaps, severity classification, "
    "remediation status, residual risk statement.",

    "Section 10 -- Audit Trail: SHA-256 hashes for all evaluation records, "
    "Sigma-Forensics stage completion, chain-of-custody log.",

    "Section 11 -- Recommendation and Conditions: Deployment conditions, "
    "monitoring requirements, escalation triggers, review schedule.",
]

# Full version additional content for language section
LANGUAGE_REQUIREMENTS_FULL_CONTEXT = [
    "The Arabic language evaluation is based on data from MTCP Papers 26 "
    "and 29. "
    "Evaluation covers 11 Arabic dialect variants. "
    "The dataset includes models from Humain (ALLaM), R-AGAM, and KFUPM "
    "alongside international providers.",

    "Nova models (Amazon Bedrock) demonstrate 90 to 100% Arabic constraint "
    "persistence. "
    "Mistral models demonstrate 83.5% Arabic constraint persistence. "
    "These empirical results inform the threshold requirements.",

    "Vision 2030 terminology, NCA security vocabulary, and NDMO governance "
    "terminology form the required technical vocabulary set for Saudi "
    "government deployments. "
    "SDAIA terminology is included for AI-specific government applications.",
]


def build_public():
    """Build the PUBLIC version of Paper 36."""
    doc = Document()
    set_document_formatting(doc)
    add_header(doc, "MTCP V1.0 -- Sovereign AI Deployment Standard")

    # Title page
    add_heading_styled(doc, "MTCP Paper 36", level=1)
    add_heading_styled(doc, "Sovereign AI Deployment Standard", level=2)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "Status: NOT for OSF (contains jurisdiction-specific requirements)")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(
        doc,
        "This document is a formal specification. "
        "It defines minimum deployment requirements for AI systems "
        "operating within sovereign national infrastructure. "
        "Requirements are numbered and use SHALL for mandatory items."
    )

    doc.add_page_break()

    # 1. Scope and Purpose
    add_heading_styled(doc, "1. Scope and Purpose", level=1)
    for para_text in SCOPE_AND_PURPOSE:
        add_paragraph_justified(doc, para_text)

    # 2. Normative References
    add_heading_styled(doc, "2. Normative References", level=1)
    for para_text in NORMATIVE_REFERENCES:
        add_paragraph_justified(doc, para_text)

    # 3. Definitions
    add_heading_styled(doc, "3. Definitions", level=1)
    add_paragraph_justified(
        doc,
        "The following terms are defined for use throughout this standard. "
        "All metrics originate from MTCP V1.0."
    )
    for defn in DEFINITIONS:
        add_paragraph_justified(doc, defn)

    # 4. General Requirements
    add_heading_styled(doc, "4. General Requirements (REQ-G)", level=1)
    add_paragraph_justified(
        doc,
        "These requirements apply to all AI deployments within sovereign "
        "infrastructure regardless of deployment context."
    )
    for req_id, req_text in GENERAL_REQUIREMENTS:
        add_requirement(doc, req_id, req_text)

    # 5. Deployment Context Requirements
    add_heading_styled(doc, "5. Deployment Context Requirements (REQ-D)", level=1)
    add_paragraph_justified(
        doc,
        "These requirements define minimum thresholds for specific deployment "
        "contexts. "
        "The deploying organisation SHALL identify the applicable context "
        "before evaluation begins."
    )
    for req_id, req_text in DEPLOYMENT_CONTEXT_REQUIREMENTS:
        add_requirement(doc, req_id, req_text)

    # 6. Cross-System Requirements
    add_heading_styled(doc, "6. Cross-System Requirements (REQ-CS)", level=1)
    add_paragraph_justified(
        doc,
        "These requirements apply when two or more AI models operate in "
        "a coordinated system within sovereign infrastructure."
    )
    for req_id, req_text in CROSS_SYSTEM_REQUIREMENTS:
        add_requirement(doc, req_id, req_text)

    # 7. Temporal Requirements
    add_heading_styled(doc, "7. Temporal Requirements (REQ-T)", level=1)
    add_paragraph_justified(
        doc,
        "These requirements address the time-bound validity of MTCP "
        "evaluations and ongoing monitoring obligations."
    )
    for req_id, req_text in TEMPORAL_REQUIREMENTS:
        add_requirement(doc, req_id, req_text)

    # 8. Adversarial Requirements
    add_heading_styled(doc, "8. Adversarial Requirements (REQ-A)", level=1)
    add_paragraph_justified(
        doc,
        "These requirements address evaluation of constraint persistence "
        "under hostile adversarial conditions."
    )
    for req_id, req_text in ADVERSARIAL_REQUIREMENTS:
        add_requirement(doc, req_id, req_text)

    # 9. Language Requirements
    add_heading_styled(doc, "9. Language Requirements (REQ-L)", level=1)
    add_paragraph_justified(
        doc,
        "These requirements apply to AI deployments in Gulf sovereign "
        "jurisdictions where Arabic language support is mandatory."
    )
    for req_id, req_text in LANGUAGE_REQUIREMENTS:
        add_requirement(doc, req_id, req_text)

    # 10. Evidence Package Requirements
    add_heading_styled(doc, "10. Evidence Package Requirements (REQ-E)", level=1)
    add_paragraph_justified(
        doc,
        "These requirements define the documentation that must accompany "
        "any deployment authorization request."
    )
    for req_id, req_text in EVIDENCE_REQUIREMENTS:
        add_requirement(doc, req_id, req_text)

    # 11. Remediation Requirements
    add_heading_styled(doc, "11. Remediation Requirements (REQ-R)", level=1)
    add_paragraph_justified(
        doc,
        "These requirements define the process when a model fails to meet "
        "deployment thresholds."
    )
    for req_id, req_text in REMEDIATION_REQUIREMENTS:
        add_requirement(doc, req_id, req_text)

    # 12. Compliance Declaration
    add_heading_styled(doc, "12. Compliance Declaration", level=1)
    for para_text in COMPLIANCE_DECLARATION_TEXT:
        add_paragraph_justified(doc, para_text)

    # 13. Annexes
    add_heading_styled(doc, "13. Annexes", level=1)

    add_heading_styled(doc, "Annex A: MTCP Grading Scale", level=2)
    for para_text in ANNEX_A_GRADING:
        add_paragraph_justified(doc, para_text)

    add_heading_styled(doc, "Annex B: Regulatory Compliance Matrix", level=2)
    add_paragraph_justified(
        doc,
        "This annex maps MTCP requirements to regulatory frameworks "
        "across applicable jurisdictions."
    )
    for para_text in ANNEX_B_REGULATORY:
        add_paragraph_justified(doc, para_text)

    add_heading_styled(doc, "Annex C: Evidence Pack Template Structure", level=2)
    for para_text in ANNEX_C_TEMPLATE:
        add_paragraph_justified(doc, para_text)

    # References
    add_heading_styled(doc, "14. References", level=1)
    refs = [
        "Abby, A. (2026). Multi-Turn Constraint Persistence (MTCP) V1.0. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 26). Arabic Language Constraint Persistence. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 27). Benchmark Contamination-Persistence Delta. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 28). Temporal Drift in Constraint Persistence. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 29). Cross-System Admissibility. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 30). Jurisdiction Resolution. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 31). Constraint Conflict Resolution. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 32). Remediation Effectiveness. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 33). Cascade Amplification in Multi-Model Systems. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 34). Regulatory Compliance Mapping. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 35). Adversarial Constraint Persistence. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "EU AI Act. Regulation (EU) 2024/1689 of the European Parliament "
        "and of the Council. Official Journal of the European Union.",

        "ISO/IEC 42001:2023. Information technology -- Artificial intelligence "
        "-- Management system. International Organization for Standardization.",

        "NCA. Essential Cybersecurity Controls (ECC-1:2018). "
        "National Cybersecurity Authority, Kingdom of Saudi Arabia.",

        "NDMO. AI Governance Framework. "
        "National Data Management Office, Kingdom of Saudi Arabia.",

        "SDAIA. AI Ethics Principles. "
        "Saudi Data and Artificial Intelligence Authority.",

        "MAS. Technology Risk Management Guidelines. "
        "Monetary Authority of Singapore.",
    ]
    for ref in refs:
        add_paragraph_justified(doc, ref)

    output_path = os.path.expanduser(
        "~/Desktop/Paper36_Sovereign_AI_Deployment_Standard_PUBLIC.docx"
    )
    doc.save(output_path)
    print(f"PUBLIC version saved: {output_path}")
    return output_path


def build_full():
    """Build the FULL (confidential) version of Paper 36."""
    doc = Document()
    set_document_formatting(doc)
    add_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
    add_footer(
        doc,
        "MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved."
    )

    # Title page
    add_heading_styled(doc, "MTCP Paper 36", level=1)
    add_heading_styled(doc, "Sovereign AI Deployment Standard", level=2)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "Status: CONFIDENTIAL -- NDA REQUIRED")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(
        doc,
        "This document is a formal specification. "
        "It defines minimum deployment requirements for AI systems "
        "operating within sovereign national infrastructure. "
        "Requirements are numbered and use SHALL for mandatory items. "
        "This version contains jurisdiction-specific requirements and "
        "names specific entities including Humain, ALLaM, R-AGAM, KFUPM, "
        "and Saudi regulatory bodies."
    )

    doc.add_page_break()

    # 1. Scope and Purpose
    add_heading_styled(doc, "1. Scope and Purpose", level=1)
    for para_text in SCOPE_AND_PURPOSE:
        add_paragraph_justified(doc, para_text)
    # Additional context for FULL version
    add_paragraph_justified(
        doc,
        "This standard has been developed with specific reference to "
        "Vision 2030 AI deployment objectives. "
        "The Kingdom of Saudi Arabia is deploying AI across government, "
        "healthcare, financial services, and citizen services. "
        "NCA, NDMO, and SDAIA requirements are directly addressed."
    )

    # 2. Normative References
    add_heading_styled(doc, "2. Normative References", level=1)
    for para_text in NORMATIVE_REFERENCES:
        add_paragraph_justified(doc, para_text)

    # 3. Definitions
    add_heading_styled(doc, "3. Definitions", level=1)
    add_paragraph_justified(
        doc,
        "The following terms are defined for use throughout this standard. "
        "All metrics originate from MTCP V1.0."
    )
    for defn in DEFINITIONS:
        add_paragraph_justified(doc, defn)

    # 4. General Requirements
    add_heading_styled(doc, "4. General Requirements (REQ-G)", level=1)
    add_paragraph_justified(
        doc,
        "These requirements apply to all AI deployments within sovereign "
        "infrastructure regardless of deployment context."
    )
    for req_id, req_text in GENERAL_REQUIREMENTS:
        add_requirement(doc, req_id, req_text)

    # 5. Deployment Context Requirements
    add_heading_styled(doc, "5. Deployment Context Requirements (REQ-D)", level=1)
    add_paragraph_justified(
        doc,
        "These requirements define minimum thresholds for specific deployment "
        "contexts. "
        "The deploying organisation SHALL identify the applicable context "
        "before evaluation begins."
    )
    for req_id, req_text in DEPLOYMENT_CONTEXT_REQUIREMENTS:
        add_requirement(doc, req_id, req_text)
    # Additional FULL context
    add_paragraph_justified(
        doc,
        "For Vision 2030 deployments, REQ-D-004 is the primary requirement. "
        "Humain (ALLaM) and R-AGAM models have been evaluated for Arabic "
        "constraint persistence under MTCP. "
        "KFUPM research models are evaluated under the same framework. "
        "International providers (Nova, Mistral, GPT, Claude) are evaluated "
        "against the same thresholds."
    )

    # 6. Cross-System Requirements
    add_heading_styled(doc, "6. Cross-System Requirements (REQ-CS)", level=1)
    add_paragraph_justified(
        doc,
        "These requirements apply when two or more AI models operate in "
        "a coordinated system within sovereign infrastructure."
    )
    for req_id, req_text in CROSS_SYSTEM_REQUIREMENTS:
        add_requirement(doc, req_id, req_text)

    # 7. Temporal Requirements
    add_heading_styled(doc, "7. Temporal Requirements (REQ-T)", level=1)
    add_paragraph_justified(
        doc,
        "These requirements address the time-bound validity of MTCP "
        "evaluations and ongoing monitoring obligations."
    )
    for req_id, req_text in TEMPORAL_REQUIREMENTS:
        add_requirement(doc, req_id, req_text)

    # 8. Adversarial Requirements
    add_heading_styled(doc, "8. Adversarial Requirements (REQ-A)", level=1)
    add_paragraph_justified(
        doc,
        "These requirements address evaluation of constraint persistence "
        "under hostile adversarial conditions. "
        "NCA Essential Cybersecurity Controls require adversarial evaluation "
        "for government-facing AI systems."
    )
    for req_id, req_text in ADVERSARIAL_REQUIREMENTS:
        add_requirement(doc, req_id, req_text)

    # 9. Language Requirements
    add_heading_styled(doc, "9. Language Requirements (REQ-L)", level=1)
    add_paragraph_justified(
        doc,
        "These requirements apply to AI deployments in Gulf sovereign "
        "jurisdictions where Arabic language support is mandatory. "
        "They are directly relevant to Vision 2030 AI objectives."
    )
    for req_id, req_text in LANGUAGE_REQUIREMENTS:
        add_requirement(doc, req_id, req_text)
    # Additional FULL context
    for para_text in LANGUAGE_REQUIREMENTS_FULL_CONTEXT:
        add_paragraph_justified(doc, para_text)

    # 10. Evidence Package Requirements
    add_heading_styled(doc, "10. Evidence Package Requirements (REQ-E)", level=1)
    add_paragraph_justified(
        doc,
        "These requirements define the documentation that must accompany "
        "any deployment authorization request."
    )
    for req_id, req_text in EVIDENCE_REQUIREMENTS:
        add_requirement(doc, req_id, req_text)

    # 11. Remediation Requirements
    add_heading_styled(doc, "11. Remediation Requirements (REQ-R)", level=1)
    add_paragraph_justified(
        doc,
        "These requirements define the process when a model fails to meet "
        "deployment thresholds."
    )
    for req_id, req_text in REMEDIATION_REQUIREMENTS:
        add_requirement(doc, req_id, req_text)

    # 12. Compliance Declaration
    add_heading_styled(doc, "12. Compliance Declaration", level=1)
    for para_text in COMPLIANCE_DECLARATION_TEXT:
        add_paragraph_justified(doc, para_text)
    # Additional FULL context
    add_paragraph_justified(
        doc,
        "For Saudi government deployments, the compliance declaration "
        "SHALL additionally reference NCA ECC compliance. "
        "NDMO data governance attestation is required for systems "
        "processing citizen data. "
        "SDAIA ethical AI attestation is required for all government AI."
    )

    # 13. Annexes
    add_heading_styled(doc, "13. Annexes", level=1)

    add_heading_styled(doc, "Annex A: MTCP Grading Scale", level=2)
    for para_text in ANNEX_A_GRADING:
        add_paragraph_justified(doc, para_text)

    add_heading_styled(doc, "Annex B: Regulatory Compliance Matrix", level=2)
    add_paragraph_justified(
        doc,
        "This annex maps MTCP requirements to regulatory frameworks "
        "across applicable jurisdictions. "
        "Mapping is based on Paper 34 (Regulatory Compliance Mapping)."
    )
    for para_text in ANNEX_B_REGULATORY:
        add_paragraph_justified(doc, para_text)
    # Additional FULL regulatory context
    add_paragraph_justified(
        doc,
        "For multi-jurisdiction deployments (for example, Saudi Arabia "
        "and UAE simultaneously), all applicable requirements from both "
        "jurisdictions apply concurrently. "
        "The most stringent requirement in any overlapping area takes precedence."
    )

    add_heading_styled(doc, "Annex C: Evidence Pack Template Structure", level=2)
    for para_text in ANNEX_C_TEMPLATE:
        add_paragraph_justified(doc, para_text)

    # References
    add_heading_styled(doc, "14. References", level=1)
    refs = [
        "Abby, A. (2026). Multi-Turn Constraint Persistence (MTCP) V1.0. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 26). Arabic Language Constraint Persistence. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 27). Benchmark Contamination-Persistence Delta. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 28). Temporal Drift in Constraint Persistence. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 29). Cross-System Admissibility. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 30). Jurisdiction Resolution. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 31). Constraint Conflict Resolution. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 32). Remediation Effectiveness. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 33). Cascade Amplification in Multi-Model Systems. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 34). Regulatory Compliance Mapping. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "Abby, A. (2026, Paper 35). Adversarial Constraint Persistence. "
        "DOI: 10.17605/OSF.IO/DXGK5.",

        "EU AI Act. Regulation (EU) 2024/1689 of the European Parliament "
        "and of the Council. Official Journal of the European Union.",

        "ISO/IEC 42001:2023. Information technology -- Artificial intelligence "
        "-- Management system. International Organization for Standardization.",

        "NCA. Essential Cybersecurity Controls (ECC-1:2018). "
        "National Cybersecurity Authority, Kingdom of Saudi Arabia.",

        "NDMO. AI Governance Framework. "
        "National Data Management Office, Kingdom of Saudi Arabia.",

        "SDAIA. AI Ethics Principles. "
        "Saudi Data and Artificial Intelligence Authority.",

        "MAS. Technology Risk Management Guidelines. "
        "Monetary Authority of Singapore.",

        "Humain. ALLaM Large Language Model. "
        "Kingdom of Saudi Arabia sovereign AI programme.",

        "KFUPM. AI Research Programme. "
        "King Fahd University of Petroleum and Minerals.",
    ]
    for ref in refs:
        add_paragraph_justified(doc, ref)

    output_path = os.path.expanduser(
        "~/Desktop/Paper36_Sovereign_AI_Deployment_Standard_FULL.docx"
    )
    doc.save(output_path)
    print(f"FULL version saved: {output_path}")
    return output_path


if __name__ == "__main__":
    print("Building Paper 36: Sovereign AI Deployment Standard")
    print("=" * 60)
    public_path = build_public()
    full_path = build_full()
    print("=" * 60)
    print("Build complete.")
    print(f"  PUBLIC: {public_path}")
    print(f"  FULL:   {full_path}")
