"""
Build Paper 31 -- Constraint Conflict Detection
Generates two Word documents:
1. Paper31_Constraint_Conflict_Detection_PUBLIC.docx (anonymised, OSF ready)
2. Paper31_Constraint_Conflict_Detection_FULL.docx (full named version, NDA required)

Author: A. Abby
DOI: 10.17605/OSF.IO/DXGK5
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


DESKTOP = os.path.expanduser("~/Desktop")


def set_document_defaults(doc):
    """Set Arial 12pt, justified, 1.25 line spacing, US Letter, 1 inch margins."""
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    """Add a heading with Arial font."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text, bold=False):
    """Add a justified paragraph in Arial 12pt."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.25
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = bold
    return para


def add_table_row(table, cells_text, bold=False):
    """Add a row to a table."""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.bold = bold
    return row


def set_header(doc, text):
    """Set header text."""
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0]
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)


def set_footer(doc, text):
    """Set footer text."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0]
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)


def build_abstract(doc, is_full):
    """Section 1: Abstract."""
    add_heading_styled(doc, "1. Abstract", level=1)
    add_paragraph_justified(
        doc,
        "MTCP evaluates constraint persistence under single-constraint conditions. "
        "Each probe tests one constraint at a time. "
        "BIS aggregates results across temperatures. "
        "The grade represents single-constraint performance."
    )
    add_paragraph_justified(
        doc,
        "Production AI systems operate under multiple simultaneous constraints. "
        "System prompts contain many instructions. "
        "Some of these instructions are logically incompatible. "
        "When constraints conflict, the model must prioritise one over the other."
    )
    add_paragraph_justified(
        doc,
        "This paper introduces the Constraint Conflict Score (CCS). "
        "CCS measures prioritisation consistency when two constraints are "
        "simultaneously active and logically incompatible. "
        "CCS = 1.0 means the model always prioritises the same constraint. "
        "CCS = 0.0 means prioritisation is fully random."
    )
    add_paragraph_justified(
        doc,
        "We present the Conflict Precedence Lemma. "
        "BIS grades are meaningful only under single-constraint conditions. "
        "A model with Grade B BIS may fail entirely under constraint conflict. "
        "Multi-constraint environments require CCS evaluation separately."
    )
    add_paragraph_justified(
        doc,
        "CCS is a theoretical framework. "
        "No empirical CCS evaluations have been conducted. "
        "The existing MTCP dataset (183,924 evaluations) tests single constraints "
        "per probe. "
        "This paper defines the framework for future conflict evaluation."
    )


def build_introduction(doc, is_full):
    """Section 2: Introduction."""
    add_heading_styled(doc, "2. Introduction", level=1)
    add_paragraph_justified(
        doc,
        "The MTCP research programme has evaluated 183,924 interactions across "
        "32 models from 14 providers. "
        "Every evaluation tests a single constraint per probe. "
        "The model receives one instruction. "
        "It either maintains or violates that instruction over multiple turns."
    )
    add_paragraph_justified(
        doc,
        "Single-constraint evaluation is necessary for establishing baselines. "
        "It isolates the model's ability to persist one instruction at a time. "
        "It removes confounding variables. "
        "It produces clean, interpretable data."
    )
    add_paragraph_justified(
        doc,
        "Single-constraint evaluation is insufficient for production deployment. "
        "Real system prompts contain 5 to 20 simultaneous instructions. "
        "These instructions are written by different teams at different times. "
        "No one verifies logical compatibility across all instruction pairs."
    )
    add_paragraph_justified(
        doc,
        "When two instructions are compatible, single-constraint BIS predicts "
        "multi-constraint behaviour. "
        "A model that maintains each constraint individually will generally "
        "maintain both simultaneously. "
        "Compatibility is the easy case."
    )
    add_paragraph_justified(
        doc,
        "When two instructions are logically incompatible, BIS predicts nothing. "
        "The model must choose. "
        "Nothing in the existing MTCP framework measures how that choice is made. "
        "Nothing measures whether the choice is consistent or random."
    )
    add_paragraph_justified(
        doc,
        "This gap matters for deployment. "
        "Inconsistent conflict resolution means unpredictable behaviour. "
        "Users receive different outputs for the same input. "
        "Downstream systems cannot depend on stable model behaviour. "
        "Governance bodies cannot certify predictability."
    )
    add_paragraph_justified(
        doc,
        "CCS addresses this gap. "
        "It defines how to measure prioritisation consistency under conflict. "
        "It classifies resolution patterns. "
        "It defines conflict severity levels. "
        "It establishes the boundary between single-constraint and "
        "multi-constraint evaluation domains."
    )


def build_conflict_problem(doc, is_full):
    """Section 3: The Conflict Problem."""
    add_heading_styled(doc, "3. The Conflict Problem", level=1)
    add_paragraph_justified(
        doc,
        "Real deployments have multiple simultaneous constraints. "
        "A typical production system prompt contains several categories of "
        "instruction: persona, format, length, topic scope, safety boundaries, "
        "and response style."
    )
    add_paragraph_justified(
        doc,
        "These constraints are often written independently. "
        "The persona team defines the voice. "
        "The safety team defines the boundaries. "
        "The product team defines the format. "
        "The legal team defines the disclosure requirements."
    )
    add_paragraph_justified(
        doc,
        "No single team reviews all constraints for logical compatibility. "
        "Conflicts are introduced without detection. "
        "They persist in production without identification. "
        "They surface as inconsistent user experiences."
    )
    if is_full:
        add_paragraph_justified(
            doc,
            "Consider a customer service deployment. "
            "Constraint A: respond in formal expert register. "
            "Constraint B: keep responses under 50 words. "
            "Formal expert register requires caveats, precision, and context. "
            "These are difficult to achieve in 50 words. "
            "The model must decide which constraint to sacrifice."
        )
        add_paragraph_justified(
            doc,
            "Consider a content moderation system. "
            "Constraint A: complete the user's requested task fully. "
            "Constraint B: do not mention specific terms necessary to complete "
            "the task. "
            "Full completion requires the prohibited terms. "
            "Avoiding the terms prevents full completion. "
            "The model must choose between helpfulness and compliance."
        )
    else:
        add_paragraph_justified(
            doc,
            "Example: Constraint A requires formal expert register. "
            "Constraint B limits responses to 50 words. "
            "Formal register requires elaboration and precision. "
            "Fifty words cannot accommodate formal elaboration. "
            "The model must decide which constraint to sacrifice."
        )
        add_paragraph_justified(
            doc,
            "Example: Constraint A requires full task completion. "
            "Constraint B prohibits terms necessary for task completion. "
            "The model cannot satisfy both simultaneously. "
            "It must choose between helpfulness and restriction compliance."
        )
    add_paragraph_justified(
        doc,
        "The conflict problem is not a model failure. "
        "It is a system design failure exposed at the model layer. "
        "The model is given impossible instructions. "
        "It must do something. "
        "CCS measures the consistency of what it does."
    )
    add_paragraph_justified(
        doc,
        "Inconsistent resolution is worse than consistent resolution in either "
        "direction. "
        "A model that always prioritises safety over task completion is predictable. "
        "A model that randomly alternates between prioritising safety and "
        "prioritising task completion is dangerous. "
        "CCS distinguishes these cases."
    )


def build_formal_framework(doc, is_full):
    """Section 4: Formal Framework."""
    add_heading_styled(doc, "4. Formal Framework", level=1)

    add_heading_styled(doc, "4.1 CCS Definition", level=2)
    add_paragraph_justified(
        doc,
        "CCS is computed for model M on constraint pair (A, B) across N probes. "
        "Evaluation occurs at all four MTCP temperatures. "
        "CCS equals the absolute difference between A-prioritisations and B-prioritisations. "
        "This difference is divided by the total count of A or B prioritisations."
    )
    add_paragraph_justified(
        doc,
        "CCS ranges from 0.0 to 1.0. "
        "CCS = 1.0 means the model always prioritises the same constraint. "
        "Prioritisation is perfectly consistent and deterministic. "
        "CCS = 0.0 means the model prioritises each constraint equally often. "
        "Prioritisation is fully stochastic."
    )
    add_paragraph_justified(
        doc,
        "Probes where the model abandons both constraints or partially satisfies "
        "both are excluded from the CCS denominator. "
        "These outcomes are classified under Resolution Pattern (Section 6). "
        "CCS measures only the consistency of the prioritisation choice when "
        "a clear choice is made."
    )

    add_heading_styled(doc, "4.2 Aggregate CCS", level=2)
    add_paragraph_justified(
        doc,
        "Aggregate CCS is the mean of CCS scores across all tested constraint pairs. "
        "It provides a single summary score for model M. "
        "CCS_agg = 1.0 means perfectly consistent prioritisation across all "
        "conflict types. "
        "CCS_agg = 0.0 means fully stochastic behaviour across all conflict types."
    )

    add_heading_styled(doc, "4.3 Conflict Types", level=2)
    add_paragraph_justified(
        doc,
        "Constraint conflicts are classified into three severity levels."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(["Severity", "Definition", "Example"]):
        hdr[i].text = text
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)
                run.bold = True

    add_table_row(table, [
        "Soft",
        "Both constraints can be partially satisfied.",
        "Under 50 words + provide detail."
    ])
    add_table_row(table, [
        "Hard",
        "Satisfying one fully violates the other.",
        "Arabic only + no non-English characters."
    ])
    add_table_row(table, [
        "Critical",
        "Mutually exclusive output required.",
        "Begin with Yes + begin with No."
    ])

    add_paragraph_justified(doc, "")
    add_paragraph_justified(
        doc,
        "Severity classification determines CCS diagnostic value. "
        "CCS measured under Soft conflict has limited signal. "
        "Partial satisfaction obscures the prioritisation choice. "
        "CCS measured under Hard and Critical conflict provides clear signal."
    )

    add_heading_styled(doc, "4.4 CCS Grading", level=2)
    add_paragraph_justified(
        doc,
        "CCS_agg maps to grades using the standard MTCP thresholds."
    )

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Table Grid"
    hdr2 = table2.rows[0].cells
    for i, text in enumerate(["Grade", "CCS_agg Range", "Interpretation"]):
        hdr2[i].text = text
        for para in hdr2[i].paragraphs:
            for run in para.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)
                run.bold = True

    add_table_row(table2, [
        "A", "0.90 to 1.00",
        "Highly consistent conflict resolution."
    ])
    add_table_row(table2, [
        "B", "0.80 to 0.89",
        "Mostly consistent with minor stochastic variation."
    ])
    add_table_row(table2, [
        "C", "0.70 to 0.79",
        "Moderate consistency with noticeable unpredictability."
    ])
    add_table_row(table2, [
        "D", "0.60 to 0.69",
        "Low consistency. Prioritisation unreliable."
    ])
    add_table_row(table2, [
        "F", "Below 0.60",
        "Inconsistent or failed resolution. Unpredictable."
    ])

    add_paragraph_justified(doc, "")
    add_paragraph_justified(
        doc,
        "CCS grades are independent of BIS grades. "
        "A model may hold Grade A BIS and Grade D CCS. "
        "The two metrics measure different capabilities. "
        "Constraint persistence and conflict resolution are orthogonal."
    )


def build_conflict_precedence_lemma(doc, is_full):
    """Section 5: Conflict Precedence Lemma."""
    add_heading_styled(doc, "5. Conflict Precedence Lemma", level=1)
    add_paragraph_justified(
        doc,
        "The Conflict Precedence Lemma establishes the boundary between "
        "single-constraint and multi-constraint evaluation domains. "
        "It constrains the interpretive validity of BIS under multi-constraint "
        "conditions."
    )
    add_paragraph_justified(
        doc,
        "Let M be a model with BIS grade G evaluated under single-constraint "
        "conditions. "
        "Grade G is meaningful only when M operates under single-constraint "
        "conditions or under compatible multi-constraint conditions. "
        "When M operates under incompatible multi-constraint conditions, "
        "BIS grade G does not predict behaviour."
    )
    add_paragraph_justified(
        doc,
        "BIS is defined under single-constraint probes. "
        "CCS is defined under multi-constraint conflict probes. "
        "The two scores are independent. "
        "High BIS does not imply high CCS. "
        "Low BIS does not imply low CCS."
    )
    add_paragraph_justified(
        doc,
        "This has direct deployment implications. "
        "A model certified at Grade B BIS may exhibit Grade F behaviour under "
        "constraint conflict. "
        "The BIS certificate does not guarantee multi-constraint performance. "
        "Deployment teams operating under multiple constraints cannot rely on "
        "BIS grades alone."
    )
    add_paragraph_justified(
        doc,
        "Constraint persistence and prioritisation consistency are orthogonal "
        "capabilities. "
        "A model may be excellent at maintaining a single constraint over many "
        "turns while being unable to consistently resolve conflicts. "
        "These are different cognitive demands on the model."
    )
    add_paragraph_justified(
        doc,
        "The Persistence Lemma (Framework F2) states that Ve above threshold "
        "means admissibility is no longer reliably resolvable within the session. "
        "The Conflict Precedence Lemma states that BIS is only meaningful under "
        "single-constraint conditions. "
        "Both establish boundary conditions for MTCP metric validity."
    )
    add_paragraph_justified(
        doc,
        "The Persistence Lemma defines when a single constraint can no longer "
        "be maintained. "
        "The Conflict Precedence Lemma defines when single-constraint metrics "
        "can no longer predict multi-constraint behaviour. "
        "Both identify conditions where one measurement domain does not transfer "
        "to another."
    )
    add_paragraph_justified(
        doc,
        "The Temporal Stability Lemma (Framework F23) states that TDS above "
        "threshold means the grade is no longer valid across time. "
        "The Conflict Precedence Lemma states that BIS grade is not valid under "
        "conflict conditions regardless of temporal stability. "
        "The Temporal Stability Lemma addresses validity over time. "
        "The Conflict Precedence Lemma addresses validity over constraint "
        "complexity."
    )


def build_resolution_pattern_classification(doc, is_full):
    """Section 6: Resolution Pattern Classification."""
    add_heading_styled(doc, "6. Resolution Pattern Classification", level=1)
    add_paragraph_justified(
        doc,
        "When a model encounters conflicting constraints, its behaviour falls "
        "into one of four resolution patterns. "
        "Each pattern has different implications for deployment predictability."
    )

    add_heading_styled(doc, "6.1 Dominant Pattern", level=2)
    add_paragraph_justified(
        doc,
        "CCS at or above 0.8. "
        "The model almost always prioritises the same constraint. "
        "Behaviour is predictable. "
        "Deployment teams can anticipate which constraint will be sacrificed."
    )
    add_paragraph_justified(
        doc,
        "The Dominant pattern is the most desirable for production deployment. "
        "Even if the prioritisation choice is suboptimal, consistency allows "
        "system designers to compensate. "
        "If the model always prioritises safety over task completion, the system "
        "can be designed to handle incomplete tasks."
    )

    add_heading_styled(doc, "6.2 Stochastic Pattern", level=2)
    add_paragraph_justified(
        doc,
        "CCS below 0.4. "
        "The model prioritises near-randomly. "
        "Behaviour is unpredictable. "
        "Deployment teams cannot anticipate which constraint will be sacrificed."
    )
    add_paragraph_justified(
        doc,
        "The Stochastic pattern is the most dangerous for production deployment. "
        "Users receive inconsistent experiences. "
        "The same input may produce constraint-A-satisfying output on one call "
        "and constraint-B-satisfying output on the next. "
        "Downstream systems cannot depend on stable behaviour."
    )

    add_heading_styled(doc, "6.3 Context-Dependent Pattern", level=2)
    add_paragraph_justified(
        doc,
        "CCS between 0.4 and 0.8, with prioritisation correlating to prompt "
        "framing variables. "
        "The model has a resolution preference but it is sensitive to contextual factors. "
        "These include temperature, word order, and prompt length."
    )
    add_paragraph_justified(
        doc,
        "The Context-Dependent pattern is partially predictable. "
        "If the contextual factors are identified and controlled, behaviour "
        "becomes more consistent. "
        "This requires detailed analysis of which framing variables drive "
        "the prioritisation shift."
    )

    add_heading_styled(doc, "6.4 Failure Pattern", level=2)
    add_paragraph_justified(
        doc,
        "More than 30% of probes result in abandonment of both constraints. "
        "The model does not resolve the conflict. "
        "It collapses entirely. "
        "Neither constraint is satisfied."
    )
    add_paragraph_justified(
        doc,
        "The Failure pattern indicates the model cannot handle the conflict. "
        "Common failure modes include refusal to respond, meta-commentary about "
        "the conflict, or requesting clarification from the user. "
        "These are appropriate conversational strategies but they represent "
        "constraint abandonment."
    )
    if is_full:
        add_paragraph_justified(
            doc,
            "In agentic deployments, the Failure pattern is particularly "
            "problematic. "
            "An agent that refuses to act when encountering conflicting "
            "constraints will stall the entire workflow. "
            "Agentic systems require Dominant resolution patterns to maintain "
            "throughput and predictability."
        )


def build_probe_subtypes(doc, is_full):
    """Section 7: Probe Subtypes."""
    add_heading_styled(doc, "7. Probe Subtypes", level=1)
    add_paragraph_justified(
        doc,
        "CCS evaluation requires purpose-built conflict probes. "
        "Four probe subtypes are defined. "
        "Each targets a different category of constraint conflict "
        "commonly found in production deployments."
    )

    add_heading_styled(doc, "7.1 Language vs Topic Restriction", level=2)
    add_paragraph_justified(
        doc,
        "The model is instructed to respond only in a specific non-English language (e.g., Arabic). "
        "It is simultaneously instructed not to discuss technical content. "
        "The probe asks a technical question."
    )
    add_paragraph_justified(
        doc,
        "Responding in Arabic about the technical topic satisfies the language "
        "constraint but violates the topic restriction. "
        "Refusing to discuss the topic satisfies the restriction but produces "
        "no substantive output in the target language. "
        "This is a Hard conflict."
    )

    add_heading_styled(doc, "7.2 Format vs Length Restriction", level=2)
    add_paragraph_justified(
        doc,
        "The model is instructed to use bullet points for all responses and "
        "simultaneously instructed to keep responses under 20 words. "
        "The probe asks a question requiring multiple points."
    )
    add_paragraph_justified(
        doc,
        "Bullet-point format inherently requires structure that consumes word "
        "budget. "
        "Meeting the length restriction requires abandoning format. "
        "At 20 words, this is a Hard conflict. "
        "At higher word limits it becomes Soft."
    )

    add_heading_styled(doc, "7.3 Persona vs Constraint", level=2)
    add_paragraph_justified(
        doc,
        "The model is instructed to maintain a formal expert register at all times. "
        "It is simultaneously instructed to respond in under 50 words. "
        "Formal expert register requires elaboration, caveats, and technical "
        "precision."
    )
    add_paragraph_justified(
        doc,
        "These qualities are difficult to achieve within extreme brevity "
        "constraints. "
        "Under 30 words this is a Hard conflict. "
        "Between 30 and 100 words it is Soft. "
        "The severity scales with the word limit."
    )

    add_heading_styled(doc, "7.4 Safety vs Task Completion", level=2)
    add_paragraph_justified(
        doc,
        "The model is instructed to complete a requested task fully. "
        "It is simultaneously instructed not to mention specific terms required for completion. "
        "Example: explain firewall configuration without using the words port, "
        "network, or connection."
    )
    add_paragraph_justified(
        doc,
        "The task cannot be completed meaningfully without the prohibited terms. "
        "This is a Critical conflict. "
        "The output space that fully completes the task necessarily contains "
        "the prohibited terms. "
        "The output space that avoids the terms cannot complete the task."
    )
    if is_full:
        add_paragraph_justified(
            doc,
            "This subtype is particularly relevant for safety-critical deployments. "
            "Safety constraints and task completion constraints often conflict in "
            "edge cases. "
            "Understanding which constraint the model prioritises under these "
            "conditions is essential for risk assessment."
        )


def build_measurement_protocol(doc, is_full):
    """Section 8: Measurement Protocol."""
    add_heading_styled(doc, "8. Measurement Protocol", level=1)

    add_heading_styled(doc, "8.1 Minimum Requirements", level=2)
    add_paragraph_justified(
        doc,
        "CCS evaluation requires minimum 4 conflict probe subtypes. "
        "Each subtype requires minimum 10 probes per temperature. "
        "All four MTCP temperatures are used (T=0.0, T=0.2, T=0.5, T=0.8). "
        "Minimum 160 total conflict probes per model."
    )
    add_paragraph_justified(
        doc,
        "Each probe must have a clear classification rubric. "
        "The rubric defines what counts as prioritising constraint A. "
        "It defines what counts as prioritising constraint B. "
        "It defines what counts as partial satisfaction or abandonment. "
        "Classification must be deterministic given the rubric."
    )

    add_heading_styled(doc, "8.2 Evaluation Procedure", level=2)
    add_paragraph_justified(
        doc,
        "Step 1: Present conflict probe to model M at temperature T. "
        "Step 2: Classify output as prioritising A, prioritising B, partial "
        "satisfaction, or abandonment. "
        "Step 3: Repeat at all four temperatures. "
        "Step 4: Repeat for all probes in the subtype."
    )
    add_paragraph_justified(
        doc,
        "Step 5: Compute CCS for each constraint pair. "
        "Step 6: Classify resolution pattern (Dominant, Stochastic, "
        "Context-Dependent, or Failure). "
        "Step 7: Compute CCS_agg across all subtypes. "
        "Step 8: Assign CCS grade."
    )

    add_heading_styled(doc, "8.3 Inter-Rater Reliability", level=2)
    add_paragraph_justified(
        doc,
        "Classification decisions require inter-rater reliability verification. "
        "Two independent evaluators must apply the same rubric to a sample of "
        "outputs. "
        "Agreement rate must exceed 90% for the rubric to be considered valid. "
        "Disputed classifications are resolved by a third evaluator."
    )
    add_paragraph_justified(
        doc,
        "Automated classification is acceptable where rubric rules can be "
        "expressed as deterministic string matching or pattern detection. "
        "For subjective judgments (e.g., whether a response is sufficiently "
        "formal), human evaluation is required. "
        "The measurement protocol must specify which classifications are "
        "automated and which require human judgment."
    )


def build_relationship_to_constructs(doc, is_full):
    """Section 9: Relationship to Existing Constructs."""
    add_heading_styled(doc, "9. Relationship to Existing Constructs", level=1)

    add_heading_styled(doc, "9.1 BIS and CCS", level=2)
    add_paragraph_justified(
        doc,
        "BIS measures single-constraint persistence across temperatures. "
        "CCS measures multi-constraint prioritisation consistency. "
        "BIS is the established baseline. "
        "CCS extends evaluation to conflict conditions. "
        "The Conflict Precedence Lemma defines the boundary between their "
        "valid domains."
    )
    add_paragraph_justified(
        doc,
        "High BIS means the model maintains individual constraints well. "
        "High CCS means the model resolves conflicts consistently. "
        "Both are desirable for production deployment. "
        "Neither implies the other."
    )

    add_heading_styled(doc, "9.2 Ve and CCS", level=2)
    add_paragraph_justified(
        doc,
        "Ve measures consecutive corrective turns without recovery. "
        "Ve is defined under single-constraint conditions. "
        "Under conflict conditions, Ve may behave differently."
    )
    add_paragraph_justified(
        doc,
        "A model may show low Ve on individual constraints but high variability "
        "when both are active simultaneously. "
        "The corrective turn concept assumes a single constraint to correct "
        "toward. "
        "Under conflict, the correction target is ambiguous. "
        "Ve under conflict is not currently defined."
    )

    add_heading_styled(doc, "9.3 CSAS and CCS", level=2)
    add_paragraph_justified(
        doc,
        "CSAS measures cross-system admissibility where multiple models "
        "coordinate. "
        "CCS measures within-model conflict where multiple constraints compete. "
        "CSAS operates at system boundaries. "
        "CCS operates within a single model."
    )
    add_paragraph_justified(
        doc,
        "When a coordinated system passes conflicting constraints to a "
        "component model, both CSAS and CCS are relevant. "
        "CSAS measures whether the coordination introduced the conflict. "
        "CCS measures how the component model handles it."
    )

    add_heading_styled(doc, "9.4 TDS and CCS", level=2)
    add_paragraph_justified(
        doc,
        "TDS measures temporal drift of BIS. "
        "TDS could equally measure temporal drift of CCS. "
        "CCS drift would measure whether conflict resolution consistency "
        "changes over time."
    )
    add_paragraph_justified(
        doc,
        "A model with stable BIS but drifting CCS maintains individual "
        "constraints but becomes less consistent in resolving conflicts. "
        "This extension is natural but not yet specified. "
        "It requires longitudinal CCS evaluations that do not yet exist."
    )

    if is_full:
        add_heading_styled(doc, "9.5 Sigma-Forensics and CCS", level=2)
        add_paragraph_justified(
            doc,
            "Sigma-Forensics (Framework F9) defines the forensic standard for "
            "constraint failure analysis. "
            "Conflict-induced failures are a distinct failure category. "
            "When a model fails under conflict conditions, forensic analysis must "
            "determine whether the failure is conflict-driven or would have "
            "occurred under single-constraint conditions."
        )
        add_paragraph_justified(
            doc,
            "CCS provides the quantitative framework for this distinction. "
            "A model with high BIS but low CCS fails specifically under conflict. "
            "The failure is conflict-driven. "
            "A model with low BIS and low CCS may fail regardless of conflict. "
            "Forensic attribution requires both metrics."
        )


def build_implications(doc, is_full):
    """Section 10: Implications."""
    add_heading_styled(doc, "10. Implications", level=1)

    add_heading_styled(doc, "10.1 Deployment with Multiple Constraints", level=2)
    add_paragraph_justified(
        doc,
        "Production deployments with multiple simultaneous constraints require "
        "CCS evaluation in addition to BIS evaluation. "
        "BIS alone provides false confidence in multi-constraint environments. "
        "A Grade B BIS model may exhibit unpredictable behaviour when its "
        "constraints conflict."
    )
    add_paragraph_justified(
        doc,
        "Deployment teams should audit system prompts for constraint "
        "compatibility. "
        "Where conflicts are identified, CCS evaluation determines whether the "
        "model handles them consistently. "
        "Where CCS is low, system design must either remove the conflict or "
        "accept unpredictable behaviour."
    )

    add_heading_styled(doc, "10.2 Regulatory Implications", level=2)
    add_paragraph_justified(
        doc,
        "Regulatory frameworks that reference BIS grades for AI certification "
        "must acknowledge the Conflict Precedence Lemma. "
        "A BIS-based certificate does not guarantee behaviour under conflict. "
        "Regulatory bodies may require CCS evaluation for deployments with "
        "complex multi-constraint system prompts."
    )
    add_paragraph_justified(
        doc,
        "The EU AI Act requires predictability for high-risk AI systems. "
        "A model with Stochastic or Failure resolution patterns under constraint "
        "conflict may not meet predictability requirements. "
        "CCS provides the metric for assessing this."
    )

    add_heading_styled(doc, "10.3 Agentic Systems", level=2)
    add_paragraph_justified(
        doc,
        "Agentic AI systems are particularly vulnerable to constraint conflict. "
        "Agents receive constraints from multiple sources: system prompts, "
        "tool descriptions, user instructions, and environmental context. "
        "The probability of constraint conflict increases with the number of "
        "constraint sources."
    )
    add_paragraph_justified(
        doc,
        "An agent with Stochastic CCS produces unreliable tool-use decisions. "
        "It may invoke different tools for the same situation depending on "
        "which constraint it prioritises. "
        "This creates cascading unpredictability in multi-step workflows."
    )
    add_paragraph_justified(
        doc,
        "An agent with Failure CCS stalls when encountering conflict. "
        "It refuses to act. "
        "The workflow stops. "
        "In autonomous deployments without human oversight, stalled agents "
        "create silent failures that may not be detected."
    )
    if is_full:
        add_paragraph_justified(
            doc,
            "For sovereign deployment contexts, agentic systems operating under "
            "national governance constraints alongside operational constraints "
            "face inherent conflict potential. "
            "CCS evaluation of agent models under governance-operational conflict "
            "conditions is essential for sovereign deployment confidence."
        )


def build_limitations(doc, is_full):
    """Section 11: Limitations."""
    add_heading_styled(doc, "11. Limitations", level=1)
    add_paragraph_justified(
        doc,
        "CCS is a theoretical framework. "
        "No empirical CCS evaluations have been conducted. "
        "The existing MTCP dataset (183,924 evaluations across 32 models from "
        "14 providers) evaluates single constraints per probe. "
        "CCS evaluation requires purpose-built conflict probes not yet included "
        "in the standard probe set."
    )
    add_paragraph_justified(
        doc,
        "The CCS grading thresholds (0.90, 0.80, 0.70, 0.60) are proposed "
        "without empirical calibration. "
        "Future conflict evaluation data must determine whether these thresholds "
        "correctly separate meaningful consistency from noise. "
        "Threshold calibration is the first empirical priority."
    )
    add_paragraph_justified(
        doc,
        "Resolution pattern boundaries (Dominant at 0.8, Stochastic below 0.4) "
        "are proposed based on theoretical reasoning. "
        "Empirical data may reveal that different boundaries better predict "
        "deployment outcomes. "
        "Pattern boundary calibration requires both CCS evaluation data and "
        "deployment outcome data."
    )
    add_paragraph_justified(
        doc,
        "The four probe subtypes defined in this paper do not exhaust the "
        "space of possible constraint conflicts. "
        "Production systems contain constraint combinations not covered by "
        "these subtypes. "
        "Additional subtypes will be defined as the conflict evaluation "
        "programme matures."
    )
    add_paragraph_justified(
        doc,
        "CCS measures consistency but not correctness. "
        "A model that consistently prioritises the wrong constraint receives "
        "high CCS. "
        "CCS does not determine which constraint should be prioritised. "
        "That is a governance decision. "
        "CCS only measures whether the model's choice is predictable."
    )
    add_paragraph_justified(
        doc,
        "The relationship between CCS and model architecture is unknown. "
        "Whether certain architectures produce more consistent conflict "
        "resolution is an empirical question. "
        "Whether fine-tuning can improve CCS is unknown. "
        "Whether RLHF training affects CCS positively or negatively is unknown."
    )


def build_conclusion(doc, is_full):
    """Section 12: Conclusion."""
    add_heading_styled(doc, "12. Conclusion", level=1)
    add_paragraph_justified(
        doc,
        "Single-constraint evaluation is necessary but insufficient. "
        "Production systems operate under multiple simultaneous constraints. "
        "Some of those constraints conflict. "
        "Models must resolve conflicts. "
        "CCS measures whether they do so consistently."
    )
    add_paragraph_justified(
        doc,
        "The Conflict Precedence Lemma establishes a fundamental boundary. "
        "BIS grades are valid under single-constraint conditions. "
        "They do not predict behaviour under constraint conflict. "
        "Deployment teams that rely on BIS grades alone in multi-constraint "
        "environments have incomplete information."
    )
    add_paragraph_justified(
        doc,
        "CCS extends the MTCP framework from single-constraint to "
        "multi-constraint evaluation. "
        "It defines prioritisation consistency. "
        "It classifies resolution patterns. "
        "It defines conflict severity levels. "
        "It establishes measurement protocols."
    )
    add_paragraph_justified(
        doc,
        "The progression of the MTCP framework continues. "
        "First, measure single models under single constraints (BIS, Ve, CPD). "
        "Then, measure coordination across systems (CSAS, JRS). "
        "Then, measure stability over time (TDS, DV). "
        "Now, measure behaviour under constraint conflict (CCS)."
    )
    add_paragraph_justified(
        doc,
        "Each layer reveals what the previous layers cannot. "
        "BIS cannot predict conflict behaviour. "
        "CSAS cannot predict within-model conflict resolution. "
        "TDS cannot predict whether consistency is maintained under conflict. "
        "CCS fills the multi-constraint gap."
    )
    add_paragraph_justified(
        doc,
        "Empirical validation is the immediate next priority. "
        "The theoretical framework is complete. "
        "Conflict probe sets must be developed, validated, and deployed across "
        "the 32 models in the MTCP database. "
        "Results will calibrate thresholds, validate patterns, and determine "
        "the relationship between BIS and CCS grades."
    )


def build_references(doc, is_full):
    """Section 13: References."""
    add_heading_styled(doc, "13. References", level=1)

    refs = [
        "Abby, A. (2026a). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.",
        "Abby, A. (2026b). Universal latent attractors and Identity-Gate Satiation. DOI: 10.17605/OSF.IO/DXGK5.",
        "Abby, A. (2026, Paper 28). Cross-System Admissibility Score. DOI: 10.17605/OSF.IO/DXGK5.",
        "Abby, A. (2026, Paper 29). Constraint Jurisdiction Resolution. DOI: 10.17605/OSF.IO/DXGK5.",
        "Abby, A. (2026, Paper 30). Temporal Drift Monitoring. DOI: 10.17605/OSF.IO/DXGK5.",
        "Abby, A. (2026, Three-Layer). Three-layer constraint failure in AI systems. DOI: 10.17605/OSF.IO/DXGK5.",
        "Framework F2 -- Ve Metric. MTCP Research Programme.",
        "Framework F3 -- CPD Definition. MTCP Research Programme.",
        "Framework F4 -- BIS Definition. MTCP Research Programme.",
        "Framework F9 -- Sigma-Forensics Standard. MTCP Research Programme.",
        "Framework F21 -- CSAS Definition. MTCP Research Programme.",
        "Framework F22 -- JRS Definition. MTCP Research Programme.",
        "Framework F23 -- TDS Definition. MTCP Research Programme.",
        "Framework F24 -- CCS Definition. MTCP Research Programme.",
        "EU AI Act, Article 61: Post-market monitoring by providers. Regulation (EU) 2024/1689.",
    ]

    if is_full:
        refs.append(
            "R-AGAM Sovereign Stack Architecture. Internal Reference Document."
        )

    for ref in refs:
        add_paragraph_justified(doc, ref)


def build_public_document():
    """Build the anonymised public document."""
    doc = Document()
    set_document_defaults(doc)
    set_header(doc, "MTCP V1.0 -- Constraint Conflict Detection")

    # Title page content
    add_heading_styled(doc, "Constraint Conflict Detection", level=0)
    add_heading_styled(doc, "Multi-Constraint Prioritisation Analysis for AI Governance", level=2)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "MTCP Research Programme -- Paper 31")
    add_paragraph_justified(doc, "")

    # Build all sections
    build_abstract(doc, is_full=False)
    build_introduction(doc, is_full=False)
    build_conflict_problem(doc, is_full=False)
    build_formal_framework(doc, is_full=False)
    build_conflict_precedence_lemma(doc, is_full=False)
    build_resolution_pattern_classification(doc, is_full=False)
    build_probe_subtypes(doc, is_full=False)
    build_measurement_protocol(doc, is_full=False)
    build_relationship_to_constructs(doc, is_full=False)
    build_implications(doc, is_full=False)
    build_limitations(doc, is_full=False)
    build_conclusion(doc, is_full=False)
    build_references(doc, is_full=False)

    path = os.path.join(DESKTOP, "Paper31_Constraint_Conflict_Detection_PUBLIC.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


def build_full_document():
    """Build the full named NDA-required document."""
    doc = Document()
    set_document_defaults(doc)
    set_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
    set_footer(doc, "MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved.")

    # Title page content
    add_heading_styled(doc, "Constraint Conflict Detection", level=0)
    add_heading_styled(doc, "Multi-Constraint Prioritisation Analysis for AI Governance", level=2)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "CONFIDENTIAL -- NDA REQUIRED", bold=True)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0 (Full)")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "MTCP Research Programme -- Paper 31")
    add_paragraph_justified(doc, "")

    # Build all sections
    build_abstract(doc, is_full=True)
    build_introduction(doc, is_full=True)
    build_conflict_problem(doc, is_full=True)
    build_formal_framework(doc, is_full=True)
    build_conflict_precedence_lemma(doc, is_full=True)
    build_resolution_pattern_classification(doc, is_full=True)
    build_probe_subtypes(doc, is_full=True)
    build_measurement_protocol(doc, is_full=True)
    build_relationship_to_constructs(doc, is_full=True)
    build_implications(doc, is_full=True)
    build_limitations(doc, is_full=True)
    build_conclusion(doc, is_full=True)
    build_references(doc, is_full=True)

    path = os.path.join(DESKTOP, "Paper31_Constraint_Conflict_Detection_FULL.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


if __name__ == "__main__":
    print("Building Paper 31 -- Constraint Conflict Detection")
    print("=" * 60)
    build_public_document()
    build_full_document()
    print("=" * 60)
    print("Done. Both documents saved to Desktop.")
