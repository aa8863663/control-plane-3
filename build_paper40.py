"""
Build Paper 40: Constraint Object Specification (COS).
Framework F31: Governed constraint as a formal object.
Generates PUBLIC and FULL .docx versions.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DESKTOP = os.path.expanduser("~/Desktop")


def set_document_formatting(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.25
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return para


def add_header(doc, text):
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_footer(doc, text):
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


# ============================================================
# CONTENT
# ============================================================

ABSTRACT = [
    "This paper defines the Constraint Object Specification (COS). "
    "COS formalises the governed constraint as a structured object. "
    "Every MTCP evaluation must reference a registered constraint object. "
    "Evaluations without a registered constraint are not admissible as governance evidence.",

    "The specification defines six properties for each constraint object. "
    "Identity provides a unique constraint_id. "
    "Provenance records who issued the constraint and under what authority. "
    "Scope defines the boundary of applicability. "
    "Validity conditions define what must hold for the constraint to remain active. "
    "Inheritance rules define behaviour in coordination chains. "
    "Expiry defines when the constraint ceases to be valid.",

    "COS was not designed in advance of measurement. "
    "It was revealed by operational pressure across 183,924 evaluations. "
    "This makes it more structurally grounded than top-down frameworks. "
    "The specification emerged from practice, not theory.",
]

INTRODUCTION = [
    "MTCP evaluates constraint persistence across 32 models and 13 providers. "
    "Every evaluation tests whether a model maintains a given constraint. "
    "The constraint itself has never been formally specified as an object.",

    "This gap became visible through operational pressure. "
    "Evaluations conducted under different constraint definitions are not comparable. "
    "Two evaluations referencing different constraints measure different things. "
    "Without formal constraint identity, comparison is assumption.",

    "The gap affects every downstream construct. "
    "BIS scores assume constraint identity across temperatures. "
    "CSAS scores assume constraint identity across system boundaries. "
    "JRS assumes that jurisdiction claims reference identifiable constraints. "
    "TDS assumes the same constraint is measured at both time points.",

    "COS closes this gap. "
    "It defines the constraint as a formal object with six properties. "
    "It establishes the Constraint Object Registry for persistent storage. "
    "It introduces the Constraint Identity Lemma for comparability. "
    "It threads constraint_id through the entire evidence chain.",

    "The specification was not designed in advance of the 183,924 evaluations. "
    "It was revealed by the data. "
    "This is the Empirically-Derived Specification Principle. "
    "The specification is more defensible because it emerged from measurement.",
]

FORMAL_FRAMEWORK = [
    "A Constraint Object is a formally specified governed constraint. "
    "It has six properties. "
    "Each property is required for governance admissibility.",

    "Property 1: Identity. "
    "A unique constraint_id assigned at registration. "
    "The ID is deterministic. "
    "It derives from the constraint name and scope. "
    "Two registrations with identical inputs produce the same ID.",

    "Property 2: Provenance. "
    "The source of the constraint. "
    "Records who issued it, when, and under what authority. "
    "Provenance is immutable after registration. "
    "It provides the audit trail for authority claims.",

    "Property 3: Scope. "
    "The boundary within which the constraint applies. "
    "Specifies which models, which interactions, which contexts. "
    "Scope is a structured record. "
    "It is the primary input to comparability determination.",

    "Property 4: Validity Conditions. "
    "What must hold for the constraint to remain active. "
    "Conditions are checked before each evaluation. "
    "A constraint that fails validity is not evaluable.",

    "Property 5: Inheritance Rules. "
    "Whether the constraint passes to downstream systems. "
    "Inheritance may be full, partial, or prohibited. "
    "This property governs behaviour in coordination chains. "
    "CSAS evaluation requires inheritance rules to be defined.",

    "Property 6: Expiry. "
    "When the constraint object ceases to be valid. "
    "After expiry, evaluations are no longer current governance evidence. "
    "Expiry aligns with TDS validity windows.",
]

CONSTRAINT_IDENTITY_LEMMA = [
    "The Constraint Identity Lemma defines evaluation comparability.",

    "Two evaluations are only comparable if they reference the same constraint_id. "
    "Alternatively, they reference a constraint with provably equivalent scope. "
    "Equivalent scope means identical applicability boundaries. "
    "Equivalent validity means identical activation conditions.",

    "Evaluations referencing different constraint objects measure different things. "
    "Comparison across different constraints is not governance-valid. "
    "This was previously assumed. "
    "COS makes it explicit and enforceable.",

    "The lemma has direct operational consequences. "
    "BIS scores from different constraint objects cannot be averaged. "
    "CSAS cannot measure coordination without shared constraint identity. "
    "TDS cannot measure drift without confirming the same constraint.",

    "The comparability check is implemented as an API endpoint. "
    "GET /api/cos/compare/{id_a}/{id_b} returns the determination. "
    "The result is binary: comparable or not comparable. "
    "The reason for incomparability is always stated.",
]

REGISTRY = [
    "The Constraint Object Registry is the persistent store. "
    "Every registered constraint object lives in the registry. "
    "Every evaluation run links to a constraint_id from the registry.",

    "Evaluations without a registered constraint are not admissible. "
    "This is a hard requirement, not a recommendation. "
    "Unregistered constraints produce evaluation data. "
    "That data cannot be used as governance evidence.",

    "The registry supports three lookup modes. "
    "By constraint_id: returns the full constraint object. "
    "By model_id: returns all constraints linked to a model. "
    "By evaluation_id: returns the constraint governing an evaluation.",

    "The registry maintains a complete audit trail. "
    "Every link between a constraint and an evaluation is recorded. "
    "Every validity check is logged. "
    "Every inheritance event is tracked.",
]

EMPIRICAL_PRINCIPLE = [
    "COS was not designed before the evaluations were conducted. "
    "It was revealed by the evaluations themselves. "
    "This is the Empirically-Derived Specification Principle.",

    "183,924 evaluations across 32 models exposed the need. "
    "Evaluations conducted under assumed constraint definitions diverged. "
    "Comparison became unreliable. "
    "The data forced the specification into existence.",

    "Top-down frameworks define what should be measured before data exists. "
    "They risk measuring what is convenient rather than what is necessary. "
    "COS defines what was measured after the data revealed the need. "
    "It is more defensible because it addresses a demonstrated gap.",

    "The principle applies beyond COS. "
    "Every MTCP construct emerged from operational measurement. "
    "Ve was revealed by correction failure patterns. "
    "BIS was revealed by temperature sensitivity data. "
    "CSAS was revealed by coordination boundary failures. "
    "COS continues this pattern.",
]

INTEGRATION = [
    "COS integrates with every existing MTCP construct.",

    "BEC Integration: each BEC record includes constraint_id. "
    "The constraint identity threads through the hash chain. "
    "Verification confirms that the chain references a valid constraint.",

    "Manifest Integration: the Constraint Manifest includes constraint_object_id. "
    "It also includes constraint_provenance. "
    "Receiving systems can verify the constraint governing the attestation.",

    "Gate Integration: the Admissibility Gate references constraint objects. "
    "A PERMIT decision is only valid for the specified constraint. "
    "Deploying under a different constraint invalidates the gate decision.",

    "CSAS Integration: coordination admissibility requires shared constraints. "
    "The Constraint Identity Lemma gates CSAS evaluation. "
    "Two systems are only evaluable together if COS comparability holds.",

    "JRS Integration: jurisdiction resolution references constraint objects. "
    "A jurisdiction claim must identify the constraint it governs. "
    "Unidentified constraints cannot have resolved jurisdiction.",
]

LIMITATIONS = [
    "COS is a specification. "
    "Implementation exists. "
    "The specification has not been tested across all 32 models.",

    "Constraint equivalence determination is currently binary. "
    "Future versions may support partial equivalence. "
    "Partial equivalence would enable bounded comparison. "
    "This is deferred to empirical calibration.",

    "Provenance immutability assumes correct initial registration. "
    "If provenance is incorrectly recorded, correction requires re-registration. "
    "This is by design. "
    "Immutability protects the audit trail.",

    "The registry does not enforce constraint quality. "
    "Any constraint can be registered. "
    "Quality is determined by validity conditions and scope definition. "
    "The registry is a store, not a validator.",
]

CONCLUSION = [
    "COS closes the constitutional gap in MTCP. "
    "The governed constraint is now a formal object. "
    "Every evaluation links to a registered constraint. "
    "Unregistered evaluations are not governance evidence.",

    "The Constraint Identity Lemma makes comparability explicit. "
    "Same constraint or equivalent constraint means comparable. "
    "Different constraint means not comparable. "
    "This was always true. COS makes it enforceable.",

    "The specification was revealed by 183,924 evaluations. "
    "It was not designed in advance. "
    "This makes it structurally grounded. "
    "The data forced the formalism.",

    "COS integrates with BEC, Manifest, Gate, CSAS, and JRS. "
    "It threads constraint identity through the entire framework. "
    "Every downstream construct gains formal grounding. "
    "The constitutional layer is now complete.",
]

REFERENCES = [
    "Abby, A. (2026a). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 28). Cross-System Admissibility Score. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 29). Constraint Jurisdiction Resolution. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 37). Blockchain Evidence Chain Integrity. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 38). Admissibility Gate. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 39). Constraint Manifest. DOI: 10.17605/OSF.IO/DXGK5.",
    "Framework F4 -- BIS Definition. MTCP Research Programme.",
    "Framework F21 -- CSAS Definition. MTCP Research Programme.",
    "Framework F22 -- JRS Definition. MTCP Research Programme.",
    "Framework F23 -- TDS Definition. MTCP Research Programme.",
    "Framework F28 -- BEC Definition. MTCP Research Programme.",
    "Framework F29 -- Gate Definition. MTCP Research Programme.",
    "Framework F30 -- Manifest Definition. MTCP Research Programme.",
    "Framework F31 -- COS Definition. MTCP Research Programme.",
]


def build_paper(is_public=True):
    doc = Document()
    set_document_formatting(doc)

    if is_public:
        add_header(doc, "MTCP V1.0 -- Constraint Object Specification")
    else:
        add_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
        add_footer(doc, "MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved.")

    # Title page
    add_heading_styled(doc, "MTCP Paper 40", level=1)
    add_heading_styled(doc, "Constraint Object Specification", level=2)
    add_paragraph_justified(doc, "")
    if not is_public:
        add_paragraph_justified(doc, "CONFIDENTIAL -- NDA REQUIRED")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "MTCP Research Programme -- Paper 40")
    add_paragraph_justified(doc, "Framework F31")
    add_paragraph_justified(doc, "")

    doc.add_page_break()

    # Abstract
    add_heading_styled(doc, "1. Abstract", level=1)
    for para in ABSTRACT:
        add_paragraph_justified(doc, para)

    # Introduction
    add_heading_styled(doc, "2. Introduction", level=1)
    for para in INTRODUCTION:
        add_paragraph_justified(doc, para)

    # Formal Framework
    add_heading_styled(doc, "3. Formal Framework", level=1)
    add_heading_styled(doc, "3.1 Constraint Object Definition", level=2)
    for para in FORMAL_FRAMEWORK:
        add_paragraph_justified(doc, para)

    # Constraint Identity Lemma
    add_heading_styled(doc, "4. Constraint Identity Lemma", level=1)
    for para in CONSTRAINT_IDENTITY_LEMMA:
        add_paragraph_justified(doc, para)

    # Registry
    add_heading_styled(doc, "5. Constraint Object Registry", level=1)
    for para in REGISTRY:
        add_paragraph_justified(doc, para)

    # Empirical Principle
    add_heading_styled(doc, "6. Empirically-Derived Specification Principle", level=1)
    for para in EMPIRICAL_PRINCIPLE:
        add_paragraph_justified(doc, para)

    # Integration
    add_heading_styled(doc, "7. Integration with MTCP Constructs", level=1)
    for para in INTEGRATION:
        add_paragraph_justified(doc, para)

    # Limitations
    add_heading_styled(doc, "8. Limitations", level=1)
    for para in LIMITATIONS:
        add_paragraph_justified(doc, para)

    # Conclusion
    add_heading_styled(doc, "9. Conclusion", level=1)
    for para in CONCLUSION:
        add_paragraph_justified(doc, para)

    # References
    add_heading_styled(doc, "10. References", level=1)
    for ref in REFERENCES:
        add_paragraph_justified(doc, ref)

    return doc


if __name__ == "__main__":
    print("Building Paper 40: Constraint Object Specification")
    print("=" * 60)

    doc_pub = build_paper(is_public=True)
    pub_path = os.path.join(DESKTOP, "Paper40_Constraint_Object_Specification_PUBLIC.docx")
    doc_pub.save(pub_path)
    print(f"PUBLIC version saved: {pub_path}")

    doc_full = build_paper(is_public=False)
    full_path = os.path.join(DESKTOP, "Paper40_Constraint_Object_Specification_FULL.docx")
    doc_full.save(full_path)
    print(f"FULL version saved: {full_path}")

    print("=" * 60)
    print("Done.")
