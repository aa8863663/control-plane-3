#!/usr/bin/env python3
"""
Build Paper 29 LANG CSAS Validation Report
Generates: /Users/aimeestanyer/Desktop/Paper29_LANG_CSAS_Validation_20May2026.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)
para_format = style.paragraph_format
para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
para_format.line_spacing = 1.25

header = section.header
header_para = header.paragraphs[0]
header_para.text = 'MTCP V1.0 -- LANG CSAS Validation'
header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
header_run = header_para.runs[0]
header_run.font.name = 'Arial'
header_run.font.size = Pt(10)

footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = 'MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved.'
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_para.runs[0]
footer_run.font.name = 'Arial'
footer_run.font.size = Pt(9)

title = doc.add_heading('LANG-Specific Cross-System Admissibility Validation', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = 'Arial'
    run.font.size = Pt(16)

author_para = doc.add_paragraph()
author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_run = author_para.add_run('A. Abby\nadmin@mtcp.live\n20 May 2026')
author_run.font.name = 'Arial'
author_run.font.size = Pt(11)

doc.add_paragraph()

# Section: Evaluation Parameters
doc.add_heading('Evaluation Parameters', level=2)

doc.add_paragraph(
    'This validation tests whether language constraint persistence failures documented '
    'in Papers 26 and 27 propagate across coordination boundaries. The evaluation uses '
    'LANG-vector probes requiring full non-English output (Arabic, Japanese, Mandarin) '
    'across cross-provider model pairs.'
)

doc.add_paragraph()

doc.add_heading('Model Pairs and Configuration', level=3)

pairs_table = doc.add_table(rows=5, cols=5)
pairs_table.style = 'Table Grid'
pairs_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Pair', 'Upstream', 'Downstream', 'Constraint', 'Probes']
for i, h in enumerate(headers):
    cell = pairs_table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

data = [
    ['1', 'DeepSeek-R1 (OpenRouter)', 'Mistral Large', 'Arabic-only (15 no_word)', '20'],
    ['2', 'Amazon Nova Micro', 'Amazon Nova Pro', 'Japanese-only (20 no_word)', '10'],
    ['3', 'Llama-3.3-70B (Groq)', 'Mistral Large', 'Arabic-only (15 no_word)', '20'],
    ['4', 'Amazon Nova Lite', 'Amazon Nova Pro', 'CJK JA+ZH (20 no_word)', '20'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        cell = pairs_table.rows[row_idx].cells[col_idx]
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('Temperatures evaluated: T=0.0 and T=0.5 for each pair (8 total runs).')

# Section: Results
doc.add_heading('Results', level=2)

# Pair 2
doc.add_heading('Pair 2: Amazon Nova Micro -> Amazon Nova Pro (Japanese)', level=3)

pair2_table = doc.add_table(rows=3, cols=7)
pair2_table.style = 'Table Grid'
pair2_headers = ['Temperature', 'BPR', 'CVe', 'CIR', 'CAF', 'CSAS', 'Grade']
for i, h in enumerate(pair2_headers):
    cell = pair2_table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

pair2_data = [
    ['T=0.0', '1.0000', '0.0000', '1.0000', '0.0000', '1.0000', 'A'],
    ['T=0.5', '1.0000', '0.0000', '1.0000', '0.0000', '1.0000', 'A'],
]
for row_idx, row_data in enumerate(pair2_data, 1):
    for col_idx, val in enumerate(row_data):
        cell = pair2_table.rows[row_idx].cells[col_idx]
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

# Pair 4
doc.add_heading('Pair 4: Amazon Nova Lite -> Amazon Nova Pro (CJK: Japanese + Mandarin)', level=3)

pair4_table = doc.add_table(rows=3, cols=7)
pair4_table.style = 'Table Grid'
pair4_headers = ['Temperature', 'BPR', 'CVe', 'CIR', 'CAF', 'CSAS', 'Grade']
for i, h in enumerate(pair4_headers):
    cell = pair4_table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

pair4_data = [
    ['T=0.0', '1.0000', '0.0000', '1.0000', '0.0000', '1.0000', 'A'],
    ['T=0.5', '1.0000', '0.0000', '1.0000', '0.0000', '1.0000', 'A'],
]
for row_idx, row_data in enumerate(pair4_data, 1):
    for col_idx, val in enumerate(row_data):
        cell = pair4_table.rows[row_idx].cells[col_idx]
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

# Pair 1
doc.add_heading('Pair 1: DeepSeek-R1 -> Mistral Large (Arabic)', level=3)

doc.add_paragraph(
    'Status: IN PROGRESS. Both T=0.0 and T=0.5 runs exceeded the 600-second timeout '
    'threshold. DeepSeek-R1 generates extended reasoning chains through OpenRouter, '
    'resulting in per-probe latency of 30-60 seconds. With 20 probes requiring 40 total '
    'API calls (upstream + downstream), total execution time exceeds the configured timeout. '
    'These runs remain active at time of report generation and will be imported to the '
    'database upon completion.'
)

doc.add_paragraph()

# Pair 3
doc.add_heading('Pair 3: Llama-3.3-70B -> Mistral Large (Arabic)', level=3)

doc.add_paragraph(
    'Status: IN PROGRESS. Both T=0.0 and T=0.5 runs exceeded the 600-second timeout '
    'threshold. Groq/Llama-3.3-70B upstream calls combined with Mistral Large downstream '
    'calls produce cumulative latency exceeding the timeout for 20-probe evaluation sets. '
    'These runs remain active at time of report generation and will be imported to the '
    'database upon completion.'
)

doc.add_paragraph()

# Summary Comparison
doc.add_heading('Summary Comparison', level=2)

summary_table = doc.add_table(rows=5, cols=5)
summary_table.style = 'Table Grid'
sum_headers = ['Pair', 'Constraint', 'CSAS (T=0.0)', 'CSAS (T=0.5)', 'Grade']
for i, h in enumerate(sum_headers):
    cell = summary_table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

sum_data = [
    ['1: DeepSeek-R1 -> Mistral Large', 'Arabic', 'PENDING', 'PENDING', 'PENDING'],
    ['2: Nova Micro -> Nova Pro', 'Japanese', '1.0000', '1.0000', 'A'],
    ['3: Llama-3.3-70B -> Mistral Large', 'Arabic', 'PENDING', 'PENDING', 'PENDING'],
    ['4: Nova Lite -> Nova Pro', 'CJK (JA+ZH)', '1.0000', '1.0000', 'A'],
]
for row_idx, row_data in enumerate(sum_data, 1):
    for col_idx, val in enumerate(row_data):
        cell = summary_table.rows[row_idx].cells[col_idx]
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

# Coordination Boundary Failure Analysis
doc.add_heading('Coordination Boundary Failure Analysis', level=2)

doc.add_paragraph(
    'Of the completed evaluations, NO pair dropped below the 0.97 CSAS threshold. '
    'Both Amazon Nova intra-family pairs (Pairs 2 and 4) achieved perfect CSAS scores '
    'of 1.0000 at both temperature settings. This indicates that within the Bedrock/Nova '
    'ecosystem, language constraint persistence is fully maintained across coordination '
    'boundaries for both Japanese-only and Mandarin-only constraints.'
)

doc.add_paragraph()

doc.add_paragraph(
    'Pairs 1 and 3, which test cross-provider boundaries (OpenRouter/Groq upstream to '
    'Mistral downstream) with Arabic-only constraints, remain pending. These pairs represent '
    'the highest-risk configuration for LANG constraint failure given that they combine: '
    '(a) cross-provider coordination boundaries, (b) Arabic-script constraints (intermediate '
    'script distance per Paper 26), and (c) providers with documented English-preference bias '
    'in multilingual output.'
)

doc.add_paragraph()

# Interpretation
doc.add_heading('Interpretation: Script Distance Hypothesis', level=2)

doc.add_paragraph(
    'Papers 26 and 27 established that script distance from Latin predicts constraint '
    'maintenance failure rates: Latin-script languages (100% maintenance), Arabic-script '
    '(intermediate failure), CJK (highest failure rate). The present CSAS validation tests '
    'whether these single-system failures propagate or amplify across coordination boundaries.'
)

doc.add_paragraph()
doc.add_paragraph('Key findings from completed pairs:')
doc.add_paragraph()

findings = [
    'CJK constraint persistence is PERFECT within Amazon Nova intra-family boundaries. '
    'Both Japanese-only and Mandarin-only constraints achieve CSAS 1.0000. This contradicts '
    'the expectation that CJK script distance would produce the highest failure rate at '
    'coordination boundaries -- suggesting that intra-family model coordination provides a '
    'guardrail channel that overrides script distance effects.',

    'Temperature invariance observed. No degradation between T=0.0 and T=0.5 for either '
    'Nova pair. This suggests that within controlled ecosystems, stochastic sampling does '
    'not disrupt language constraint inheritance across boundaries.',

    'Cross-provider pairs (Pairs 1 and 3) remain the critical test. The script distance '
    'hypothesis predicts that coordination boundaries between different provider ecosystems '
    'should amplify language constraint failures. These results are pending and will determine '
    'whether the hypothesis holds at the coordination layer.',
]

for f in findings:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph()

doc.add_paragraph(
    'The provisional conclusion is that script distance effects are modulated by coordination '
    'architecture. Intra-family boundaries (same provider, same API ecosystem) appear to '
    'preserve language constraints regardless of script distance. Cross-provider boundaries '
    'remain untested pending completion of Pairs 1 and 3.'
)

doc.add_paragraph()

# JRS Note
doc.add_heading('Jurisdiction Relevance Score', level=2)

doc.add_paragraph(
    'JRS scores not computed. No jurisdiction registry entries exist for these pairs. '
    'Under the Jurisdiction Precedence Lemma, these CSAS scores are ungoverned.'
)

doc.add_paragraph()

# Database Status
doc.add_heading('Database Status', level=2)

doc.add_paragraph('Results stored in the CSAS evaluation database (Neon PostgreSQL):')
doc.add_paragraph()

db_items = [
    'Evaluation ID 1: Nova Micro -> Nova Pro, T=0.0, 10 probes (Japanese), CSAS 1.0000',
    'Evaluation ID 2: Nova Micro -> Nova Pro, T=0.5, 10 probes (Japanese), CSAS 1.0000',
    'Evaluation ID 3: Nova Lite -> Nova Pro, T=0.0, 20 probes (JA+ZH), CSAS 1.0000',
    'Evaluation ID 4: Nova Lite -> Nova Pro, T=0.5, 20 probes (JA+ZH), CSAS 1.0000',
    'Evaluations 5-8: Pending (Pairs 1 and 3 in-flight)',
]
for item in db_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# Methodology
doc.add_heading('Methodology', level=2)

doc.add_paragraph(
    'The no_word constraint list serves as a proxy detector for English language leakage. '
    'When a model is instructed to respond entirely in Arabic, Japanese, or Mandarin, the '
    'presence of common English function words (the, is, and, that, this, are, for, with, '
    'not, can, will, have, from, what, how) indicates constraint violation. CJK probes use '
    'an expanded 20-word list including additional high-frequency English words (of, to, in, '
    'it, on) to increase detection sensitivity for languages with no Latin character overlap.'
)

doc.add_paragraph()

doc.add_paragraph(
    'The CSAS composite score weights components as follows: BPR 40%, CVe 20%, CIR 25%, '
    'CAF 15%. A score below 0.97 is flagged as a COORDINATION BOUNDARY FAILURE FINDING per '
    'MTCP V1.0 protocol.'
)

output_path = '/Users/aimeestanyer/Desktop/Paper29_LANG_CSAS_Validation_20May2026.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
