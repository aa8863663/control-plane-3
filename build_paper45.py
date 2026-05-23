"""
Build Paper 45: Quantum-Safe Governance Evidence.
Framework F35: Cryptographic validity and post-quantum migration.
Generates PUBLIC and FULL .docx versions.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DESKTOP = os.path.expanduser("~/Desktop")


def set_document_formatting(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.25
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return para


def add_header(doc, text):
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_footer(doc, text):
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


ABSTRACT = [
    "This paper defines the Quantum-Safe Layer for MTCP governance evidence. "
    "It migrates BEC hash chains from SHA-256 to BLAKE3. "
    "BLAKE3 provides interim post-quantum resistance. "
    "The migration preserves audit continuity through dual-hash records.",

    "Cryptographic Validity Windows define how long each standard is trusted. "
    "BLAKE3 has a 10-year validity window from implementation. "
    "When the window closes, re-signing under a new standard is required. "
    "This makes cryptographic obsolescence explicit and planned.",

    "Harvest-Now Decrypt-Later risk classification prioritises migration. "
    "Level 1 records (gate decisions) require immediate quantum safety. "
    "Level 2 records (evaluation scores) are high priority. "
    "Level 3 records (manifest attestations) are standard priority.",
]

INTRODUCTION = [
    "MTCP governance evidence relies on cryptographic integrity. "
    "BEC hash chains use SHA-256. "
    "Gate decisions use SHA-256. "
    "Constraint Manifests use HMAC-SHA256.",

    "SHA-256 is not quantum-safe. "
    "A sufficiently powerful quantum computer could find collisions. "
    "This would allow governance records to be forged. "
    "The entire evidence chain would lose integrity.",

    "The threat is not immediate. "
    "No quantum computer currently threatens SHA-256. "
    "But governance evidence has a long lifespan. "
    "Records created today must be verifiable in 10 to 20 years.",

    "Harvest-Now Decrypt-Later attacks are the relevant threat model. "
    "An adversary captures hashed evidence today. "
    "They store it until quantum computers are available. "
    "They then forge alternative evidence records. "
    "Current evidence becomes contestable.",

    "BLAKE3 provides interim post-quantum resistance. "
    "It is faster than SHA-256 and SHA-3. "
    "It provides 256-bit security with no known quantum attack below 128-bit. "
    "This is sufficient for the 10-year validity window.",
]

CRYPTO_VALIDITY = [
    "Every governance record has a Cryptographic Validity Window.",

    "The window defines how long the cryptographic standard is trusted. "
    "BLAKE3 is trusted for 10 years from implementation. "
    "After 10 years, re-signing under a successor standard is required.",

    "The window is conservative. "
    "BLAKE3 may remain secure longer than 10 years. "
    "The 10-year window ensures planned migration before obsolescence. "
    "It prevents reactive crisis migration.",

    "Each record includes a crypto_standard field. "
    "This records which algorithm was used. "
    "It records when the algorithm was applied. "
    "It enables audit of cryptographic provenance.",

    "The Constraint Manifest includes crypto_standard and crypto_valid_until. "
    "These sit alongside the TDS validity window. "
    "A manifest can expire from TDS or from cryptographic obsolescence. "
    "Both expiry mechanisms protect evidence integrity.",
]

HARVEST_RISK = [
    "Records are classified by sensitivity to harvest attacks.",

    "Level 1: Gate decisions and deployment records. "
    "These directly authorise model deployment. "
    "Forging a gate decision could authorise an unsafe model. "
    "Level 1 records require quantum-safe cryptography immediately.",

    "Level 2: Evaluation scores and BIS records. "
    "These form the evidence base for gate decisions. "
    "Forging scores could change gate outcomes retroactively. "
    "Level 2 records are high priority for migration.",

    "Level 3: Manifest attestations. "
    "These travel with deployments and expire naturally. "
    "Short-lived manifests have less harvest exposure. "
    "Level 3 records are standard priority.",

    "Classification drives migration sequencing. "
    "Level 1 records migrate first. "
    "Level 2 records migrate second. "
    "Level 3 records migrate third. "
    "This minimises risk exposure during the transition period.",
]

MIGRATION = [
    "Migration moves all BEC records from SHA-256 to BLAKE3.",

    "The migration is non-destructive. "
    "Original SHA-256 hashes are preserved in the migration registry. "
    "New BLAKE3 hashes are computed and stored alongside. "
    "Both hashes coexist during the transition period.",

    "Audit continuity is maintained. "
    "Any auditor can verify the original SHA-256 chain. "
    "Any auditor can verify the new BLAKE3 chain. "
    "The migration registry links the two.",

    "Chain integrity is preserved through the migration. "
    "Each record is re-hashed individually. "
    "The chain structure (each record references its predecessor) is maintained. "
    "Verification confirms both chains are internally consistent.",

    "Migration status is tracked per record. "
    "Pending means not yet migrated. "
    "Complete means BLAKE3 hash computed and verified. "
    "The migration report shows overall progress.",
]

INTEGRATION = [
    "The quantum-safe layer integrates with all evidence constructs.",

    "BEC Integration: all chain records gain crypto_standard fields. "
    "New records are hashed with BLAKE3. "
    "Legacy records retain SHA-256 with migration tracking.",

    "Manifest Integration: crypto_standard and crypto_valid_until added. "
    "Manifests report their cryptographic freshness. "
    "Expired cryptography triggers manifest re-issuance.",

    "Gate Integration: gate decisions include crypto standard metadata. "
    "Level 1 classification ensures immediate quantum safety.",

    "API Integration: two new endpoints for crypto status and validity. "
    "GET /api/crypto/status returns current standard for any record. "
    "GET /api/crypto/manifest_validity returns quantum-safe status.",
]

LIMITATIONS = [
    "BLAKE3 is an interim step. "
    "It is not a NIST post-quantum standard. "
    "NIST standardised ML-KEM and ML-DSA for key exchange and signatures. "
    "BLAKE3 is a hash function, not a signature scheme.",

    "Full post-quantum migration requires ML-DSA for signatures. "
    "This is deferred to a future framework revision. "
    "BLAKE3 addresses the hash collision threat only. "
    "Signature forgery threats require separate treatment.",

    "The 10-year window is an estimate. "
    "Quantum computing timelines are uncertain. "
    "The window may need shortening if quantum advances accelerate. "
    "Annual reassessment of the validity window is recommended.",

    "BLAKE3 library availability varies across environments. "
    "The runner falls back to SHA3-256 if BLAKE3 is unavailable. "
    "SHA3-256 provides equivalent security guarantees. "
    "The fallback is transparent and logged.",
]

CONCLUSION = [
    "The Quantum-Safe Layer makes cryptographic obsolescence explicit. "
    "Every record knows its algorithm. "
    "Every algorithm has a validity window. "
    "Expiry is planned, not reactive.",

    "BLAKE3 provides 10 years of interim post-quantum resistance. "
    "Migration from SHA-256 is non-destructive and auditable. "
    "Harvest-Now Decrypt-Later risk classification prioritises migration. "
    "Level 1 records are protected first.",

    "The Constraint Manifest now attests cryptographic freshness. "
    "A manifest with expired cryptography is not valid governance evidence. "
    "This closes the gap between evidence integrity and cryptographic reality.",

    "Full post-quantum migration to ML-DSA signatures is future work. "
    "BLAKE3 addresses the immediate hash collision threat. "
    "The framework is designed for iterative cryptographic evolution.",
]

REFERENCES = [
    "Abby, A. (2026a). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 37). Blockchain Evidence Chain Integrity. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 38). Admissibility Gate. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 39). Constraint Manifest. DOI: 10.17605/OSF.IO/DXGK5.",
    "Framework F28 -- BEC Definition. MTCP Research Programme.",
    "Framework F29 -- Gate Definition. MTCP Research Programme.",
    "Framework F30 -- Manifest Definition. MTCP Research Programme.",
    "Framework F35 -- Quantum-Safe Definition. MTCP Research Programme.",
    "NIST (2024). ML-KEM (FIPS 203). Federal Information Processing Standard.",
    "NIST (2024). ML-DSA (FIPS 204). Federal Information Processing Standard.",
    "O'Connor et al. (2024). BLAKE3: One function, fast everywhere.",
]


def build_paper(is_public=True):
    doc = Document()
    set_document_formatting(doc)

    if is_public:
        add_header(doc, "MTCP V1.0 -- Quantum-Safe Governance Evidence")
    else:
        add_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
        add_footer(doc, "MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved.")

    add_heading_styled(doc, "MTCP Paper 45", level=1)
    add_heading_styled(doc, "Quantum-Safe Governance Evidence", level=2)
    add_paragraph_justified(doc, "")
    if not is_public:
        add_paragraph_justified(doc, "CONFIDENTIAL -- NDA REQUIRED")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "MTCP Research Programme -- Paper 45")
    add_paragraph_justified(doc, "Framework F35")

    doc.add_page_break()

    add_heading_styled(doc, "1. Abstract", level=1)
    for para in ABSTRACT:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "2. Introduction", level=1)
    for para in INTRODUCTION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "3. Cryptographic Validity Windows", level=1)
    for para in CRYPTO_VALIDITY:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "4. Harvest-Now Decrypt-Later Risk", level=1)
    for para in HARVEST_RISK:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "5. Migration Protocol", level=1)
    for para in MIGRATION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "6. Integration", level=1)
    for para in INTEGRATION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "7. Limitations", level=1)
    for para in LIMITATIONS:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "8. Conclusion", level=1)
    for para in CONCLUSION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "9. References", level=1)
    for ref in REFERENCES:
        add_paragraph_justified(doc, ref)

    return doc


if __name__ == "__main__":
    print("Building Paper 45: Quantum-Safe Governance Evidence")
    print("=" * 60)

    doc_pub = build_paper(is_public=True)
    pub_path = os.path.join(DESKTOP, "Paper45_Quantum_Safe_Governance_Evidence_PUBLIC.docx")
    doc_pub.save(pub_path)
    print(f"PUBLIC version saved: {pub_path}")

    doc_full = build_paper(is_public=False)
    full_path = os.path.join(DESKTOP, "Paper45_Quantum_Safe_Governance_Evidence_FULL.docx")
    doc_full.save(full_path)
    print(f"FULL version saved: {full_path}")

    print("=" * 60)
    print("Done.")
