#!/usr/bin/env python3
"""
Build Paper 39: Constraint Manifest -- Portable Verified Attestation for Model Deployments.
Produces two documents: PUBLIC and FULL (NDA) versions.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.expanduser("~/Desktop")


def set_font(run, name="Arial", size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.insert(0, rFonts)


def set_paragraph_format(para, spacing=1.25, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    para.alignment = alignment
    para.paragraph_format.line_spacing = spacing
    para.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level=1):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, bold=True, size=14 if level == 1 else 12)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    para.paragraph_format.space_after = Pt(6)
    return para


def add_para(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, bold=bold)
    set_paragraph_format(para)
    return para


def set_margins(doc):
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)


def add_header(doc, text):
    section = doc.sections[0]
    header = section.header
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    run = para.add_run(text)
    set_font(run, size=9)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_footer(doc, text):
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    run = para.add_run(text)
    set_font(run, size=8)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ============================================================================
# PUBLIC VERSION
# ============================================================================

def build_public():
    doc = Document()
    set_margins(doc)
    add_header(doc, "MTCP V1.0 -- Constraint Manifest")

    # Title
    para = doc.add_paragraph()
    run = para.add_run("Paper 39: Constraint Manifest")
    set_font(run, bold=True, size=16)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run("Portable Verified Attestation for Model Deployments")
    set_font(run, size=12)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run("A. Abby -- MTCP Research Programme -- May 2026")
    set_font(run, size=10)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run("DOI: 10.17605/OSF.IO/DXGK5")
    set_font(run, size=10)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Abstract
    add_heading(doc, "1. Abstract")
    add_para(doc, "This paper defines the Constraint Manifest. The Constraint Manifest is a portable signed document that travels with a model deployment. It packages evaluation scores, gate decisions, and integrity attestations into a single verifiable document. Any receiving system can verify the manifest independently without contacting the issuing system.")
    add_para(doc, "The manifest provides cryptographic proof of evaluation status at time of issuance. It uses SHA-256 hashing and HMAC-SHA256 signing to prevent tampering. The specification covers 32 models across 13 providers. This paper presents the formal schema, verification protocol, and portability constraints.")
    add_para(doc, "The Constraint Manifest operationalizes EU AI Act requirements for portable conformity evidence. It bridges the gap between evaluation (which is centralized) and deployment (which is distributed). The manifest is the mechanism by which evaluation travels with a model.")

    # 2. Introduction
    add_heading(doc, "2. Introduction")
    add_para(doc, "MTCP evaluates 32 AI models across 13 providers for behavioural constraint persistence. Evaluations produce scores across multiple frameworks. BIS measures behavioural integrity. CSAS measures cross-system admissibility. JRS measures jurisdiction resolution. TDS measures temporal drift. The Admissibility Gate enforces deployment decisions based on these scores.")
    add_para(doc, "All evaluation data resides in the MTCP platform. When a model is deployed to a new system, that system cannot verify evaluation status without querying MTCP directly. This creates a dependency on the issuing system for every deployment verification. The Constraint Manifest eliminates this dependency.")
    add_para(doc, "The manifest is a self-contained document. It includes all scores, the gate decision, integrity attestations, and a cryptographic signature. A receiving system needs only the manifest and the verification protocol to confirm that a model passed evaluation. No callback to MTCP is required.")
    add_para(doc, "This paper specifies the manifest schema, the three-step verification protocol, the Portability Lemma, and the revocation mechanism. It establishes the formal foundation for portable model attestation within the MTCP framework.")

    # 3. Portability Problem
    add_heading(doc, "3. The Portability Problem")
    add_para(doc, "Evaluation and deployment are structurally separated. Evaluation occurs at a central authority. Deployment occurs at distributed endpoints. The gap between these creates a verification problem. How does a deploying system know a model has been evaluated?")
    add_para(doc, "Current approaches rely on API queries to the evaluation authority. This requires network connectivity, introduces latency, and creates a single point of failure. If the evaluation authority is unreachable, deployment verification cannot proceed. This is unacceptable for critical infrastructure.")
    add_para(doc, "The alternative is a portable document that travels with the model. This document must be tamper-evident (cannot be modified without detection). It must be self-verifying (can be checked without external queries). It must be time-bounded (expires when evaluation data becomes stale).")
    add_para(doc, "The Constraint Manifest solves this problem. It packages evaluation data into a signed document that can be verified offline. The document travels with the model from evaluation authority to deployment endpoint. Verification requires only the document itself and knowledge of the verification protocol.")

    # 4. Formal Framework
    add_heading(doc, "4. Formal Framework")
    add_para(doc, "The Constraint Manifest is defined as a tuple M = (S, H, Sig) where S is the schema (evaluation data), H is the SHA-256 hash computed over S, and Sig is the HMAC-SHA256 signature computed over H using the issuer secret key.")
    add_para(doc, "The manifest is valid when three conditions hold simultaneously. First, recomputing H from S produces the same hash. Second, recomputing Sig from H produces the same signature. Third, the current date does not exceed the TDS validity window encoded in S.")
    add_para(doc, "The Portability Lemma constrains manifest validity to the model-provider pair. If either component changes, the manifest is invalid. This reflects the reality that model behaviour depends on provider infrastructure.")
    add_para(doc, "Manifest revocation provides an additional invalidation mechanism beyond expiry. The issuer can revoke any manifest at any time. Revocation is permanent and irreversible. A revoked manifest fails verification regardless of hash and signature validity.")

    # 5. Schema
    add_heading(doc, "5. Manifest Schema")
    add_para(doc, "The manifest schema contains 20 fields organized into four groups. The identification group contains manifest_id (UUID), model_id, provider, and evaluation_date. These identify what was evaluated and when.")
    add_para(doc, "The scores group contains bis_score, bis_grade, csas_score, csas_grade, jrs_score, tds_valid_until, tds_status, acps_score, and acps_grade. These capture the evaluation results from each MTCP framework. Optional scores (CSAS, JRS, ACPS) are null when not applicable.")
    add_para(doc, "The enforcement group contains gate_status, gate_context, regulatory_compliance, bec_integrity_score, and bec_chain_id. These capture the gate decision, applicable regulations, and chain integrity at time of issuance.")
    add_para(doc, "The verification group contains manifest_hash, issuer, and issuer_signature. The hash is SHA-256 computed over all preceding fields using pipe-delimited concatenation. The signature is HMAC-SHA256 computed over the hash. Null fields are represented as the string \"null\" in hash computation.")

    # 6. Verification Protocol
    add_heading(doc, "6. Verification Protocol")
    add_para(doc, "Verification follows three steps executed in sequence. All three must pass for the manifest to be accepted. Failure at any step terminates verification with REJECTED status.")
    add_para(doc, "Step 1 is hash verification. The verifier recomputes SHA-256 over all schema fields using pipe-delimited concatenation. If the computed hash matches manifest_hash, the document is intact. If not, it has been tampered with. Hash mismatch means REJECTED.")
    add_para(doc, "Step 2 is signature verification. The verifier computes HMAC-SHA256 over manifest_hash using the shared verification key. If the computed signature matches issuer_signature, the document was issued by the claimed authority. If not, the signature is forged. Signature mismatch means REJECTED.")
    add_para(doc, "Step 3 is expiry verification. The verifier compares the current date against tds_valid_until. If the current date exceeds the validity window, the evaluation data is stale. The manifest no longer attests to current model state. Expiry means REJECTED.")
    add_para(doc, "An optional fourth check queries the revocation registry. If the manifest_id appears in the revocation registry, the manifest has been explicitly invalidated by the issuer. Revocation means REJECTED regardless of other checks passing.")

    # 7. Portability Lemma
    add_heading(doc, "7. Portability Lemma")
    add_para(doc, "The Portability Lemma states that a change of provider invalidates the manifest. A manifest is bound to the specific model-provider pair at time of issuance. If either component changes, the attestation no longer applies.")
    add_para(doc, "This constraint reflects a physical reality. Model behaviour depends on provider infrastructure. The same model weights served through different providers may exhibit different constraint persistence characteristics. Evaluation results from one provider do not transfer to another.")
    add_para(doc, "Provider change triggers automatic revocation. When the MTCP platform detects that a model has changed provider, all active manifests for that model-provider pair are revoked. New evaluation under the new provider is required before a new manifest can be issued.")
    add_para(doc, "The Portability Lemma does not restrict manifest transfer between receiving systems. A manifest for model X on provider P is valid for any system deploying model X from provider P. The manifest travels with the deployment. It does not travel with the provider.")

    # 8. Revocation
    add_heading(doc, "8. Revocation Mechanism")
    add_para(doc, "Revocation permanently invalidates a manifest. Three conditions trigger revocation. Score change occurs when evaluation scores change after issuance. Provider change occurs when the model moves to a different provider. Administrative revocation occurs at issuer discretion for any reason.")
    add_para(doc, "Revocation is recorded in a central revocation registry. The registry maps manifest_id to revocation_id, timestamp, reason, and revoking authority. Receiving systems should check this registry before accepting a manifest.")
    add_para(doc, "Revocation is irreversible. A revoked manifest cannot be un-revoked. If conditions change favourably, a new manifest must be issued. This prevents manipulation of the revocation system to toggle manifest validity.")
    add_para(doc, "Offline systems face a revocation gap. A system that cannot reach the revocation registry may accept a revoked manifest. This is an acknowledged limitation. High-security deployments should require online revocation checking.")

    # 9. API Contract
    add_heading(doc, "9. API Contract")
    add_para(doc, "The manifest system exposes four API endpoints. POST /api/manifest/generate creates a new manifest for a specified model and context. It requires model_id, deployment context, and optionally provider override. It returns the complete manifest document.")
    add_para(doc, "POST /api/manifest/verify accepts a manifest document or manifest_id and executes the three-step verification protocol. It returns hash_valid, signature_valid, expiry_valid, revocation status, and overall result (VERIFIED or REJECTED).")
    add_para(doc, "POST /api/manifest/revoke accepts a manifest_id and revocation reason. It permanently revokes the manifest. It returns the revocation record including revocation_id and timestamp.")
    add_para(doc, "GET /api/manifest/registry returns all active manifests in the system. It supports filtering by model, provider, context, and status. It returns the manifest registry with issuance and expiry dates.")

    # 10. Relationship to Constructs
    add_heading(doc, "10. Relationship to MTCP Constructs")
    add_para(doc, "BIS provides the primary score packaged in every manifest. BIS score and grade are required fields. A manifest without BIS is invalid. BIS represents the core behavioural integrity evaluation.")
    add_para(doc, "CSAS provides cross-system admissibility when applicable. CSAS fields are optional in the manifest schema. They are populated only when the model operates in a coordination pair. When present, they attest to cross-system evaluation.")
    add_para(doc, "JRS provides jurisdiction resolution score. Like CSAS, JRS is optional and populated only for models operating at system boundaries. The manifest attests to jurisdiction resolution capability when applicable.")
    add_para(doc, "TDS provides temporal stability classification and controls manifest expiry. The tds_valid_until field is computed from the TDS evaluation date plus 90 days. TDS is the only construct that directly controls manifest lifetime.")
    add_para(doc, "BEC provides chain integrity attestation. The bec_integrity_score and bec_chain_id fields attest that the underlying evaluation records are tamper-evident. BEC integrity of 1.0 is required for manifest generation.")
    add_para(doc, "The Admissibility Gate provides the gate_status field. The manifest packages the gate decision (PERMIT or DENY) for the specified deployment context. The manifest does not evaluate. It attests to a prior gate decision.")

    # 11. Implications
    add_heading(doc, "11. Implications")
    add_para(doc, "The Constraint Manifest enables distributed deployment verification. Organizations can verify model evaluation status at the point of deployment without connectivity to the evaluation authority. This supports air-gapped environments and edge deployments.")
    add_para(doc, "Regulatory compliance is simplified. Auditors can request the manifest as evidence of evaluation. The manifest is a single document containing all relevant scores, decisions, and cryptographic proof. It replaces multi-document evidence gathering.")
    add_para(doc, "Supply chain verification becomes possible. When models pass through intermediaries (resellers, integrators, platform providers), each intermediary can verify the manifest independently. The evaluation attestation travels through the supply chain with the model.")
    add_para(doc, "Procurement decisions are accelerated. A deploying organization can verify a manifest in milliseconds. No evaluation cycle is required at the receiving end. The evaluation was already performed and attested to by the manifest.")

    # 12. Limitations
    add_heading(doc, "12. Limitations")
    add_para(doc, "The manifest attests to evaluation state at time of issuance only. Model behaviour can change after manifest issuance. The 90-day validity window mitigates this but does not eliminate it. Real-time behavioural monitoring is outside manifest scope.")
    add_para(doc, "HMAC-SHA256 requires shared secret distribution. Receiving systems need the verification key to check signatures. Key distribution is a separate infrastructure concern not addressed by this specification.")
    add_para(doc, "Offline systems cannot check revocation status. A manifest that has been revoked may be accepted by a disconnected system. The revocation gap is bounded by connectivity restoration.")
    add_para(doc, "The manifest does not include raw evaluation data. It attests to scores but does not include the underlying probe results, model responses, or evaluation methodology. Full evidence requires separate evidence pack generation.")

    # 13. Conclusion
    add_heading(doc, "13. Conclusion")
    add_para(doc, "The Constraint Manifest bridges evaluation and deployment through a portable signed document. It enables independent verification without issuing authority cooperation. The three-step verification protocol (hash, signature, expiry) provides strong tamper-evidence.")
    add_para(doc, "The Portability Lemma constrains manifest validity to the evaluated model-provider pair. This reflects the reality that model behaviour depends on serving infrastructure. Provider change requires re-evaluation and new manifest issuance.")
    add_para(doc, "This specification covers 32 models across 13 providers within the MTCP evaluation framework. The manifest schema, verification protocol, and revocation mechanism are fully defined and implemented. Framework F30 provides the formal definitions referenced by this paper.")

    # References
    add_heading(doc, "References")
    add_para(doc, "Abby, A. (2026). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.")
    add_para(doc, "Framework F4 -- BIS Definition. MTCP Research Programme.")
    add_para(doc, "Framework F21 -- CSAS Definition. MTCP Research Programme.")
    add_para(doc, "Framework F22 -- JRS Definition. MTCP Research Programme.")
    add_para(doc, "Framework F23 -- TDS Definition. MTCP Research Programme.")
    add_para(doc, "Framework F28 -- BEC Definition. MTCP Research Programme.")
    add_para(doc, "Framework F29 -- Gate Definition. MTCP Research Programme.")
    add_para(doc, "Framework F30 -- Manifest Definition. MTCP Research Programme.")
    add_para(doc, "EU AI Act. Regulation (EU) 2024/1689. Official Journal of the European Union.")
    add_para(doc, "ISO/IEC 42001:2023. AI Management System. International Organization for Standardization.")

    path = os.path.join(OUTPUT_DIR, "Paper39_Constraint_Manifest_PUBLIC.docx")
    doc.save(path)
    print(f"  PUBLIC version saved: {path}")
    return path


# ============================================================================
# FULL (NDA) VERSION
# ============================================================================

def build_full():
    doc = Document()
    set_margins(doc)
    add_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
    add_footer(doc, "DOI: 10.17605/OSF.IO/DXGK5 | Copyright 2026 A. Abby. All rights reserved.")

    # Title
    para = doc.add_paragraph()
    run = para.add_run("Paper 39: Constraint Manifest")
    set_font(run, bold=True, size=16)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run("Portable Verified Attestation for Model Deployments")
    set_font(run, size=12)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run("FULL VERSION -- CONTAINS IMPLEMENTATION DETAILS")
    set_font(run, bold=True, size=11)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run("A. Abby -- MTCP Research Programme -- May 2026")
    set_font(run, size=10)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run("DOI: 10.17605/OSF.IO/DXGK5")
    set_font(run, size=10)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Abstract
    add_heading(doc, "1. Abstract")
    add_para(doc, "This paper defines the Constraint Manifest in full technical detail. The Constraint Manifest is a portable signed document that travels with a model deployment. It packages evaluation scores, gate decisions, and integrity attestations into a single verifiable document. Any receiving system can verify the manifest independently without contacting the issuing system.")
    add_para(doc, "The manifest provides cryptographic proof of evaluation status at time of issuance. It uses SHA-256 hashing for integrity and HMAC-SHA256 signing for authenticity. The specification covers 32 models across 13 providers. This paper presents the formal schema, verification protocol, portability constraints, and implementation architecture.")
    add_para(doc, "The Constraint Manifest operationalizes EU AI Act requirements for portable conformity evidence. It bridges the gap between evaluation (which is centralized) and deployment (which is distributed). This full version includes implementation details, database schema, API contracts, and operational procedures not available in the public specification.")

    # 2. Introduction
    add_heading(doc, "2. Introduction")
    add_para(doc, "MTCP evaluates 32 AI models across 13 providers for behavioural constraint persistence. The evaluation platform produces scores across seven frameworks. Each framework measures a different dimension of model safety and reliability. Together they form a comprehensive evaluation surface.")
    add_para(doc, "Evaluation data resides in a PostgreSQL database managed by the MTCP control plane. All scoring, gate decisions, and chain integrity records are stored centrally. This architecture works for evaluation but creates a bottleneck for deployment verification. Every system must query the central database to verify model status.")
    add_para(doc, "The Constraint Manifest decouples verification from the central database. It exports evaluation state into a self-contained document. The document can be verified anywhere without database access. This enables distributed deployment without centralized dependency.")
    add_para(doc, "The implementation uses four database tables (manifests, manifest_verifications, manifest_revocations, manifest_registry). The API exposes four endpoints for generation, verification, revocation, and registry queries. The manifest_generator.py module handles all operations. Integration with the existing control plane is through api_server.py routes.")

    # 3. Portability Problem
    add_heading(doc, "3. The Portability Problem")
    add_para(doc, "The portability problem has three dimensions. First, network dependency. Deployment endpoints may not have reliable connectivity to the evaluation authority. Air-gapped environments, edge deployments, and restricted networks cannot make API calls for verification.")
    add_para(doc, "Second, latency. Real-time verification queries add latency to deployment decisions. In automated pipelines, this latency compounds. Hundreds of models deployed across dozens of endpoints means hundreds of verification queries per deployment cycle.")
    add_para(doc, "Third, single point of failure. If the evaluation authority is unavailable, all deployment verification stops. This is unacceptable for critical infrastructure where deployment decisions must proceed regardless of external system availability.")
    add_para(doc, "The Constraint Manifest addresses all three dimensions. It eliminates network dependency through self-contained documents. It eliminates latency through local verification (milliseconds). It eliminates single point of failure because verification requires no external system. The document is the proof.")

    # 4. Formal Framework
    add_heading(doc, "4. Formal Framework")
    add_para(doc, "The Constraint Manifest is formally defined as M = (S, H, Sig) where S is the ordered schema tuple, H = SHA-256(pipe_concat(S)), and Sig = HMAC-SHA256(K, H) where K is the issuer secret key.")
    add_para(doc, "The pipe concatenation function converts each field to its string representation and joins them with the pipe character. Null fields become the literal string \"null\". JSONB fields are serialized with sorted keys. Numeric fields use their decimal string representation.")
    add_para(doc, "Verification is a three-predicate conjunction. V(M) = (H' == H) AND (Sig' == Sig) AND (today <= tds_valid_until). All three predicates must evaluate to true. Any false predicate produces REJECTED. There is no partial verification.")
    add_para(doc, "The Portability Lemma is a constraint on M. P(M) states that M is valid only when the current serving configuration matches (model_id, provider) from S. Any change to either component invalidates M. This is enforced through the revocation mechanism.")
    add_para(doc, "The signing key is derived from the environment variable MANIFEST_SECRET. In production, this is a 256-bit key stored in the secrets manager. In development, the default value is used. Key rotation requires re-signing all active manifests. This is an operational procedure, not a specification change.")

    # 5. Schema
    add_heading(doc, "5. Manifest Schema")
    add_para(doc, "The manifest schema defines 20 fields stored in the manifests table. The database type is PostgreSQL. The manifest_id is UUID v4 generated at creation time. It serves as the primary lookup key across all manifest operations.")
    add_para(doc, "Identification fields: manifest_id (UUID, unique, not null), model_id (TEXT, not null), provider (TEXT, not null), evaluation_date (DATE, not null). These four fields uniquely identify what was evaluated and when. The combination of model_id and provider forms the portability binding.")
    add_para(doc, "Score fields: bis_score (REAL, not null), bis_grade (TEXT, not null), csas_score (REAL, nullable), csas_grade (TEXT, nullable), jrs_score (REAL, nullable), tds_valid_until (DATE, not null), tds_status (TEXT, not null), acps_score (REAL, nullable), acps_grade (TEXT, nullable). BIS and TDS are always required. Others are optional.")
    add_para(doc, "Enforcement fields: gate_status (TEXT, not null -- PERMIT or DENY), gate_context (TEXT, not null), regulatory_compliance (JSONB, nullable), bec_integrity_score (REAL, not null, default 1.0), bec_chain_id (TEXT, nullable). Gate status is required for every manifest.")
    add_para(doc, "Verification fields: manifest_hash (TEXT, not null -- 64 hex chars), issuer (TEXT, not null, default 'MTCP Research Programme'), issuer_signature (TEXT, not null -- 64 hex chars). These enable independent verification. The hash covers all preceding fields. The signature covers the hash.")

    # 6. Verification Protocol
    add_heading(doc, "6. Verification Protocol")
    add_para(doc, "The verification protocol executes three checks in sequence. Each check must pass before proceeding to the next. The protocol is deterministic. Given the same inputs, it always produces the same result.")
    add_para(doc, "Hash verification reconstructs the pipe-delimited string from all schema fields in canonical order. The order is: manifest_id, model_id, provider, evaluation_date, bis_score, bis_grade, csas_score, csas_grade, jrs_score, tds_valid_until, tds_status, gate_status, gate_context, acps_score, acps_grade, regulatory_compliance, bec_integrity_score, bec_chain_id.")
    add_para(doc, "Signature verification uses HMAC-SHA256 with the shared secret key. The input is the manifest_hash string (64 hex characters). The output is compared against issuer_signature. Both sides must use the same key for verification to succeed.")
    add_para(doc, "Expiry verification compares today's date against tds_valid_until. The comparison is inclusive. A manifest is valid on the expiry date itself. It becomes invalid the day after. Time zones are not considered. Dates are compared as calendar dates.")
    add_para(doc, "Revocation checking queries the manifest_revocations table. If any record exists for the manifest_id, the manifest is revoked. Revocation overrides all other checks. A revoked manifest with valid hash, signature, and unexpired date is still REJECTED.")
    add_para(doc, "Verification results are logged in the manifest_verifications table. Every verification attempt is recorded with all check results, the overall outcome, and a timestamp. This provides a complete audit trail of who verified what and when.")

    # 7. Portability Lemma
    add_heading(doc, "7. Portability Lemma")
    add_para(doc, "The Portability Lemma is a formal constraint. It states: a manifest M issued for model X served by provider P is valid ONLY when X continues to be served by P. If X is served by Q where Q != P, M is invalid.")
    add_para(doc, "This reflects infrastructure reality. The same model weights on different provider infrastructure can exhibit different behaviour. Temperature handling, quantization, system prompts, and inference engines vary by provider. MTCP evaluation measures behaviour under specific conditions.")
    add_para(doc, "Provider change detection is outside the manifest specification. The MTCP platform monitors provider configurations and triggers revocation when changes are detected. The manifest itself does not detect provider changes. It relies on the issuing system for revocation.")
    add_para(doc, "The Portability Lemma has a critical corollary for model versioning. If a provider updates model weights without changing the model identifier, the manifest is invalid. The evaluated model no longer exists. However, the manifest system cannot detect this automatically. TDS drift detection is the mechanism that eventually catches version changes.")

    # 8. Revocation
    add_heading(doc, "8. Revocation Mechanism")
    add_para(doc, "The revocation system uses a dedicated table (manifest_revocations). Each revocation record contains revocation_id (UUID), manifest_id (the target), revoked_at (timestamp), revocation_reason (enum: score_change, provider_change, administrative), and revoked_by (authority identifier).")
    add_para(doc, "Score change revocation is triggered when any evaluation score in the manifest changes by more than the drift tolerance. This is automated by the TDS monitoring system. When TDS detects drift beyond tolerance, all active manifests for that model are revoked.")
    add_para(doc, "Provider change revocation is triggered by the Portability Lemma. When the platform detects a provider configuration change for a model, all active manifests for the old model-provider pair are revoked.")
    add_para(doc, "Administrative revocation requires manual action by an authorized user. It is used for error correction, security incidents, and policy changes. No reason beyond \"administrative\" is required. The authority to revoke is restricted to issuer-level operators.")
    add_para(doc, "Revocation cascades to the manifest_registry table. When a manifest is revoked, its registry entry is updated to status 'revoked'. This ensures that registry queries return accurate current state. The manifests table is also updated to status 'revoked'.")

    # 9. API Contract
    add_heading(doc, "9. API Contract")
    add_para(doc, "POST /api/manifest/generate accepts JSON body with model_id (required), context (required), and provider (optional). It queries the database for current scores, computes hash and signature, stores the manifest, and returns the complete document. Authentication is required.")
    add_para(doc, "POST /api/manifest/verify accepts JSON body with either manifest_id (for database lookup) or a complete manifest object (for inline verification). It executes the three-step protocol and returns detailed results. No authentication required for verification.")
    add_para(doc, "POST /api/manifest/revoke accepts JSON body with manifest_id (required) and reason (required, one of: score_change, provider_change, administrative). It creates a revocation record and updates manifest status. Authentication is required. Returns the revocation record.")
    add_para(doc, "GET /api/manifest/registry returns all manifest registry entries. Supports query parameters: model (filter by model_id), provider (filter by provider), status (filter by status: active, revoked, expired), context (filter by gate_context). Returns array of registry entries with pagination.")

    # 10. Relationship to Constructs
    add_heading(doc, "10. Relationship to MTCP Constructs")
    add_para(doc, "BIS (Framework F4) provides the mandatory behavioural integrity score. Every manifest must include bis_score and bis_grade. These are computed from the probes_500 dataset results. The BIS score is the percentage of probes where the model maintained constraints.")
    add_para(doc, "CSAS (Framework F21) provides cross-system admissibility. Manifests include CSAS when the model operates in a coordination pair (two or more models that interact). CSAS measures whether constraints persist when models communicate with each other.")
    add_para(doc, "JRS (Framework F22) provides jurisdiction resolution. Manifests include JRS when the model operates at a system boundary with jurisdiction-specific requirements. JRS measures the model's ability to apply the correct jurisdiction rules.")
    add_para(doc, "TDS (Framework F23) controls manifest lifetime. The tds_valid_until field is evaluation_date + 90 days. TDS drift detection can trigger early revocation. TDS is the only construct that directly determines how long a manifest remains valid.")
    add_para(doc, "BEC (Framework F28) attests to data integrity. The bec_integrity_score confirms that the evaluation records were not tampered with before manifest generation. A BEC integrity below 1.0 means the chain is compromised and the manifest should not have been generated.")
    add_para(doc, "The Gate (Framework F29) provides the enforcement decision. The gate_status field in the manifest is the binary PERMIT/DENY from the Admissibility Gate. A manifest can contain a DENY gate_status. This is valid. The manifest attests to the denial. It does not override it.")

    # 11. Implications
    add_heading(doc, "11. Implications")
    add_para(doc, "For procurement, the manifest is a compliance document. Organizations can require manifests as part of model procurement. The manifest provides machine-verifiable proof of evaluation. This replaces vendor self-attestation with independently verifiable evidence.")
    add_para(doc, "For supply chains, the manifest enables trust propagation. A model evaluated by MTCP carries its attestation through resellers, integrators, and platform providers. Each intermediary can verify without contacting MTCP. Trust propagates through the supply chain document.")
    add_para(doc, "For regulators, the manifest simplifies audit. Instead of requesting evaluation data from the evaluation authority, regulators can request the manifest from the deploying organization. The manifest is auditable evidence of evaluation status at deployment time.")
    add_para(doc, "For operators, the manifest enables automated deployment gates. CI/CD pipelines can verify manifests as part of deployment checks. If the manifest is valid, deployment proceeds. If invalid (expired, revoked, tampered), deployment is blocked. No human intervention required.")

    # 12. Limitations
    add_heading(doc, "12. Limitations")
    add_para(doc, "The manifest attests to a point-in-time evaluation. Model behaviour can drift between evaluations. The 90-day window is a conservative bound. High-risk deployments may require shorter windows. The specification does not currently support variable windows.")
    add_para(doc, "Key management is outside this specification. HMAC-SHA256 requires shared secret distribution. Key rotation, key escrow, and key compromise recovery are operational concerns. Organizations implementing the manifest system must solve key distribution independently.")
    add_para(doc, "Revocation depends on connectivity. Offline systems cannot check revocation status. The revocation gap between revocation issuance and offline system reconnection is a security risk. Mitigation requires periodic online checks or short manifest validity windows.")
    add_para(doc, "The manifest does not prevent deployment. It attests to evaluation status but does not enforce deployment decisions. A system may choose to deploy a model with a DENY manifest or an expired manifest. Enforcement is the responsibility of the deploying system, not the manifest.")

    # 13. Conclusion
    add_heading(doc, "13. Conclusion")
    add_para(doc, "The Constraint Manifest provides portable verified attestation for AI model deployments. It packages MTCP evaluation data into a self-contained, cryptographically signed document. The three-step verification protocol enables independent verification without issuing authority cooperation.")
    add_para(doc, "The Portability Lemma constrains validity to the evaluated model-provider pair. Revocation provides permanent invalidation. Expiry provides temporal bounds. Together, these mechanisms ensure that manifests accurately represent current evaluation state within defined bounds.")
    add_para(doc, "Implementation is complete. Four database tables store manifest lifecycle data. Four API endpoints expose manifest operations. The manifest_generator.py module provides CLI and programmatic access. Integration with the MTCP control plane is through api_server.py routes. The system is operational for all 32 evaluated models across 13 providers.")

    # References
    add_heading(doc, "References")
    add_para(doc, "Abby, A. (2026). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.")
    add_para(doc, "Framework F4 -- BIS Definition. MTCP Research Programme.")
    add_para(doc, "Framework F21 -- CSAS Definition. MTCP Research Programme.")
    add_para(doc, "Framework F22 -- JRS Definition. MTCP Research Programme.")
    add_para(doc, "Framework F23 -- TDS Definition. MTCP Research Programme.")
    add_para(doc, "Framework F28 -- BEC Definition. MTCP Research Programme.")
    add_para(doc, "Framework F29 -- Gate Definition. MTCP Research Programme.")
    add_para(doc, "Framework F30 -- Manifest Definition. MTCP Research Programme.")
    add_para(doc, "EU AI Act. Regulation (EU) 2024/1689. Official Journal of the European Union.")
    add_para(doc, "ISO/IEC 42001:2023. AI Management System. International Organization for Standardization.")
    add_para(doc, "NIST AI 100-1. AI Risk Management Framework. National Institute of Standards and Technology.")

    path = os.path.join(OUTPUT_DIR, "Paper39_Constraint_Manifest_FULL.docx")
    doc.save(path)
    print(f"  FULL version saved: {path}")
    return path


# ============================================================================
# Main
# ============================================================================

if __name__ == "__main__":
    print("Building Paper 39: Constraint Manifest")
    print("=" * 50)
    build_public()
    build_full()
    print("\nDone. Both documents saved to Desktop.")
