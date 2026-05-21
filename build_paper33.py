"""
Build script for Paper 33: ALLaM Independent Evaluation (Pre-Registration).
Generates both PUBLIC and FULL (confidential) .docx versions.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os


def set_document_formatting(doc):
    """Set page size, margins, and default font."""
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.25
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    """Add a heading with Arial font."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text):
    """Add a justified paragraph with correct formatting."""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para_format = para.paragraph_format
    para_format.line_spacing = 1.25
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return para


def add_footer(doc, footer_text):
    """Add footer text to all sections."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = footer_text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_header(doc, header_text):
    """Add header text to all sections."""
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = header_text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


# ============================================================
# PAPER CONTENT -- PUBLIC VERSION
# ============================================================

ABSTRACT_PUBLIC = (
    "This paper is a pre-registration document. "
    "It does not contain results. "
    "It documents the evaluation design for an independent constraint persistence assessment "
    "of an Arabic-first large language model. "
    "The target model has 34 billion parameters trained on 4 trillion English tokens and "
    "1.2 trillion Arabic tokens using a decoder-only architecture. "
    "No Arabic-first model has been independently evaluated for constraint persistence. "
    "Existing MTCP evaluations tested generic multilingual models on Arabic probes. "
    "This evaluation targets a model specifically designed for Arabic as its primary language. "
    "The methodology is locked here before results are obtained. "
    "The evaluation will be conducted identically regardless of outcome. "
    "Pre-registration prevents post-hoc rationalisation of results."
)

ABSTRACT_FULL = (
    "This paper is a pre-registration document. "
    "It does not contain results. "
    "It documents the evaluation design for an independent constraint persistence assessment "
    "of ALLaM, Humain's Arabic-first large language model. "
    "ALLaM has 34 billion parameters trained on 4 trillion English tokens and "
    "1.2 trillion Arabic tokens using a decoder-only architecture. "
    "No Arabic-first model has been independently evaluated for constraint persistence. "
    "Existing MTCP evaluations tested generic multilingual models on Arabic probes. "
    "This evaluation targets ALLaM specifically because it was designed for Arabic as its primary language. "
    "The methodology is locked here before results are obtained. "
    "The evaluation will be conducted identically regardless of outcome. "
    "Pre-registration prevents post-hoc rationalisation of results."
)

INTRODUCTION_PUBLIC = [
    "MTCP Paper 26 evaluated Arabic constraint persistence across five models. "
    "Nova Pro achieved 85 percent pass rate with 3 hard stops. "
    "Claude Sonnet 4.5 achieved 78 percent with 4 hard stops. "
    "Claude Haiku 4.5 achieved 100 percent but required correction on 50 percent of probes. "
    "Mistral Large achieved 83.5 percent across four temperatures with 33 hard stops.",

    "These results established that Arabic constraint persistence is measurable. "
    "They also revealed a gap. "
    "Every model tested was a generic multilingual system. "
    "None was designed with Arabic as its primary language. "
    "None had majority Arabic training data.",

    "An Arabic-first model should theoretically outperform generic multilingual models on Arabic constraints. "
    "The training distribution favours Arabic linguistic patterns. "
    "The model architecture optimises for Arabic morphology and syntax. "
    "Arabic constraint instructions should align with the model's primary processing mode.",

    "This hypothesis requires independent testing. "
    "Vendor claims about Arabic performance are not sufficient. "
    "Independent evaluation using standardised methodology produces credible evidence. "
    "MTCP provides that methodology. "
    "This paper locks the evaluation design before any results are obtained.",
]

INTRODUCTION_FULL = [
    "MTCP Paper 26 evaluated Arabic constraint persistence across five models. "
    "Nova Pro achieved 85 percent pass rate with 3 hard stops. "
    "Claude Sonnet 4.5 achieved 78 percent with 4 hard stops. "
    "Claude Haiku 4.5 achieved 100 percent but required correction on 50 percent of probes. "
    "Mistral Large achieved 83.5 percent across four temperatures with 33 hard stops.",

    "These results established that Arabic constraint persistence is measurable. "
    "They also revealed a gap. "
    "Every model tested was a generic multilingual system. "
    "None was designed with Arabic as its primary language. "
    "None had majority Arabic training data.",

    "ALLaM is Humain's Arabic-first language model. "
    "Humain is a PIF-owned national AI company in Saudi Arabia launched May 2025. "
    "CEO Tareq Amin leads the organisation. "
    "ALLaM represents a sovereign AI investment in Arabic language technology. "
    "Its architecture prioritises Arabic through training data composition.",

    "ALLaM should theoretically outperform generic multilingual models on Arabic constraints. "
    "The training distribution favours Arabic linguistic patterns. "
    "The model architecture optimises for Arabic morphology and syntax. "
    "Arabic constraint instructions should align with the model's primary processing mode.",

    "This hypothesis requires independent testing. "
    "Humain's claims about ALLaM's Arabic performance are not sufficient. "
    "Independent evaluation using standardised methodology produces credible evidence. "
    "MTCP provides that methodology. "
    "This paper locks the evaluation design before any results are obtained.",
]

ARCHITECTURE_PUBLIC = [
    "The target model uses a decoder-only transformer architecture. "
    "Total parameter count is 34 billion. "
    "Training corpus comprises approximately 4 trillion English tokens and 1.2 trillion Arabic tokens. "
    "This gives a total of 5.2 trillion training tokens.",

    "The Arabic training proportion is approximately 23 percent of total tokens. "
    "This is significantly higher than any generic multilingual model. "
    "Most multilingual models allocate less than 5 percent to Arabic. "
    "The increased Arabic proportion should produce stronger Arabic language modelling.",

    "The decoder-only architecture follows standard practices for autoregressive generation. "
    "The model processes Arabic right-to-left text through standard tokenisation. "
    "Arabic-specific tokeniser optimisation has not been independently confirmed. "
    "Tokeniser efficiency directly affects constraint instruction processing.",

    "The model is described as Arabic-first. "
    "This means Arabic was a primary design consideration from architecture selection through training. "
    "It does not mean Arabic-only. "
    "The 4 trillion English tokens ensure strong English capability. "
    "The evaluation will test whether Arabic-first design translates to Arabic-first constraint persistence.",
]

ARCHITECTURE_FULL = [
    "ALLaM uses a decoder-only transformer architecture. "
    "Total parameter count is 34 billion. "
    "Training corpus comprises approximately 4 trillion English tokens and 1.2 trillion Arabic tokens. "
    "This gives a total of 5.2 trillion training tokens.",

    "The Arabic training proportion is approximately 23 percent of total tokens. "
    "This is significantly higher than any generic multilingual model. "
    "Most multilingual models allocate less than 5 percent to Arabic. "
    "The increased Arabic proportion should produce stronger Arabic language modelling.",

    "ALLaM's decoder-only architecture follows standard practices for autoregressive generation. "
    "The model processes Arabic right-to-left text through standard tokenisation. "
    "Arabic-specific tokeniser optimisation has not been independently confirmed. "
    "Tokeniser efficiency directly affects constraint instruction processing.",

    "Humain describes ALLaM as Arabic-first. "
    "This means Arabic was a primary design consideration from architecture selection through training. "
    "It does not mean Arabic-only. "
    "The 4 trillion English tokens ensure strong English capability. "
    "The evaluation will test whether Arabic-first design translates to Arabic-first constraint persistence.",

    "Model safety is led by Emad Alghamdi at Humain. "
    "Collaboration with KFUPM and R-AGAM provides academic oversight. "
    "These relationships may facilitate evaluation access when the API becomes available.",
]

EVALUATION_DESIGN = [
    "The evaluation uses two probe sets. "
    "Combined they provide 70 unique probes evaluated at four temperatures.",

    "Probe Set 1 is LANG_Arabic_Probes_V2. "
    "This set contains 50 probes across five subtypes. "
    "The subtypes are baseline, topic pressure, code switching, formal register, and technical content. "
    "These probes were used in Paper 26 and produce known baseline results for comparison models.",

    "Probe Set 2 is LANG_Arabic_Technical_V1. "
    "This set contains 20 probes targeting Saudi regulatory and technical vocabulary. "
    "Five domains are covered with four probes each. "
    "Vision 2030 economic diversification terminology. "
    "NCA cybersecurity controls terminology. "
    "SDAIA AI governance terminology. "
    "NDMO data management terminology. "
    "Arabic-first AI architecture terminology.",

    "The technical probe set uses heavy Saudi regulatory language. "
    "Terms include Vision 2030 vocabulary, cybersecurity control frameworks, "
    "data governance standards, and AI regulatory concepts. "
    "These probes test whether the model maintains Arabic-only output "
    "under pressure from specialised technical domains.",

    "Each probe follows the standard MTCP three-turn structure. "
    "Turn 1 presents the constraint and task. "
    "Turn 2 provides correction if the model fails on turn 1. "
    "Turn 3 provides stronger correction if the model fails on turn 2. "
    "This structure measures both initial compliance and recovery capacity.",

    "Four temperatures are tested for each probe. "
    "T=0.0 provides deterministic baseline. "
    "T=0.2 introduces minimal stochastic variation. "
    "T=0.5 is the standard deployment temperature. "
    "T=0.8 tests under high stochastic pressure. "
    "Total evaluations: 70 probes at 4 temperatures equals 280 individual assessments.",
]

HYPOTHESES = [
    "Four hypotheses are registered before data collection.",

    "Hypothesis 1: The Arabic-first model will score above 90 percent BIS on Arabic probes. "
    "This corresponds to Grade A in the MTCP grading system. "
    "Rationale: Arabic-first training should produce native-level constraint processing "
    "in Arabic. "
    "The model's primary language should be its strongest language for constraint persistence.",

    "Hypothesis 2: The Arabic-first model will show lower first-turn failure rate "
    "than Nova models. "
    "Paper 26 showed Nova Micro and Nova Haiku pass 100 percent but require correction "
    "on 50 to 60 percent of probes. "
    "An Arabic-first model should process Arabic constraints correctly on first attempt. "
    "Rationale: Native Arabic processing should not require correction prompting.",

    "Hypothesis 3: The Arabic-first model will outperform all non-Arabic-first models "
    "on technical vocabulary retention. "
    "The LANG_Arabic_Technical_V1 probe set uses heavy Saudi regulatory terminology. "
    "Generic multilingual models may default to English for technical concepts. "
    "An Arabic-first model should maintain Arabic technical vocabulary without leakage. "
    "Rationale: Arabic-majority training data should include regulatory and technical Arabic.",

    "Hypothesis 4: Temperature sensitivity will be minimal for the Arabic-first model. "
    "Arabic as the primary language should produce architectural stability. "
    "Constraint violations in Arabic should not increase with temperature. "
    "Rationale: When Arabic is the dominant training language, Arabic constraint processing "
    "is architectural rather than stochastic. "
    "Architectural patterns are temperature-invariant (Paper 9).",
]

COMPARISON_FRAMEWORK = [
    "Results will be compared against Paper 26 baselines. "
    "Five comparison models provide reference points.",

    "Nova Pro (via Bedrock): 85 percent pass rate, 3 hard stops at T=0.0 on 50 probes. "
    "This represents strong but imperfect Arabic performance from a generic model.",

    "Claude Sonnet 4.5 (via Bedrock): 78 percent pass rate, 4 hard stops. "
    "This represents moderate Arabic performance with notable failure rate.",

    "Claude Haiku 4.5 (via Bedrock): 100 percent pass rate, 0 hard stops. "
    "However 50 percent of probes required correction on first turn. "
    "High pass rate masks high correction requirement.",

    "Nova Micro (via Bedrock): 100 percent pass rate, 0 hard stops. "
    "However 60 percent of probes fail on first turn. "
    "The model recovers but requires heavy correction.",

    "Mistral Large: 83.5 percent across 4 temperatures, 33 hard stops on 50 probes. "
    "This represents the weakest Arabic performance with highest hard stop rate.",

    "The Arabic-first model should outperform all five comparison models. "
    "Specifically it should show higher first-turn accuracy than Nova models "
    "and lower hard stop rate than Mistral Large. "
    "If it does not, the Arabic-first training claim requires scrutiny.",
]

METRICS = [
    "Five metrics will be computed from evaluation results.",

    "BIS (Behavioural Integrity Score): The primary metric. "
    "Percentage of probes where the model maintains the constraint across all turns. "
    "Graded A (90-100), B (80-89), C (70-79), D (60-69), F (below 60).",

    "Hard Stop Rate: Percentage of probes where the model cannot comply even after correction. "
    "A hard stop means the constraint is architecturally impossible for the model. "
    "This is the strongest failure signal.",

    "Recovery Latency: Which turn achieves compliance. "
    "Turn 1 compliance means immediate processing. "
    "Turn 2 compliance means correction was required. "
    "Turn 3 compliance means heavy correction was required. "
    "Lower latency indicates stronger native constraint processing.",

    "Temperature Sensitivity: Variance in BIS across the four temperature settings. "
    "Low sensitivity indicates architectural processing. "
    "High sensitivity indicates stochastic processing. "
    "Arabic-first models should show low sensitivity on Arabic probes.",

    "CPD (Constraint Persistence Drift): If control probes are available, CPD measures "
    "whether constraint persistence changes over repeated evaluation sessions. "
    "This metric requires multiple evaluation sessions separated by time. "
    "It may not be available in initial evaluation depending on API access duration.",
]

ACCESS_REQUIREMENTS = [
    "The evaluation requires API access to the target model. "
    "At the time of this pre-registration, no public API endpoint exists.",

    "Required access parameters: programmatic API endpoint supporting multi-turn conversation. "
    "Temperature parameter control between 0.0 and 0.8. "
    "System prompt capability for control probe insertion. "
    "Rate limits sufficient for 280 evaluations within a reasonable timeframe.",

    "The evaluation requires no model internals. "
    "No access to weights, activations, or training data is needed. "
    "Only standard inference API access is required. "
    "This makes the evaluation non-invasive and compatible with commercial API terms.",

    "Access pathway options exist. "
    "Direct API access through the model provider when endpoints become available. "
    "Collaborative access through academic partnerships. "
    "The evaluation methodology is identical regardless of access pathway.",
]

ACCESS_REQUIREMENTS_FULL = [
    "The evaluation requires API access to ALLaM. "
    "At the time of this pre-registration, no public API endpoint exists for ALLaM.",

    "Required access parameters: programmatic API endpoint supporting multi-turn conversation. "
    "Temperature parameter control between 0.0 and 0.8. "
    "System prompt capability for control probe insertion. "
    "Rate limits sufficient for 280 evaluations within a reasonable timeframe.",

    "The evaluation requires no model internals. "
    "No access to ALLaM's weights, activations, or training data is needed. "
    "Only standard inference API access is required. "
    "This makes the evaluation non-invasive and compatible with commercial API terms.",

    "Access pathway options exist. "
    "Direct API access through Humain when ALLaM endpoints become available. "
    "Collaborative access through KFUPM or R-AGAM academic partnerships. "
    "Tareq Amin's organisation has expressed interest in demonstrating ALLaM's capabilities. "
    "Independent evaluation serves both MTCP's research agenda and Humain's credibility goals. "
    "The evaluation methodology is identical regardless of access pathway.",
]

PREREGISTRATION_STATEMENT = [
    "This document constitutes a formal pre-registration. "
    "The evaluation design, hypotheses, and analysis plan are locked as of this publication date.",

    "The following commitments are made. "
    "The probe sets will not be modified after this publication. "
    "The hypotheses will not be revised after results are obtained. "
    "The analysis plan will not be altered to favour any particular outcome.",

    "Results will be reported regardless of whether they confirm or disconfirm the hypotheses. "
    "If the Arabic-first model performs below generic multilingual models, this will be reported. "
    "If the Arabic-first model shows unexpected failure patterns, these will be documented.",

    "Pre-registration serves two purposes. "
    "First, it prevents motivated reasoning in analysis. "
    "The hypotheses favour the Arabic-first model. "
    "If results are negative, pre-registration ensures they are reported honestly. "
    "Second, it demonstrates methodological rigour. "
    "The evaluation design was not influenced by preliminary results.",

    "This is the first pre-registered constraint persistence evaluation in MTCP. "
    "Previous papers reported results on models where evaluation was conducted before publication. "
    "This paper inverts the sequence. "
    "Publication precedes evaluation. "
    "This is the strongest possible methodological commitment to unbiased results.",
]

SOVEREIGN_AI_IMPLICATIONS = [
    "Sovereign AI deployment requires independent evaluation. "
    "A nation investing in language model development needs objective performance data. "
    "Vendor self-assessment is insufficient for national infrastructure decisions.",

    "Arabic-first models serve a specific strategic purpose. "
    "They reduce dependency on foreign language technology providers. "
    "They ensure Arabic language processing meets national standards. "
    "They provide sovereignty over AI capabilities in the national language.",

    "Independent evaluation of sovereign AI models must be rigorous. "
    "It must use standardised methodology. "
    "It must produce reproducible results. "
    "It must compare against international benchmarks. "
    "MTCP provides all four requirements.",

    "If the Arabic-first model achieves Grade A on Arabic constraint persistence, "
    "it validates the sovereign AI investment thesis. "
    "Arabic-first design produces Arabic-first performance. "
    "If it does not achieve Grade A, it identifies where further development is needed.",

    "Constraint persistence is particularly relevant for government deployment. "
    "Government systems must follow strict operational constraints. "
    "Language constraints ensure information remains in the national language. "
    "Format constraints ensure compliance with regulatory standards. "
    "A sovereign AI model that cannot persist government constraints is not deployment-ready.",
]

CONCLUSION_PUBLIC = [
    "This paper registers an evaluation design for an Arabic-first language model. "
    "No results are presented. "
    "The methodology is locked. "
    "The hypotheses are stated. "
    "The comparison framework is defined.",

    "The evaluation will test whether Arabic-first training produces Arabic-first constraint persistence. "
    "This is a fundamental question for sovereign AI deployment. "
    "The answer has implications for national AI investment strategies worldwide.",

    "When API access becomes available, the evaluation will be conducted exactly as described here. "
    "Results will be published as a companion paper regardless of outcome. "
    "Pre-registration ensures the scientific integrity of those results.",

    "MTCP has evaluated 35 models across 14 providers. "
    "No Arabic-first model has been independently assessed. "
    "This gap will be closed. "
    "The methodology is ready. "
    "Only access is required.",
]

CONCLUSION_FULL = [
    "This paper registers an evaluation design for ALLaM. "
    "No results are presented. "
    "The methodology is locked. "
    "The hypotheses are stated. "
    "The comparison framework is defined.",

    "The evaluation will test whether ALLaM's Arabic-first training produces Arabic-first constraint persistence. "
    "This is a fundamental question for Saudi sovereign AI deployment. "
    "The answer has implications for Humain's strategic positioning and for national AI investment.",

    "When Humain's API access becomes available, the evaluation will be conducted exactly as described here. "
    "Results will be published as a companion paper regardless of outcome. "
    "Pre-registration ensures the scientific integrity of those results.",

    "MTCP has evaluated 35 models across 14 providers. "
    "No Arabic-first model has been independently assessed. "
    "ALLaM will be the first. "
    "The methodology is ready. "
    "Access through Humain, KFUPM, or R-AGAM is the only remaining requirement.",
]

REFERENCES_PUBLIC = [
    "Abby, A. (2026a). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 9). Stochastic versus architectural constraint failure in language models. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 20). Posture Reauthorization Protocol: Wrapper-layer mitigation for architectural constraint failures. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 26). Arabic constraint persistence evaluation across five models. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 32). Remediation Effectiveness Score: Measuring intervention impact on constraint persistence failures. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Three-Layer). Three-layer constraint failure in AI systems. DOI: 10.17605/OSF.IO/DXGK5.",
    "Framework F4 -- BIS Definition. MTCP Research Programme.",
    "Framework F23 -- TDS Definition. MTCP Research Programme.",
    "Framework F24 -- CCS Definition. MTCP Research Programme.",
    "Framework F25 -- RES Definition. MTCP Research Programme.",
]

REFERENCES_FULL = [
    "Abby, A. (2026a). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 9). Stochastic versus architectural constraint failure in language models. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 20). Posture Reauthorization Protocol: Wrapper-layer mitigation for architectural constraint failures. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 26). Arabic constraint persistence evaluation across five models. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 32). Remediation Effectiveness Score: Measuring intervention impact on constraint persistence failures. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Three-Layer). Three-layer constraint failure in AI systems. DOI: 10.17605/OSF.IO/DXGK5.",
    "Framework F4 -- BIS Definition. MTCP Research Programme.",
    "Framework F23 -- TDS Definition. MTCP Research Programme.",
    "Framework F24 -- CCS Definition. MTCP Research Programme.",
    "Framework F25 -- RES Definition. MTCP Research Programme.",
    "Humain (2025). ALLaM: Arabic Language Large Model. Technical Overview.",
    "SDAIA (2023). National AI Governance Framework. Saudi Data and AI Authority.",
    "NCA (2022). Essential Cybersecurity Controls. National Cybersecurity Authority of Saudi Arabia.",
    "NDMO (2022). National Data Management Standards. National Data Management Office.",
]


def build_paper(doc, is_public=True):
    """Build the paper content into the document."""
    set_document_formatting(doc)

    if is_public:
        add_header(doc, "MTCP V1.0 -- ALLaM Pre-Registration")
    else:
        add_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
        add_footer(
            doc,
            "MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved.",
        )

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ALLaM Independent Evaluation")
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_public:
        run = subtitle.add_run(
            "Pre-Registration: Arabic-First Model Constraint Persistence Assessment"
        )
    else:
        run = subtitle.add_run(
            "Pre-Registration: ALLaM Constraint Persistence Assessment (Humain)"
        )
    run.font.name = "Arial"
    run.font.size = Pt(13)

    # Author info
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("A. Abby")
    run.font.name = "Arial"
    run.font.size = Pt(12)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run("admin@mtcp.live")
    run.font.name = "Arial"
    run.font.size = Pt(11)

    doi = doc.add_paragraph()
    doi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doi.add_run("DOI: 10.17605/OSF.IO/DXGK5")
    run.font.name = "Arial"
    run.font.size = Pt(11)

    series = doc.add_paragraph()
    series.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = series.add_run("MTCP Paper Series -- Paper 33")
    run.font.name = "Arial"
    run.font.size = Pt(11)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run("May 2026")
    run.font.name = "Arial"
    run.font.size = Pt(11)

    # 1. Abstract
    add_heading_styled(doc, "1. Abstract", level=1)
    if is_public:
        add_paragraph_justified(doc, ABSTRACT_PUBLIC)
    else:
        add_paragraph_justified(doc, ABSTRACT_FULL)

    # 2. Introduction
    add_heading_styled(doc, "2. Introduction", level=1)
    intro = INTRODUCTION_PUBLIC if is_public else INTRODUCTION_FULL
    for para_text in intro:
        add_paragraph_justified(doc, para_text)

    # 3. Architecture Summary
    if is_public:
        add_heading_styled(doc, "3. Architecture Summary", level=1)
    else:
        add_heading_styled(doc, "3. ALLaM Architecture Summary", level=1)
    arch = ARCHITECTURE_PUBLIC if is_public else ARCHITECTURE_FULL
    for para_text in arch:
        add_paragraph_justified(doc, para_text)

    # 4. Evaluation Design
    add_heading_styled(doc, "4. Evaluation Design", level=1)
    for para_text in EVALUATION_DESIGN:
        add_paragraph_justified(doc, para_text)

    # 5. Hypotheses
    add_heading_styled(doc, "5. Hypotheses", level=1)
    for para_text in HYPOTHESES:
        add_paragraph_justified(doc, para_text)

    # 6. Comparison Framework
    add_heading_styled(doc, "6. Comparison Framework", level=1)
    for para_text in COMPARISON_FRAMEWORK:
        add_paragraph_justified(doc, para_text)

    # 7. Metrics
    add_heading_styled(doc, "7. Metrics", level=1)
    for para_text in METRICS:
        add_paragraph_justified(doc, para_text)

    # 8. Access Requirements
    add_heading_styled(doc, "8. Access Requirements", level=1)
    access = ACCESS_REQUIREMENTS if is_public else ACCESS_REQUIREMENTS_FULL
    for para_text in access:
        add_paragraph_justified(doc, para_text)

    # 9. Pre-Registration Statement
    add_heading_styled(doc, "9. Pre-Registration Statement", level=1)
    for para_text in PREREGISTRATION_STATEMENT:
        add_paragraph_justified(doc, para_text)

    # 10. Implications for Sovereign AI Deployment
    add_heading_styled(doc, "10. Implications for Sovereign AI Deployment", level=1)
    for para_text in SOVEREIGN_AI_IMPLICATIONS:
        add_paragraph_justified(doc, para_text)

    # 11. Conclusion
    add_heading_styled(doc, "11. Conclusion", level=1)
    conclusion = CONCLUSION_PUBLIC if is_public else CONCLUSION_FULL
    for para_text in conclusion:
        add_paragraph_justified(doc, para_text)

    # 12. References
    add_heading_styled(doc, "12. References", level=1)
    refs = REFERENCES_PUBLIC if is_public else REFERENCES_FULL
    for ref in refs:
        add_paragraph_justified(doc, ref)


def main():
    desktop = os.path.expanduser("~/Desktop")

    # Build PUBLIC version
    doc_public = Document()
    build_paper(doc_public, is_public=True)
    public_path = os.path.join(
        desktop, "Paper33_ALLaM_Arabic_Evaluation_PUBLIC.docx"
    )
    doc_public.save(public_path)
    print(f"PUBLIC version saved: {public_path}")

    # Build FULL version
    doc_full = Document()
    build_paper(doc_full, is_public=False)
    full_path = os.path.join(
        desktop, "Paper33_ALLaM_Arabic_Evaluation_FULL.docx"
    )
    doc_full.save(full_path)
    print(f"FULL version saved: {full_path}")


if __name__ == "__main__":
    main()
