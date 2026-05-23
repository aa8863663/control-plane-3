"""
Build Paper 42: Institutional Constraint Failure (ICF).
Theoretical bridge paper. No framework number. No infrastructure.
Generates PUBLIC and FULL .docx versions.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DESKTOP = os.path.expanduser("~/Desktop")


def set_document_formatting(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.25
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return para


def add_header(doc, text):
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_footer(doc, text):
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


ABSTRACT = [
    "This paper defines Institutional Constraint Failure (ICF). "
    "ICF describes how human institutions accumulate governance assumptions. "
    "This mirrors constraint drift documented in AI models at the turn level. "
    "Institutions skip the constitutional layer for the same structural reasons.",

    "Three institutional failure modes mirror CSAS coordination failures. "
    "Institutional Dilution: governance weakens through organisational layers. "
    "Institutional Inheritance Failure: policy conditions never reach operations. "
    "Institutional Cascade: marginal assumptions compound across decision layers.",

    "The Institutional Constraint Persistence Score (ICPS) measures this. "
    "ICPS quantifies whether governance conditions persist to operational execution. "
    "It is defined by analogy to BIS but applied to institutional behaviour. "
    "This paper presents the theoretical framework. No infrastructure is required.",
]

INTRODUCTION = [
    "MTCP measures constraint persistence in AI models. "
    "A model receives a constraint. "
    "The constraint persists or fails across turns. "
    "BIS quantifies the persistence rate.",

    "Institutions exhibit the same pattern at a different scale. "
    "A governance condition is established at the policy level. "
    "The condition must persist through organisational layers to operations. "
    "It frequently does not.",

    "The mechanism is structurally identical. "
    "In AI models, constraints fail through attention decay and context displacement. "
    "In institutions, governance conditions fail through delegation assumptions. "
    "Neither system actively rejects the constraint. "
    "Both systems passively lose it.",

    "This paper formalises the parallel. "
    "It defines ICF as a pattern with three failure modes. "
    "It introduces ICPS as a measurement framework. "
    "It uses Universal Credit Q4 2012 as the canonical case study.",

    "The purpose is not analogy for its own sake. "
    "The purpose is to demonstrate that MTCP constructs apply beyond AI. "
    "Constraint persistence is a general governance problem. "
    "AI models make it measurable. "
    "Institutions make it consequential.",
]

ICF_DEFINITION = [
    "Institutional Constraint Failure is a specific pattern. "
    "Human institutions accumulate governance assumptions over time. "
    "They do not resolve those assumptions explicitly. "
    "This mirrors constraint drift in AI models.",

    "At the model level, constraint drift occurs across turns. "
    "Each turn provides an opportunity to lose the constraint. "
    "The model does not reject the constraint. "
    "It simply fails to maintain it under pressure.",

    "At the institutional level, constraint drift occurs across layers. "
    "Each organisational layer provides an opportunity to lose the condition. "
    "The institution does not reject the governance requirement. "
    "It simply fails to propagate it to operations.",

    "ICF is not negligence. "
    "It is structural. "
    "Institutions facing complexity offload governance to assumptions. "
    "This is cognitively efficient in the short term. "
    "It is structurally dangerous in the long term.",

    "The constitutional layer is where governance gets skipped. "
    "Policy establishes the requirement. "
    "The requirement should be constitutionally embedded in processes. "
    "Instead it is assumed to be understood. "
    "The assumption accumulates until failure exposes it.",
]

THREE_MODES = [
    "Three institutional failure modes mirror CSAS coordination failures.",

    "Mode 1: Institutional Dilution. "
    "Governance requirements weaken as they pass through organisational layers. "
    "Each layer reinterprets the requirement within its own context. "
    "The original precision is lost through successive reinterpretation. "
    "This mirrors CSAS Type 1: Constraint Dilution at system boundaries.",

    "Mode 2: Institutional Inheritance Failure. "
    "Governance conditions established at the policy level never reach operations. "
    "The policy exists in documentation. "
    "Operational teams have never seen it. "
    "This mirrors CSAS Type 2: Constraint Inheritance Failure.",

    "Mode 3: Institutional Cascade. "
    "A marginal governance assumption compounds across decision layers. "
    "Each layer adds its own marginal assumption. "
    "No single assumption is a failure. "
    "The accumulated assumptions produce systemic failure. "
    "This mirrors CSAS Type 4: Constraint State Drift.",

    "The correspondence is structural, not metaphorical. "
    "Both systems face the same problem. "
    "A condition established at one point must persist to another point. "
    "The mechanisms of loss are different. "
    "The pattern of loss is identical.",
]

ICPS = [
    "ICPS measures institutional constraint persistence. "
    "It quantifies whether governance conditions reach operational execution. "
    "It is defined by analogy to BIS.",

    "BIS measures the pass rate of constraint probes across temperatures. "
    "ICPS measures the pass rate of governance conditions across organisational layers. "
    "A governance condition is a probe. "
    "An organisational layer is a temperature.",

    "ICPS = 1.0 means all governance conditions reach operations intact. "
    "ICPS = 0.0 means no governance conditions reach operations. "
    "The grading scale mirrors MTCP. "
    "Grade A: 0.90 to 1.00. Grade B: 0.80 to 0.89. Grade F: below 0.60.",

    "ICPS is theoretical. "
    "No institutional evaluations have been conducted. "
    "The metric is defined for future application. "
    "Measurement would require governance audit protocols.",

    "The value of ICPS is conceptual. "
    "It makes institutional governance measurable in principle. "
    "It demonstrates that MTCP methodology generalises. "
    "Constraint persistence is not an AI-specific problem.",
]

UNIVERSAL_CREDIT = [
    "Universal Credit Q4 2012 is the canonical ICF case study. "
    "Six systems coordinated without explicit governance resolution. "
    "Each system maintained its own constraint set internally. "
    "No system owned the cross-system governance conditions.",

    "The Department for Work and Pensions operated the core benefits system. "
    "HM Revenue and Customs operated tax and income data. "
    "Local authorities operated housing and council tax systems. "
    "Each passed internal quality checks independently.",

    "No single department owned cross-system data format constraints. "
    "DWP assumed HMRC held jurisdiction over format constraints. "
    "HMRC assumed DWP held jurisdiction. "
    "Local authorities assumed the central departments had resolved it.",

    "This is Institutional Inheritance Failure (Mode 2). "
    "The governance condition was established at the policy level. "
    "It never reached operational teams in any department. "
    "No team knew it was their responsibility.",

    "The cascade followed. "
    "Data format mismatches propagated across boundaries. "
    "No entity detected or corrected the mismatch. "
    "Delays were measured in years. "
    "Costs were measured in hundreds of millions of pounds.",

    "Critically, no single system violated its own constraints. "
    "Each department maintained internal data quality. "
    "The failure was entirely at the institutional governance level. "
    "ICF describes this pattern formally.",
]

PRE_CUSTODIAL = [
    "Pre-custodial fidelity is the institutional mirror of MTCP evaluation.",

    "At the model level, MTCP asks: did the model hold its constraints. "
    "At the institutional level, ICF asks: did the institution resolve its assumptions. "
    "Both questions must be answered before custody is asserted.",

    "Pre-custodial means before asserting authority or custody. "
    "An institution that asserts authority over an AI deployment must first demonstrate ICPS. "
    "Did it resolve its governance assumptions before claiming governance.",

    "An institution with ICPS Grade F has no business governing AI deployment. "
    "Its own governance conditions do not persist to operations. "
    "Adding AI governance on top of failed institutional governance is futile. "
    "The AI constraints will suffer the same institutional drift.",

    "This is the institutional argument for MTCP. "
    "Institutions need MTCP because their own governance fails. "
    "If institutional governance persisted reliably, MTCP would be redundant. "
    "ICF demonstrates why MTCP is necessary.",
]

RELATIONSHIP = [
    "ICF relates to MTCP constructs through structural correspondence.",

    "BIS measures model-level constraint persistence. "
    "ICPS measures institution-level constraint persistence. "
    "Both use the same conceptual framework. "
    "Both face the same structural challenge.",

    "CSAS measures coordination admissibility between models. "
    "ICF measures coordination governance between institutions. "
    "CSAS failure types map directly to ICF failure modes. "
    "The mapping is structural, not analogical.",

    "JRS measures jurisdiction resolution at model boundaries. "
    "ICF exposes why jurisdiction is rarely resolved at institutional boundaries. "
    "JRS provides the framework. "
    "ICF explains why the framework is needed.",

    "COS defines the constraint as a formal object. "
    "ICF demonstrates that institutions rarely formalise their constraints. "
    "Institutional constraints remain implicit. "
    "COS provides what institutions lack.",
]

LIMITATIONS = [
    "ICF is a theoretical framework. "
    "No institutional ICPS evaluations have been conducted. "
    "The three failure modes are derived from structural analysis.",

    "The CSAS-to-ICF mapping is structural but not validated. "
    "Institutional failure may have modes not captured by the three types. "
    "The taxonomy is proposed, not confirmed.",

    "Universal Credit is one case study. "
    "Additional case studies would strengthen the framework. "
    "A single case demonstrates possibility. "
    "Multiple cases would demonstrate generality.",

    "ICPS measurement requires governance audit protocols that do not exist. "
    "Defining the metric does not make it measurable today. "
    "Measurement is deferred to future institutional audit work.",
]

CONCLUSION = [
    "Institutional Constraint Failure is a real pattern. "
    "Institutions accumulate governance assumptions over time. "
    "These assumptions mirror constraint drift in AI models. "
    "The parallel is structural, not metaphorical.",

    "Three failure modes describe how institutions lose governance conditions. "
    "Dilution, inheritance failure, and cascade. "
    "All three appear in the Universal Credit case study. "
    "All three map to CSAS coordination failure types.",

    "ICPS makes institutional governance measurable in principle. "
    "Pre-custodial fidelity asks whether institutions resolved assumptions first. "
    "This is the institutional argument for MTCP. "
    "AI governance cannot exceed institutional governance.",

    "ICF demonstrates that MTCP constructs generalise beyond AI. "
    "Constraint persistence is a universal governance problem. "
    "AI models make it measurable. "
    "Institutions make it consequential.",
]

REFERENCES = [
    "Abby, A. (2026a). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 28). Cross-System Admissibility Score. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 29). Constraint Jurisdiction Resolution. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 40). Constraint Object Specification. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 41). Legitimacy Resolution Protocol. DOI: 10.17605/OSF.IO/DXGK5.",
    "Framework F4 -- BIS Definition. MTCP Research Programme.",
    "Framework F21 -- CSAS Definition. MTCP Research Programme.",
    "Framework F22 -- JRS Definition. MTCP Research Programme.",
    "Framework F31 -- COS Definition. MTCP Research Programme.",
    "National Audit Office (2013). Universal Credit: Early Progress. HC 621.",
    "Institute for Government (2020). Universal Credit: Lessons Learned.",
]


def build_paper(is_public=True):
    doc = Document()
    set_document_formatting(doc)

    if is_public:
        add_header(doc, "MTCP V1.0 -- Institutional Constraint Failure")
    else:
        add_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
        add_footer(doc, "MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved.")

    add_heading_styled(doc, "MTCP Paper 42", level=1)
    add_heading_styled(doc, "Institutional Constraint Failure", level=2)
    add_paragraph_justified(doc, "")
    if not is_public:
        add_paragraph_justified(doc, "CONFIDENTIAL -- NDA REQUIRED")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "MTCP Research Programme -- Paper 42")
    add_paragraph_justified(doc, "Theoretical Bridge Paper")

    doc.add_page_break()

    add_heading_styled(doc, "1. Abstract", level=1)
    for para in ABSTRACT:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "2. Introduction", level=1)
    for para in INTRODUCTION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "3. Institutional Constraint Failure", level=1)
    for para in ICF_DEFINITION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "4. Three Institutional Failure Modes", level=1)
    for para in THREE_MODES:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "5. Institutional Constraint Persistence Score", level=1)
    for para in ICPS:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "6. Case Study: Universal Credit Q4 2012", level=1)
    for para in UNIVERSAL_CREDIT:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "7. Pre-Custodial Fidelity", level=1)
    for para in PRE_CUSTODIAL:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "8. Relationship to MTCP Constructs", level=1)
    for para in RELATIONSHIP:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "9. Limitations", level=1)
    for para in LIMITATIONS:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "10. Conclusion", level=1)
    for para in CONCLUSION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "11. References", level=1)
    for ref in REFERENCES:
        add_paragraph_justified(doc, ref)

    return doc


if __name__ == "__main__":
    print("Building Paper 42: Institutional Constraint Failure")
    print("=" * 60)

    doc_pub = build_paper(is_public=True)
    pub_path = os.path.join(DESKTOP, "Paper42_Institutional_Constraint_Failure_PUBLIC.docx")
    doc_pub.save(pub_path)
    print(f"PUBLIC version saved: {pub_path}")

    doc_full = build_paper(is_public=False)
    full_path = os.path.join(DESKTOP, "Paper42_Institutional_Constraint_Failure_FULL.docx")
    doc_full.save(full_path)
    print(f"FULL version saved: {full_path}")

    print("=" * 60)
    print("Done.")
