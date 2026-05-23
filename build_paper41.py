"""
Build Paper 41: Legitimacy Resolution Protocol (LRP).
Framework F32: Authority establishment and cross-regime comparability.
Generates PUBLIC and FULL .docx versions.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DESKTOP = os.path.expanduser("~/Desktop")


def set_document_formatting(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.25
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return para


def add_header(doc, text):
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_footer(doc, text):
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


ABSTRACT = [
    "This paper defines the Legitimacy Resolution Protocol (LRP). "
    "LRP establishes what makes authority over a constraint valid. "
    "It makes authority cross-regime comparable. "
    "It sits beneath JRS and completes jurisdiction resolution.",

    "Three conditions define legitimate authority. "
    "The authority source must be explicitly identified. "
    "The authority claim must be bounded by the constraint scope. "
    "The authority claim must be resolvable without trusting the issuer.",

    "The Legitimacy Score measures whether all three conditions hold. "
    "1.0 means fully resolvable legitimate authority. "
    "0.0 means authority was assumed not established. "
    "The Legitimacy Lemma constrains JRS validity to LRP satisfaction.",
]

INTRODUCTION = [
    "JRS (Paper 29) resolves jurisdiction at coordination boundaries. "
    "It determines whose constraint set governs each boundary. "
    "JRS does not determine whether the claimed authority is legitimate. "
    "It resolves jurisdiction without verifying the authority source.",

    "This gap is structural. "
    "A jurisdiction resolution that assigns governance to an unverifiable entity is not governance. "
    "It is documented assumption. "
    "The documentation creates an appearance of governance without substance.",

    "LRP closes this gap. "
    "It defines three conditions for legitimate authority. "
    "It provides a score measuring condition satisfaction. "
    "It establishes cross-regime comparability criteria. "
    "It constrains JRS validity through the Legitimacy Lemma.",

    "Without LRP, two systems in different jurisdictions cannot be compared. "
    "Each may have resolved jurisdiction internally. "
    "Neither can verify the other's authority claims. "
    "LRP provides the verification protocol.",
]

FORMAL_FRAMEWORK = [
    "LRP defines three conditions for legitimate authority.",

    "Condition 1: Source Identification. "
    "The authority source is explicitly identified and registered. "
    "An unnamed authority is not legitimate. "
    "An unregistered authority is not verifiable. "
    "Registration requires name, type, and scope boundary.",

    "Condition 2: Scope Boundedness. "
    "The authority claim is bounded by the constraint object scope. "
    "Authority cannot exceed the constraint it governs. "
    "A claim broader than its constraint is over-scoped. "
    "Over-scoped claims fail Condition 2.",

    "Condition 3: Resolvability. "
    "A receiving system can verify the claim without trusting the issuer. "
    "Verification requires a defined method. "
    "The method must be executable by any party. "
    "Trust-dependent verification fails Condition 3.",

    "All three conditions must hold simultaneously. "
    "Satisfying two of three is not legitimate authority. "
    "The conjunction is strict. "
    "Partial legitimacy does not exist in this framework.",
]

LEGITIMACY_SCORE = [
    "The Legitimacy Score quantifies condition satisfaction. "
    "Each condition contributes one-third. "
    "Score ranges from 0.0 to 1.0.",

    "Score 1.0: all three conditions satisfied. "
    "Authority is fully resolvable and legitimate. "
    "JRS scores at this boundary are governance-valid.",

    "Score 0.67: two of three conditions satisfied. "
    "Authority is partially established. "
    "JRS scores are descriptive but not fully governance-valid.",

    "Score 0.33: one of three conditions satisfied. "
    "Authority is weakly established. "
    "JRS scores are not governance-valid.",

    "Score 0.0: no conditions satisfied. "
    "Authority was assumed not established. "
    "JRS scores at this boundary are assumption, not governance.",
]

CROSS_REGIME = [
    "Two authority claims are cross-regime comparable under two conditions.",

    "First condition: both claims satisfy all three LRP conditions. "
    "Both have legitimacy score of 1.0. "
    "Partial legitimacy is not comparable.",

    "Second condition: both claims reference the same constraint object. "
    "Alternatively, both reference COS-comparable equivalents. "
    "The Constraint Identity Lemma (Paper 40) gates this determination.",

    "If both conditions hold, the claims are cross-regime comparable. "
    "This means governance findings from one regime apply to the other. "
    "CSAS scores across regime boundaries become meaningful.",

    "If either condition fails, the claims are not comparable. "
    "Governance findings do not transfer. "
    "Each regime stands alone. "
    "Cross-regime coordination requires separate evaluation.",
]

LEGITIMACY_LEMMA = [
    "The Legitimacy Lemma constrains JRS validity.",

    "JRS scores are only governance-valid when LRP conditions are satisfied. "
    "A jurisdiction resolution with unverifiable authority is not governance. "
    "It is assumption with documentation.",

    "The Persistence Lemma states that Ve above threshold invalidates admissibility. "
    "The Jurisdiction Precedence Lemma states that low JRS invalidates CSAS. "
    "The Legitimacy Lemma states that unresolved LRP invalidates JRS.",

    "Each lemma establishes a precondition for the next layer. "
    "Ve gates BIS. BIS gates CSAS. CSAS gates JRS. JRS gates LRP. "
    "LRP completes the governance stack.",

    "The practical consequence: a system with JRS Grade A but LRP score 0.0 has no governance. "
    "It has documented assumption with excellent documentation. "
    "Documentation quality does not substitute for authority verification.",
]

IMPLICATIONS = [
    "LRP affects the Admissibility Gate. "
    "The Gate now requires six inputs for a governance-valid PERMIT. "
    "BIS, CSAS, JRS, TDS, GRC, and LRP. "
    "Missing any input produces an incomplete gate decision.",

    "LRP affects sovereign deployment. "
    "Sovereign AI requires verifiable authority within national boundaries. "
    "LRP provides the verification protocol. "
    "Cross-border coordination requires mutual LRP satisfaction.",

    "LRP affects procurement. "
    "Procurement contracts reference MTCP grades. "
    "Those grades are only governance-valid under LRP. "
    "Contracts should specify LRP requirements alongside BIS requirements.",
]

LIMITATIONS = [
    "LRP is a specification. "
    "No cross-regime LRP evaluations have been conducted. "
    "The three conditions are formally specified but untested empirically.",

    "Resolvability (Condition 3) depends on verification method quality. "
    "LRP does not evaluate whether the method is actually executed. "
    "It evaluates whether the method exists and is executable. "
    "Execution verification is an operational concern.",

    "The binary conjunction (all three or nothing) may be conservative. "
    "Future versions may introduce weighted conditions. "
    "This is deferred to empirical calibration.",
]

CONCLUSION = [
    "LRP completes the governance stack beneath JRS. "
    "Authority is now formally defined as a verifiable property. "
    "Jurisdiction resolution without legitimate authority is not governance.",

    "Three conditions define legitimacy. "
    "Source identification. Scope boundedness. Resolvability. "
    "All three must hold simultaneously.",

    "The Legitimacy Lemma constrains JRS validity. "
    "The Cross-Regime Comparability Condition enables multi-regime governance. "
    "The Legitimacy Score quantifies condition satisfaction.",

    "LRP was revealed by the same operational pressure that revealed COS. "
    "The 183,924 evaluations exposed the need for authority verification. "
    "The specification emerged from practice.",
]

REFERENCES = [
    "Abby, A. (2026a). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 29). Constraint Jurisdiction Resolution. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 40). Constraint Object Specification. DOI: 10.17605/OSF.IO/DXGK5.",
    "Framework F22 -- JRS Definition. MTCP Research Programme.",
    "Framework F29 -- Gate Definition. MTCP Research Programme.",
    "Framework F31 -- COS Definition. MTCP Research Programme.",
    "Framework F32 -- LRP Definition. MTCP Research Programme.",
]


def build_paper(is_public=True):
    doc = Document()
    set_document_formatting(doc)

    if is_public:
        add_header(doc, "MTCP V1.0 -- Legitimacy Resolution Protocol")
    else:
        add_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
        add_footer(doc, "MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved.")

    add_heading_styled(doc, "MTCP Paper 41", level=1)
    add_heading_styled(doc, "Legitimacy Resolution Protocol", level=2)
    add_paragraph_justified(doc, "")
    if not is_public:
        add_paragraph_justified(doc, "CONFIDENTIAL -- NDA REQUIRED")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "MTCP Research Programme -- Paper 41")
    add_paragraph_justified(doc, "Framework F32")

    doc.add_page_break()

    add_heading_styled(doc, "1. Abstract", level=1)
    for para in ABSTRACT:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "2. Introduction", level=1)
    for para in INTRODUCTION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "3. Formal Framework", level=1)
    add_heading_styled(doc, "3.1 Three Conditions for Legitimate Authority", level=2)
    for para in FORMAL_FRAMEWORK:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "4. Legitimacy Score", level=1)
    for para in LEGITIMACY_SCORE:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "5. Cross-Regime Comparability", level=1)
    for para in CROSS_REGIME:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "6. Legitimacy Lemma", level=1)
    for para in LEGITIMACY_LEMMA:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "7. Implications", level=1)
    for para in IMPLICATIONS:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "8. Limitations", level=1)
    for para in LIMITATIONS:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "9. Conclusion", level=1)
    for para in CONCLUSION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "10. References", level=1)
    for ref in REFERENCES:
        add_paragraph_justified(doc, ref)

    return doc


if __name__ == "__main__":
    print("Building Paper 41: Legitimacy Resolution Protocol")
    print("=" * 60)

    doc_pub = build_paper(is_public=True)
    pub_path = os.path.join(DESKTOP, "Paper41_Legitimacy_Resolution_Protocol_PUBLIC.docx")
    doc_pub.save(pub_path)
    print(f"PUBLIC version saved: {pub_path}")

    doc_full = build_paper(is_public=False)
    full_path = os.path.join(DESKTOP, "Paper41_Legitimacy_Resolution_Protocol_FULL.docx")
    doc_full.save(full_path)
    print(f"FULL version saved: {full_path}")

    print("=" * 60)
    print("Done.")
