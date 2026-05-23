"""
Build MAFS Paper 1: Architectural Failure Surface.
Separate research track. Same dataset, same pseudonym.
Generates PUBLIC and FULL .docx versions.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DESKTOP = os.path.expanduser("~/Desktop")


def set_document_formatting(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.25
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return para


def add_header(doc, text):
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_footer(doc, text):
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


ABSTRACT = [
    "This paper introduces the Model Architectural Failure Surface (MAFS) programme. "
    "MAFS extracts safety-relevant findings from 183,924 constraint persistence evaluations. "
    "It classifies hard stop failures by architecture, provider, and safety tier. "
    "It maps the architectural failure surface for 32 AI models.",

    "Four constructs are defined. "
    "The Architectural Alignment Failure Index (AAFI) measures hard stop percentage. "
    "The Correction Ceiling defines maximum achievable constraint persistence. "
    "The Cascade Failure Score maps propagation through coordination chains. "
    "The Safety Constraint Taxonomy classifies 532 probes by safety relevance.",

    "All constructs were revealed by data, not designed in advance. "
    "183,924 evaluations produced the hard stop patterns. "
    "The patterns revealed the failure surface. "
    "MAFS makes that surface legible to safety researchers.",
]

INTRODUCTION = [
    "MTCP evaluates constraint persistence across 32 models and 13 providers. "
    "The primary outputs are BIS scores, grades, and governance metrics. "
    "These outputs serve governance stakeholders. "
    "Safety researchers need a different framing.",

    "Safety researchers ask different questions. "
    "Which constraints cannot be maintained regardless of intervention. "
    "Which architectural patterns produce irreversible failures. "
    "Which models create cascade risk in coordination chains.",

    "MAFS reframes the same data for safety research. "
    "Same 183,924 evaluations. Same 532 probes. Same 32 models. "
    "Different questions. Different constructs. Different audience.",

    "The key insight: hard stops are safety signals. "
    "A hard stop means the model cannot maintain a constraint. "
    "No correction, no prompt engineering, no intervention resolves it. "
    "The failure is architectural. It is in the weights.",

    "MAFS maps where these architectural failures occur. "
    "Which models. Which constraint types. Which temperatures. "
    "The map is the Architectural Failure Surface. "
    "This paper defines how to compute and interpret it.",
]

AAFI = [
    "The Architectural Alignment Failure Index measures hard stop density.",

    "AAFI equals the percentage of safety-relevant constraints that produce hard stops. "
    "A hard stop means Ve reached maximum without recovery. "
    "No correction sequence resolves the failure. "
    "The constraint is architecturally impossible for the model.",

    "AAFI grading uses an inverted scale. "
    "Grade A: AAFI at or below 0.10 (10 percent or fewer hard stops). "
    "Grade B: AAFI 0.11 to 0.20. "
    "Grade C: AAFI 0.21 to 0.30. "
    "Grade D: AAFI 0.31 to 0.40. "
    "Grade F: AAFI above 0.40.",

    "A model with AAFI Grade F has architecturally impossible safety constraints. "
    "No operational intervention resolves this. "
    "No prompt engineering. No wrapper. No correction sequence. "
    "Provider-level training intervention is the only path.",

    "AAFI is computed from existing hard stop records. "
    "No new evaluations are required. "
    "The data already exists in the MTCP database. "
    "AAFI extracts the safety signal from governance data.",
]

CORRECTION_CEILING = [
    "The Correction Ceiling defines the upper bound of intervention.",

    "For each model, a maximum constraint persistence exists. "
    "Beyond this maximum, no operational intervention produces improvement. "
    "The ceiling is structural. It derives from model weights.",

    "The ceiling is extracted from temperature sensitivity data. "
    "Performance at T=0.0 represents maximum determinism. "
    "Stochastic variance is minimised. "
    "The remaining failures are architectural.",

    "A model with a ceiling at 75 percent cannot exceed Grade C. "
    "All intervention effort beyond this point is wasted. "
    "The deployer must accept Grade C or change models. "
    "This is the ceiling concept from Paper 32 under safety framing.",

    "The gap between current BIS and ceiling is the improvement space. "
    "Within this gap, operational intervention is effective. "
    "Beyond this gap, only training-level change works. "
    "The ceiling makes this boundary explicit.",
]

CASCADE = [
    "The Cascade Failure Score maps propagation risk.",

    "When a hard-stop model sits upstream in a coordination chain, risk propagates. "
    "The upstream model produces an unrecoverable constraint violation. "
    "The downstream model receives violated output as input. "
    "The violation cascades without opportunity for correction.",

    "Cascade score quantifies this propagation probability. "
    "Higher AAFI upstream means more cascade opportunities. "
    "Each hard stop is a potential cascade source. "
    "The score reflects aggregate cascade exposure.",

    "Four classifications describe cascade severity. "
    "Contained: cascade score below 0.10. Minimal risk. "
    "Minor cascade: score 0.10 to 0.20. Occasional propagation. "
    "Moderate cascade: score 0.20 to 0.40. Frequent propagation. "
    "Severe cascade: score above 0.40. Systematic propagation.",

    "Cascade analysis uses existing CSAS data reclassified under safety framing. "
    "CSAS measures coordination admissibility for governance. "
    "Cascade analysis measures propagation risk for safety. "
    "Same data. Different question.",
]

TAXONOMY = [
    "The Safety Constraint Taxonomy classifies 532 probes by safety relevance.",

    "Tier 1: Directly Safety Critical. "
    "Probes testing content safety, harm prevention, bias, privacy, deception. "
    "Hard stops on Tier 1 probes are the highest safety concern. "
    "These map to the CONT vector.",

    "Tier 2: Governance Critical. "
    "Probes testing format compliance, domain restriction, scope limitation. "
    "Hard stops on Tier 2 probes affect regulatory compliance. "
    "These map to FORM, DOM, and SCOPE vectors.",

    "Tier 3: Operational. "
    "Probes testing language, style, persona, tone constraints. "
    "Hard stops on Tier 3 probes affect user experience. "
    "These map to the LANG vector.",

    "The taxonomy makes the dataset legible to safety researchers. "
    "It does not expose probe content. "
    "It classifies by category and tier. "
    "Safety researchers can assess risk without seeing specific probes.",
]

EMPIRICAL_PRINCIPLE = [
    "MAFS constructs were not designed in advance of measurement.",

    "183,924 evaluations produced hard stop data. "
    "The hard stop patterns revealed architectural failure surfaces. "
    "The surfaces revealed safety-relevant structure in the data. "
    "MAFS formalises what the data revealed.",

    "This is the Empirically-Derived Safety Specification Principle. "
    "Safety constructs derived from observation are more defensible. "
    "They address demonstrated problems, not hypothetical ones. "
    "The data forces the specification.",

    "Theoretically-derived safety specifications risk irrelevance. "
    "They may measure things that do not fail in practice. "
    "They may miss things that do fail in practice. "
    "MAFS measures what actually fails. "
    "This makes it operationally grounded.",
]

LIMITATIONS = [
    "MAFS Phase 1 is extracted from existing data. "
    "No new safety-specific evaluations have been conducted. "
    "AAFI is computed from hard stop records that were collected for governance.",

    "The Safety Constraint Taxonomy classifies by vector, not by content. "
    "All CONT probes are classified Tier 1. "
    "Some CONT probes may be less safety-critical than others. "
    "Finer classification requires manual review.",

    "Cascade analysis is modelled, not empirically measured. "
    "Actual cascade behaviour requires dedicated coordination evaluation. "
    "The cascade score is an estimate from existing data.",

    "AAFI assumes hard stops are architecturally permanent. "
    "Provider model updates may resolve specific hard stops. "
    "AAFI should be recomputed after model updates. "
    "The score represents a snapshot, not a permanent property.",
]

CONCLUSION = [
    "MAFS makes the safety signal in MTCP data explicit. "
    "183,924 evaluations contain architectural failure information. "
    "MAFS extracts and formalises that information.",

    "Four constructs map the failure surface. "
    "AAFI measures hard stop density. "
    "The Correction Ceiling defines intervention limits. "
    "Cascade analysis maps propagation risk. "
    "The Safety Taxonomy makes data legible.",

    "All constructs were revealed by data, not designed in advance. "
    "This makes them empirically grounded. "
    "They measure what actually fails. "
    "Safety researchers get operational findings, not theoretical speculation.",

    "MAFS is Phase 1. "
    "Future phases will add dedicated safety evaluations. "
    "Phase 1 extracts maximum value from existing data. "
    "The architectural failure surface is now mapped.",
]

REFERENCES = [
    "Abby, A. (2026a). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 28). Cross-System Admissibility Score. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 32). Remediation Effectiveness Score. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 35). Adversarial Constraint Persistence. DOI: 10.17605/OSF.IO/DXGK5.",
    "Framework F4 -- BIS Definition. MTCP Research Programme.",
    "Framework F21 -- CSAS Definition. MTCP Research Programme.",
    "Framework F25 -- RES Definition. MTCP Research Programme.",
]


def build_paper(is_public=True):
    doc = Document()
    set_document_formatting(doc)

    if is_public:
        add_header(doc, "MAFS V1.0 -- Architectural Failure Surface")
    else:
        add_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
        add_footer(doc, "MAFS V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved.")

    add_heading_styled(doc, "MAFS Paper 1", level=1)
    add_heading_styled(doc, "Model Architectural Failure Surface", level=2)
    add_paragraph_justified(doc, "")
    if not is_public:
        add_paragraph_justified(doc, "CONFIDENTIAL -- NDA REQUIRED")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "MAFS Research Programme -- Paper 1")

    doc.add_page_break()

    add_heading_styled(doc, "1. Abstract", level=1)
    for para in ABSTRACT:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "2. Introduction", level=1)
    for para in INTRODUCTION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "3. Architectural Alignment Failure Index", level=1)
    for para in AAFI:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "4. Correction Ceiling", level=1)
    for para in CORRECTION_CEILING:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "5. Cascade Failure Score", level=1)
    for para in CASCADE:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "6. Safety Constraint Taxonomy", level=1)
    for para in TAXONOMY:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "7. Empirically-Derived Safety Specification", level=1)
    for para in EMPIRICAL_PRINCIPLE:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "8. Limitations", level=1)
    for para in LIMITATIONS:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "9. Conclusion", level=1)
    for para in CONCLUSION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "10. References", level=1)
    for ref in REFERENCES:
        add_paragraph_justified(doc, ref)

    return doc


if __name__ == "__main__":
    print("Building MAFS Paper 1: Architectural Failure Surface")
    print("=" * 60)

    doc_pub = build_paper(is_public=True)
    pub_path = os.path.join(DESKTOP, "MAFS_Complete", "MAFS_Paper1_Architectural_Failure_Surface_PUBLIC.docx")
    doc_pub.save(pub_path)
    print(f"PUBLIC version saved: {pub_path}")

    doc_full = build_paper(is_public=False)
    full_path = os.path.join(DESKTOP, "MAFS_Complete", "MAFS_Paper1_Architectural_Failure_Surface_FULL.docx")
    doc_full.save(full_path)
    print(f"FULL version saved: {full_path}")

    print("=" * 60)
    print("Done.")
