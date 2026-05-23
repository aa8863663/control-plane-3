"""
Build script for Paper 37: Blockchain Evidence Chain (BEC) Integrity.
Generates both PUBLIC and FULL (confidential) .docx versions.
Framework F28: Cryptographic integrity for MTCP evaluation records.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_document_formatting(doc):
    """Set page size, margins, and default font."""
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.25
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    """Add a heading with Arial font."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text):
    """Add a justified paragraph with correct formatting."""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para_format = para.paragraph_format
    para_format.line_spacing = 1.25
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return para


def add_footer(doc, footer_text):
    """Add footer text to all sections."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = footer_text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_header(doc, header_text):
    """Add header text to all sections."""
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = header_text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


# ============================================================
# SECTION CONTENT
# ============================================================

ABSTRACT = [
    "This paper defines the Blockchain Evidence Chain (BEC) architecture "
    "for MTCP evaluation records. "
    "BEC provides cryptographic integrity guarantees through SHA-256 hash chaining. "
    "Each evaluation record includes a hash of the previous record. "
    "Any modification to a historical record invalidates all subsequent hashes.",

    "BEC is not a distributed blockchain. "
    "It is a local hash chain providing tamper evidence. "
    "It does not prevent modification. "
    "It makes modification detectable. "
    "This distinction is critical for correct interpretation.",

    "The Integrity Lemma establishes the admissibility condition for MTCP scores. "
    "Scores are admissible as compliance evidence only when BEC integrity equals 1.0. "
    "This applies to the relevant chain segment. "
    "A broken chain invalidates all scores in the affected segment.",

    "BEC integrates with the existing Sigma-Forensics framework by "
    "extending individual record hashes into a chain structure. "
    "It addresses EU AI Act Articles 13 and 17 requirements for "
    "traceable and verifiable evaluation records.",
]

INTRODUCTION = [
    "AI model evaluation produces records that inform deployment "
    "decisions worth millions of dollars. "
    "These records must be trustworthy. "
    "If evaluation grades can be modified after issuance without "
    "detection then the evaluation framework lacks credibility.",

    "MTCP has produced 183,924 evaluations across 32 models and "
    "13 providers. "
    "These evaluations generate BIS grades, CSAS scores, TDS "
    "status reports, and ACPS results. "
    "Each result informs procurement, compliance, and deployment decisions.",

    "Regulators require verifiable records. "
    "EU AI Act Article 13 mandates transparency and traceability "
    "for high-risk AI systems. "
    "Article 17 mandates quality management with documented and "
    "verifiable processes. "
    "Both articles require that evaluation results can be verified "
    "at any future point by authorized parties.",

    "Current integrity measures hash individual records using SHA-256. "
    "This detects modification of individual records. "
    "However individual hashes do not detect deletion of records "
    "or insertion of fabricated records or reordering. "
    "A gap exists between individual record integrity and complete "
    "chain integrity.",

    "BEC fills this gap. "
    "It extends individual hashes into a hash chain. "
    "Each record includes the hash of its predecessor. "
    "Modification of any record breaks all subsequent hashes. "
    "This paper defines the BEC architecture and its integration "
    "with MTCP evaluation infrastructure.",
]

INTEGRITY_PROBLEM = [
    "Evaluation records can be modified after issuance. "
    "A model that received Grade C could have its record changed "
    "to show Grade A. "
    "Without chain integrity verification this modification is "
    "undetectable if the individual record hash is also updated.",

    "The threat model includes the following scenarios. "
    "First: an internal actor modifies a grade to favor a provider. "
    "Second: an external attacker compromises the database and "
    "alters historical records. "
    "Third: a provider pressures for grade modification and the "
    "modification is applied retroactively.",

    "Individual SHA-256 hashes detect modification of a single "
    "record. "
    "But if the attacker updates both the record content and its "
    "hash then the individual hash verification passes. "
    "The modification is invisible at the single-record level.",

    "Chain integrity solves this. "
    "If record R_i is modified then its hash changes. "
    "Record R_{i+1} contains the original hash as previous_hash. "
    "The linkage between R_i and R_{i+1} breaks. "
    "The attacker must also modify R_{i+1} to fix the linkage. "
    "But modifying R_{i+1} breaks R_{i+2}. "
    "The cascade continues to the end of the chain.",

    "To falsify a single historical record the attacker must "
    "recompute hashes for every subsequent record. "
    "This is computationally feasible but operationally detectable. "
    "Periodic external verification creates checkpoints that make "
    "wholesale chain recomputation detectable.",
]

FORMAL_FRAMEWORK = [
    "A Blockchain Evidence Chain is an ordered sequence of records "
    "R_1 through R_n. "
    "Each record R_i contains a field previous_hash equal to the "
    "record_hash of R_{i-1}. "
    "The genesis record R_1 has previous_hash equal to 64 zeros.",

    "The BEC record format contains the following fields. "
    "record_id: unique identifier. "
    "chain_id: identifies which chain this record belongs to. "
    "sequence_number: position in the chain starting at 1. "
    "previous_hash: SHA-256 hash of the preceding record. "
    "model_id: identifier of the evaluated model. "
    "evaluation_type: BIS or CSAS or TDS or CCS or RES or ACPS or JRS. "
    "evaluation_data: JSON object with evaluation results. "
    "timestamp: ISO 8601 creation time. "
    "evaluator_id: identifier of the evaluator. "
    "record_hash: SHA-256 of all fields concatenated with pipe delimiters.",

    "The hash computation uses pipe-delimited concatenation of all "
    "fields converted to string representation. "
    "The hash is computed as SHA-256 of the concatenated string. "
    "The result is a 64-character lowercase hexadecimal string. "
    "This is consistent with Sigma-Forensics SHA-256 usage.",

    "The chain is append-only. "
    "Records are added to the end. "
    "Records are never modified or deleted in normal operation. "
    "The append-only property combined with hash chaining creates "
    "a tamper-evident log.",

    "Provider chains are isolated. "
    "Each provider has an independent chain. "
    "A compromise of one chain does not affect another. "
    "Cross-provider comparisons reference both chains but do not merge them.",
]

INTEGRITY_LEMMA = [
    "The Integrity Lemma defines admissibility for MTCP evaluation scores. "
    "Scores are admissible as compliance evidence only when BEC integrity equals 1.0. "
    "A broken chain invalidates all scores in the affected segment "
    "regardless of their individual validity.",

    "For evaluation record R_i to be admissible all of the following "
    "must hold. "
    "First: R_i record_hash verifies correctly against its fields. "
    "Second: R_i previous_hash equals R_{i-1} record_hash. "
    "Third: the chain is continuous from genesis through R_i. "
    "Fourth: BEC integrity for the segment from R_1 through R_i equals 1.0.",

    "If any condition fails then R_i is inadmissible. "
    "The evaluation result may still be correct. "
    "The model may still perform as recorded. "
    "But the record cannot be presented as verified evidence.",

    "A single tampered record invalidates all subsequent records. "
    "Even if subsequent records have valid individual hashes the "
    "chain is broken. "
    "Re-establishing admissibility requires re-evaluation and "
    "creation of a new chain segment.",

    "This is strict by design. "
    "Regulatory compliance requires absolute certainty about "
    "record integrity. "
    "Partial chain integrity is not acceptable for compliance evidence. "
    "The binary nature of the Integrity Lemma (1.0 or inadmissible) "
    "eliminates ambiguity.",
]

PROVIDER_ISOLATION = [
    "Evaluation records from different providers are maintained "
    "in separate chains. "
    "Each provider chain operates independently. "
    "A compromise of one chain does not affect another.",

    "This design decision reflects the operational reality of "
    "multi-provider evaluation. "
    "MTCP evaluates 32 models from 13 providers. "
    "Provider-specific chains ensure that a breach affecting one "
    "provider does not contaminate the records of all providers.",

    "Cross-provider comparisons reference records in both chains "
    "by record_id and chain_id. "
    "The comparison does not merge the chains. "
    "If Provider A chain is compromised then Provider A grades "
    "are inadmissible. "
    "Provider B grades remain fully admissible from the intact "
    "Provider B chain.",

    "Chain identifiers follow a naming convention. "
    "The primary chain is named main. "
    "Evaluation-specific chains are named by type such as csas "
    "or tds or acps. "
    "Provider-specific chains use the format provider_{name}.",
]

VERIFICATION_PROTOCOL = [
    "BEC verification follows three steps. "
    "All three must pass for the chain to be considered intact.",

    "Step 1: Hash Recomputation. "
    "For each record R_i recompute the record_hash from the "
    "fields using the SHA-256 concatenation formula. "
    "Compare the recomputed hash against the stored record_hash. "
    "Flag any mismatch.",

    "Step 2: Chain Linkage Verification. "
    "For each record R_i where i > 1 verify that previous_hash "
    "equals R_{i-1} record_hash. "
    "This confirms each record correctly references its predecessor. "
    "Flag any mismatch.",

    "Step 3: Continuity Verification. "
    "Verify sequence numbers are continuous from 1 to n with "
    "no gaps and no duplicates. "
    "Verify timestamps are non-decreasing. "
    "Flag any gaps or duplicates or timestamp inversions.",

    "A chain passes verification only if all three steps produce "
    "zero flags. "
    "Any flag means the chain is compromised at that point. "
    "The first flagged record is the earliest point of compromise.",

    "Verification results are recorded as BEC verification records. "
    "These include chain_id, verification timestamp, records checked, "
    "records valid, integrity score, and first invalid sequence.",
]

BEC_INTEGRITY_SCORE = [
    "The BEC Integrity Score is the proportion of records in the "
    "chain that pass hash verification. "
    "Score of 1.0 means the full chain is intact. "
    "Any score below 1.0 indicates tampering or corruption.",

    "The score is computed as records passing verification divided "
    "by total records in the chain. "
    "For a chain of 1000 records where 998 pass verification the "
    "integrity score is 0.998.",

    "Interpretation: 1.0 is INTACT. "
    "The chain is fully verified and all records are admissible. "
    "0.95 to 0.99 is COMPROMISED with isolated failures. "
    "Below 0.95 is BROKEN with significant failures.",

    "For compliance purposes only 1.0 is acceptable. "
    "The Integrity Lemma requires a perfect score for admissibility. "
    "The graded scale exists for diagnostic purposes. "
    "It helps identify the severity of a breach but does not "
    "create partial admissibility.",
]

RELATIONSHIP_EXISTING = [
    "BEC integrates with Sigma-Forensics by extending individual "
    "record hashes into a chain. "
    "Sigma-Forensics Stage 3 produces SHA-256 hashes for each "
    "evaluation record. "
    "BEC takes these hashes and chains them sequentially. "
    "The two systems are complementary layers.",

    "EU AI Act Article 13 requires transparency and traceability. "
    "BEC provides the traceability mechanism. "
    "Any regulator can verify chain integrity using the three-step "
    "protocol with only the SHA-256 algorithm and the stored records. "
    "No proprietary tools are required for verification.",

    "EU AI Act Article 17 requires quality management with verifiable "
    "records. "
    "BEC provides the verification layer. "
    "Records protected by BEC satisfy the requirement for tamper-evident "
    "documentation of AI system evaluation.",

    "TDS validity depends on historical record integrity. "
    "Temporal drift analysis compares current results against historical "
    "baselines. "
    "If historical records are compromised then drift analysis is "
    "unreliable. "
    "BEC ensures historical records remain trustworthy for drift comparison.",

    "CSAS cross-system evaluations reference multiple provider chains. "
    "BEC provider isolation ensures that a compromise in one chain is contained. "
    "Cross-system evaluations referencing intact chains from other providers remain valid.",
]

IMPLICATIONS = [
    "Regulatory admissibility: evaluation results presented to "
    "regulators for conformity assessment carry BEC verification. "
    "Regulators can independently verify that results have not "
    "been modified since the original evaluation.",

    "Audit trail: BEC provides a complete and verifiable audit "
    "trail for all evaluation activity. "
    "Auditors can verify not only individual records but the "
    "complete sequence of evaluations.",

    "Procurement evidence: when organizations use MTCP grades "
    "for procurement decisions BEC ensures those grades are "
    "the original unmodified evaluation results. "
    "This protects both buyer and provider.",

    "Legal defensibility: in disputes over model performance, "
    "BEC-verified records provide cryptographic evidence. "
    "They prove what was recorded at the time of evaluation. "
    "This is stronger evidence than unverified database records.",

    "Provider accountability: providers cannot claim their model "
    "was unfairly graded if the evaluation chain is intact. "
    "The chain proves the grade was recorded as evaluated "
    "and has not been modified since.",
]

LIMITATIONS = [
    "BEC is not distributed. "
    "It runs on a single system. "
    "It does not provide the tamper resistance of a distributed "
    "ledger with multiple independent validators.",

    "BEC provides local integrity only. "
    "A sufficiently privileged administrator could theoretically "
    "rebuild the entire chain from scratch with false data. "
    "External checkpointing mitigates this risk but does not "
    "eliminate it.",

    "BEC does not prevent initial falsification. "
    "If an evaluation result is false when first recorded then "
    "BEC faithfully protects that false result. "
    "BEC guarantees records have not been changed. "
    "It does not guarantee records were correct when created.",

    "Verification is computationally O(n) in chain length. "
    "Very long chains require proportional verification time. "
    "Periodic checkpointing can reduce verification scope "
    "but introduces trust in the checkpoint.",

    "BEC does not address availability. "
    "If chain records are deleted entirely rather than modified "
    "then BEC cannot verify what is no longer present. "
    "Regular backups and external replication address this.",
]

CONCLUSION = [
    "BEC provides the integrity layer beneath all MTCP evaluation "
    "constructs. "
    "It extends existing SHA-256 hashing from individual records "
    "into a chain structure that detects any modification to "
    "any historical record.",

    "The Integrity Lemma establishes the compliance standard. "
    "Chain integrity score must equal 1.0 for records to be admissible. "
    "This is strict and binary by design.",

    "BEC is a specification. "
    "Implementation exists as private intellectual property. "
    "No distributed ledger is required. "
    "The architecture requires only SHA-256 hashing and "
    "sequential append-only storage.",

    "The three-step verification protocol enables independent "
    "verification by any party with access to the chain data. "
    "No proprietary tools or algorithms are required. "
    "This supports regulatory independence in conformity assessment.",
]

REFERENCES = [
    "Abby, A. (2026). Multi-Turn Constraint Persistence (MTCP) V1.0. "
    "DOI: 10.17605/OSF.IO/DXGK5.",

    "Abby, A. (2026, Framework F4). Behavioural Integrity Score Definition. "
    "MTCP Research Programme.",

    "Abby, A. (2026, Framework F9). Sigma-Forensics Standard. "
    "MTCP Research Programme.",

    "Abby, A. (2026, Framework F21). Cross-System Admissibility Score. "
    "MTCP Research Programme.",

    "Abby, A. (2026, Framework F23). Temporal Drift Status. "
    "MTCP Research Programme.",

    "Abby, A. (2026, Framework F27). Adversarial Constraint Persistence Score. "
    "MTCP Research Programme.",

    "Abby, A. (2026, Framework F28). Blockchain Evidence Chain Definition. "
    "MTCP Research Programme.",

    "EU AI Act. Regulation (EU) 2024/1689 of the European Parliament "
    "and of the Council. Official Journal of the European Union.",

    "NIST. Secure Hash Standard (SHS). FIPS PUB 180-4. "
    "National Institute of Standards and Technology.",

    "ISO/IEC 42001:2023. Information technology -- Artificial intelligence "
    "-- Management system. International Organization for Standardization.",

    "Merkle, R. (1979). A certified digital signature. Advances in "
    "Cryptology -- CRYPTO 89.",
]


def build_public():
    """Build the PUBLIC version of Paper 37."""
    doc = Document()
    set_document_formatting(doc)
    add_header(doc, "MTCP V1.0 -- Blockchain Evidence Chain")

    # Title page
    add_heading_styled(doc, "MTCP Paper 37", level=1)
    add_heading_styled(doc, "Blockchain Evidence Chain (BEC) Integrity", level=2)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(
        doc,
        "This paper defines the Blockchain Evidence Chain architecture "
        "for MTCP evaluation records. "
        "BEC provides cryptographic integrity through SHA-256 hash chaining. "
        "It is a specification. "
        "Implementation exists as private intellectual property. "
        "No distributed ledger is required."
    )

    doc.add_page_break()

    # 1. Abstract
    add_heading_styled(doc, "1. Abstract", level=1)
    for para_text in ABSTRACT:
        add_paragraph_justified(doc, para_text)

    # 2. Introduction
    add_heading_styled(doc, "2. Introduction", level=1)
    for para_text in INTRODUCTION:
        add_paragraph_justified(doc, para_text)

    # 3. The Integrity Problem
    add_heading_styled(doc, "3. The Integrity Problem", level=1)
    for para_text in INTEGRITY_PROBLEM:
        add_paragraph_justified(doc, para_text)

    # 4. Formal Framework
    add_heading_styled(doc, "4. Formal Framework", level=1)
    for para_text in FORMAL_FRAMEWORK:
        add_paragraph_justified(doc, para_text)

    # 5. Integrity Lemma
    add_heading_styled(doc, "5. The Integrity Lemma", level=1)
    for para_text in INTEGRITY_LEMMA:
        add_paragraph_justified(doc, para_text)

    # 6. Provider Isolation Guarantee
    add_heading_styled(doc, "6. Provider Isolation Guarantee", level=1)
    for para_text in PROVIDER_ISOLATION:
        add_paragraph_justified(doc, para_text)

    # 7. Verification Protocol
    add_heading_styled(doc, "7. Verification Protocol", level=1)
    for para_text in VERIFICATION_PROTOCOL:
        add_paragraph_justified(doc, para_text)

    # 8. BEC Integrity Score
    add_heading_styled(doc, "8. BEC Integrity Score", level=1)
    for para_text in BEC_INTEGRITY_SCORE:
        add_paragraph_justified(doc, para_text)

    # 9. Relationship to Existing Constructs
    add_heading_styled(doc, "9. Relationship to Existing Constructs", level=1)
    for para_text in RELATIONSHIP_EXISTING:
        add_paragraph_justified(doc, para_text)

    # 10. Implications
    add_heading_styled(doc, "10. Implications", level=1)
    for para_text in IMPLICATIONS:
        add_paragraph_justified(doc, para_text)

    # 11. Limitations
    add_heading_styled(doc, "11. Limitations", level=1)
    for para_text in LIMITATIONS:
        add_paragraph_justified(doc, para_text)

    # 12. Conclusion
    add_heading_styled(doc, "12. Conclusion", level=1)
    for para_text in CONCLUSION:
        add_paragraph_justified(doc, para_text)

    # 13. References
    add_heading_styled(doc, "13. References", level=1)
    for ref in REFERENCES:
        add_paragraph_justified(doc, ref)

    output_path = os.path.expanduser(
        "~/Desktop/Paper37_BEC_Integrity_PUBLIC.docx"
    )
    doc.save(output_path)
    print(f"PUBLIC version saved: {output_path}")
    return output_path


def build_full():
    """Build the FULL (confidential) version of Paper 37."""
    doc = Document()
    set_document_formatting(doc)
    add_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
    add_footer(
        doc,
        "MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved."
    )

    # Title page
    add_heading_styled(doc, "MTCP Paper 37", level=1)
    add_heading_styled(doc, "Blockchain Evidence Chain (BEC) Integrity", level=2)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "Status: CONFIDENTIAL -- NDA REQUIRED")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(
        doc,
        "This paper defines the Blockchain Evidence Chain architecture "
        "for MTCP evaluation records. "
        "BEC provides cryptographic integrity through SHA-256 hash chaining. "
        "This version names specific models and providers. "
        "Implementation exists as private intellectual property. "
        "No distributed ledger is required."
    )

    doc.add_page_break()

    # 1. Abstract
    add_heading_styled(doc, "1. Abstract", level=1)
    for para_text in ABSTRACT:
        add_paragraph_justified(doc, para_text)
    add_paragraph_justified(
        doc,
        "BEC protects the integrity of 183,924 evaluations across 35 "
        "models from 14 providers including OpenAI, Anthropic, Google, "
        "Meta, Mistral, Cohere, Amazon, xAI, Cerebras, Fireworks, "
        "NVIDIA, Moonshot, and IBM."
    )

    # 2. Introduction
    add_heading_styled(doc, "2. Introduction", level=1)
    for para_text in INTRODUCTION:
        add_paragraph_justified(doc, para_text)
    add_paragraph_justified(
        doc,
        "The full model roster includes GPT-4o, GPT-4o-mini, GPT-3.5, "
        "Claude Sonnet 4, Claude Haiku 4.5, Gemini 2.0 Flash, "
        "Gemini 2.5, Gemini 2.5 Flash, Gemma 2 27B, Gemma 3 27B, "
        "Llama 3.1 8B, Llama 3.1 70B, Llama 3.3 70B, Llama 4 Scout, "
        "Llama 4 Maverick, Mistral Large, Mistral Small, "
        "Magistral Medium, Command R+, DeepSeek R1, Grok 3 Mini, "
        "Nova Micro, Nova Lite, Nova Pro, NVIDIA Llama 3.1, "
        "Phi-4, Kimi, Qwen 2.5, Qwen 3 8B, Qwen 3 32B, and Granite. "
        "Each model has evaluation records protected by BEC chain integrity."
    )

    # 3. The Integrity Problem
    add_heading_styled(doc, "3. The Integrity Problem", level=1)
    for para_text in INTEGRITY_PROBLEM:
        add_paragraph_justified(doc, para_text)
    add_paragraph_justified(
        doc,
        "A concrete example: GPT-4o received BIS Grade B (87.3%). "
        "If an attacker modifies this record to show Grade A (92.1%) "
        "a procurement decision based on that grade is compromised. "
        "BEC makes this modification detectable."
    )

    # 4. Formal Framework
    add_heading_styled(doc, "4. Formal Framework", level=1)
    for para_text in FORMAL_FRAMEWORK:
        add_paragraph_justified(doc, para_text)
    add_paragraph_justified(
        doc,
        "The implementation uses PostgreSQL with JSONB storage for "
        "evaluation_data. "
        "Chains are identified by text string (main, csas, tds, etc). "
        "The bec_records table stores all chain records. "
        "The bec_chains table stores chain metadata and integrity scores. "
        "The bec_verifications table stores verification audit logs."
    )

    # 5. Integrity Lemma
    add_heading_styled(doc, "5. The Integrity Lemma", level=1)
    for para_text in INTEGRITY_LEMMA:
        add_paragraph_justified(doc, para_text)
    add_paragraph_justified(
        doc,
        "In practice this means that if any record in the main chain "
        "is tampered with then all subsequent BIS, CSAS, TDS, CCS, "
        "RES, and ACPS scores are inadmissible until the chain is "
        "re-established through re-evaluation. "
        "For 183,924 records this is a significant operational consequence."
    )

    # 6. Provider Isolation Guarantee
    add_heading_styled(doc, "6. Provider Isolation Guarantee", level=1)
    for para_text in PROVIDER_ISOLATION:
        add_paragraph_justified(doc, para_text)
    add_paragraph_justified(
        doc,
        "Provider chains are defined for: OpenAI (GPT models), "
        "Anthropic (Claude models), Google (Gemini and Gemma models), "
        "Meta (Llama models), Mistral (Mistral and Magistral models), "
        "Cohere (Command R+), Amazon (Nova models), xAI (Grok models), "
        "Cerebras (hosted Llama), Fireworks (DeepSeek R1), "
        "NVIDIA (hosted models), Moonshot (Kimi), and IBM (Granite). "
        "A breach of the OpenAI chain does not affect Claude grades."
    )

    # 7. Verification Protocol
    add_heading_styled(doc, "7. Verification Protocol", level=1)
    for para_text in VERIFICATION_PROTOCOL:
        add_paragraph_justified(doc, para_text)

    # 8. BEC Integrity Score
    add_heading_styled(doc, "8. BEC Integrity Score", level=1)
    for para_text in BEC_INTEGRITY_SCORE:
        add_paragraph_justified(doc, para_text)

    # 9. Relationship to Existing Constructs
    add_heading_styled(doc, "9. Relationship to Existing Constructs", level=1)
    for para_text in RELATIONSHIP_EXISTING:
        add_paragraph_justified(doc, para_text)

    # 10. Implications
    add_heading_styled(doc, "10. Implications", level=1)
    for para_text in IMPLICATIONS:
        add_paragraph_justified(doc, para_text)
    add_paragraph_justified(
        doc,
        "For Sirar by stc and other procurement partners the BEC "
        "verification provides confidence that grades presented in "
        "evidence packs are the original evaluation results. "
        "This is critical for board-level deployment decisions."
    )

    # 11. Limitations
    add_heading_styled(doc, "11. Limitations", level=1)
    for para_text in LIMITATIONS:
        add_paragraph_justified(doc, para_text)

    # 12. Conclusion
    add_heading_styled(doc, "12. Conclusion", level=1)
    for para_text in CONCLUSION:
        add_paragraph_justified(doc, para_text)

    # 13. References
    add_heading_styled(doc, "13. References", level=1)
    for ref in REFERENCES:
        add_paragraph_justified(doc, ref)

    output_path = os.path.expanduser(
        "~/Desktop/Paper37_BEC_Integrity_FULL.docx"
    )
    doc.save(output_path)
    print(f"FULL version saved: {output_path}")
    return output_path


if __name__ == "__main__":
    print("Building Paper 37: Blockchain Evidence Chain (BEC) Integrity")
    print("=" * 60)
    pub_path = build_public()
    full_path = build_full()
    print("=" * 60)
    print("Build complete.")
    print(f"  PUBLIC: {pub_path}")
    print(f"  FULL:   {full_path}")
