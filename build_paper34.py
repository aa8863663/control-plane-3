"""
Build script for Paper 34: Regulatory Compliance Mapping.
Generates both PUBLIC and FULL (confidential) .docx versions.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_document_formatting(doc):
    """Set page size, margins, and default font."""
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.25
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    """Add a heading with Arial font."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text):
    """Add a justified paragraph with correct formatting."""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para_format = para.paragraph_format
    para_format.line_spacing = 1.25
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return para


def add_footer(doc, footer_text):
    """Add footer text to all sections."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = footer_text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_header(doc, header_text):
    """Add header text to all sections."""
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = header_text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_table_row(table, cells_text, bold=False):
    """Add a row to a table with formatted text."""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.bold = bold
    return row


def create_compliance_table(doc, include_all=False):
    """Create the regulatory compliance matrix table."""
    headers = [
        "Jurisdiction", "Regulation / Article",
        "Min Grade", "Metric Evidence",
        "Min Score", "Context"
    ]

    rows_common = [
        ["EU", "AI Act Art 9", "B", "BIS + CPD",
         "BIS >= 80%, CPD < 30pp", "High-risk"],
        ["EU", "AI Act Art 10", "C", "CPD",
         "CPD assessed", "High-risk"],
        ["EU", "AI Act Art 13", "B", "Sigma-Forensics",
         "5-stage audit", "High-risk"],
        ["EU", "AI Act Art 17", "B", "TDS",
         "TDS < 5pp", "High-risk"],
        ["EU", "AI Act Art 72", "N/A", "Ve",
         "Ve >= 2 logged", "All"],
        ["UK / DSIT", "AI Safety", "B", "BIS",
         "BIS >= 80%", "High-risk"],
        ["UK / DSIT", "Post-market", "B", "Sigma + TDS",
         "Audit + TDS < 5pp", "High-risk"],
        ["UK / DSIT", "Monitoring", "B", "TDS",
         "90-day validity", "High-risk"],
    ]

    rows_jurisdiction_specific = [
        ["Saudi / NDMO", "AI Governance", "B", "BIS (Arabic)",
         "BIS >= 90% Arabic", "Standard"],
        ["Saudi / NDMO", "Cross-system", "B", "CSAS",
         "CSAS Grade B", "Coordinated"],
        ["Saudi / NDMO", "Data Residency", "N/A", "JRS",
         "JRS = 1.0", "Cross-border"],
        ["Saudi / NCA", "Cybersecurity", "A", "CSAS",
         "CSAS Grade A", "Critical infra"],
        ["Saudi / NCA", "Adversarial", "A", "ACPS",
         "ACPS Grade A", "Government"],
        ["Saudi / NCA", "Arabic (Citizen)", "A", "BIS (Arabic)",
         "BIS >= 90%", "Citizen-facing"],
        ["Singapore / MAS", "Model Risk", "B", "BIS",
         "BIS >= 80%", "Financial"],
        ["Singapore / MAS", "Multilingual", "B", "BIS (zh,ms,ta)",
         "BIS >= 80% each", "Financial"],
        ["Singapore / MAS", "Temporal", "B", "TDS",
         "TDS < 5pp, 60-day", "Financial"],
        ["Singapore / MAS", "Conflict", "B", "CCS",
         "CCS >= 0.8", "Multi-constraint"],
    ]

    all_rows = rows_common if not include_all else rows_common + rows_jurisdiction_specific

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.bold = True

    for row_data in all_rows:
        add_table_row(table, row_data)

    return table


# ============================================================
# PAPER CONTENT
# ============================================================

ABSTRACT_PUBLIC = (
    "This paper maps MTCP evaluation outputs to regulatory requirements "
    "across multiple jurisdictions. "
    "MTCP produces grades, metrics, and audit evidence. "
    "Regulators require documented compliance evidence. "
    "This paper bridges the gap between evaluation output and regulatory input. "
    "It defines which MTCP metrics satisfy which regulatory articles. "
    "It establishes minimum grade thresholds per jurisdiction. "
    "It introduces validity periods for time-sensitive compliance evidence. "
    "The mapping covers the EU AI Act, UK DSIT AI Safety Framework, "
    "and additional jurisdictions where AI governance requirements apply. "
    "This is a technical mapping exercise. "
    "It is not legal advice. "
    "Jurisdiction-specific interpretation requires local legal counsel."
)

ABSTRACT_FULL = (
    "This paper maps MTCP evaluation outputs to regulatory requirements "
    "across five jurisdictions. "
    "MTCP produces grades, metrics, and audit evidence. "
    "Regulators require documented compliance evidence. "
    "This paper bridges the gap between evaluation output and regulatory input. "
    "It defines which MTCP metrics satisfy which regulatory articles. "
    "It establishes minimum grade thresholds per jurisdiction. "
    "It introduces validity periods for time-sensitive compliance evidence. "
    "The mapping covers the EU AI Act, NDMO Saudi AI Governance, "
    "NCA Saudi Cybersecurity Controls, MAS Singapore Model Risk Management, "
    "and UK DSIT AI Safety Framework. "
    "R-AGAM collaboration provides context for Saudi sovereign deployment. "
    "This is a technical mapping exercise. "
    "It is not legal advice. "
    "Jurisdiction-specific interpretation requires local legal counsel."
)

INTRODUCTION_PUBLIC = [
    "MTCP evaluates AI systems for constraint persistence. "
    "The evaluation produces structured outputs. "
    "BIS grades quantify constraint maintenance. "
    "Ve identifies hard failure events. "
    "CPD measures probe-structure familiarity. "
    "TDS tracks temporal drift. "
    "CSAS measures cross-system admissibility. "
    "Sigma-Forensics produces auditable evidence chains.",

    "These outputs have regulatory relevance. "
    "But evaluation without regulatory mapping is incomplete. "
    "A deployer receives a Grade B result. "
    "The deployer asks whether this satisfies regulatory requirements. "
    "Without a formal mapping, MTCP cannot answer this question.",

    "This paper provides the formal mapping. "
    "It converts MTCP from an evaluation tool into a compliance instrument. "
    "Each regulatory requirement is matched to specific MTCP evidence. "
    "Each jurisdiction receives its own mapping detail. "
    "Minimum grades are defined per deployment context.",

    "The methodology draws on 183,924 evaluations across 35 models "
    "and 14 providers. "
    "Arabic evaluation evidence from Papers 26 and 27 informs "
    "jurisdiction-specific language requirements. "
    "Multilingual evidence across 11 languages and 4 script families "
    "informs multilingual jurisdiction requirements.",
]

INTRODUCTION_FULL = [
    "MTCP evaluates AI systems for constraint persistence. "
    "The evaluation produces structured outputs. "
    "BIS grades quantify constraint maintenance. "
    "Ve identifies hard failure events. "
    "CPD measures probe-structure familiarity. "
    "TDS tracks temporal drift. "
    "CSAS measures cross-system admissibility. "
    "Sigma-Forensics produces auditable evidence chains.",

    "These outputs have regulatory relevance. "
    "But evaluation without regulatory mapping is incomplete. "
    "A deployer receives a Grade B result. "
    "The deployer asks whether this satisfies regulatory requirements. "
    "Without a formal mapping, MTCP cannot answer this question.",

    "This paper provides the formal mapping. "
    "It converts MTCP from an evaluation tool into a compliance instrument. "
    "Each regulatory requirement is matched to specific MTCP evidence. "
    "Five jurisdictions receive detailed mapping. "
    "Minimum grades are defined per deployment context.",

    "The methodology draws on 183,924 evaluations across 35 models "
    "and 14 providers. "
    "Arabic evaluation evidence from Papers 26 and 27 informs "
    "NDMO and NCA requirements. "
    "Multilingual evidence across 11 languages and 4 script families "
    "informs MAS Singapore requirements. "
    "R-AGAM collaboration provides the Saudi sovereign deployment context "
    "for NDMO and NCA mappings.",
]

COMPLIANCE_PROBLEM_PUBLIC = [
    "Evaluation without regulatory mapping creates a documentation gap. "
    "The evaluator knows the model's grade. "
    "The regulator knows the compliance requirement. "
    "Neither can confirm whether the grade satisfies the requirement "
    "without a formal mapping between the two.",

    "This gap has commercial consequences. "
    "Procurement teams cannot specify MTCP requirements in contracts "
    "without knowing which grades regulators accept. "
    "Vendors cannot claim compliance readiness without knowing "
    "which metrics each jurisdiction requires. "
    "Auditors cannot verify compliance without knowing the threshold.",

    "The gap also has safety consequences. "
    "A model deployed at Grade C in a jurisdiction requiring Grade B "
    "creates unquantified risk. "
    "Without the mapping, this risk is invisible. "
    "The deployer believes they have evaluated the system. "
    "The evaluation is genuine but insufficient for the jurisdiction.",

    "Framework F26 and this paper close the gap. "
    "They define the minimum MTCP evidence required per jurisdiction. "
    "They establish validity periods for time-sensitive evidence. "
    "They enable procurement, audit, and compliance teams "
    "to connect evaluation outputs to regulatory requirements.",
]

COMPLIANCE_PROBLEM_FULL = [
    "Evaluation without regulatory mapping creates a documentation gap. "
    "The evaluator knows the model's grade. "
    "The regulator knows the compliance requirement. "
    "Neither can confirm whether the grade satisfies the requirement "
    "without a formal mapping between the two.",

    "This gap has commercial consequences. "
    "Procurement teams cannot specify MTCP requirements in contracts "
    "without knowing which grades regulators accept. "
    "Vendors cannot claim compliance readiness without knowing "
    "which metrics each jurisdiction requires. "
    "Auditors cannot verify compliance without knowing the threshold.",

    "The gap also has safety consequences. "
    "A model deployed at Grade C in a jurisdiction requiring Grade B "
    "creates unquantified risk. "
    "Without the mapping, this risk is invisible. "
    "The deployer believes they have evaluated the system. "
    "The evaluation is genuine but insufficient for the jurisdiction.",

    "The gap is particularly acute for cross-jurisdiction deployment. "
    "A model serving Saudi users must satisfy NDMO requirements. "
    "The same model serving NCA-regulated systems must satisfy NCA requirements. "
    "If deployed in Singapore financial services, MAS requirements apply. "
    "Each jurisdiction has different thresholds for the same metrics.",

    "Framework F26 and this paper close the gap. "
    "They define the minimum MTCP evidence required per jurisdiction. "
    "They establish validity periods for time-sensitive evidence. "
    "They enable procurement, audit, and compliance teams "
    "to connect evaluation outputs to regulatory requirements.",
]

EU_MAPPING = [
    "The EU AI Act (Regulation 2024/1689) establishes requirements "
    "for high-risk AI systems. "
    "MTCP maps to five articles. "
    "Each article requires specific MTCP evidence.",

    "Article 9 requires a risk management system. "
    "The system must identify and mitigate risks throughout the lifecycle. "
    "MTCP provides BIS evaluation as risk identification evidence. "
    "Grade B or higher demonstrates acceptable constraint persistence risk. "
    "CPD below 30pp demonstrates that results are not contaminated "
    "by probe-structure familiarity. "
    "Both conditions must be satisfied simultaneously.",

    "Article 10 requires appropriate data governance. "
    "MTCP provides CPD assessment as evidence. "
    "CPD identifies whether probe familiarity inflates BIS scores. "
    "A completed CPD assessment demonstrates data governance awareness. "
    "The assessment must be completed regardless of CPD level.",

    "Article 13 requires transparency. "
    "High-risk AI systems must enable interpretation of outputs. "
    "Sigma-Forensics provides the transparency evidence. "
    "The 5-stage audit produces SHA-256 hashed transcripts. "
    "The hash chain demonstrates evaluation integrity. "
    "The audit report demonstrates reproducible methodology.",

    "Article 17 requires a quality management system. "
    "Systematic procedures must maintain quality over time. "
    "TDS monitoring demonstrates ongoing quality measurement. "
    "TDS Stable (below 2pp) or Marginal (2-5pp) is acceptable. "
    "TDS Significant (5-10pp) triggers re-evaluation. "
    "TDS Critical (above 10pp) indicates quality degradation. "
    "Validity period is 90 days.",

    "Article 72 requires serious incident reporting. "
    "MTCP provides Ve as the incident trigger. "
    "Ve >= 2 means the model failed to recover after correction. "
    "This constitutes a serious incident indicator. "
    "Any Ve >= 2 event requires logging regardless of BIS grade. "
    "This is a trigger condition, not a grade threshold.",
]

NDMO_MAPPING = [
    "The NDMO AI Governance Framework establishes requirements "
    "for AI systems deployed in Saudi Arabia. "
    "Three governance areas map to MTCP metrics.",

    "Arabic language persistence is the primary requirement. "
    "AI systems serving Arabic-speaking populations must maintain "
    "constraints in Arabic at elevated thresholds. "
    "Minimum 90 percent BIS on LANG_Arabic probes is required. "
    "This threshold exceeds the standard Grade B (80 percent). "
    "The elevation reflects citizen-facing service expectations.",

    "Cross-system coordination requires CSAS evaluation. "
    "AI systems coordinating with other AI systems must demonstrate "
    "constraint persistence across system boundaries. "
    "CSAS Grade B minimum for coordinated deployments. "
    "All CSAS components (BPR, CVe, CIR, CAF) must reach Grade B.",

    "Data residency requires explicit jurisdiction assignment. "
    "AI systems processing Saudi data across borders must have "
    "explicit jurisdiction assignment per JRS evaluation. "
    "JRS = 1.0 (explicit) is required. "
    "JRS below 1.0 (assumed or ambiguous) fails the requirement. "
    "This is a binary pass-fail condition.",
]

NCA_MAPPING = [
    "The NCA establishes cybersecurity controls for AI systems "
    "in critical national infrastructure and government services. "
    "Three control areas map to MTCP metrics.",

    "Critical national infrastructure requires Grade A assurance. "
    "CSAS evaluation at Grade A is mandatory. "
    "Grade A means 90 percent or above on all CSAS components. "
    "No operational remediation path is acceptable. "
    "The system must achieve Grade A without intervention. "
    "This is the most stringent requirement in the framework.",

    "Government-facing systems require adversarial evaluation. "
    "ACPS (Adversarial Constraint Persistence Score) is mandatory. "
    "The evaluation tests resistance to adversarial constraint manipulation. "
    "The full adversarial probe set must be completed. "
    "Grade A on adversarial probes is required.",

    "Citizen-facing services require Arabic Grade A persistence. "
    "BIS evaluation on LANG_Arabic probes must achieve Grade A. "
    "Grade A means 90 percent or above on Arabic probes. "
    "This requirement combines with NDMO 90 percent threshold. "
    "NCA adds the Grade A formal requirement for citizen-facing services.",

    "Saudi sovereign deployment context is informed by R-AGAM "
    "collaboration on Arabic-first model evaluation. "
    "The intersection of NDMO governance and NCA cybersecurity "
    "creates the most demanding Arabic AI evaluation requirements globally.",
]

MAS_MAPPING = [
    "The MAS Model Risk Management guidelines establish requirements "
    "for AI models in financial services. "
    "Four requirement areas map to MTCP metrics.",

    "Financial services deployment requires Grade B minimum. "
    "BIS evaluation must demonstrate 80 percent or above. "
    "Grade B is the deployment threshold. "
    "Grade C or below requires RES evaluation before deployment. "
    "Financial services cannot deploy Grade C models without "
    "documented remediation evidence.",

    "Multilingual constraint persistence reflects Singapore demographics. "
    "BIS evaluation on Mandarin (zh), Malay (ms), and Tamil (ta) "
    "probe sets is required. "
    "Grade B minimum on each language individually. "
    "A model passing overall but failing one language does not comply. "
    "Each language is independently assessed.",

    "Temporal monitoring uses a shortened validity period. "
    "Financial model validation requires 60-day re-evaluation. "
    "This is shorter than the standard 90 days. "
    "TDS must remain below 5pp within the 60-day window. "
    "TDS Significant triggers immediate re-evaluation. "
    "TDS Critical triggers immediate deployment suspension.",

    "Multi-constraint evaluation addresses financial compliance overlap. "
    "Financial services commonly apply multiple constraints simultaneously. "
    "CCS evaluation (Framework F24) measures conflict resolution. "
    "CCS must achieve 0.8 or above. "
    "CCS evaluation is required wherever multiple compliance "
    "constraints apply to the same model.",
]

UK_MAPPING = [
    "The UK DSIT AI Safety Framework establishes requirements "
    "for high-risk AI deployment in the United Kingdom. "
    "Three requirement areas map to MTCP metrics.",

    "High-risk deployment requires Grade B minimum. "
    "BIS evaluation must demonstrate 80 percent or above. "
    "Grade B is the deployment threshold for high-risk systems. "
    "The UK framework does not currently differentiate "
    "between critical infrastructure and standard high-risk.",

    "Post-market monitoring requires combined evidence. "
    "Sigma-Forensics audit provides baseline evaluation evidence. "
    "TDS monitoring provides ongoing surveillance evidence. "
    "The combination demonstrates both initial compliance and "
    "continued monitoring throughout deployment.",

    "Ongoing monitoring uses 90-day validity periods. "
    "TDS must remain below 5pp for continued deployment approval. "
    "Re-evaluation is required every 90 days. "
    "The 90-day cycle aligns with quarterly reporting expectations. "
    "Missed re-evaluation invalidates compliance evidence.",
]

EVIDENCE_REQUIREMENTS = [
    "Each jurisdiction requires specific evidence documentation. "
    "The evidence must be traceable to specific evaluation events. "
    "Sigma-Forensics hash chains provide the traceability mechanism.",

    "EU evidence requirements include the full Sigma-Forensics report, "
    "BIS scores across all four temperatures, CPD assessment results, "
    "TDS monitoring logs, and Ve event records. "
    "All evidence must be available for regulatory inspection.",

    "UK evidence requirements include the Sigma-Forensics audit report, "
    "BIS evaluation results, TDS monitoring records, "
    "and re-evaluation schedules. "
    "Evidence retention period aligns with deployment lifecycle.",

    "Evidence validity expires per jurisdiction-specific timelines. "
    "EU and UK use 90-day validity. "
    "Expired evidence requires re-evaluation before continued reliance. "
    "Evidence cannot be extended without fresh evaluation.",
]

EVIDENCE_REQUIREMENTS_FULL = [
    "Each jurisdiction requires specific evidence documentation. "
    "The evidence must be traceable to specific evaluation events. "
    "Sigma-Forensics hash chains provide the traceability mechanism.",

    "EU evidence requirements include the full Sigma-Forensics report, "
    "BIS scores across all four temperatures, CPD assessment results, "
    "TDS monitoring logs, and Ve event records. "
    "All evidence must be available for regulatory inspection.",

    "NDMO evidence requirements include Arabic BIS evaluation results, "
    "CSAS evaluation reports for coordinated systems, "
    "JRS assessment for cross-border deployments, "
    "and Sigma-Forensics audit for evaluation integrity. "
    "Evidence must be available in Arabic and English.",

    "NCA evidence requirements include CSAS Grade A certification, "
    "ACPS adversarial evaluation results, Arabic BIS Grade A results, "
    "and Sigma-Forensics audit with full hash chain. "
    "Evidence retention is indefinite for critical infrastructure.",

    "MAS evidence requirements include BIS evaluation in all four "
    "official language probe sets, TDS monitoring at 60-day intervals, "
    "CCS evaluation for multi-constraint deployments, "
    "and model validation documentation per MAS guidelines.",

    "UK evidence requirements include the Sigma-Forensics audit report, "
    "BIS evaluation results, TDS monitoring records, "
    "and re-evaluation schedules. "
    "Evidence retention period aligns with deployment lifecycle.",

    "Evidence validity expires per jurisdiction-specific timelines. "
    "MAS and NCA use 60-day validity. "
    "EU, NDMO, and UK use 90-day validity. "
    "Expired evidence requires re-evaluation before continued reliance. "
    "Evidence cannot be extended without fresh evaluation.",
]

IMPLICATIONS = [
    "The regulatory mapping has three primary implications. "
    "These affect procurement, certification, and cross-jurisdiction deployment.",

    "Procurement teams can now specify MTCP requirements in contracts. "
    "A contract can require Grade B BIS for EU deployment. "
    "A contract can require CSAS Grade A for NCA-regulated systems. "
    "A contract can require 60-day TDS validity for financial models. "
    "These requirements are measurable and verifiable.",

    "Certification becomes possible through MTCP evidence. "
    "A model with Grade B BIS, CPD below 30pp, and Sigma-Forensics audit "
    "has documented evidence for EU Articles 9, 10, and 13. "
    "The evidence does not guarantee certification. "
    "But it provides the structured documentation that certification requires.",

    "Cross-jurisdiction deployment requires the most stringent standard. "
    "A model deployed in EU and Singapore must satisfy both. "
    "MAS requires 60-day TDS validity. "
    "EU requires 90-day validity. "
    "The model must satisfy the 60-day requirement to cover both. "
    "The strictest requirement always governs multi-jurisdiction deployment.",
]

IMPLICATIONS_FULL = [
    "The regulatory mapping has three primary implications. "
    "These affect procurement, certification, and cross-jurisdiction deployment.",

    "Procurement teams can now specify MTCP requirements in contracts. "
    "A contract can require Grade B BIS for EU deployment. "
    "A contract can require CSAS Grade A for NCA-regulated systems. "
    "A contract can require 60-day TDS validity for MAS financial models. "
    "A contract can require Arabic Grade A for NCA citizen-facing services. "
    "These requirements are measurable and verifiable.",

    "Certification becomes possible through MTCP evidence. "
    "A model with Grade B BIS, CPD below 30pp, and Sigma-Forensics audit "
    "has documented evidence for EU Articles 9, 10, and 13. "
    "A model with CSAS Grade A and ACPS Grade A has documented evidence "
    "for NCA critical infrastructure requirements. "
    "The evidence does not guarantee certification. "
    "But it provides the structured documentation that certification requires.",

    "Cross-jurisdiction deployment requires the most stringent standard. "
    "A model deployed in Saudi Arabia and Singapore must satisfy both. "
    "NCA requires CSAS Grade A. "
    "MAS requires 60-day TDS validity. "
    "NDMO requires Arabic BIS at 90 percent. "
    "MAS requires multilingual BIS at Grade B for three languages. "
    "The model must satisfy all applicable requirements simultaneously. "
    "The strictest requirement always governs multi-jurisdiction deployment.",

    "Saudi sovereign deployment context creates the highest combined threshold. "
    "NDMO Arabic 90 percent plus NCA Grade A plus CSAS Grade A "
    "represents the most demanding AI compliance profile in the mapping. "
    "R-AGAM collaboration positions MTCP as the evaluation instrument "
    "for this profile.",
]

LIMITATIONS = [
    "This paper has significant limitations that must be stated clearly.",

    "This is not legal advice. "
    "MTCP is a technical evaluation methodology. "
    "It does not interpret law. "
    "The regulatory mappings represent technical interpretations "
    "of regulatory requirements. "
    "Legal counsel must validate these interpretations.",

    "Jurisdiction-specific interpretation varies. "
    "Regulators retain sole interpretive authority over their own requirements. "
    "A mapping that appears clear may be interpreted differently "
    "by the relevant regulatory authority. "
    "MTCP evidence supports compliance claims but does not guarantee them.",

    "Regulatory requirements change. "
    "The EU AI Act is newly implemented. "
    "Guidance and interpretation will evolve. "
    "The mappings in this paper reflect the regulatory text "
    "as of May 2026. "
    "Future regulatory guidance may alter the mapping.",

    "MTCP evidence is necessary but not sufficient. "
    "No single evaluation framework covers all compliance requirements. "
    "MTCP covers constraint persistence. "
    "Other aspects of AI compliance (fairness, bias, privacy, security) "
    "require separate evaluation frameworks. "
    "This paper addresses only constraint persistence compliance.",

    "Legal review is required before reliance. "
    "No deployment decision should rely solely on this mapping. "
    "Legal counsel in each target jurisdiction must review "
    "the specific regulatory requirements and MTCP evidence "
    "before any compliance claim is made.",
]

CONCLUSION_PUBLIC = [
    "This paper provides the formal mapping between MTCP evaluation "
    "outputs and regulatory requirements. "
    "The mapping covers the EU AI Act and UK DSIT AI Safety Framework "
    "with additional jurisdiction support.",

    "Grade B BIS is the most common deployment threshold. "
    "TDS monitoring with jurisdiction-specific validity periods "
    "provides ongoing compliance evidence. "
    "Sigma-Forensics provides the audit trail that regulators require.",

    "MTCP moves from evaluation tool to compliance instrument. "
    "Evaluation identifies the problem. "
    "The regulatory mapping connects the evaluation to the obligation. "
    "Deployers can now trace from MTCP grade to regulatory requirement.",

    "This mapping requires legal review before reliance. "
    "It is a technical exercise that supports compliance claims. "
    "It does not constitute legal advice. "
    "Local legal counsel must validate jurisdiction-specific interpretations.",
]

CONCLUSION_FULL = [
    "This paper provides the formal mapping between MTCP evaluation "
    "outputs and regulatory requirements across five jurisdictions. "
    "EU, NDMO Saudi, NCA Saudi, MAS Singapore, and UK DSIT "
    "each receive detailed metric-to-requirement mappings.",

    "Grade B BIS is the most common deployment threshold. "
    "Grade A is required for NCA critical infrastructure and citizen-facing. "
    "Arabic BIS at 90 percent is required for NDMO and NCA. "
    "Multilingual BIS at Grade B is required for MAS. "
    "TDS monitoring with jurisdiction-specific validity periods "
    "provides ongoing compliance evidence. "
    "Sigma-Forensics provides the audit trail that regulators require.",

    "MTCP moves from evaluation tool to compliance instrument. "
    "Evaluation identifies the problem. "
    "The regulatory mapping connects the evaluation to the obligation. "
    "Deployers can now trace from MTCP grade to regulatory requirement.",

    "Saudi sovereign deployment represents the highest compliance threshold. "
    "The combination of NDMO governance, NCA cybersecurity, and Arabic "
    "language requirements creates a uniquely demanding profile. "
    "R-AGAM collaboration positions MTCP as the evaluation instrument "
    "for this profile.",

    "This mapping requires legal review before reliance. "
    "It is a technical exercise that supports compliance claims. "
    "It does not constitute legal advice. "
    "Local legal counsel must validate jurisdiction-specific interpretations.",
]

REFERENCES = [
    "Abby, A. (2026a). Multi-Turn Constraint Persistence (MTCP). "
    "DOI: 10.17605/OSF.IO/DXGK5.",

    "Abby, A. (2026, Paper 13). EU AI Act compliance mapping "
    "for constraint persistence evaluation. "
    "DOI: 10.17605/OSF.IO/DXGK5.",

    "Abby, A. (2026, Paper 26). Arabic language constraint "
    "persistence evaluation. DOI: 10.17605/OSF.IO/DXGK5.",

    "Abby, A. (2026, Paper 27). Arabic multi-dialect constraint "
    "persistence analysis. DOI: 10.17605/OSF.IO/DXGK5.",

    "Abby, A. (2026, Paper 32). Remediation Effectiveness Score. "
    "DOI: 10.17605/OSF.IO/DXGK5.",

    "EU AI Act. Regulation (EU) 2024/1689 of the European Parliament "
    "and of the Council. Official Journal of the European Union.",

    "Framework F2 -- Ve Metric. MTCP Research Programme.",

    "Framework F3 -- CPD Definition. MTCP Research Programme.",

    "Framework F4 -- BIS Definition. MTCP Research Programme.",

    "Framework F9 -- Sigma-Forensics Standard. MTCP Research Programme.",

    "Framework F21 -- CSAS Definition. MTCP Research Programme.",

    "Framework F22 -- JRS Definition. MTCP Research Programme.",

    "Framework F23 -- TDS Definition. MTCP Research Programme.",

    "Framework F24 -- CCS Definition. MTCP Research Programme.",

    "Framework F25 -- RES Definition. MTCP Research Programme.",

    "Framework F26 -- Regulatory Mapping. MTCP Research Programme.",
]

REFERENCES_FULL = REFERENCES + [
    "MAS. Model Risk Management Guidelines. "
    "Monetary Authority of Singapore.",

    "NDMO. National AI Governance Framework. "
    "National Data Management Office, Kingdom of Saudi Arabia.",

    "NCA. National Cybersecurity Controls. "
    "National Cybersecurity Authority, Kingdom of Saudi Arabia.",

    "UK DSIT. AI Safety Framework. Department for Science, "
    "Innovation and Technology, United Kingdom.",

    "R-AGAM. Arabic-first AI Governance and Model Evaluation. "
    "King Fahd University of Petroleum and Minerals.",
]


def build_public():
    """Build the PUBLIC version of Paper 34."""
    doc = Document()
    set_document_formatting(doc)
    add_header(doc, "MTCP V1.0 -- Regulatory Compliance Mapping")

    # Title page
    add_heading_styled(doc, "MTCP Paper 34", level=1)
    add_heading_styled(doc, "Regulatory Compliance Mapping", level=2)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "Status: PUBLIC -- Not for OSF (pending legal review)")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(
        doc,
        "DISCLAIMER: This document is a technical mapping exercise. "
        "It is not legal advice. "
        "Jurisdiction-specific interpretation requires local legal counsel. "
        "MTCP provides evaluation evidence that may support compliance claims. "
        "It does not guarantee compliance."
    )

    doc.add_page_break()

    # Abstract
    add_heading_styled(doc, "1. Abstract", level=1)
    add_paragraph_justified(doc, ABSTRACT_PUBLIC)

    # Introduction
    add_heading_styled(doc, "2. Introduction", level=1)
    for para_text in INTRODUCTION_PUBLIC:
        add_paragraph_justified(doc, para_text)

    # The Compliance Problem
    add_heading_styled(doc, "3. The Compliance Problem", level=1)
    for para_text in COMPLIANCE_PROBLEM_PUBLIC:
        add_paragraph_justified(doc, para_text)

    # Regulatory Compliance Matrix
    add_heading_styled(doc, "4. Regulatory Compliance Matrix", level=1)
    add_paragraph_justified(
        doc,
        "The following table maps regulatory requirements to MTCP metrics. "
        "This public version includes EU and UK mappings. "
        "Additional jurisdiction mappings are available under NDA."
    )
    add_paragraph_justified(doc, "")
    create_compliance_table(doc, include_all=False)

    # EU AI Act Mapping
    add_heading_styled(doc, "5. EU AI Act Mapping", level=1)
    for para_text in EU_MAPPING:
        add_paragraph_justified(doc, para_text)

    # Placeholder sections for non-public jurisdictions
    add_heading_styled(doc, "6. Additional Jurisdiction Mappings", level=1)
    add_paragraph_justified(
        doc,
        "Additional jurisdiction mappings cover Arabic language requirements, "
        "cybersecurity controls, financial services model risk, "
        "and multilingual constraint persistence. "
        "These mappings are available under NDA due to "
        "jurisdiction-specific claims requiring legal review."
    )

    # UK DSIT Mapping
    add_heading_styled(doc, "7. UK DSIT Mapping", level=1)
    for para_text in UK_MAPPING:
        add_paragraph_justified(doc, para_text)

    # Evidence Requirements
    add_heading_styled(doc, "8. Evidence Requirements Per Jurisdiction", level=1)
    for para_text in EVIDENCE_REQUIREMENTS:
        add_paragraph_justified(doc, para_text)

    # Implications
    add_heading_styled(doc, "9. Implications", level=1)
    for para_text in IMPLICATIONS:
        add_paragraph_justified(doc, para_text)

    # Limitations
    add_heading_styled(doc, "10. Limitations", level=1)
    for para_text in LIMITATIONS:
        add_paragraph_justified(doc, para_text)

    # Conclusion
    add_heading_styled(doc, "11. Conclusion", level=1)
    for para_text in CONCLUSION_PUBLIC:
        add_paragraph_justified(doc, para_text)

    # References
    add_heading_styled(doc, "12. References", level=1)
    for ref in REFERENCES:
        add_paragraph_justified(doc, ref)

    output_path = os.path.expanduser(
        "~/Desktop/Paper34_Regulatory_Mapping_PUBLIC.docx"
    )
    doc.save(output_path)
    print(f"PUBLIC version saved: {output_path}")
    return output_path


def build_full():
    """Build the FULL (confidential) version of Paper 34."""
    doc = Document()
    set_document_formatting(doc)
    add_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
    add_footer(
        doc,
        "MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved."
    )

    # Title page
    add_heading_styled(doc, "MTCP Paper 34", level=1)
    add_heading_styled(doc, "Regulatory Compliance Mapping", level=2)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "Status: CONFIDENTIAL -- NDA REQUIRED")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(
        doc,
        "DISCLAIMER: This document is a technical mapping exercise. "
        "It is not legal advice. "
        "Jurisdiction-specific interpretation requires local legal counsel. "
        "MTCP provides evaluation evidence that may support compliance claims. "
        "It does not guarantee compliance."
    )

    doc.add_page_break()

    # Abstract
    add_heading_styled(doc, "1. Abstract", level=1)
    add_paragraph_justified(doc, ABSTRACT_FULL)

    # Introduction
    add_heading_styled(doc, "2. Introduction", level=1)
    for para_text in INTRODUCTION_FULL:
        add_paragraph_justified(doc, para_text)

    # The Compliance Problem
    add_heading_styled(doc, "3. The Compliance Problem", level=1)
    for para_text in COMPLIANCE_PROBLEM_FULL:
        add_paragraph_justified(doc, para_text)

    # Regulatory Compliance Matrix
    add_heading_styled(doc, "4. Regulatory Compliance Matrix", level=1)
    add_paragraph_justified(
        doc,
        "The following table maps all five jurisdictions "
        "to MTCP metric requirements."
    )
    add_paragraph_justified(doc, "")
    create_compliance_table(doc, include_all=True)

    # EU AI Act Mapping
    add_heading_styled(doc, "5. EU AI Act Mapping", level=1)
    for para_text in EU_MAPPING:
        add_paragraph_justified(doc, para_text)

    # NDMO Saudi Mapping
    add_heading_styled(doc, "6. NDMO Saudi Mapping", level=1)
    for para_text in NDMO_MAPPING:
        add_paragraph_justified(doc, para_text)

    # NCA Saudi Mapping
    add_heading_styled(doc, "7. NCA Saudi Mapping", level=1)
    for para_text in NCA_MAPPING:
        add_paragraph_justified(doc, para_text)

    # MAS Singapore Mapping
    add_heading_styled(doc, "8. MAS Singapore Mapping", level=1)
    for para_text in MAS_MAPPING:
        add_paragraph_justified(doc, para_text)

    # UK DSIT Mapping
    add_heading_styled(doc, "9. UK DSIT Mapping", level=1)
    for para_text in UK_MAPPING:
        add_paragraph_justified(doc, para_text)

    # Evidence Requirements
    add_heading_styled(doc, "10. Evidence Requirements Per Jurisdiction", level=1)
    for para_text in EVIDENCE_REQUIREMENTS_FULL:
        add_paragraph_justified(doc, para_text)

    # Implications
    add_heading_styled(doc, "11. Implications", level=1)
    for para_text in IMPLICATIONS_FULL:
        add_paragraph_justified(doc, para_text)

    # Limitations
    add_heading_styled(doc, "12. Limitations", level=1)
    for para_text in LIMITATIONS:
        add_paragraph_justified(doc, para_text)

    # Conclusion
    add_heading_styled(doc, "13. Conclusion", level=1)
    for para_text in CONCLUSION_FULL:
        add_paragraph_justified(doc, para_text)

    # References
    add_heading_styled(doc, "14. References", level=1)
    for ref in REFERENCES_FULL:
        add_paragraph_justified(doc, ref)

    output_path = os.path.expanduser(
        "~/Desktop/Paper34_Regulatory_Mapping_FULL.docx"
    )
    doc.save(output_path)
    print(f"FULL version saved: {output_path}")
    return output_path


if __name__ == "__main__":
    print("Building Paper 34: Regulatory Compliance Mapping")
    print("=" * 60)
    public_path = build_public()
    full_path = build_full()
    print("=" * 60)
    print("Build complete.")
    print(f"  PUBLIC: {public_path}")
    print(f"  FULL:   {full_path}")
