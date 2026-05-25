"""Build Paper 47: Permanence Architecture. Framework F37."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
DESKTOP = os.path.expanduser("~/Desktop")

def setup(doc):
    for s in doc.sections:
        s.page_width=Inches(8.5);s.page_height=Inches(11)
        s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Inches(1)
    st=doc.styles["Normal"];st.font.name="Arial";st.font.size=Pt(12)
    st.paragraph_format.line_spacing=1.25;st.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

def h(doc,t,l=1):
    hd=doc.add_heading(t,level=l)
    for r in hd.runs: r.font.name="Arial"

def p(doc,t):
    pa=doc.add_paragraph(t);pa.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    pa.paragraph_format.line_spacing=1.25
    for r in pa.runs: r.font.name="Arial";r.font.size=Pt(12)

def hdr(doc,t):
    for s in doc.sections:
        pa=s.header.paragraphs[0] if s.header.paragraphs else s.header.add_paragraph()
        pa.text=t;pa.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in pa.runs: r.font.name="Arial";r.font.size=Pt(9)

def ftr(doc,t):
    for s in doc.sections:
        pa=s.footer.paragraphs[0] if s.footer.paragraphs else s.footer.add_paragraph()
        pa.text=t;pa.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in pa.runs: r.font.name="Arial";r.font.size=Pt(9)

C = [
 ("1. Abstract",[
  "This paper defines Permanence Architecture as a formal property. "
  "Most governance stacks re-derive constraints at each application boundary. "
  "Re-derivation is where semantic drift and provenance breaks occur. "
  "Permanence Architecture binds constraint, provenance, and evidence to the payload.",
  "Three guarantees define permanence. "
  "Constraint Binding prevents separation of constraint from evaluation record. "
  "Provenance Continuity preserves the full custody chain in the payload. "
  "Crypto-Agility names the algorithm as a field value not a hardcoded primitive.",
  "The Permanence Lemma establishes that re-derived records are not governance evidence. "
  "Only records persisting intact across boundaries are admissible. "
  "The Boundary Persistence Score measures this property per deployment.",
 ]),
 ("2. Introduction",[
  "Governance stacks face a structural problem at system boundaries. "
  "A constraint is defined in System A. "
  "An evaluation is conducted in System A. "
  "The result must be consumed in System B.",
  "Most architectures re-derive the constraint and provenance at System B. "
  "They query the originating system or reconstruct from application state. "
  "This re-derivation introduces three failure modes. "
  "Semantic drift: the reconstructed constraint differs from the original. "
  "Provenance break: the chain of custody is not verifiable at B. "
  "Crypto coupling: the verification algorithm is embedded in the schema.",
  "MTCP solves all three through Permanence Architecture. "
  "The BEC hash chain binds constraint, provenance, and evidence to each record. "
  "The record travels with the payload across boundaries. "
  "No re-derivation is required. No reconstruction is needed.",
 ]),
 ("3. Three Permanence Guarantees",[
  "Guarantee 1: Constraint Binding.",
  "The constraint object is bound to the evaluation record at registration time. "
  "The binding is cryptographic. "
  "The payload_binding_hash includes the constraint_id. "
  "Separating constraint from record breaks the hash. "
  "Any receiving system can verify the binding independently.",
  "Guarantee 2: Provenance Continuity.",
  "The full chain of custody lives in the payload. "
  "From constraint issuance through evaluation through attestation. "
  "The BEC chain preserves every step with previous_hash links. "
  "No step can be removed without breaking the chain. "
  "The payload is self-proving.",
  "Guarantee 3: Crypto-Agility.",
  "The cryptographic algorithm is a field value. "
  "The crypto_standard field names which algorithm was used. "
  "Migrating from SHA-256 to BLAKE3 changes a field value. "
  "It does not change the schema. "
  "The governance semantics are untouched by algorithm evolution.",
  "This is the difference between hardcoded and declared cryptography. "
  "Hardcoded systems must be re-engineered when algorithms change. "
  "Declared systems update a field value. "
  "Permanence Architecture uses declared cryptography by design.",
 ]),
 ("4. The Permanence Lemma",[
  "Records requiring re-derivation at boundaries are not governance evidence.",
  "They are governance assumptions. "
  "An assumption can be wrong without detection. "
  "Evidence cannot be wrong without breaking verification.",
  "A constraint manifest reconstructed from application state at System B proves nothing. "
  "It proves that System B believes the constraint existed. "
  "It does not prove the constraint was enforced. "
  "It does not prove the evaluation was conducted. "
  "It does not prove the chain of custody is intact.",
  "Only records persisting intact across boundaries are admissible. "
  "Intact means: hash verifiable, constraint bound, provenance preserved. "
  "Any break in any of these three disqualifies the record. "
  "The Permanence Lemma is a binary determination.",
 ]),
 ("5. Boundary Persistence Score",[
  "The score measures deployment-level permanence compliance.",
  "For a deployment with N evaluation records across M system boundaries. "
  "Boundary Persistence Score equals verified records divided by total records. "
  "Verified means payload_binding_hash is correct and permanence_verified is true.",
  "Score 1.0: every record persists intact across all boundaries. "
  "The deployment has full permanence. "
  "All governance evidence is admissible.",
  "Score below 1.0: some records were reconstructed. "
  "Those records are governance assumptions not evidence. "
  "The deployment has partial permanence. "
  "Only verified records are admissible.",
 ]),
 ("6. Implementation",[
  "Permanence Architecture is implemented through three infrastructure changes.",
  "First: payload_binding_hash added to every BEC record. "
  "Computed as SHA-256 of record_hash, model_id, evaluation_type, and chain_id. "
  "Binds the evaluation to its context cryptographically.",
  "Second: permanence_verified flag on each record. "
  "Set to true when the binding hash is computed and stored. "
  "Records without verification are legacy and not permanence-compliant.",
  "Third: crypto_standard field in the Constraint Manifest. "
  "Names the algorithm used for all hashes in the manifest. "
  "Enables algorithm migration without schema change.",
 ]),
 ("7. Relationship to Existing Constructs",[
  "Permanence Architecture extends BEC (Framework F28). "
  "BEC provides chain integrity. Permanence provides boundary persistence. "
  "Both are required. Chain integrity without boundary persistence is local-only.",
  "Permanence Architecture extends the Constraint Manifest (Framework F30). "
  "The manifest is a permanent record by design. "
  "Permanence Architecture formalises why this matters.",
  "Permanence Architecture extends COS (Framework F31). "
  "COS defines the constraint object. "
  "Permanence Architecture binds that object to evaluation records permanently.",
 ]),
 ("8. Limitations",[
  "Permanence Architecture requires infrastructure support. "
  "Systems that cannot store or forward payload-bound records cannot participate. "
  "Legacy systems may require adapter layers.",
  "The payload_binding_hash increases record size. "
  "Each record gains 64 characters of hash data. "
  "This is negligible for modern systems but non-zero.",
 ]),
 ("9. Conclusion",[
  "Most governance breaks at system boundaries. "
  "Re-derivation is the structural cause. "
  "Permanence Architecture eliminates re-derivation.",
  "Three guarantees hold: binding, continuity, crypto-agility. "
  "The Permanence Lemma disqualifies re-derived records. "
  "The Boundary Persistence Score measures deployment compliance. "
  "MTCP now provides formally permanent governance evidence.",
 ]),
 ("10. References",[
  "Abby, A. (2026). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.",
  "Abby, A. (2026, Paper 37). Blockchain Evidence Chain. DOI: 10.17605/OSF.IO/DXGK5.",
  "Abby, A. (2026, Paper 39). Constraint Manifest. DOI: 10.17605/OSF.IO/DXGK5.",
  "Abby, A. (2026, Paper 40). Constraint Object Specification. DOI: 10.17605/OSF.IO/DXGK5.",
  "Framework F28 -- BEC Definition. MTCP Research Programme.",
  "Framework F35 -- Quantum-Safe Definition. MTCP Research Programme.",
  "Framework F37 -- Permanence Architecture Definition. MTCP Research Programme.",
 ]),
]

def build(is_public=True):
    doc=Document();setup(doc)
    if is_public: hdr(doc,"MTCP V1.0 -- Permanence Architecture")
    else: hdr(doc,"CONFIDENTIAL -- NDA REQUIRED");ftr(doc,"MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby.")
    h(doc,"MTCP Paper 47",1);h(doc,"Permanence Architecture",2)
    p(doc,"")
    if not is_public: p(doc,"CONFIDENTIAL -- NDA REQUIRED")
    p(doc,"Author: A. Abby");p(doc,"Contact: admin@mtcp.live")
    p(doc,"DOI: 10.17605/OSF.IO/DXGK5");p(doc,"Date: May 2026");p(doc,"Version: 1.0")
    p(doc,"");p(doc,"MTCP Research Programme -- Paper 47");p(doc,"Framework F37")
    doc.add_page_break()
    for title,paras in C:
        h(doc,title,1)
        for t in paras: p(doc,t)
    return doc

if __name__=="__main__":
    print("Building Paper 47: Permanence Architecture")
    build(True).save(os.path.join(DESKTOP,"Paper47_Permanence_Architecture_PUBLIC.docx"))
    build(False).save(os.path.join(DESKTOP,"Paper47_Permanence_Architecture_FULL.docx"))
    print("Done.")
