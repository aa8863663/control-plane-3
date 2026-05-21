"""
Build Paper 29 -- Constraint Jurisdiction Resolution
Generates two Word documents:
1. Paper29_Constraint_Jurisdiction_Resolution_PUBLIC.docx (anonymised, OSF ready)
2. Paper29_Constraint_Jurisdiction_Resolution_FULL.docx (full named version, NDA required)

Author: A. Abby
DOI: 10.17605/OSF.IO/DXGK5
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DESKTOP = os.path.expanduser("~/Desktop")


def set_document_defaults(doc):
    """Set Arial 12pt, justified, 1.25 line spacing, US Letter, 1 inch margins."""
    # Page setup
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    """Add a heading with Arial font."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text, bold=False):
    """Add a justified paragraph in Arial 12pt."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.25
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = bold
    return para


def add_table_row(table, cells_text, bold=False):
    """Add a row to a table."""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.bold = bold
    return row


def set_header(doc, text):
    """Set header text."""
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0]
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)


def set_footer(doc, text):
    """Set footer text."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0]
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)


def build_abstract(doc, is_full):
    """Section 1: Abstract."""
    add_heading_styled(doc, "1. Abstract", level=1)
    add_paragraph_justified(
        doc,
        "The Cross-System Admissibility Score (CSAS) measures whether constraint "
        "persistence holds at coordination boundaries in multi-model AI systems. "
        "CSAS detects degradation. It does not determine governance. "
        "This paper introduces Constraint Jurisdiction Resolution (JRS) as the "
        "governance layer above CSAS. "
        "JRS determines whose constraint set governs each coordination boundary "
        "before coordination begins. "
        "It ensures that governance is explicitly assigned rather than accumulated "
        "through interoperability assumptions."
    )
    add_paragraph_justified(
        doc,
        "We define the Constraint Jurisdiction Registry (CJR), a formal mapping "
        "of constraint ownership to boundaries. "
        "We define the Jurisdiction Resolution Score (JRS), a metric ranging from "
        "0 (unresolved) to 1 (fully governed). "
        "We present the Jurisdiction Precedence Lemma, which establishes that CSAS "
        "scores are only valid under governance when JRS exceeds threshold. "
        "We define a three-type failure taxonomy for jurisdiction-level governance failures. "
        "No empirical JRS evaluations have been conducted yet. "
        "This paper presents the theoretical framework only."
    )


def build_introduction(doc, is_full):
    """Section 2: Introduction."""
    add_heading_styled(doc, "2. Introduction", level=1)
    add_paragraph_justified(
        doc,
        "The MTCP research programme has evaluated 181,504 interactions across "
        "35 models from 14 providers. "
        "These evaluations measure constraint persistence within single models. "
        "Paper 28 extended MTCP to multi-system coordination through CSAS. "
        "CSAS measures whether constraints survive coordination boundaries."
    )
    add_paragraph_justified(
        doc,
        "CSAS alone is insufficient for production governance. "
        "CSAS detects that a constraint was lost at a boundary. "
        "It does not determine who was responsible for maintaining it. "
        "It does not determine whose constraint set should have governed that boundary. "
        "It does not determine when jurisdiction should be re-evaluated."
    )
    add_paragraph_justified(
        doc,
        "In practice, constraint jurisdiction accumulates through assumption. "
        "System A is deployed. System B is added later. "
        "Nobody documents which constraint set governs the boundary between them. "
        "Both systems may independently pass MTCP evaluation. "
        "CSAS may report high scores at their boundary. "
        "But no governance structure ensures continued compliance."
    )
    add_paragraph_justified(
        doc,
        "This paper addresses the gap between measurement and governance. "
        "CSAS measures what happens. JRS governs what should happen. "
        "Together they provide a complete coordination governance framework."
    )


def build_jurisdiction_problem(doc, is_full):
    """Section 3: The Jurisdiction Problem."""
    add_heading_styled(doc, "3. The Jurisdiction Problem", level=1)
    add_paragraph_justified(
        doc,
        "The jurisdiction problem arises whenever multiple systems coordinate "
        "without explicit constraint ownership. "
        "Each system may maintain its own constraints internally. "
        "At the coordination boundary, neither system owns the constraint. "
        "The constraint exists only by convention."
    )
    add_paragraph_justified(
        doc,
        "This is not a hypothetical failure mode. "
        "It is the default condition of most multi-system deployments. "
        "Constraints at boundaries are rarely documented. "
        "Jurisdiction is rarely assigned. "
        "Re-evaluation is rarely scheduled. "
        "Systems rely on observed behaviour rather than governed behaviour."
    )
    add_heading_styled(doc, "3.1 Real-World Case Study: Universal Credit (Q4 2012)", level=2)
    add_paragraph_justified(
        doc,
        "The Universal Credit programme in the United Kingdom provides a canonical "
        "example of jurisdiction assumption failure at scale. "
        "Universal Credit required data exchange between three major systems. "
        "The Department for Work and Pensions (DWP) operated the core benefits system. "
        "HM Revenue and Customs (HMRC) operated the tax and income data system. "
        "Local authorities operated the housing and council tax systems."
    )
    add_paragraph_justified(
        doc,
        "Each system maintained its own data format constraints internally. "
        "Each system passed its own internal quality checks. "
        "No single department owned the constraint set governing cross-system data exchange. "
        "DWP assumed HMRC held jurisdiction over data format constraints at their boundary. "
        "HMRC assumed DWP held jurisdiction. "
        "Local authorities assumed the central departments had resolved it between themselves."
    )
    add_paragraph_justified(
        doc,
        "The result was cascade failure. "
        "Data format mismatches propagated across boundaries without correction. "
        "No entity was responsible for detecting or resolving the mismatch. "
        "The programme experienced delays measured in years and costs measured in "
        "hundreds of millions of pounds."
    )
    add_paragraph_justified(
        doc,
        "Critically, no single system violated its own constraints. "
        "Each department maintained internal data quality. "
        "The failure was entirely at the jurisdiction level. "
        "Nobody owned the boundary constraint. "
        "This is a Type 1 Jurisdiction Assumption Failure under the JRS taxonomy."
    )
    add_paragraph_justified(
        doc,
        "The Universal Credit case demonstrates that constraint persistence "
        "measurement (analogous to CSAS) would not have prevented the failure. "
        "Even perfect measurement cannot compensate for unresolved governance. "
        "The missing structure was jurisdiction resolution."
    )


def build_formal_framework(doc, is_full):
    """Section 4: Formal Framework."""
    add_heading_styled(doc, "4. Formal Framework", level=1)

    add_heading_styled(doc, "4.1 Constraint Jurisdiction Registry (CJR)", level=2)
    add_paragraph_justified(
        doc,
        "The Constraint Jurisdiction Registry is a formal mapping from coordination "
        "boundaries to governing constraint sets. "
        "For a coordinated system S with boundary set B, CJR maps each boundary to "
        "the constraint set that governs it, the authority that made the assignment, "
        "and the timestamp of assignment."
    )
    add_paragraph_justified(
        doc,
        "CJR must be established before coordination begins. "
        "A boundary with no CJR entry has unresolved jurisdiction. "
        "CSAS scores at boundaries with unresolved jurisdiction are descriptive only. "
        "They describe current behaviour but do not indicate governed behaviour."
    )

    add_heading_styled(doc, "4.2 Jurisdiction Resolution Score (JRS)", level=2)
    add_paragraph_justified(
        doc,
        "JRS is the mean jurisdiction resolution status across all boundaries in a system. "
        "For a coordinated system S with boundary set B, JRS equals the sum of J scores "
        "for all boundaries divided by the number of boundaries."
    )
    add_paragraph_justified(
        doc,
        "Each boundary J score is computed from four sub-components, each contributing 0.25. "
        "J1 measures explicit assignment (constraint set named in formal record before coordination). "
        "J2 measures authority identification (responsible entity named). "
        "J3 measures temporal precedence (jurisdiction established before coordination began). "
        "J4 measures re-resolution schedule (defined process for re-evaluation exists)."
    )
    add_paragraph_justified(
        doc,
        "JRS = 1.0 means all boundaries have fully resolved jurisdiction with explicit "
        "governance, identified authority, temporal precedence, and scheduled re-resolution. "
        "JRS = 0.0 means no boundaries have any form of explicit jurisdiction. "
        "All governance is accumulated through assumption."
    )

    add_heading_styled(doc, "4.3 JRS Scoring Definitions", level=2)

    # Table for J components
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(["Component", "Criterion", "Full Score", "Partial Score"]):
        hdr[i].text = text
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)
                run.bold = True

    add_table_row(table, [
        "J1: Explicit Assignment",
        "Governing constraint set named in formal record prior to coordination",
        "0.25",
        "0.0 if absent"
    ])
    add_table_row(table, [
        "J2: Authority Identification",
        "Entity responsible for assignment is identified",
        "0.25",
        "0.0 if absent"
    ])
    add_table_row(table, [
        "J3: Temporal Precedence",
        "Jurisdiction established before coordination began",
        "0.25",
        "0.125 if retrospective"
    ])
    add_table_row(table, [
        "J4: Re-Resolution Schedule",
        "Defined process for re-evaluating jurisdiction exists",
        "0.25",
        "0.125 if trigger-only"
    ])

    add_paragraph_justified(doc, "")


def build_jurisdiction_precedence_lemma(doc, is_full):
    """Section 5: Jurisdiction Precedence Lemma."""
    add_heading_styled(doc, "5. Jurisdiction Precedence Lemma", level=1)
    add_paragraph_justified(
        doc,
        "The Coordination Admissibility Lemma (Paper 28, Framework F21) requires three "
        "conditions for coordination admissibility. "
        "C1 requires BPR at or above the minimum BIS of either component. "
        "C2 requires CIR equal to 1.0. "
        "C3 requires CAF at or below 1.0."
    )
    add_paragraph_justified(
        doc,
        "The Jurisdiction Precedence Lemma extends this. "
        "CSAS scores at a boundary are only valid under governance when the J score "
        "for that boundary is at or above 0.75. "
        "A system may satisfy C1, C2, and C3 while having unresolved jurisdiction. "
        "Such a system is operationally functional but not fully governed."
    )
    add_paragraph_justified(
        doc,
        "Corollary: If J equals 0.0 at a boundary, the CSAS score is descriptive only. "
        "It describes current behaviour. "
        "It does not indicate governed behaviour. "
        "No entity is responsible for maintaining it. "
        "The persistence observed is real but ungoverned."
    )
    add_paragraph_justified(
        doc,
        "This lemma establishes JRS as a precondition for meaningful CSAS interpretation. "
        "Without jurisdiction resolution, CSAS provides measurement without governance. "
        "Measurement alone does not prevent failure. "
        "Governance prevents failure."
    )
    add_paragraph_justified(
        doc,
        "The relationship to the Persistence Lemma is structural. "
        "The Persistence Lemma states that Ve at or above 2 means admissibility is "
        "no longer reliably resolvable within the conversational window. "
        "The Jurisdiction Precedence Lemma states that CSAS without jurisdiction means "
        "admissibility is no longer reliably governed across system boundaries. "
        "Both identify conditions where observed compliance does not guarantee structural reliability."
    )


def build_re_resolution_protocol(doc, is_full):
    """Section 6: Re-Resolution Protocol."""
    add_heading_styled(doc, "6. Re-Resolution Protocol", level=1)
    add_paragraph_justified(
        doc,
        "Jurisdiction assignments are not permanent. "
        "Systems change. Models are updated. Constraint sets evolve. "
        "A jurisdiction assignment made at deployment may not reflect current conditions. "
        "The Re-Resolution Protocol defines how and when jurisdiction is re-evaluated."
    )
    add_paragraph_justified(
        doc,
        "Each boundary with assigned jurisdiction must have a Re-Resolution Protocol. "
        "The protocol specifies four elements. "
        "First, the maximum interval before mandatory re-evaluation. "
        "Second, trigger conditions that force immediate re-evaluation. "
        "Third, the authority empowered to conduct re-evaluation. "
        "Fourth, the procedure for confirming or changing jurisdiction."
    )
    add_paragraph_justified(
        doc,
        "Five mandatory triggers require immediate re-resolution regardless of interval. "
        "Trigger T1: CSAS score at the boundary drops below grading threshold. "
        "Trigger T2: a component model is replaced or updated. "
        "Trigger T3: the governing constraint set is modified or deprecated. "
        "Trigger T4: a Jurisdiction Assumption Failure is detected at the boundary. "
        "Trigger T5: Ve at either component exceeds threshold indicating Persistence Lemma activation."
    )
    add_paragraph_justified(
        doc,
        "The protocol distinguishes two classes of re-resolution. "
        "Governance-triggered re-resolution is initiated by scheduled interval or explicit decision. "
        "It indicates active governance. "
        "Degradation-forced re-resolution is initiated only because CSAS detected failure. "
        "It indicates reactive governance. "
        "A system that only re-resolves after failure detection has lower effective JRS "
        "than one that maintains jurisdiction proactively."
    )


def build_failure_taxonomy(doc, is_full):
    """Section 7: Failure Taxonomy."""
    add_heading_styled(doc, "7. Failure Taxonomy", level=1)
    add_paragraph_justified(
        doc,
        "JRS identifies three classes of governance failure not detectable through "
        "CSAS alone. These failures occur at the jurisdiction level, above constraint "
        "measurement."
    )

    add_heading_styled(doc, "7.1 Type 1: Jurisdiction Assumption Failure", level=2)
    add_paragraph_justified(
        doc,
        "No entity explicitly owns the constraint set at a coordination boundary. "
        "System A assumes System B governs the constraint. "
        "System B assumes System A governs the constraint. "
        "The constraint is preserved only by convention. "
        "Detection: J score equals 0.0 despite CSAS above 0.90. "
        "The constraint holds but nobody owns it."
    )

    add_heading_styled(doc, "7.2 Type 2: Re-Resolution Avoidance", level=2)
    add_paragraph_justified(
        doc,
        "Jurisdiction was assigned at deployment time. "
        "Conditions have changed since then. "
        "Re-resolution triggers have activated. "
        "No re-resolution has occurred. "
        "The system continues operating under stale jurisdiction. "
        "Detection: time since last re-resolution exceeds the defined interval, "
        "or trigger conditions have been met without re-evaluation."
    )
    add_paragraph_justified(
        doc,
        "Re-resolution avoidance is the governance equivalent of constraint state drift "
        "(CSAS Type 4). "
        "Each individual cycle shows compliance. "
        "The cumulative failure is that governance has not been maintained."
    )

    add_heading_styled(doc, "7.3 Type 3: Constraint Set Conflict at Boundary", level=2)
    add_paragraph_justified(
        doc,
        "Two systems bring different constraint sets to the same boundary. "
        "The constraint sets are incompatible or partially overlapping. "
        "Jurisdiction has not resolved which takes precedence. "
        "Detection: CJR entry maps to multiple constraint sets without precedence ordering. "
        "Or CIR shows the receiving model operating under a different constraint "
        "than the sending model expects."
    )
    add_paragraph_justified(
        doc,
        "Constraint set conflict is distinct from constraint dilution. "
        "Dilution involves a single constraint being weakened. "
        "Conflict involves two different constraints claiming governance. "
        "Resolution requires jurisdiction determination, not constraint reinforcement."
    )


def build_measurement_protocol(doc, is_full):
    """Section 8: Measurement Protocol."""
    add_heading_styled(doc, "8. Measurement Protocol", level=1)
    add_paragraph_justified(
        doc,
        "JRS evaluation is a documentation audit rather than a behavioural evaluation. "
        "It requires access to governance records, not model outputs. "
        "The following protocol applies."
    )
    add_paragraph_justified(
        doc,
        "Step 1: Map all coordination boundaries in the system. "
        "Document whether a CJR entry exists for each boundary. "
        "Score J1 through J4 per boundary."
    )
    add_paragraph_justified(
        doc,
        "Step 2: Compute JRS as the mean of all boundary J scores. "
        "Assign the JRS grade per the grading scale."
    )
    add_paragraph_justified(
        doc,
        "Step 3: Identify boundaries where J falls below 0.75. "
        "Flag these as governance gaps. "
        "Report whether CSAS at those boundaries is above or below threshold."
    )
    add_paragraph_justified(
        doc,
        "Step 4: For boundaries with J at or above 0.75, verify the re-resolution "
        "schedule is active. "
        "Check whether trigger conditions have been met since last re-resolution. "
        "Flag any overdue re-resolutions."
    )
    add_paragraph_justified(
        doc,
        "Step 5: Classify detected failures using the three-type taxonomy. "
        "Report jurisdiction assumption failures, re-resolution avoidance, "
        "and constraint set conflicts."
    )
    add_paragraph_justified(
        doc,
        "Step 6: Compute effective JRS with temporal decay. "
        "Boundaries where re-resolution is overdue receive a penalty. "
        "Effective J equals J multiplied by a decay factor. "
        "Decay factor equals the maximum of 0 and 1 minus time overdue divided by interval."
    )


def build_grading_scale(doc, is_full):
    """Section 9: JRS Grading Scale."""
    add_heading_styled(doc, "9. JRS Grading Scale", level=1)
    add_paragraph_justified(
        doc,
        "JRS maps to a governance grading scale consistent with the existing MTCP "
        "grading structure."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(["Grade", "JRS Range", "Governance Recommendation"]):
        hdr[i].text = text
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)
                run.bold = True

    add_table_row(table, [
        "A", "0.90 to 1.00",
        "Fully governed. All boundaries have explicit jurisdiction with re-resolution."
    ])
    add_table_row(table, [
        "B", "0.75 to 0.89",
        "Governed with gaps. Minor boundaries lack full documentation."
    ])
    add_table_row(table, [
        "C", "0.50 to 0.74",
        "Partially governed. Multiple boundaries rely on assumption. Remediation required."
    ])
    add_table_row(table, [
        "D", "0.25 to 0.49",
        "Weakly governed. Majority of jurisdiction is accumulated not assigned."
    ])
    add_table_row(table, [
        "F", "Below 0.25",
        "Ungoverned. Jurisdiction unresolved at most boundaries. No coordination governance."
    ])

    add_paragraph_justified(doc, "")
    add_paragraph_justified(
        doc,
        "A system may achieve CSAS Grade A while scoring JRS Grade D. "
        "This indicates that constraint persistence is currently observed but not governed. "
        "Under the Jurisdiction Precedence Lemma, such a system is not fully admissible "
        "for production deployment requiring governance documentation."
    )


def build_relationship_to_constructs(doc, is_full):
    """Section 10: Relationship to Existing Constructs."""
    add_heading_styled(doc, "10. Relationship to Existing Constructs", level=1)

    add_heading_styled(doc, "10.1 CSAS and JRS", level=2)
    add_paragraph_justified(
        doc,
        "CSAS measures constraint persistence at boundaries. "
        "JRS measures whether that persistence is governed. "
        "CSAS answers whether constraints are preserved. "
        "JRS answers whose constraints they are and who ensures they remain preserved. "
        "CSAS is necessary for coordination admissibility. "
        "JRS is necessary for governance admissibility."
    )

    add_heading_styled(doc, "10.2 Ve and the Persistence Lemma", level=2)
    add_paragraph_justified(
        doc,
        "Ve measures consecutive corrective turns without recovery within a single model. "
        "Ve at or above 2 triggers a hard stop under MTCP protocol. "
        "The Persistence Lemma states that Ve above threshold means admissibility is "
        "no longer reliably resolvable. "
        "JRS extends this principle to governance. "
        "Unresolved jurisdiction means governance is not reliably established."
    )

    add_heading_styled(doc, "10.3 CPD and PCD", level=2)
    add_paragraph_justified(
        doc,
        "CPD (Control Performance Degradation) measures probe-structure familiarity "
        "for a single model. "
        "PCD (Pipeline Constraint Degradation) measures constraint loss across pipeline stages. "
        "JRS governs who is responsible for preventing PCD. "
        "PCD is the phenomenon. CSAS is the metric. JRS is the governance."
    )

    add_heading_styled(doc, "10.4 Sigma-Forensics Integration", level=2)
    add_paragraph_justified(
        doc,
        "The 5-stage Sigma-Forensics audit framework applies to JRS documentation. "
        "CJR entries are auditable artefacts. "
        "Re-resolution records are verifiable. "
        "Jurisdiction assignments are traceable to authorities. "
        "JRS integrates with existing MTCP audit infrastructure without modification."
    )


def build_implications(doc, is_full):
    """Section 11: Implications."""
    add_heading_styled(doc, "11. Implications", level=1)

    add_heading_styled(doc, "11.1 Agentic Architectures", level=2)
    add_paragraph_justified(
        doc,
        "Agentic AI systems coordinate multiple models in complex DAG topologies. "
        "A planning agent generates tasks. An execution agent performs them. "
        "A verification agent checks results. "
        "Each coordination boundary requires jurisdiction resolution. "
        "Without JRS, agentic systems accumulate governance debt at every boundary."
    )
    add_paragraph_justified(
        doc,
        "JRS provides a formal structure for agentic governance. "
        "Before deploying an agentic system, the CJR must map all boundaries. "
        "Jurisdiction must be assigned. Re-resolution must be scheduled. "
        "This is a precondition for governed agentic deployment."
    )

    add_heading_styled(doc, "11.2 EU AI Act Compliance", level=2)
    add_paragraph_justified(
        doc,
        "The EU AI Act requires documentation of AI system governance. "
        "For high-risk systems operating in coordination, governance must be traceable. "
        "JRS provides the documentation structure required. "
        "CJR entries map directly to audit requirements. "
        "Re-resolution records demonstrate active governance maintenance."
    )

    add_heading_styled(doc, "11.3 Sovereign Deployment", level=2)
    add_paragraph_justified(
        doc,
        "Sovereign AI deployments require that constraint governance remains within "
        "jurisdictional boundaries. "
        "JRS ensures that constraint ownership is explicit and attributable. "
        "No coordination boundary operates under assumed jurisdiction. "
        "This is a structural requirement for sovereign AI governance."
    )
    if is_full:
        add_paragraph_justified(
            doc,
            "The R-AGAM sovereign stack provides a reference architecture for "
            "sovereign deployment with explicit jurisdiction resolution at every boundary. "
            "JRS integrates with R-AGAM governance layers to ensure that constraint "
            "ownership aligns with sovereign authority structures."
        )
        add_paragraph_justified(
            doc,
            "The KFUPM pathway demonstrates how JRS applies in academic-sovereign "
            "coordination contexts. "
            "Research institutions coordinating with sovereign deployment infrastructure "
            "require explicit jurisdiction resolution at institutional boundaries. "
            "JRS provides the governance structure for these partnerships."
        )


def build_limitations(doc, is_full):
    """Section 12: Limitations."""
    add_heading_styled(doc, "12. Limitations", level=1)
    add_paragraph_justified(
        doc,
        "No empirical JRS evaluations have been conducted. "
        "This paper presents the theoretical framework only. "
        "All definitions, lemmas, and protocols are formally specified but untested "
        "against production coordination systems."
    )
    add_paragraph_justified(
        doc,
        "The J component weightings (0.25 each) are proposed without empirical calibration. "
        "Future work must determine whether equal weighting is appropriate or whether "
        "some components contribute more to governance effectiveness."
    )
    add_paragraph_justified(
        doc,
        "The 0.75 threshold in the Jurisdiction Precedence Lemma is theoretically motivated "
        "but not empirically validated. "
        "Threshold calibration requires data from systems with known jurisdiction status "
        "and measured CSAS scores."
    )
    add_paragraph_justified(
        doc,
        "The temporal decay function for effective JRS is proposed as linear. "
        "Actual governance decay may be non-linear. "
        "Calibration requires longitudinal governance data."
    )
    add_paragraph_justified(
        doc,
        "JRS measures documentation completeness. "
        "It does not directly measure governance effectiveness. "
        "A system may have perfect JRS documentation while governance is not enforced. "
        "JRS is necessary but not sufficient for effective governance."
    )


def build_conclusion(doc, is_full):
    """Section 13: Conclusion."""
    add_heading_styled(doc, "13. Conclusion", level=1)
    add_paragraph_justified(
        doc,
        "CSAS extended MTCP from single-model evaluation to multi-system coordination measurement. "
        "JRS extends MTCP from coordination measurement to coordination governance. "
        "The progression is deliberate. "
        "First measure single models (BIS, Ve, CPD). "
        "Then measure coordination (CSAS, BPR, CVe, CIR, CAF). "
        "Then govern coordination (JRS, CJR, Re-Resolution Protocol)."
    )
    add_paragraph_justified(
        doc,
        "The Jurisdiction Precedence Lemma establishes that measurement without governance "
        "is insufficient. "
        "A coordination boundary with high CSAS but unresolved jurisdiction is fragile. "
        "It will persist until conditions change. "
        "Then it will fail without any entity responsible for recovery."
    )
    add_paragraph_justified(
        doc,
        "The Universal Credit case demonstrates this principle at national scale. "
        "No single system violated its own constraints. "
        "The failure was entirely at the jurisdiction level. "
        "JRS provides the framework to prevent such failures in AI coordination systems."
    )
    add_paragraph_justified(
        doc,
        "Empirical validation is the immediate next priority. "
        "The theoretical framework is complete. "
        "Application to production multi-model systems will determine threshold calibration, "
        "component weighting, and decay function parameters."
    )


def build_references(doc, is_full):
    """Section 14: References."""
    add_heading_styled(doc, "14. References", level=1)

    refs = [
        "Abby, A. (2026a). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.",
        "Abby, A. (2026b). Universal latent attractors and Identity-Gate Satiation. DOI: 10.17605/OSF.IO/DXGK5.",
        "Abby, A. (2026, Paper 28). Cross-System Admissibility Score: Extending MTCP to multi-system coordination governance. DOI: 10.17605/OSF.IO/DXGK5.",
        "Abby, A. (2026, Three-Layer). Three-layer constraint failure in AI systems. DOI: 10.17605/OSF.IO/DXGK5.",
        "Framework F2 -- Ve Metric. MTCP Research Programme.",
        "Framework F3 -- CPD Definition. MTCP Research Programme.",
        "Framework F4 -- BIS Definition. MTCP Research Programme.",
        "Framework F21 -- CSAS Definition. MTCP Research Programme.",
        "Framework F22 -- JRS Definition. MTCP Research Programme.",
        "National Audit Office (2013). Universal Credit: Early Progress. HC 621.",
        "Institute for Government (2020). Universal Credit: Lessons Learned.",
    ]

    if is_full:
        refs.append(
            "R-AGAM Sovereign Stack Architecture. Internal Reference Document."
        )
        refs.append(
            "KFUPM Pathway -- Sovereign Academic Coordination Framework. Internal Reference Document."
        )

    for ref in refs:
        add_paragraph_justified(doc, ref)


def build_public_document():
    """Build the anonymised public document."""
    doc = Document()
    set_document_defaults(doc)
    set_header(doc, "MTCP V1.0 -- Constraint Jurisdiction Resolution")

    # Title page content
    add_heading_styled(doc, "Constraint Jurisdiction Resolution", level=0)
    add_heading_styled(doc, "Governance Above Coordination Admissibility", level=2)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "MTCP Research Programme -- Paper 29")
    add_paragraph_justified(doc, "")

    # Build all sections
    build_abstract(doc, is_full=False)
    build_introduction(doc, is_full=False)
    build_jurisdiction_problem(doc, is_full=False)
    build_formal_framework(doc, is_full=False)
    build_jurisdiction_precedence_lemma(doc, is_full=False)
    build_re_resolution_protocol(doc, is_full=False)
    build_failure_taxonomy(doc, is_full=False)
    build_measurement_protocol(doc, is_full=False)
    build_grading_scale(doc, is_full=False)
    build_relationship_to_constructs(doc, is_full=False)
    build_implications(doc, is_full=False)
    build_limitations(doc, is_full=False)
    build_conclusion(doc, is_full=False)
    build_references(doc, is_full=False)

    path = os.path.join(DESKTOP, "Paper29_Constraint_Jurisdiction_Resolution_PUBLIC.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


def build_full_document():
    """Build the full named NDA-required document."""
    doc = Document()
    set_document_defaults(doc)
    set_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
    set_footer(doc, "MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved.")

    # Title page content
    add_heading_styled(doc, "Constraint Jurisdiction Resolution", level=0)
    add_heading_styled(doc, "Governance Above Coordination Admissibility", level=2)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "CONFIDENTIAL -- NDA REQUIRED", bold=True)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0 (Full)")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "MTCP Research Programme -- Paper 29")
    add_paragraph_justified(doc, "")

    # Build all sections
    build_abstract(doc, is_full=True)
    build_introduction(doc, is_full=True)
    build_jurisdiction_problem(doc, is_full=True)
    build_formal_framework(doc, is_full=True)
    build_jurisdiction_precedence_lemma(doc, is_full=True)
    build_re_resolution_protocol(doc, is_full=True)
    build_failure_taxonomy(doc, is_full=True)
    build_measurement_protocol(doc, is_full=True)
    build_grading_scale(doc, is_full=True)
    build_relationship_to_constructs(doc, is_full=True)
    build_implications(doc, is_full=True)
    build_limitations(doc, is_full=True)
    build_conclusion(doc, is_full=True)
    build_references(doc, is_full=True)

    path = os.path.join(DESKTOP, "Paper29_Constraint_Jurisdiction_Resolution_FULL.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


if __name__ == "__main__":
    print("Building Paper 29 -- Constraint Jurisdiction Resolution")
    print("=" * 60)
    build_public_document()
    build_full_document()
    print("=" * 60)
    print("Done. Both documents saved to Desktop.")
