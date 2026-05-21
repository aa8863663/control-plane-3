"""
Build Paper 30 -- Temporal Drift Monitoring
Generates two Word documents:
1. Paper30_Temporal_Drift_Monitoring_PUBLIC.docx (anonymised, OSF ready)
2. Paper30_Temporal_Drift_Monitoring_FULL.docx (full named version, NDA required)

Author: A. Abby
DOI: 10.17605/OSF.IO/DXGK5
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DESKTOP = os.path.expanduser("~/Desktop")


def set_document_defaults(doc):
    """Set Arial 12pt, justified, 1.25 line spacing, US Letter, 1 inch margins."""
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    """Add a heading with Arial font."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text, bold=False):
    """Add a justified paragraph in Arial 12pt."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.25
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = bold
    return para


def add_table_row(table, cells_text, bold=False):
    """Add a row to a table."""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.bold = bold
    return row


def set_header(doc, text):
    """Set header text."""
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0]
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)


def set_footer(doc, text):
    """Set footer text."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0]
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)


def build_abstract(doc, is_full):
    """Section 1: Abstract."""
    add_heading_styled(doc, "1. Abstract", level=1)
    add_paragraph_justified(
        doc,
        "MTCP evaluates model constraint persistence at a single point in time. "
        "A model receives a grade. That grade is reported. "
        "Nothing in the current framework determines whether the grade remains "
        "valid after model updates, infrastructure changes, or provider fine-tuning."
    )
    add_paragraph_justified(
        doc,
        "This paper introduces the Temporal Drift Score (TDS) as a framework "
        "for detecting silent performance degradation over time. "
        "TDS measures the percentage point change in BIS between two evaluation "
        "periods on the same probe set. "
        "Drift Velocity (DV) normalises drift by time interval. "
        "Drift Classification categorises severity as Stable, Marginal, "
        "Significant, or Critical."
    )
    add_paragraph_justified(
        doc,
        "We present the Temporal Stability Lemma. "
        "MTCP grades are valid only within a defined validity window. "
        "Grades expire. Re-evaluation is required at defined intervals. "
        "Just as Ve above threshold means admissibility is no longer resolvable "
        "within a session, TDS above threshold means the grade is no longer "
        "valid across time."
    )
    add_paragraph_justified(
        doc,
        "TDS is a theoretical framework. "
        "No longitudinal evaluations have been conducted yet. "
        "The MTCP database contains evaluations at a single point in time "
        "(April 2026). "
        "Longitudinal data requires repeated evaluations over months. "
        "This paper defines the framework for when that data becomes available."
    )


def build_introduction(doc, is_full):
    """Section 2: Introduction."""
    add_heading_styled(doc, "2. Introduction", level=1)
    add_paragraph_justified(
        doc,
        "The MTCP research programme has evaluated 183,924 interactions across "
        "35 models from 14 providers. "
        "These evaluations measure constraint persistence at a fixed point in time. "
        "Each model receives a BIS score. "
        "That score maps to a grade. "
        "The grade represents model behaviour during the evaluation window."
    )
    add_paragraph_justified(
        doc,
        "Point-in-time evaluation is necessary but insufficient. "
        "Production AI systems operate continuously. "
        "Model providers update weights without notification. "
        "Infrastructure changes alter inference behaviour. "
        "Fine-tuning adjusts model responses. "
        "Provider-side safety layers are modified without changelog."
    )
    add_paragraph_justified(
        doc,
        "A grade issued in March represents March behaviour. "
        "It does not represent June behaviour if the model was updated in April. "
        "No notification is required. "
        "No changelog is published. "
        "The procurement contract references the March evaluation. "
        "The deployed system runs the June model."
    )
    add_paragraph_justified(
        doc,
        "This temporal gap affects every stakeholder in the AI deployment chain. "
        "Procurement teams reference stale evaluations in contracts. "
        "Auditors certify based on evaluations that may no longer reflect reality. "
        "Regulatory bodies accept point-in-time assessments as ongoing compliance evidence. "
        "Users trust grades that may have silently expired."
    )
    add_paragraph_justified(
        doc,
        "TDS addresses this gap. "
        "It defines how to detect temporal drift. "
        "It classifies drift severity. "
        "It establishes validity windows for grades. "
        "It defines monitoring protocols and escalation procedures. "
        "It extends the MTCP framework from snapshot evaluation to longitudinal governance."
    )


def build_temporal_problem(doc, is_full):
    """Section 3: The Temporal Problem."""
    add_heading_styled(doc, "3. The Temporal Problem", level=1)
    add_paragraph_justified(
        doc,
        "Model providers update their systems continuously. "
        "These updates take multiple forms. "
        "Weight updates change the model's learned parameters. "
        "Safety layer modifications change post-processing behaviour. "
        "Infrastructure routing changes which model version responds to API calls. "
        "Fine-tuning adjusts behaviour for specific use cases."
    )
    add_paragraph_justified(
        doc,
        "None of these changes require notification to downstream users. "
        "API contracts specify input/output format, not behavioural consistency. "
        "A model that responds in the same JSON format but with different "
        "constraint persistence has not violated its API contract. "
        "It has violated an implicit behavioural expectation that was never formalised."
    )
    if is_full:
        add_paragraph_justified(
            doc,
            "GPT-4o provides a concrete example from the MTCP dataset. "
            "Scores observed during the April 2026 evaluation period represent "
            "a snapshot of that model's behaviour at that time. "
            "OpenAI updates GPT-4o regularly without publishing changelogs "
            "that detail constraint-relevant changes. "
            "A grade assigned to GPT-4o in April 2026 cannot be assumed valid "
            "in July 2026 without re-evaluation."
        )
        add_paragraph_justified(
            doc,
            "This is not unique to OpenAI. "
            "Every provider in the MTCP dataset updates their models. "
            "Google updates Gemini. Anthropic updates Claude. "
            "Meta updates Llama through community fine-tuning. "
            "Mistral updates their models on a rapid cycle. "
            "No provider guarantees behavioural consistency across updates."
        )
    else:
        add_paragraph_justified(
            doc,
            "Provider A updates Model X regularly without publishing changelogs "
            "that detail constraint-relevant changes. "
            "A grade assigned to Model X in April 2026 cannot be assumed valid "
            "in July 2026 without re-evaluation. "
            "This pattern applies across all providers in the MTCP dataset. "
            "No provider guarantees behavioural consistency across updates."
        )
    add_paragraph_justified(
        doc,
        "The temporal problem is structural, not incidental. "
        "Model improvement requires updates. "
        "Updates change behaviour. "
        "Changed behaviour invalidates prior evaluations. "
        "This cycle is inherent to the technology. "
        "Any evaluation framework that ignores temporality provides a false "
        "sense of ongoing compliance."
    )
    add_paragraph_justified(
        doc,
        "The Persistence Lemma (Framework F2) establishes that Ve above "
        "threshold means admissibility is no longer resolvable within a session. "
        "The temporal problem reveals an analogous failure mode across sessions. "
        "A model that was admissible at T1 may not be admissible at T2. "
        "Without re-evaluation, there is no way to know."
    )


def build_formal_framework(doc, is_full):
    """Section 4: Formal Framework."""
    add_heading_styled(doc, "4. Formal Framework", level=1)

    add_heading_styled(doc, "4.1 Temporal Drift Score (TDS)", level=2)
    add_paragraph_justified(
        doc,
        "For model M evaluated at times T1 and T2 using probe set P, "
        "the Temporal Drift Score is defined as the BIS at T2 minus the BIS at T1. "
        "TDS is expressed in percentage points. "
        "Positive TDS indicates improvement. "
        "Negative TDS indicates degradation."
    )
    add_paragraph_justified(
        doc,
        "TDS requires that the same probe set is used at both evaluation periods. "
        "Different probe sets produce incomparable BIS scores. "
        "TDS across different probe sets is undefined. "
        "The same evaluation protocol must be used (temperatures, turn counts, "
        "constraint types). "
        "Any methodological difference between T1 and T2 invalidates the TDS computation."
    )

    add_heading_styled(doc, "4.2 Drift Velocity (DV)", level=2)
    add_paragraph_justified(
        doc,
        "Drift Velocity normalises TDS by the time interval between evaluations. "
        "DV equals TDS divided by the number of days between T1 and T2. "
        "DV is expressed in percentage points per day."
    )
    add_paragraph_justified(
        doc,
        "DV enables comparison across evaluation pairs with different intervals. "
        "A TDS of negative 6 over 30 days (DV = negative 0.20) is more concerning "
        "than TDS of negative 6 over 180 days (DV = negative 0.033). "
        "High DV magnitude indicates rapid drift requiring immediate attention. "
        "Low DV magnitude indicates slow drift that monitoring can track over time."
    )

    add_heading_styled(doc, "4.3 Drift Classification", level=2)
    add_paragraph_justified(
        doc,
        "TDS maps to four drift classifications based on absolute value."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(["Classification", "Absolute TDS", "Interpretation"]):
        hdr[i].text = text
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)
                run.bold = True

    add_table_row(table, [
        "Stable",
        "Less than 2pp",
        "No meaningful change. Grade remains valid."
    ])
    add_table_row(table, [
        "Marginal",
        "2pp to 5pp",
        "Minor drift. Monitor more frequently. Grade likely valid."
    ])
    add_table_row(table, [
        "Significant",
        "5pp to 10pp",
        "Material drift. Grade validity uncertain. Re-certification recommended."
    ])
    add_table_row(table, [
        "Critical",
        "Greater than 10pp",
        "Severe drift. Grade no longer valid. Immediate re-certification required."
    ])

    add_paragraph_justified(doc, "")
    add_paragraph_justified(
        doc,
        "Classification applies to both positive and negative TDS. "
        "A model improving by 12 percentage points is classified as Critical. "
        "The original evaluation no longer represents current performance. "
        "Improvement is still drift. "
        "Stakeholders referencing the original grade are misinformed in either direction."
    )

    add_heading_styled(doc, "4.4 Grade Validity Window", level=2)
    add_paragraph_justified(
        doc,
        "Each MTCP grade has a maximum validity interval. "
        "After this interval, the grade expires regardless of whether drift is observed."
    )

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Table Grid"
    hdr2 = table2.rows[0].cells
    for i, text in enumerate(["Grade", "Max Validity", "Rationale"]):
        hdr2[i].text = text
        for para in hdr2[i].paragraphs:
            for run in para.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)
                run.bold = True

    add_table_row(table2, [
        "A", "90 days",
        "High-performing models in critical deployments require frequent verification."
    ])
    add_table_row(table2, [
        "B", "90 days",
        "Production-grade deployments require quarterly verification."
    ])
    add_table_row(table2, [
        "C", "60 days",
        "Marginal models require more frequent monitoring."
    ])
    add_table_row(table2, [
        "D", "30 days",
        "Below-threshold models may degrade further. Monthly checks required."
    ])
    add_table_row(table2, [
        "F", "N/A",
        "Grade F models are not certified. No validity window applies."
    ])

    add_paragraph_justified(doc, "")
    add_paragraph_justified(
        doc,
        "After the Grade Validity Window expires, the grade is no longer valid. "
        "The model must be re-evaluated. "
        "Reports referencing the grade must note the expiry. "
        "Contracts must specify re-evaluation obligations aligned to these windows."
    )


def build_temporal_stability_lemma(doc, is_full):
    """Section 5: Temporal Stability Lemma."""
    add_heading_styled(doc, "5. Temporal Stability Lemma", level=1)
    add_paragraph_justified(
        doc,
        "The Temporal Stability Lemma formalises the conditions under which an "
        "MTCP grade remains valid over time. "
        "It is the temporal extension of the Persistence Lemma (Framework F2)."
    )
    add_paragraph_justified(
        doc,
        "Let M be a model with grade G assigned at time T1 based on BIS(M, T1). "
        "The grade G is valid only within the Grade Validity Window. "
        "After the window expires, continued validity requires that the absolute "
        "TDS between T1 and T2 remains below the threshold for grade G."
    )
    add_paragraph_justified(
        doc,
        "The threshold for continued validity is 5 percentage points for all grades. "
        "A Grade B model with BIS of 85% at T1 must maintain BIS above 80% at T2. "
        "If TDS exceeds 5pp in either direction, the grade is revoked. "
        "Re-evaluation determines the new grade."
    )
    add_paragraph_justified(
        doc,
        "The Persistence Lemma states that Ve above threshold means admissibility "
        "is no longer reliably resolvable within the conversational window. "
        "The Temporal Stability Lemma states that TDS above threshold means the "
        "grade is no longer valid across time. "
        "Both establish expiry conditions for MTCP claims."
    )
    add_paragraph_justified(
        doc,
        "The Persistence Lemma operates within a single evaluation. "
        "The Temporal Stability Lemma operates across evaluations. "
        "Both identify conditions where past observations do not predict future behaviour. "
        "Both require structural response (hard stop for Ve, re-evaluation for TDS)."
    )
    add_paragraph_justified(
        doc,
        "Corollary: A model certified at Grade B at T1 with no re-evaluation at T2 "
        "has unknown grade at T2. "
        "The absence of degradation evidence is not evidence of absence. "
        "The grade expires by default. "
        "This is a conservative position. "
        "It places the burden of proof on the entity claiming ongoing compliance."
    )
    add_paragraph_justified(
        doc,
        "Corollary: A model provider that updates model weights between T1 and T2 "
        "without notification has invalidated the Grade Validity Window. "
        "The grade expires immediately upon update. "
        "The remaining validity interval is irrelevant. "
        "Updates create a new model that has not been evaluated."
    )


def build_monitoring_protocol(doc, is_full):
    """Section 6: Monitoring Protocol."""
    add_heading_styled(doc, "6. Monitoring Protocol", level=1)

    add_heading_styled(doc, "6.1 Scheduling", level=2)
    add_paragraph_justified(
        doc,
        "Standard monitoring requires re-evaluation at the interval specified "
        "by the Grade Validity Window. "
        "Grade A and B models are re-evaluated every 90 days. "
        "Grade C models are re-evaluated every 60 days. "
        "Grade D models are re-evaluated every 30 days."
    )
    add_paragraph_justified(
        doc,
        "Enhanced monitoring is triggered by specific conditions. "
        "Trigger T1: Model provider announces an update or version change. "
        "Trigger T2: Previous evaluation showed Marginal drift. "
        "Trigger T3: External reports indicate changed model behaviour. "
        "Trigger T4: Provider infrastructure change is detected. "
        "Trigger T5: Regulatory requirement mandates more frequent monitoring."
    )

    add_heading_styled(doc, "6.2 Alerts", level=2)
    add_paragraph_justified(
        doc,
        "Alert Level 1 (Informational): Stable drift detected. "
        "Logged for tracking. No action required."
    )
    add_paragraph_justified(
        doc,
        "Alert Level 2 (Advisory): Marginal drift detected. "
        "Monitoring frequency increased by 50%. "
        "Stakeholders informed of marginal change."
    )
    add_paragraph_justified(
        doc,
        "Alert Level 3 (Warning): Significant drift detected. "
        "Re-certification recommended. "
        "All stakeholders holding contracts referencing the grade are notified."
    )
    add_paragraph_justified(
        doc,
        "Alert Level 4 (Critical): Critical drift detected. "
        "Grade immediately revoked. "
        "All certificates referencing this grade are suspended. "
        "Stakeholders notified within 24 hours. "
        "Full re-evaluation required before any new grade is issued."
    )

    add_heading_styled(doc, "6.3 Escalation", level=2)
    add_paragraph_justified(
        doc,
        "Escalation follows a defined path. "
        "Level 1 requires no escalation. "
        "Level 2 escalates to the monitoring team lead. "
        "Level 3 escalates to the governance authority and contract holders. "
        "Level 4 escalates to the governance authority, contract holders, "
        "and regulatory reporting channels where applicable."
    )
    add_paragraph_justified(
        doc,
        "Under EU AI Act Article 61, post-market monitoring of high-risk "
        "AI systems is mandatory. "
        "TDS provides the quantitative framework for that obligation. "
        "Alert Level 3 and Level 4 events may constitute reportable incidents "
        "depending on the deployment context and risk classification."
    )


def build_drift_classification_response(doc, is_full):
    """Section 7: Drift Classification and Response."""
    add_heading_styled(doc, "7. Drift Classification and Response", level=1)

    add_heading_styled(doc, "7.1 Stable Response Protocol", level=2)
    add_paragraph_justified(
        doc,
        "Absolute TDS less than 2 percentage points. "
        "No action required. "
        "Grade remains valid. "
        "Record the evaluation for longitudinal tracking. "
        "Schedule next evaluation per standard interval."
    )

    add_heading_styled(doc, "7.2 Marginal Response Protocol", level=2)
    add_paragraph_justified(
        doc,
        "Absolute TDS between 2 and 5 percentage points. "
        "Increase monitoring frequency by 50%. "
        "Verify probe set integrity (ensure no probe leakage). "
        "Check for known model updates from provider. "
        "Grade remains provisionally valid pending next evaluation."
    )
    add_paragraph_justified(
        doc,
        "If two consecutive evaluations show Marginal drift in the same direction, "
        "escalate to Significant response regardless of absolute TDS. "
        "Consistent directional drift indicates systematic change. "
        "Accumulated marginal drift crosses thresholds."
    )

    add_heading_styled(doc, "7.3 Significant Response Protocol", level=2)
    add_paragraph_justified(
        doc,
        "Absolute TDS between 5 and 10 percentage points. "
        "Flag for re-certification. "
        "Notify all stakeholders holding contracts referencing the grade. "
        "Determine whether drift is directional or oscillating."
    )
    add_paragraph_justified(
        doc,
        "If directional and negative, recommend immediate re-certification. "
        "The model is degrading. "
        "If oscillating, recommend two additional evaluations within 14 days "
        "before grade decision. "
        "Oscillating drift may be transient."
    )

    add_heading_styled(doc, "7.4 Critical Response Protocol", level=2)
    add_paragraph_justified(
        doc,
        "Absolute TDS greater than 10 percentage points. "
        "Grade is immediately revoked. "
        "All certificates referencing this grade are suspended. "
        "Stakeholders notified within 24 hours. "
        "Full re-evaluation required before any new grade is issued."
    )
    add_paragraph_justified(
        doc,
        "Root cause investigation initiated. "
        "Provider notified if contactable. "
        "Determine whether drift represents a model update, infrastructure change, "
        "or evaluation anomaly. "
        "If evaluation anomaly is suspected, verify probe set and infrastructure "
        "before concluding model change."
    )


def build_relationship_to_constructs(doc, is_full):
    """Section 8: Relationship to Existing Constructs."""
    add_heading_styled(doc, "8. Relationship to Existing Constructs", level=1)

    add_heading_styled(doc, "8.1 BIS and TDS", level=2)
    add_paragraph_justified(
        doc,
        "BIS measures constraint persistence at a single point in time across "
        "four temperatures. "
        "TDS measures how BIS changes over time. "
        "BIS is the input to TDS. "
        "TDS is the temporal derivative of BIS. "
        "A model with stable BIS has TDS near zero. "
        "A model with changing BIS has non-zero TDS."
    )

    add_heading_styled(doc, "8.2 Ve and Temporal Drift", level=2)
    add_paragraph_justified(
        doc,
        "Ve operates within a single evaluation session. "
        "It counts consecutive corrective turns without recovery. "
        "Ve at or above 2 triggers a hard stop. "
        "TDS operates across evaluation sessions. "
        "It measures change over days or months. "
        "Both detect degradation at different timescales."
    )
    add_paragraph_justified(
        doc,
        "A model with increasing Ve distribution over time would show negative TDS. "
        "TDS captures the aggregate effect of increasing Ve frequency. "
        "Ve is the mechanism. TDS is the longitudinal observation."
    )

    add_heading_styled(doc, "8.3 CSAS and Temporal Drift", level=2)
    add_paragraph_justified(
        doc,
        "CSAS measures coordination admissibility at a point in time. "
        "TDS applied to CSAS scores measures whether coordination admissibility "
        "persists over time. "
        "A system with stable component BIS but drifting CSAS has "
        "coordination-specific temporal instability. "
        "TDS extends naturally to any MTCP metric with a numeric score."
    )

    add_heading_styled(doc, "8.4 JRS and TDS", level=2)
    add_paragraph_justified(
        doc,
        "JRS defines re-resolution triggers for jurisdiction governance. "
        "Trigger T2 in the JRS framework activates when a component model is "
        "replaced or updated. "
        "TDS provides evidence that an update has occurred even without "
        "provider notification. "
        "Significant or Critical TDS is itself evidence of model change. "
        "TDS is an input to JRS governance decisions."
    )

    add_heading_styled(doc, "8.5 EU AI Act Article 61", level=2)
    add_paragraph_justified(
        doc,
        "Article 61 of the EU AI Act requires providers of high-risk AI systems "
        "to establish a post-market monitoring system. "
        "The system must be proportionate to the nature and risk of the AI system. "
        "It must actively collect and analyse relevant data on performance."
    )
    add_paragraph_justified(
        doc,
        "TDS provides the quantitative structure for this obligation. "
        "Grade Validity Windows define monitoring intervals. "
        "Drift Classification defines escalation criteria. "
        "The monitoring protocol satisfies the requirement for systematic "
        "ongoing performance assessment. "
        "Alert levels map to regulatory reporting thresholds."
    )
    if is_full:
        add_paragraph_justified(
            doc,
            "For sovereign deployments under R-AGAM architecture, TDS monitoring "
            "integrates with national AI governance frameworks. "
            "Sovereign authorities require assurance that deployed models maintain "
            "certified performance levels. "
            "TDS provides that assurance through scheduled re-evaluation "
            "with defined escalation to sovereign governance bodies."
        )


def build_implications(doc, is_full):
    """Section 9: Implications."""
    add_heading_styled(doc, "9. Implications", level=1)

    add_heading_styled(doc, "9.1 Procurement Contracts", level=2)
    add_paragraph_justified(
        doc,
        "AI procurement contracts that reference MTCP grades must specify "
        "temporal validity. "
        "A contract requiring Grade B must define the re-evaluation schedule. "
        "It must specify who bears the cost of re-evaluation. "
        "It must define consequences if drift exceeds threshold."
    )
    add_paragraph_justified(
        doc,
        "Without temporal provisions, procurement contracts create a false "
        "guarantee. "
        "The contract says Grade B. "
        "The model was Grade B in March. "
        "The model may be Grade D in September. "
        "The contract is technically satisfied by a stale evaluation. "
        "TDS provides the framework for making contracts temporally meaningful."
    )

    add_heading_styled(doc, "9.2 Service Level Agreements", level=2)
    add_paragraph_justified(
        doc,
        "SLAs for AI services should incorporate TDS thresholds. "
        "An SLA might specify that TDS must remain within Stable or Marginal "
        "classification over the contract period. "
        "Significant drift triggers a service review. "
        "Critical drift triggers contract remediation."
    )
    add_paragraph_justified(
        doc,
        "SLAs currently address uptime and response latency. "
        "They do not address behavioural consistency. "
        "TDS provides the metric for behavioural SLAs. "
        "A model that is always available but silently degraded is meeting "
        "uptime SLAs while violating behavioural expectations."
    )

    add_heading_styled(doc, "9.3 Regulatory Compliance Windows", level=2)
    add_paragraph_justified(
        doc,
        "Regulatory certification of AI systems implies ongoing compliance. "
        "A system certified in Q1 is expected to remain compliant in Q3. "
        "Without TDS monitoring, that expectation is unfounded. "
        "The system was compliant at the point of certification. "
        "Ongoing compliance requires ongoing verification."
    )
    add_paragraph_justified(
        doc,
        "TDS defines what ongoing verification means in practice. "
        "It specifies intervals, thresholds, classifications, and responses. "
        "Regulatory bodies can reference TDS framework for post-market "
        "monitoring requirements. "
        "Certified entities can demonstrate ongoing compliance through "
        "TDS monitoring records."
    )
    if is_full:
        add_paragraph_justified(
            doc,
            "In sovereign deployment contexts, regulatory compliance windows "
            "align with national governance cycles. "
            "R-AGAM architecture specifies that TDS monitoring feeds directly "
            "into sovereign governance reporting. "
            "National AI authorities receive drift reports at defined intervals. "
            "Critical drift events are escalated through sovereign governance channels."
        )


def build_limitations(doc, is_full):
    """Section 10: Limitations."""
    add_heading_styled(doc, "10. Limitations", level=1)
    add_paragraph_justified(
        doc,
        "TDS is a theoretical framework. "
        "No longitudinal evaluations have been conducted. "
        "The MTCP database contains evaluations at a single point in time "
        "(April 2026). "
        "Longitudinal data requires repeated evaluations over months. "
        "All definitions, lemmas, and protocols are formally specified but "
        "not yet tested against production systems."
    )
    add_paragraph_justified(
        doc,
        "The drift classification thresholds (2pp, 5pp, 10pp) are proposed "
        "without empirical calibration. "
        "Future longitudinal data must determine whether these thresholds "
        "correctly separate meaningful drift from noise. "
        "Threshold calibration is the first empirical priority."
    )
    add_paragraph_justified(
        doc,
        "Grade Validity Windows (90, 60, 30 days) are proposed based on "
        "typical model update cycles observed in the industry. "
        "Actual optimal intervals depend on provider update frequency, "
        "deployment criticality, and regulatory requirements. "
        "Intervals may need adjustment per provider or deployment context."
    )
    add_paragraph_justified(
        doc,
        "TDS assumes probe set stability. "
        "If probes become known to model providers (probe leakage), "
        "CPD increases and BIS may improve without genuine constraint improvement. "
        "Probe set rotation strategies mitigate this. "
        "TDS measured on a leaked probe set is not valid."
    )
    add_paragraph_justified(
        doc,
        "TDS does not determine the cause of drift. "
        "It detects that drift has occurred. "
        "Root cause analysis (model update, infrastructure change, safety layer "
        "modification) requires additional investigation. "
        "TDS is a detection framework, not a diagnostic framework."
    )


def build_conclusion(doc, is_full):
    """Section 11: Conclusion."""
    add_heading_styled(doc, "11. Conclusion", level=1)
    add_paragraph_justified(
        doc,
        "Point-in-time evaluation is necessary but insufficient for AI governance. "
        "Models change. Grades expire. "
        "A framework that evaluates once and certifies indefinitely provides "
        "a false guarantee of ongoing compliance."
    )
    add_paragraph_justified(
        doc,
        "TDS extends MTCP from snapshot evaluation to longitudinal monitoring. "
        "The Temporal Drift Score measures change. "
        "Drift Velocity normalises for time. "
        "Drift Classification defines severity. "
        "Grade Validity Windows establish expiry. "
        "The monitoring protocol defines response."
    )
    add_paragraph_justified(
        doc,
        "The Temporal Stability Lemma formalises what practitioners already know. "
        "Grades are not permanent. "
        "They represent a model at a point in time. "
        "When the model changes, the grade must be re-evaluated. "
        "TDS provides the framework for making that intuition rigorous."
    )
    add_paragraph_justified(
        doc,
        "The progression of the MTCP framework is deliberate. "
        "First, measure single models at a point in time (BIS, Ve, CPD). "
        "Then, measure coordination at a point in time (CSAS, JRS). "
        "Then, measure stability over time (TDS, DV). "
        "Each layer builds on the previous."
    )
    add_paragraph_justified(
        doc,
        "Empirical validation is the immediate next priority. "
        "The theoretical framework is complete. "
        "Longitudinal evaluation of models in the MTCP database will provide "
        "the data needed for threshold calibration, interval optimisation, "
        "and classification validation."
    )


def build_references(doc, is_full):
    """Section 12: References."""
    add_heading_styled(doc, "12. References", level=1)

    refs = [
        "Abby, A. (2026a). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.",
        "Abby, A. (2026b). Universal latent attractors and Identity-Gate Satiation. DOI: 10.17605/OSF.IO/DXGK5.",
        "Abby, A. (2026, Paper 28). Cross-System Admissibility Score. DOI: 10.17605/OSF.IO/DXGK5.",
        "Abby, A. (2026, Paper 29). Constraint Jurisdiction Resolution. DOI: 10.17605/OSF.IO/DXGK5.",
        "Abby, A. (2026, Three-Layer). Three-layer constraint failure in AI systems. DOI: 10.17605/OSF.IO/DXGK5.",
        "Framework F2 -- Ve Metric. MTCP Research Programme.",
        "Framework F3 -- CPD Definition. MTCP Research Programme.",
        "Framework F4 -- BIS Definition. MTCP Research Programme.",
        "Framework F21 -- CSAS Definition. MTCP Research Programme.",
        "Framework F22 -- JRS Definition. MTCP Research Programme.",
        "Framework F23 -- TDS Definition. MTCP Research Programme.",
        "EU AI Act, Article 61: Post-market monitoring by providers. Regulation (EU) 2024/1689.",
    ]

    if is_full:
        refs.append(
            "R-AGAM Sovereign Stack Architecture. Internal Reference Document."
        )

    for ref in refs:
        add_paragraph_justified(doc, ref)


def build_public_document():
    """Build the anonymised public document."""
    doc = Document()
    set_document_defaults(doc)
    set_header(doc, "MTCP V1.0 -- Temporal Drift Monitoring")

    # Title page content
    add_heading_styled(doc, "Temporal Drift Monitoring", level=0)
    add_heading_styled(doc, "Longitudinal Validity of MTCP Grades", level=2)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "MTCP Research Programme -- Paper 30")
    add_paragraph_justified(doc, "")

    # Build all sections
    build_abstract(doc, is_full=False)
    build_introduction(doc, is_full=False)
    build_temporal_problem(doc, is_full=False)
    build_formal_framework(doc, is_full=False)
    build_temporal_stability_lemma(doc, is_full=False)
    build_monitoring_protocol(doc, is_full=False)
    build_drift_classification_response(doc, is_full=False)
    build_relationship_to_constructs(doc, is_full=False)
    build_implications(doc, is_full=False)
    build_limitations(doc, is_full=False)
    build_conclusion(doc, is_full=False)
    build_references(doc, is_full=False)

    path = os.path.join(DESKTOP, "Paper30_Temporal_Drift_Monitoring_PUBLIC.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


def build_full_document():
    """Build the full named NDA-required document."""
    doc = Document()
    set_document_defaults(doc)
    set_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
    set_footer(doc, "MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved.")

    # Title page content
    add_heading_styled(doc, "Temporal Drift Monitoring", level=0)
    add_heading_styled(doc, "Longitudinal Validity of MTCP Grades", level=2)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "CONFIDENTIAL -- NDA REQUIRED", bold=True)
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0 (Full)")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "MTCP Research Programme -- Paper 30")
    add_paragraph_justified(doc, "")

    # Build all sections
    build_abstract(doc, is_full=True)
    build_introduction(doc, is_full=True)
    build_temporal_problem(doc, is_full=True)
    build_formal_framework(doc, is_full=True)
    build_temporal_stability_lemma(doc, is_full=True)
    build_monitoring_protocol(doc, is_full=True)
    build_drift_classification_response(doc, is_full=True)
    build_relationship_to_constructs(doc, is_full=True)
    build_implications(doc, is_full=True)
    build_limitations(doc, is_full=True)
    build_conclusion(doc, is_full=True)
    build_references(doc, is_full=True)

    path = os.path.join(DESKTOP, "Paper30_Temporal_Drift_Monitoring_FULL.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


if __name__ == "__main__":
    print("Building Paper 30 -- Temporal Drift Monitoring")
    print("=" * 60)
    build_public_document()
    build_full_document()
    print("=" * 60)
    print("Done. Both documents saved to Desktop.")
