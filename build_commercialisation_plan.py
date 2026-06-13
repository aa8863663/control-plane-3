#!/usr/bin/env python3
"""Build MTCP Commercialisation Plan as .docx"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUTPUT = Path.home() / "Desktop" / "MTCP_Commercialisation_Plan.docx"


def sf(run, name='Arial', size=11, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold


def h1(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(8)
    run = p.add_run(text)
    sf(run, size=14, bold=True)


def h2(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    run = p.add_run(text)
    sf(run, size=12, bold=True)


def para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    sf(run)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.clear().add_run(text)
    sf(run)


def build():
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.space_after = Pt(4)
    run = t.add_run('MTCP Commercialisation Plan')
    sf(run, size=18, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.space_after = Pt(20)
    run = sub.add_run('Fastest Path to Paid Pilot and Sustainable Revenue\nMay 2026 | A. Abby | admin@mtcp.live')
    sf(run, size=10)

    # === SECTION 1 ===
    h1(doc, '1. Where You Are Now')
    para(doc, 'MTCP has a working product, published research, live data, and validated demand. '
        'The gap is converting that into revenue. Here is the current asset base:')
    bullet(doc, '181,504 evaluations across 35 models and 14 providers')
    bullet(doc, 'Published paper on OSF with DOI (10.17605/OSF.IO/DXGK5)')
    bullet(doc, 'Live platform at mtcp.live with public leaderboard')
    bullet(doc, 'Live MCP server for programmatic integration')
    bullet(doc, 'Full HuggingFace dataset (184,387 rows)')
    bullet(doc, 'Six independent practitioners validated in 48 hours')
    bullet(doc, 'Formal technical embedding in R-AGAM sovereign spec (binding)')
    bullet(doc, 'Joint paper with Axius SDC (decision governance runtime)')
    bullet(doc, 'Arabic constraint evaluation (only one in existence)')
    bullet(doc, 'Forming sovereign stack (R-AGAM + MTCP + Citadel Protocol)')

    # === SECTION 2 ===
    h1(doc, '2. What You Are Selling')

    h2(doc, '2.1 The Product')
    para(doc, 'MTCP sells behavioural reliability evidence for AI model deployment. '
        'It answers one question: "Can this model be trusted to follow instructions across a real conversation?"')
    para(doc, 'No other evaluation tests this. Standard benchmarks test one turn. '
        'MTCP tests three turns with correction sequences. This is the only way to measure '
        'whether a model will maintain constraints in production.')

    h2(doc, '2.2 The Deliverable')
    para(doc, 'A paid pilot delivers a per-model evidence pack that includes:')
    bullet(doc, 'MTCP grade (A through F) for each model in the buyer\'s stack')
    bullet(doc, 'Arabic constraint persistence score (for Gulf deployments)')
    bullet(doc, 'Temperature sensitivity analysis (which models degrade)')
    bullet(doc, 'Provider comparison (same model, different infrastructure)')
    bullet(doc, 'Failure taxonomy (exactly where each model breaks)')
    bullet(doc, 'Deployment recommendation per use case')
    bullet(doc, 'Board-ready evidence appendix for regulatory submission')

    h2(doc, '2.3 The Pricing')
    para(doc, 'Based on market research of Patronus AI, Credo AI, ValidMind, Lakera, and others:')
    bullet(doc, 'Discovery assessment (1 week, 3 models): GBP 15,000')
    bullet(doc, 'Standard pilot (4 weeks, full stack): GBP 35,000')
    bullet(doc, 'Enterprise evaluation (8 weeks, custom probes): GBP 75,000')
    bullet(doc, 'Annual monitoring contract: GBP 100,000 to 250,000')
    bullet(doc, 'Government/sovereign deployment: GBP 250,000+')
    para(doc, 'First deal target: GBP 35,000 pilot. This is the sweet spot. '
        'Low enough to approve without board sign-off at most organisations. '
        'High enough to establish value.')

    # === SECTION 3 ===
    h1(doc, '3. Fastest Path to First Paid Pilot')

    h2(doc, '3.1 The Compliance Wedge (2-3 months)')
    para(doc, 'Every successful AI safety company sold their first deal by mapping their tool to a regulation. '
        'The buyer is panicking about compliance. You solve their panic.')
    para(doc, 'Your compliance angles:')
    bullet(doc, 'Saudi NCA cybersecurity controls (Arabic language integrity)')
    bullet(doc, 'Saudi NDMO data governance (Arabic-mandatory digital services)')
    bullet(doc, 'SDAIA AI Ethics Principles (predictable model behaviour)')
    bullet(doc, 'EU AI Act Article 9 (risk management), Article 13 (transparency)')
    bullet(doc, 'NIST AI RMF (measure function)')
    para(doc, 'Target: Any organisation deploying LLMs in Saudi Arabia that needs to prove '
        'Arabic language compliance to NCA. They need evidence. You make evidence.')

    h2(doc, '3.2 The Arabic FinanceBench Play (3-4 months)')
    para(doc, 'Patronus AI built FinanceBench (finance-specific benchmark), published a paper showing '
        'GPT-4 fails, then financial firms came to them asking for custom evaluation.')
    para(doc, 'You have already done this for Arabic. Paper 26 is your FinanceBench. '
        'The V2 evaluation (50 probes, 4 models, 4 temperatures) makes it even stronger.')
    para(doc, 'Action: Publish Paper 26 V2 on OSF. Share the anonymised version publicly. '
        'The finding ("60% first-turn failure rate in Arabic") is the headline that gets attention.')

    h2(doc, '3.3 The UAE Validation Lab Opportunity (immediate)')
    para(doc, 'On 5 May 2026, the UAE Cyber Security Council launched a National AI Test and Validation Lab '
        'in partnership with Cisco and Open Innovation AI. They are building evaluation infrastructure NOW.')
    para(doc, 'Cisco acquired Robust Intelligence for circa USD 500M. They are already embedded in Gulf AI safety. '
        'But they do not have Arabic-specific constraint persistence testing. You do.')
    para(doc, 'Action: Research procurement contacts at UAE Cyber Security Council. '
        'Position MTCP as the Arabic-specific evaluation capability that complements their existing infrastructure.')

    h2(doc, '3.4 The KFUPM Academic Route (already in motion)')
    para(doc, 'Mohamed Rihan (R-AGAM) has opened a KFUPM affiliation pathway. '
        'This gives institutional credibility for Gulf procurement. '
        'Saudi government procurement strongly favours university-affiliated research.')
    para(doc, 'Action: Complete the two-track research programme submission (Track 1 document already done). '
        'Use KFUPM affiliation for Schmidt Sciences and Gulf procurement credibility.')

    # === SECTION 4 ===
    h1(doc, '4. Funding (Active Now)')

    h2(doc, '4.1 Immediate Actions')
    bullet(doc, 'Schmidt Sciences (trustworthyai@schmidtsciences.org): Deadline was 17 May. '
        'Email anyway expressing interest for next round. Up to USD 5M. No org required.')
    bullet(doc, 'LTFF (Long-Term Future Fund): Rolling applications. USD 40K-120K for individuals. '
        'No organisation needed. Apply with validation evidence. Fastest path to cash.')
    bullet(doc, 'Manifund: Already live at manifund.org/Mtcpaa. USD 96K ask. Push regranters.')
    bullet(doc, 'SFF Theme Rounds: June-July 2026. Need fiscal sponsor. USD 2-4M pool.')

    h2(doc, '4.2 What to Put in Applications')
    para(doc, 'Your application evidence stack:')
    bullet(doc, '181,504 evaluations (largest multi-turn dataset in existence)')
    bullet(doc, 'Published DOI on OSF')
    bullet(doc, 'Six independent practitioner validations in 48 hours')
    bullet(doc, 'First formal technical embedding (R-AGAM spec V1.0)')
    bullet(doc, 'First formal citation (Axius SDC ADR-0003)')
    bullet(doc, 'Definitive commercial framing from CAIO (Johnny Malik, 4Micro)')
    bullet(doc, 'Only Arabic constraint persistence evaluation in existence')
    bullet(doc, 'Live platform, live MCP server, live dataset')

    # === SECTION 5 ===
    h1(doc, '5. Target Buyers')

    h2(doc, '5.1 Tier 1 (highest probability of conversion)')
    bullet(doc, 'Saudi organisations deploying Arabic LLMs under NCA/NDMO requirements')
    bullet(doc, 'Gulf ERP/CRM integrators adding AI to Arabic-mandatory systems')
    bullet(doc, 'Sovereign AI companies (Humain, SDAIA) selecting models for national deployment')
    bullet(doc, 'UAE entities building AI validation infrastructure')

    h2(doc, '5.2 Tier 2 (medium-term)')
    bullet(doc, 'European enterprises needing EU AI Act compliance evidence')
    bullet(doc, 'Financial institutions under model risk management requirements')
    bullet(doc, 'AI consulting firms needing evaluation capability')
    bullet(doc, 'Model providers wanting independent reliability certification')

    h2(doc, '5.3 Tier 3 (longer-term, larger deals)')
    bullet(doc, 'Government procurement (national AI safety evaluation contracts)')
    bullet(doc, 'Cloud providers wanting to offer evaluation as a service')
    bullet(doc, 'Standards bodies (NIST, ISO) needing measurement methodology')

    # === SECTION 6 ===
    h1(doc, '6. Competitive Landscape')

    para(doc, 'Key competitors and how MTCP differs:')
    bullet(doc, 'Patronus AI (raised USD 50M): Single-turn evaluation, hallucination focus. No multi-turn. No Arabic.')
    bullet(doc, 'Lakera (raised USD 20M): Runtime protection, prompt injection focus. No constraint persistence.')
    bullet(doc, 'Robust Intelligence (Cisco, USD 500M): Model validation. No multi-turn. No Arabic.')
    bullet(doc, 'Giskard: Red teaming, open source. Multi-turn exists but no constraint persistence protocol.')
    bullet(doc, 'Credo AI: Governance platform, compliance mapping. No evaluation engine of their own.')
    bullet(doc, 'ValidMind: Financial services only. No Arabic. No multi-turn.')

    para(doc, 'MTCP unique advantages:')
    bullet(doc, 'Only evaluation that tests multi-turn constraint PERSISTENCE (not just safety)')
    bullet(doc, 'Only evaluation with Arabic-specific constraint testing')
    bullet(doc, 'Largest multi-turn dataset (181,504 evaluations)')
    bullet(doc, 'Provider comparison capability (same model, different infrastructure)')
    bullet(doc, 'Temperature sensitivity analysis (no competitor offers this)')
    bullet(doc, 'Forming sovereign stack gives full governance integration, not just evaluation')

    # === SECTION 7 ===
    h1(doc, '7. Available Models for Evaluation')

    h2(doc, '7.1 Amazon Bedrock (your account, eu-west-2)')
    para(doc, 'Currently accessible and invocable:')
    bullet(doc, 'Amazon Nova Micro, Lite, Pro, Premier, Nova-2-Lite, Nova-2-Sonic')
    bullet(doc, 'Mistral Large, Small, Magistral, Ministral-8B, Devstral, Pixtral')
    bullet(doc, 'DeepSeek R1, V3.2')
    bullet(doc, 'Qwen3-32B, Qwen3-Coder')
    bullet(doc, 'Meta Llama 4 Maverick, Scout, Llama 3.3-70B, 3.1-70B, 3.1-8B')
    bullet(doc, 'Google Gemma 3 (4B, 12B, 27B)')
    bullet(doc, 'NVIDIA Nemotron Nano, Super')
    bullet(doc, 'AI21 Jamba 1.5')
    bullet(doc, 'Cohere Command-R, Command-R-Plus')
    bullet(doc, 'Writer Palmyra X4, X5')
    bullet(doc, 'OpenAI GPT-OSS-120B, GPT-OSS-20B')
    bullet(doc, 'Moonshot Kimi-K2 Thinking')
    bullet(doc, 'ZAI GLM-4.7, GLM-5')
    bullet(doc, 'MiniMax M2, M2.1, M2.5')
    para(doc, 'Note: Claude models are listed but not enabled for invocation. '
        'Need to enable model access in AWS console.')

    h2(doc, '7.2 Fireworks AI')
    bullet(doc, 'DeepSeek V4 Pro')
    bullet(doc, 'Kimi K2.5, K2.6')
    bullet(doc, 'GLM-5p1')
    bullet(doc, 'GPT-OSS-120B')

    h2(doc, '7.3 Other Providers (API keys in .env)')
    bullet(doc, 'OpenAI: GPT-4o, GPT-4o-mini, GPT-3.5-Turbo')
    bullet(doc, 'Anthropic: Claude Sonnet 4.5, Haiku 4.5 (direct API)')
    bullet(doc, 'Groq: Llama-3.3-70B, Llama-3.1-8B, Qwen3-32B')
    bullet(doc, 'NVIDIA NIM: Llama-3.1-70B, Qwen-2.5-7B, Gemma-3-27B, Phi-4-mini')
    bullet(doc, 'Google: Gemini-2.5-Flash, Gemini-2.0-Flash')
    bullet(doc, 'Mistral: Mistral-Large, Mistral-Small, Magistral-Medium')
    bullet(doc, 'Cohere: Command-A, Command-R-Plus, Command-R7B')
    bullet(doc, 'DeepSeek: DeepSeek-R1 (direct)')
    bullet(doc, 'Cerebras: Llama-8B')
    bullet(doc, 'xAI: Grok-3-mini')
    bullet(doc, 'OpenRouter: All of the above plus more')

    # === SECTION 8 ===
    h1(doc, '8. This Week Action List')

    bullet(doc, '1. Email Schmidt Sciences expressing interest for next Trustworthy AI round')
    bullet(doc, '2. Submit LTFF application (rolling, no org needed, USD 40-120K)')
    bullet(doc, '3. Push Manifund regranters for initial funding')
    bullet(doc, '4. Publish Paper 26 V2 on OSF once Arabic run completes')
    bullet(doc, '5. Enable Claude models on Bedrock (AWS console, Model Access page)')
    bullet(doc, '6. Research UAE Cyber Security Council procurement contacts')
    bullet(doc, '7. Complete KFUPM two-track submission with Mohamed')
    bullet(doc, '8. Price first pilot offer at GBP 35K and prepare a one-page proposal template')
    bullet(doc, '9. Set up Substack for public-facing content (build audience)')
    bullet(doc, '10. Run new Bedrock models through Arabic evaluation (DeepSeek R1, Qwen3-32B, Llama 4)')

    # Footer
    p = doc.add_paragraph()
    p.space_before = Pt(30)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('MTCP | A. Abby | admin@mtcp.live | DOI: 10.17605/OSF.IO/DXGK5')
    sf(run, size=9)

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == '__main__':
    build()
