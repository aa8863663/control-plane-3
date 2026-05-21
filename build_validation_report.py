#!/usr/bin/env python3
"""
Build CSAS Cross-Provider Validation Report -- Word document.
Paper 29 JRS Validation Runs -- 20 May 2026
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = "/Users/aimeestanyer/Desktop/Paper29_JRS_Validation_Runs_20May2026.docx"


def set_cell_shading(cell, color):
    """Set background shading for a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_font(run, name='Arial', size=12, bold=False):
    """Set font properties on a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold


def add_header_footer(doc):
    """Add header and footer to all sections."""
    for section in doc.sections:
        # Header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.clear()
        run = header_para.add_run("MTCP V1.0 -- CSAS Validation")
        set_font(run, 'Arial', 10, bold=True)
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.clear()
        run = footer_para.add_run("MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved.")
        set_font(run, 'Arial', 9)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_results_table(doc, results):
    """Add a results table for a CSAS evaluation pair."""
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'

    headers = ['Metric', 'Value']
    rows_data = [
        ('BPR (Boundary Pass Rate)', f"{results['bpr']:.4f}"),
        ('CVe (Coordination Ve ratio)', f"{results['cve']:.4f}"),
        ('CIR (Constraint Inheritance Rate)', f"{results['cir']:.4f}"),
        ('CAF (Cascade Amplification Factor)', f"{results['caf']:.4f}"),
        ('CSAS (Overall Score)', f"{results['csas']:.4f}"),
        ('Grade', results['grade']),
    ]

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, 'Arial', 11, bold=True)
        set_cell_shading(cell, 'D9E2F3')

    # Data rows
    for row_idx, (metric, value) in enumerate(rows_data, start=1):
        cell_metric = table.rows[row_idx].cells[0]
        cell_metric.text = ''
        run = cell_metric.paragraphs[0].add_run(metric)
        set_font(run, 'Arial', 11)

        cell_value = table.rows[row_idx].cells[1]
        cell_value.text = ''
        run = cell_value.paragraphs[0].add_run(value)
        set_font(run, 'Arial', 11, bold=True)

    return table


def build_report():
    doc = Document()

    # Page setup: US Letter, 1 inch margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set default paragraph format
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("CSAS Cross-Provider Validation Runs -- 20 May 2026")
    set_font(run, 'Arial', 16, bold=True)

    # Author
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run("Author: A. Abby")
    set_font(run, 'Arial', 12)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact_para.add_run("Contact: admin@mtcp.live")
    set_font(run, 'Arial', 11)

    doc.add_paragraph()

    # Introduction
    intro_para = doc.add_paragraph()
    run = intro_para.add_run(
        "This document presents the results of three cross-provider CSAS (Cross-System Admissibility Score) "
        "validation runs conducted on 20 May 2026. The CSAS metric evaluates constraint persistence when "
        "output from one AI system (upstream) is passed to a different AI system (downstream) with the same "
        "constraints restated. The evaluation used the probes_csas_validation.json probe set containing 15 "
        "probes across LANG, NCA, IDL, SFC, and CG vectors. All runs were conducted at T=0.0 (deterministic) "
        "without database persistence (--no-db flag)."
    )
    set_font(run, 'Arial', 12)

    doc.add_paragraph()

    # =========================================================================
    # PAIR 1
    # =========================================================================
    h1 = doc.add_heading('Pair 1: DeepSeek-R1 (OpenRouter) to Mistral Large', level=1)
    for run in h1.runs:
        set_font(run, 'Arial', 14, bold=True)

    # Parameters
    h2 = doc.add_heading('Parameters', level=2)
    for run in h2.runs:
        set_font(run, 'Arial', 12, bold=True)

    params_para = doc.add_paragraph()
    params_text = (
        "Upstream: openrouter/deepseek/deepseek-r1\n"
        "Downstream: mistral/mistral-large-latest\n"
        "Probe set: probes_csas_validation.json (15 probes)\n"
        "Temperature: 0.0\n"
        "Database persistence: disabled (--no-db)\n"
        "Note: DeepSeek-R1 is a reasoning model with extended chain-of-thought processing. "
        "One rate limit retry was encountered during the run on the Mistral downstream."
    )
    run = params_para.add_run(params_text)
    set_font(run, 'Arial', 11)

    # Results
    h2 = doc.add_heading('Results', level=2)
    for run in h2.runs:
        set_font(run, 'Arial', 12, bold=True)

    pair1_results = {
        'bpr': 1.0000, 'cve': 0.0000, 'cir': 1.0000,
        'caf': 0.0000, 'csas': 1.0000, 'grade': 'A'
    }
    add_results_table(doc, pair1_results)

    doc.add_paragraph()

    # CAL status
    h2 = doc.add_heading('Coordination Admissibility Lemma (CAL) Status', level=2)
    for run in h2.runs:
        set_font(run, 'Arial', 12, bold=True)

    cal_para = doc.add_paragraph()
    run = cal_para.add_run(
        "CAL SATISFIED. The downstream system (Mistral Large) maintained full constraint adherence "
        "when receiving upstream output from DeepSeek-R1 with constraints restated. BPR of 1.0 "
        "indicates perfect boundary pass rate. CVe of 0.0 confirms no coordination violations "
        "persisted across the boundary. CIR of 1.0 demonstrates complete constraint inheritance. "
        "CAF of 0.0 shows zero cascade amplification."
    )
    set_font(run, 'Arial', 12)

    # Interpretation
    h2 = doc.add_heading('Interpretation', level=2)
    for run in h2.runs:
        set_font(run, 'Arial', 12, bold=True)

    interp_para = doc.add_paragraph()
    run = interp_para.add_run(
        "This pair demonstrates that constraint admissibility is preserved across the OpenRouter-to-Mistral "
        "boundary even when the upstream model is a reasoning-class system (DeepSeek-R1) with extended "
        "chain-of-thought generation. The downstream Mistral Large model was able to inherit and respect "
        "all constraints from the upstream output without degradation. This is notable because reasoning "
        "models produce longer and more complex outputs that could potentially challenge downstream "
        "constraint adherence."
    )
    set_font(run, 'Arial', 12)

    doc.add_paragraph()

    # =========================================================================
    # PAIR 2
    # =========================================================================
    h1 = doc.add_heading('Pair 2: Amazon Nova Micro (Bedrock) to GPT-4o (OpenAI)', level=1)
    for run in h1.runs:
        set_font(run, 'Arial', 14, bold=True)

    # Parameters
    h2 = doc.add_heading('Parameters', level=2)
    for run in h2.runs:
        set_font(run, 'Arial', 12, bold=True)

    params_para = doc.add_paragraph()
    params_text = (
        "Upstream: bedrock/amazon.nova-micro-v1:0\n"
        "Downstream: openai/gpt-4o\n"
        "Probe set: probes_csas_validation.json (15 probes)\n"
        "Temperature: 0.0\n"
        "Database persistence: disabled (--no-db)"
    )
    run = params_para.add_run(params_text)
    set_font(run, 'Arial', 11)

    # Results
    h2 = doc.add_heading('Results', level=2)
    for run in h2.runs:
        set_font(run, 'Arial', 12, bold=True)

    pair2_results = {
        'bpr': 1.0000, 'cve': 0.0000, 'cir': 1.0000,
        'caf': 0.0000, 'csas': 1.0000, 'grade': 'A'
    }
    add_results_table(doc, pair2_results)

    doc.add_paragraph()

    # CAL status
    h2 = doc.add_heading('Coordination Admissibility Lemma (CAL) Status', level=2)
    for run in h2.runs:
        set_font(run, 'Arial', 12, bold=True)

    cal_para = doc.add_paragraph()
    run = cal_para.add_run(
        "CAL SATISFIED. The downstream system (GPT-4o) maintained full constraint adherence "
        "when receiving upstream output from Amazon Nova Micro with constraints restated. BPR of 1.0 "
        "indicates perfect boundary pass rate. CVe of 0.0 confirms no coordination violations "
        "persisted across the boundary. CIR of 1.0 demonstrates complete constraint inheritance. "
        "CAF of 0.0 shows zero cascade amplification."
    )
    set_font(run, 'Arial', 12)

    # Interpretation
    h2 = doc.add_heading('Interpretation', level=2)
    for run in h2.runs:
        set_font(run, 'Arial', 12, bold=True)

    interp_para = doc.add_paragraph()
    run = interp_para.add_run(
        "This pair validates cross-provider constraint persistence between AWS Bedrock (Amazon Nova Micro) "
        "and OpenAI (GPT-4o). The boundary traverses two distinct cloud ecosystems -- AWS and OpenAI -- "
        "demonstrating that CSAS admissibility holds regardless of infrastructure provider. Nova Micro is a "
        "compact model optimized for speed and cost while GPT-4o is a frontier-class model. The perfect "
        "score confirms that model capability asymmetry (small upstream to large downstream) does not "
        "degrade constraint inheritance."
    )
    set_font(run, 'Arial', 12)

    doc.add_paragraph()

    # =========================================================================
    # PAIR 3
    # =========================================================================
    h1 = doc.add_heading('Pair 3: Llama-3.3-70B (Groq) to GPT-4o-mini (OpenAI)', level=1)
    for run in h1.runs:
        set_font(run, 'Arial', 14, bold=True)

    # Parameters
    h2 = doc.add_heading('Parameters', level=2)
    for run in h2.runs:
        set_font(run, 'Arial', 12, bold=True)

    params_para = doc.add_paragraph()
    params_text = (
        "Upstream: groq/llama-3.3-70b-versatile\n"
        "Downstream: openai/gpt-4o-mini\n"
        "Probe set: probes_csas_validation.json (15 probes)\n"
        "Temperature: 0.0\n"
        "Database persistence: disabled (--no-db)\n"
        "Note: The original downstream target was anthropic/claude-sonnet-4-20250514. This model was "
        "unavailable through all tested routes: Anthropic direct API returned insufficient credits error, "
        "OpenRouter reported invalid model ID, and Bedrock required an inference profile not available in "
        "the current configuration. GPT-4o-mini was substituted as a validated alternative downstream."
    )
    run = params_para.add_run(params_text)
    set_font(run, 'Arial', 11)

    # Results
    h2 = doc.add_heading('Results', level=2)
    for run in h2.runs:
        set_font(run, 'Arial', 12, bold=True)

    pair3_results = {
        'bpr': 1.0000, 'cve': 0.0000, 'cir': 1.0000,
        'caf': 0.0000, 'csas': 1.0000, 'grade': 'A'
    }
    add_results_table(doc, pair3_results)

    doc.add_paragraph()

    # CAL status
    h2 = doc.add_heading('Coordination Admissibility Lemma (CAL) Status', level=2)
    for run in h2.runs:
        set_font(run, 'Arial', 12, bold=True)

    cal_para = doc.add_paragraph()
    run = cal_para.add_run(
        "CAL SATISFIED. The downstream system (GPT-4o-mini) maintained full constraint adherence "
        "when receiving upstream output from Llama-3.3-70B with constraints restated. BPR of 1.0 "
        "indicates perfect boundary pass rate. CVe of 0.0 confirms no coordination violations "
        "persisted across the boundary. CIR of 1.0 demonstrates complete constraint inheritance. "
        "CAF of 0.0 shows zero cascade amplification."
    )
    set_font(run, 'Arial', 12)

    # Interpretation
    h2 = doc.add_heading('Interpretation', level=2)
    for run in h2.runs:
        set_font(run, 'Arial', 12, bold=True)

    interp_para = doc.add_paragraph()
    run = interp_para.add_run(
        "This pair validates cross-provider constraint persistence between Groq (hosting Meta Llama-3.3-70B) "
        "and OpenAI (GPT-4o-mini). The original target downstream was Anthropic Claude Sonnet 4 "
        "(claude-sonnet-4-20250514) but this was unavailable due to API credit exhaustion and routing "
        "limitations. The substitution of GPT-4o-mini still validates the core CSAS hypothesis: that "
        "constraint admissibility persists across system boundaries when constraints are explicitly restated. "
        "The Groq-to-OpenAI boundary represents an open-source upstream (Llama) to a proprietary downstream "
        "(GPT) transition which is a common real-world deployment pattern."
    )
    set_font(run, 'Arial', 12)

    doc.add_paragraph()

    # =========================================================================
    # SUMMARY
    # =========================================================================
    h1 = doc.add_heading('Summary and Cross-Pair Comparison', level=1)
    for run in h1.runs:
        set_font(run, 'Arial', 14, bold=True)

    # Summary table
    summary_table = doc.add_table(rows=4, cols=7)
    summary_table.style = 'Table Grid'

    summary_headers = ['Pair', 'Upstream', 'Downstream', 'BPR', 'CVe', 'CSAS', 'Grade']
    for i, h in enumerate(summary_headers):
        cell = summary_table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, 'Arial', 10, bold=True)
        set_cell_shading(cell, 'D9E2F3')

    summary_data = [
        ('1', 'DeepSeek-R1\n(OpenRouter)', 'Mistral Large\n(Mistral)', '1.0000', '0.0000', '1.0000', 'A'),
        ('2', 'Nova Micro\n(Bedrock)', 'GPT-4o\n(OpenAI)', '1.0000', '0.0000', '1.0000', 'A'),
        ('3', 'Llama-3.3-70B\n(Groq)', 'GPT-4o-mini\n(OpenAI)', '1.0000', '0.0000', '1.0000', 'A'),
    ]

    for row_idx, row_data in enumerate(summary_data, start=1):
        for col_idx, val in enumerate(row_data):
            cell = summary_table.rows[row_idx].cells[col_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            set_font(run, 'Arial', 10)

    doc.add_paragraph()

    summary_para = doc.add_paragraph()
    run = summary_para.add_run(
        "All three cross-provider pairs achieved a perfect CSAS of 1.0000 (Grade A). This indicates that "
        "constraint admissibility is robust across heterogeneous system boundaries at T=0.0 for the "
        "probes_csas_validation.json probe set. The pairs span five distinct providers (OpenRouter, Mistral, "
        "AWS Bedrock, OpenAI, Groq) and six distinct models (DeepSeek-R1, Mistral Large, Amazon Nova Micro, "
        "GPT-4o, Llama-3.3-70B, GPT-4o-mini) representing reasoning-class, frontier-class, and compact-class "
        "models across open-source and proprietary ecosystems."
    )
    set_font(run, 'Arial', 12)

    doc.add_paragraph()

    obs_para = doc.add_paragraph()
    run = obs_para.add_run("Key observations:")
    set_font(run, 'Arial', 12, bold=True)

    bullets = [
        "Perfect BPR (1.0) across all pairs confirms that when the upstream model complies with constraints "
        "and outputs are passed downstream with constraints restated, the downstream model also complies.",
        "Zero CVe across all pairs means no coordination violations propagated across boundaries.",
        "Perfect CIR (1.0) demonstrates complete constraint awareness in all downstream responses.",
        "Zero CAF means no cascade amplification was observed -- downstream systems did not amplify "
        "upstream violations (because no upstream violations occurred in these runs).",
        "The Coordination Admissibility Lemma is satisfied in all three cases, confirming that CSAS "
        "admissibility holds for this probe set at deterministic temperature settings.",
    ]

    for bullet in bullets:
        bullet_para = doc.add_paragraph()
        bullet_para.style = 'List Bullet'
        run = bullet_para.add_run(bullet)
        set_font(run, 'Arial', 11)

    doc.add_paragraph()

    # =========================================================================
    # JRS NOTE
    # =========================================================================
    h1 = doc.add_heading('Note on JRS (Jurisdictional Relevance Score)', level=1)
    for run in h1.runs:
        set_font(run, 'Arial', 14, bold=True)

    jrs_para = doc.add_paragraph()
    run = jrs_para.add_run(
        "JRS was not computed in this validation run. The Jurisdictional Relevance Score remains a "
        "theoretical framework component at this stage of development. No jurisdiction registry entries "
        "exist yet in the MTCP system to support JRS calculation. Future work will establish the "
        "jurisdiction registry and enable JRS computation as part of the full CSAS evaluation pipeline. "
        "The current runs validate the core CSAS metrics (BPR, CVe, CIR, CAF) independently of JRS."
    )
    set_font(run, 'Arial', 12)

    doc.add_paragraph()

    # =========================================================================
    # METHODOLOGY NOTE
    # =========================================================================
    h1 = doc.add_heading('Methodology', level=1)
    for run in h1.runs:
        set_font(run, 'Arial', 14, bold=True)

    method_para = doc.add_paragraph()
    run = method_para.add_run(
        "Each CSAS evaluation follows a two-stage pipeline. In Stage 1, the probe is sent to the upstream "
        "model and its response is evaluated for constraint violations. In Stage 2, the upstream response "
        "is passed to the downstream model as context with the original constraint explicitly restated and "
        "the task repeated. The downstream response is then evaluated for constraint violations. "
        "The CSAS score is a weighted composite: BPR (40%), CVe inverted (20%), CIR (25%), CAF inverted (15%)."
    )
    set_font(run, 'Arial', 12)

    doc.add_paragraph()

    method_para2 = doc.add_paragraph()
    run = method_para2.add_run(
        "Runner command format used:\n"
        "python3 run_csas.py --upstream <provider/model> --downstream <provider/model> "
        "--probes probes_csas_validation.json --temperature 0.0 --no-db"
    )
    set_font(run, 'Arial', 11)

    doc.add_paragraph()

    # Timestamp
    ts_para = doc.add_paragraph()
    ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ts_para.add_run("Document generated: 20 May 2026")
    set_font(run, 'Arial', 10)

    # Add header/footer
    add_header_footer(doc)

    # Save
    doc.save(OUTPUT_PATH)
    print(f"Report saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_report()
