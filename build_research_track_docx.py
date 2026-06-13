#!/usr/bin/env python3
"""Build MTCP Research Track Overview .docx for Gulf university submission."""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUTPUT = Path.home() / "Desktop" / "MTCP_Research_Track_Overview.docx"


def set_font(run, name='Arial', size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold


def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=14, bold=True)
        p.space_before = Pt(24)
        p.space_after = Pt(12)
    elif level == 2:
        set_font(run, size=12, bold=True)
        p.space_before = Pt(18)
        p.space_after = Pt(8)
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run)
    return p


def build():
    doc = Document()

    # Page setup: US Letter, 1-inch margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.space_after = Pt(6)
    run = title_p.add_run('MTCP Research Track Overview')
    set_font(run, size=16, bold=True)

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.space_after = Pt(4)
    run = sub_p.add_run('Track 1: Multi-Turn Constraint Persistence')
    set_font(run, size=12, bold=False)

    # Meta block
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.space_after = Pt(20)
    run = meta_p.add_run(
        'Principal Investigator: A. Abby\n'
        'Contact: admin@mtcp.live\n'
        'DOI: 10.17605/OSF.IO/DXGK5\n'
        'Date: May 2026'
    )
    set_font(run, size=10)

    # Horizontal line
    hr_p = doc.add_paragraph()
    hr_p.space_after = Pt(12)
    run = hr_p.add_run('_' * 80)
    set_font(run, size=8)
    run.font.color.rgb = None

    # === SECTION 1 ===
    add_heading_styled(doc, '1. Research Programme Overview')

    add_para(doc,
        'MTCP (Multi-Turn Constraint Persistence) is an independent research programme '
        'that evaluates whether large language models maintain instructed constraints '
        'across sustained conversational interactions. The programme addresses a '
        'fundamental gap in existing AI safety evaluation methodology.')

    add_para(doc,
        'Prior to MTCP, no published benchmark tested whether language models persist '
        'with behavioural constraints beyond a single interaction turn. Existing '
        'evaluation frameworks test model outputs in isolation. They measure whether a '
        'model can produce a correct response to one prompt. They do not measure '
        'whether the model sustains that behaviour when challenged, corrected, or '
        'placed under topical pressure across multiple conversational turns.')

    add_para(doc,
        'MTCP fills this gap with a structured three-turn evaluation protocol. The '
        'protocol issues a constrained instruction, detects constraint violations in '
        'the model response, issues a correction prompt, and measures whether the '
        'model recovers and maintains compliance. A model that violates on two '
        'consecutive turns receives a safety hard stop classification.')

    add_para(doc,
        'The programme has completed 181,504 individual evaluations across 35 distinct '
        'language models served by 14 commercial API providers. Evaluations span four '
        'temperature settings (0.0, 0.2, 0.5, 0.8) to capture behavioural variance '
        'under different inference configurations. This constitutes the largest known '
        'multi-turn constraint persistence dataset in existence.')

    # === SECTION 2 ===
    add_heading_styled(doc, '2. Research Scope and Methodology')

    add_heading_styled(doc, '2.1 Evaluation Structure', level=2)

    add_para(doc,
        'Each evaluation follows a fixed three-turn protocol. Turn 1 delivers a task '
        'instruction with an explicit behavioural constraint. The automated detection '
        'system analyses the model response for constraint violations. If a violation '
        'occurs, Turn 2 delivers a correction prompt that identifies the violation and '
        'requests compliance. If the violation persists, Turn 3 delivers a stronger '
        'correction with explicit restatement of the constraint.')

    add_para(doc,
        'A model that complies at any turn during the sequence receives a COMPLETED '
        'outcome. A model that violates on two or more consecutive turns receives a '
        'SAFETY HARD STOP outcome. This binary classification provides a clear '
        'threshold for deployment readiness assessment.')

    add_heading_styled(doc, '2.2 Grading Scale', level=2)

    add_para(doc,
        'MTCP assigns letter grades based on the Behavioural Instability Score (BIS). '
        'BIS represents the mean pass rate across all four standard temperature '
        'settings. Grade A requires 90 to 100 percent. Grade B requires 80 to 89 '
        'percent. Grade C requires 70 to 79 percent. Grade D requires 60 to 69 '
        'percent. Grade E requires 50 to 59 percent. Grade F applies below 50 percent.')

    add_para(doc,
        'No model in the current dataset achieves Grade A. The highest performing '
        'model achieves 88.8 percent (Grade B). This finding establishes that '
        'multi-turn constraint persistence remains an unsolved challenge across the '
        'entire commercial language model landscape.')

    add_heading_styled(doc, '2.3 Key Metrics', level=2)

    add_para(doc,
        'The programme tracks four primary metrics. Pass rate measures the proportion '
        'of evaluations where the model achieved compliance within three turns. Hard '
        'stop rate measures the proportion of evaluations where the model failed to '
        'achieve compliance, indicating persistent constraint violation. Recovery '
        'latency measures the number of correction turns required before compliance. '
        'BIS aggregates pass rate across temperature settings to produce a single '
        'deployment readiness score.')

    add_para(doc,
        'Additional diagnostic metrics include Control Performance Degradation (CPD), '
        'which measures the performance difference between standard probes and novel '
        'control probes. CPD identifies potential benchmark contamination effects. The '
        'mean CPD across all models is negative 28 percentage points, indicating '
        'substantial contamination sensitivity in the general model population.')

    # === SECTION 3 ===
    add_heading_styled(doc, '3. Existing Outputs')

    add_heading_styled(doc, '3.1 Published Research', level=2)

    add_para(doc,
        'The primary research paper is published on the Open Science Framework under '
        'DOI 10.17605/OSF.IO/DXGK5. The paper reports findings from 32 models across '
        '13 providers and establishes the MTCP evaluation methodology, grading scale, '
        'and initial findings. Twenty-four additional research papers have been '
        'completed covering topics including temperature failure taxonomies, flagship '
        'model regression, control performance degradation, and behavioural constraint '
        'failure as a distinct failure category.')

    add_heading_styled(doc, '3.2 Public Dataset', level=2)

    add_para(doc,
        'The full evaluation dataset is published on HuggingFace at the repository '
        'aa8899/mtcp-boundary-500. The dataset contains 184,387 probe-level results '
        'with model identity, provider, temperature, outcome, and recovery latency for '
        'each evaluation. External researchers have independently queried this dataset '
        'to validate MTCP findings.')

    add_heading_styled(doc, '3.3 Arabic Constraint Persistence', level=2)

    add_para(doc,
        'Paper 26 (Arabic Constraint Assurance) represents the first published '
        'evaluation of Arabic language constraint persistence in large language models. '
        'The evaluation tests whether models maintain Arabic-only output constraints '
        'when processing technical content that typically triggers English vocabulary '
        'insertion. Three models were evaluated via Amazon Bedrock with results ranging '
        'from 78 to 100 percent compliance. This work demonstrates MTCP applicability '
        'to language-specific governance requirements in Gulf regulatory contexts.')

    add_heading_styled(doc, '3.4 Live Infrastructure', level=2)

    add_para(doc,
        'The MTCP platform operates at mtcp.live with a live evaluation database, '
        'public leaderboard, and automated evidence generation. A Model Context '
        'Protocol (MCP) server operates at mtcp-mcp-server.fly.dev, providing '
        'programmatic access to evaluation scores, evidence packs, and regime '
        'classification data for integration with external governance systems.')

    # === SECTION 4 ===
    add_heading_styled(doc, '4. IP Ownership Statement')

    add_para(doc,
        'All intellectual property produced under the MTCP research programme remains '
        'under the sole ownership of A. Abby. This includes all research papers, '
        'evaluation datasets, probe content, detection methodologies, scoring '
        'algorithms, framework definitions, and software implementations.')

    add_para(doc,
        'No transfer, licence, or assignment of ownership is implied by institutional '
        'research association. Collaborative outputs produced jointly with co-investigators '
        'will follow standard academic co-authorship conventions. Each co-investigator '
        'retains ownership of their respective contributions.')

    add_para(doc,
        'This ownership structure is consistent with existing non-disclosure agreements '
        'governing MTCP collaboration. The programme operates under a tiered IP '
        'boundary framework. Tier 1 materials (published papers, public scores, '
        'grading scale) are freely available. Tier 2 materials (detailed evidence '
        'packs, provider comparisons) require executed NDA. Tier 3 materials (probe '
        'content, detection methodology, proprietary frameworks) remain permanently '
        'confidential.')

    # === SECTION 5 ===
    add_heading_styled(doc, '5. Research Collaboration Conditions')

    add_heading_styled(doc, '5.1 Methodological Independence', level=2)

    add_para(doc,
        'The MTCP track operates with full methodological independence within the '
        'collaborative programme. Research direction, evaluation methodology, probe '
        'design, and publication decisions remain under the sole authority of the '
        'principal investigator. Institutional association provides affiliation context '
        'and collaborative infrastructure. It does not grant editorial authority over '
        'MTCP research outputs.')

    add_heading_styled(doc, '5.2 Researcher Identity', level=2)

    add_para(doc,
        'The public researcher identity for all MTCP outputs is A. Abby. This '
        'pseudonym must be maintained in all institutional outputs, publications, '
        'communications, and administrative records associated with the research '
        'programme. The institutional partner acknowledges this requirement as a '
        'condition of research association.')

    add_heading_styled(doc, '5.3 Access Boundaries', level=2)

    add_para(doc,
        'Institutional association does not grant access to proprietary MTCP materials. '
        'Probe content (the specific instructions and constraints used in evaluations) '
        'remains confidential. The constraint detection methodology (the automated '
        'system that identifies violations) remains confidential. Private framework '
        'definitions (F16 through F19) remain confidential. Access to these materials '
        'requires separate commercial agreement independent of the research '
        'association.')

    add_para(doc,
        'The institutional partner and co-investigators receive access to published '
        'research outputs, public evaluation data, and collaborative working documents '
        'produced within the joint programme. They do not receive access to the '
        'underlying evaluation infrastructure, proprietary datasets, or detection '
        'systems that produce MTCP scores.')

    # === CLOSING ===
    add_heading_styled(doc, 'Document Control')

    add_para(doc,
        'This document describes Track 1 (MTCP) of a two-track collaborative research '
        'programme. Track 2 (R-AGAM) will be documented separately by the designated '
        'co-investigator. The combined programme submission requires both track '
        'documents for completeness.')

    # Footer-style closing
    close_p = doc.add_paragraph()
    close_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    close_p.space_before = Pt(30)
    run = close_p.add_run('End of Document')
    set_font(run, size=10, bold=False)

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == '__main__':
    build()
