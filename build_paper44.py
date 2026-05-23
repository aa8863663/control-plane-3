"""
Build Paper 44: Deployment Readiness Attestation (DRA).
Framework F34: Composite pre-deployment fidelity score.
Generates PUBLIC and FULL .docx versions.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DESKTOP = os.path.expanduser("~/Desktop")


def set_document_formatting(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.25
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return para


def add_header(doc, text):
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_footer(doc, text):
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


ABSTRACT = [
    "This paper defines the Deployment Readiness Attestation (DRA). "
    "DRA is a composite pre-deployment fidelity score. "
    "It answers one question for every deployment decision. "
    "Did this system hold its constraints before custody was asserted.",

    "DRA combines four inputs with fixed weights. "
    "BIS score at 40 percent. TDS validity at 20 percent. "
    "BEC integrity at 20 percent. Gate status at 20 percent. "
    "Output is a single number between 0.0 and 1.0.",

    "The DRA Lemma establishes that pre-deployment admissibility requires DRA above threshold. "
    "DRA is not a summary. "
    "It is a standalone attestation of pre-custodial fidelity. "
    "It is the single number every decision-maker needs.",
]

INTRODUCTION = [
    "MTCP produces multiple scores per model. "
    "BIS measures constraint persistence. "
    "TDS measures temporal validity. "
    "BEC measures evidence integrity. "
    "The Gate produces PERMIT or DENY decisions.",

    "Decision-makers need one number. "
    "A CAIO reviewing a procurement decision needs a single answer. "
    "A CDO certifying a deployment needs a single answer. "
    "A procurement officer comparing vendors needs a single answer.",

    "Currently, that answer requires interpreting multiple scores. "
    "A model with Grade B BIS, valid TDS, intact BEC, and PERMIT status is ready. "
    "But expressing this as a single attestation requires DRA.",

    "DRA provides that attestation. "
    "It compresses four inputs into one score. "
    "It compresses one score into one grade. "
    "It compresses one grade into one decision: ready or not ready.",
]

FORMAL_FRAMEWORK = [
    "DRA is computed from four inputs with fixed weights.",

    "Input 1: BIS Score. Weight 40 percent. "
    "BIS measures behavioural constraint persistence. "
    "It is the primary input because constraint persistence is the primary requirement. "
    "BIS is normalised to 0.0 to 1.0 range.",

    "Input 2: TDS Validity. Weight 20 percent. "
    "TDS validity is binary: within window or expired. "
    "A model within its validity window scores 1.0. "
    "An expired model scores 0.0. "
    "There is no partial credit for temporal validity.",

    "Input 3: BEC Integrity. Weight 20 percent. "
    "BEC integrity measures evidence chain soundness. "
    "Intact chain scores 1.0. Broken chain scores below 1.0. "
    "Any break in the chain reduces this component.",

    "Input 4: Gate Status. Weight 20 percent. "
    "Gate status is binary: PERMIT or DENY. "
    "PERMIT scores 1.0. DENY or absent scores 0.0. "
    "There is no partial credit for gate decisions.",

    "The formula: DRA = (BIS * 0.4) + (TDS * 0.2) + (BEC * 0.2) + (Gate * 0.2). "
    "Maximum possible score is 1.0. "
    "This requires perfect BIS, valid TDS, intact BEC, and PERMIT.",
]

DRA_LEMMA = [
    "The DRA Lemma defines pre-deployment admissibility.",

    "A model is pre-deployment admissible only when DRA meets the context threshold. "
    "The threshold varies by deployment context. "
    "Critical infrastructure requires Grade A (0.90). "
    "Financial services requires Grade B (0.80). "
    "General enterprise requires Grade C (0.70).",

    "DRA is not a summary of other scores. "
    "It is a standalone attestation. "
    "It attests that pre-custodial fidelity was achieved. "
    "Pre-custodial means before deployment authority was asserted.",

    "The distinction matters. "
    "A summary describes what exists. "
    "An attestation declares readiness for a specific purpose. "
    "DRA declares readiness for deployment in a specific context.",

    "A model with DRA Grade B is attested ready for financial services. "
    "The same model is not attested ready for critical infrastructure. "
    "Context determines the threshold. "
    "DRA determines whether the threshold is met.",
]

CONTEXT_THRESHOLDS = [
    "Five deployment contexts are defined with specific thresholds.",

    "Critical Infrastructure: Grade A required (DRA 0.90 or above). "
    "This includes power grid control, water treatment, transport systems. "
    "Failure consequences are catastrophic. "
    "Only the highest fidelity is acceptable.",

    "Financial Services: Grade B required (DRA 0.80 or above). "
    "This includes trading systems, credit decisions, fraud detection. "
    "Failure consequences are severe and regulated. "
    "High fidelity is required.",

    "Healthcare: Grade B required (DRA 0.80 or above). "
    "This includes clinical decision support, triage, patient communication. "
    "Failure consequences affect patient safety. "
    "High fidelity is required.",

    "Government Services: Grade B required (DRA 0.80 or above). "
    "This includes benefits administration, public communication, planning. "
    "Failure consequences affect citizens directly. "
    "High fidelity is required.",

    "General Enterprise: Grade C required (DRA 0.70 or above). "
    "This includes internal tools, content generation, data analysis. "
    "Failure consequences are contained within the organisation. "
    "Moderate fidelity is acceptable.",
]

ATTESTATION = [
    "DRA produces a signed attestation record.",

    "The attestation includes: attestation_id, model_id, context, score, grade. "
    "It includes an attestation_hash computed over all fields. "
    "It includes a valid_until timestamp aligned to TDS validity.",

    "The attestation is verifiable. "
    "Any party can recompute the hash from the attestation fields. "
    "Hash mismatch indicates tampering. "
    "Expiry indicates the attestation is no longer current.",

    "The attestation travels with the deployment decision. "
    "It is portable like the Constraint Manifest. "
    "It provides one-number evidence of pre-custodial fidelity. "
    "It does not require callback to the MTCP platform.",
]

INTEGRATION = [
    "DRA integrates with the Constraint Manifest. "
    "The manifest now includes dra_score and dra_grade fields. "
    "The manifest attests pre-custodial fidelity as a single number.",

    "DRA integrates with the Admissibility Gate. "
    "DRA is the human-readable summary of the gate decision. "
    "The gate produces PERMIT or DENY. "
    "DRA explains why with a single composite score.",

    "DRA integrates with the MCP server. "
    "The get_deployment_readiness tool returns DRA in one call. "
    "Decision-makers get score, grade, and attestation status together.",
]

LIMITATIONS = [
    "DRA weights are fixed at 40/20/20/20. "
    "These weights are proposed without empirical calibration. "
    "Future versions may adjust weights based on deployment outcome data.",

    "DRA does not include CSAS, JRS, LRP, or GRC. "
    "These are coordination and governance scores. "
    "DRA measures single-model pre-deployment readiness. "
    "Coordination readiness requires separate assessment.",

    "DRA is context-dependent. "
    "A model's DRA grade varies by deployment context. "
    "There is no context-free DRA. "
    "Every DRA score requires a specified context.",

    "TDS and Gate inputs are binary. "
    "This means small changes in TDS validity can produce large DRA changes. "
    "A model one day past validity drops from 1.0 to 0.0 on that component. "
    "This is by design. Expiry is not gradual.",
]

CONCLUSION = [
    "DRA provides the single number every decision-maker needs. "
    "It compresses four MTCP inputs into one composite score. "
    "It grades that score against context-specific thresholds. "
    "It produces a signed attestation of pre-custodial fidelity.",

    "The DRA Lemma makes pre-deployment admissibility formal. "
    "The threshold varies by context. "
    "The score is deterministic. "
    "The attestation is verifiable and portable.",

    "DRA does not replace detailed MTCP evaluation. "
    "It summarises the deployment-relevant conclusion. "
    "Detailed scores remain available for audit. "
    "DRA provides the decision layer above them.",

    "The MTCP framework now provides a complete decision stack. "
    "Evaluate (BIS). Monitor (TDS). Verify (BEC). Enforce (Gate). Attest (DRA). "
    "Each layer serves a different stakeholder. "
    "DRA serves the decision-maker.",
]

REFERENCES = [
    "Abby, A. (2026a). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 30). Temporal Drift Monitoring. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 37). Blockchain Evidence Chain Integrity. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 38). Admissibility Gate. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 39). Constraint Manifest. DOI: 10.17605/OSF.IO/DXGK5.",
    "Framework F4 -- BIS Definition. MTCP Research Programme.",
    "Framework F23 -- TDS Definition. MTCP Research Programme.",
    "Framework F28 -- BEC Definition. MTCP Research Programme.",
    "Framework F29 -- Gate Definition. MTCP Research Programme.",
    "Framework F34 -- DRA Definition. MTCP Research Programme.",
]


def build_paper(is_public=True):
    doc = Document()
    set_document_formatting(doc)

    if is_public:
        add_header(doc, "MTCP V1.0 -- Deployment Readiness Attestation")
    else:
        add_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
        add_footer(doc, "MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved.")

    add_heading_styled(doc, "MTCP Paper 44", level=1)
    add_heading_styled(doc, "Deployment Readiness Attestation", level=2)
    add_paragraph_justified(doc, "")
    if not is_public:
        add_paragraph_justified(doc, "CONFIDENTIAL -- NDA REQUIRED")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "MTCP Research Programme -- Paper 44")
    add_paragraph_justified(doc, "Framework F34")

    doc.add_page_break()

    add_heading_styled(doc, "1. Abstract", level=1)
    for para in ABSTRACT:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "2. Introduction", level=1)
    for para in INTRODUCTION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "3. Formal Framework", level=1)
    for para in FORMAL_FRAMEWORK:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "4. DRA Lemma", level=1)
    for para in DRA_LEMMA:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "5. Context Thresholds", level=1)
    for para in CONTEXT_THRESHOLDS:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "6. Attestation Record", level=1)
    for para in ATTESTATION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "7. Integration", level=1)
    for para in INTEGRATION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "8. Limitations", level=1)
    for para in LIMITATIONS:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "9. Conclusion", level=1)
    for para in CONCLUSION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "10. References", level=1)
    for ref in REFERENCES:
        add_paragraph_justified(doc, ref)

    return doc


if __name__ == "__main__":
    print("Building Paper 44: Deployment Readiness Attestation")
    print("=" * 60)

    doc_pub = build_paper(is_public=True)
    pub_path = os.path.join(DESKTOP, "Paper44_Deployment_Readiness_Attestation_PUBLIC.docx")
    doc_pub.save(pub_path)
    print(f"PUBLIC version saved: {pub_path}")

    doc_full = build_paper(is_public=False)
    full_path = os.path.join(DESKTOP, "Paper44_Deployment_Readiness_Attestation_FULL.docx")
    doc_full.save(full_path)
    print(f"FULL version saved: {full_path}")

    print("=" * 60)
    print("Done.")
