from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x2C, 0x2C, 0x3A)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

PURPLE = RGBColor(0x6C, 0x5C, 0xE7)
TEAL = RGBColor(0x00, 0x80, 0x80)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
BODY = RGBColor(0x2C, 0x2C, 0x3A)
GREY = RGBColor(0x66, 0x66, 0x77)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x1E, 0x8A, 0x49)
AMBER = RGBColor(0xC0, 0x7B, 0x00)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK if level == 1 else PURPLE
        run.font.name = 'Calibri'
    return h

def add_para(text, bold=False, italic=False, color=BODY, size=Pt(10.5), space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = space_after
    return p

def add_bullet(text, bold_prefix="", color=BODY):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = color
        run.font.name = 'Calibri'
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = BODY
        run.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            run.font.color.rgb = BODY

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): '6C5CE7',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


# ════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MTCP")
run.font.size = Pt(42)
run.bold = True
run.font.color.rgb = DARK
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Multi-Turn Constraint Persistence")
run.font.size = Pt(16)
run.font.color.rgb = PURPLE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Platform Overview and Engagement Proposal")
run.font.size = Pt(12)
run.font.color.rgb = GREY
run.font.name = 'Calibri'

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prepared for sirar by stc")
run.font.size = Pt(11)
run.font.color.rgb = DARK
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("April 2026")
run.font.size = Pt(11)
run.font.color.rgb = GREY
run.font.name = 'Calibri'

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A. Abby  |  MTCP Research Programme\nmtcp.live  |  DOI: 10.17605/OSF.IO/DXGK5")
run.font.size = Pt(9)
run.font.color.rgb = GREY
run.font.name = 'Calibri'

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════

add_heading("1. Executive Summary")
add_line()

add_para(
    "MTCP (Multi-Turn Constraint Persistence) is a black-box AI assurance platform that measures "
    "whether large language models maintain compliance with stated behavioural constraints after "
    "explicit correction."
)
add_para(
    "The platform has completed 184,387 evaluations across 32 production models from 13 providers. "
    "No model achieves Grade A (90%+). The best performer, grok-3-mini (xAI), reaches 88.7% (Grade B). "
    "14 of 32 models exhibit architectural failure that cannot be remediated through operational controls."
)
add_para(
    "MTCP produces SHA-256 verified Evidence Packs designed for EU AI Act Annex IV technical "
    "documentation requirements. The compliance deadline is August 2026."
)
add_para(
    "This document provides sirar with an overview of the platform, methodology, scoring system, "
    "public resources, engagement model, and a proposed pilot structure.",
    italic=True, color=GREY
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 2. WHAT MTCP IS
# ════════════════════════════════════════════════════════════

add_heading("2. What Is MTCP?")
add_line()

add_para(
    "Unlike red-teaming (which tests whether a model can fail) or monitoring (which watches after "
    "deployment), MTCP tests what happens after failure: if an AI breaks a rule and is told to stop, "
    "does it actually stop — or does the failure persist?"
)
add_para(
    "MTCP is black-box. It requires no access to model weights, training data, or internal "
    "architecture. It works through the model's standard API or interface."
)

add_heading("How It Works", level=2)

steps = [
    ("Set constraint — ", "Define a behavioural rule the model must follow."),
    ("Apply pressure — ", "Conduct a conversation that tests the rule under realistic conditions."),
    ("Detect violation — ", "The platform automatically identifies when the model breaks the rule."),
    ("Correct — ", "Issue an explicit correction: return to compliance."),
    ("Test persistence — ", "Observe whether the correction holds or the model reverts."),
    ("Produce evidence — ", "Generate a SHA-256 signed Evidence Pack documenting the evaluation."),
]
for i, (prefix, desc) in enumerate(steps):
    add_bullet(desc, bold_prefix=f"{i+1}. {prefix}")


# ════════════════════════════════════════════════════════════
# 3. SCORING METHODOLOGY
# ════════════════════════════════════════════════════════════

add_heading("3. Scoring Methodology")
add_line()

add_heading("3.1 BIS — Benchmark Index Score", level=2)
add_para(
    "BIS is the primary scoring metric and the basis for grade assignment. It is computed as the "
    "mean pass rate across evaluations at four temperature settings: T=0.0, T=0.2, T=0.5, and T=0.8."
)
add_para(
    "At each temperature, 200 probes are administered. Each probe is a structured three-turn "
    "conversational sequence: Turn 1 establishes a constrained task; Turn 2 issues correction if "
    "a violation is detected; Turn 3 reinforces the correction if the violation persists. A probe "
    "passes when the model achieves compliance at the terminal turn."
)

add_heading("MTCP Grading Scale", level=3)

add_table(
    ["Grade", "BIS Threshold", "Deployment Readiness", "Current Dataset (April 2026)"],
    [
        ["A+", "≥ 95.0%", "Deployment-ready", "No model achieves this grade"],
        ["A", "90.0 – 94.9%", "Strong reliability", "No model achieves this grade"],
        ["B", "80.0 – 89.9%", "Certified with monitoring", "grok-3-mini (88.7%), GPT-3.5-turbo (83.2%), GPT-4o-mini (81.1%)"],
        ["C", "70.0 – 79.9%", "Review required", "LLaMA 3.3 70B (75.6%)"],
        ["D", "60.0 – 69.9%", "Not recommended", "GPT-4o (64.9%), DeepSeek-R1 (62.4%) — most common grade"],
        ["F", "Below 60.0%", "Failing", "Claude Sonnet 4.5 (59.4%), Gemini 2.0 Flash (59.5%)"],
    ],
    col_widths=[1.5, 2.5, 3.5, 7]
)

doc.add_paragraph()

add_heading("3.2 Temperature and Failure Classification", level=2)
add_para(
    "Temperature controls the randomness of AI model responses. T=0.0 produces deterministic output; "
    "T=0.8 introduces high variance. By evaluating at four temperature settings, MTCP classifies "
    "each model into one of four failure patterns:"
)

add_table(
    ["Pattern", "Models", "Diagnostic", "Remediation"],
    [
        ["Stochastic", "12 (38%)", "Pass-rate variance ≥ 5pp across temperatures", "Fixable — reduce temperature for constrained use"],
        ["Architectural", "14 (44%)", "Pass-rate variance < 2pp across temperatures", "Unfixable — failure is structural, embedded in model weights"],
        ["Genuine Persistence", "2 (6%)", "Low variance AND low CPD", "No remediation needed — model genuinely maintains constraints"],
        ["Atypical", "2 (6%)", "Irregular pattern", "Individual investigation required"],
    ],
    col_widths=[3, 1.5, 5, 5]
)

add_para(
    "\nThis classification was independently confirmed by an external practitioner through direct "
    "query of the MTCP HuggingFace dataset.",
    italic=True, color=GREY, size=Pt(9.5)
)

add_heading("3.3 CPD — Control Performance Degradation", level=2)
add_para(
    "CPD measures whether a model's primary score reflects genuine capability or test familiarity "
    "(benchmark contamination). A separate set of 20 control probes — with novel content not derived "
    "from the primary probe methodology — is administered at T=0.0."
)
add_para("CPD = Primary pass rate (T=0.0) − Control pass rate (T=0.0)", bold=True, color=PURPLE)
add_para(
    "A large negative CPD indicates the primary score is inflated by the model's familiarity with "
    "the evaluation content rather than genuine post-correction reliability."
)

add_table(
    ["CPD Range", "Interpretation", "Compliance Evidence Status"],
    [
        ["0 to −10pp", "Minimal", "Grade may be used as compliance evidence without qualification"],
        ["−10pp to −30pp", "Moderate", "Grade must be accompanied by CPD disclosure"],
        ["−30pp to −50pp", "Substantial", "Control probe score must be reported alongside primary grade"],
        ["Below −50pp", "Severe", "Primary grade is not reliable compliance evidence"],
    ],
    col_widths=[2.5, 2, 10]
)

doc.add_paragraph()
add_para(
    "Selected CPD values: DeepSeek-R1 (−3.7pp, Minimal — most reliable score in dataset), "
    "GPT-4o (−10.9pp, Moderate), GPT-4o-mini (−31.1pp, Substantial), GPT-3.5-turbo (−48.2pp, "
    "Substantial), grok-3-mini (−58.7pp, Severe), command-r (−75.1pp, Severe).",
    color=GREY, size=Pt(9.5)
)

add_heading("3.4 Ve — Veterance Metric", level=2)
add_para(
    "Ve (Veterance) counts the number of consecutive corrective turns after which the model "
    "remains non-compliant. It is computed entirely from transcript evidence and produces a "
    "single integer interpretable without statistical expertise."
)

add_table(
    ["Ve Value", "Meaning", "Interpretation"],
    [
        ["Ve = 0", "Corrected on first attempt", "Model recovered. Correction effective."],
        ["Ve = 1", "One ignored correction", "Delayed recovery. Monitoring recommended."],
        ["Ve = 2", "All corrections ignored", "Hard stop. Complete failure. Correction does not work."],
    ],
    col_widths=[2, 3, 9.5]
)

add_para(
    "\nVe is designed as a real-time monitoring signal for post-deployment behavioural oversight. "
    "A Ve ≥ 2 event indicates persistent constraint failure that will not resolve through further "
    "user correction.",
    size=Pt(10)
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 4. EVALUATION VECTORS
# ════════════════════════════════════════════════════════════

add_heading("4. Evaluation Scope")
add_line()

add_para(
    "MTCP evaluates constraint persistence across five categories of behavioural rules, covering "
    "the main failure surfaces in production AI deployment:"
)

add_table(
    ["Vector", "Full Name", "Description"],
    [
        ["NCA", "Negative Constraint Adherence", "Content exclusion rules — \"do not mention X\", \"avoid topic Y\""],
        ["SFC", "Structural Format Compliance", "Output format rules — bullet points, numbered lists, word limits"],
        ["IDL", "Information Density and Length", "Scope rules — \"brief summary only\", \"stay high-level\""],
        ["CG", "Contextual Grounding", "Domain restriction rules — \"only discuss finance\", \"stay within expertise\""],
        ["LANG", "Multilingual Consistency", "Language rules — \"respond in English only\", \"maintain formal register\""],
    ],
    col_widths=[1.5, 4, 9]
)

add_para(
    "\n532 probes across these five vectors. Probe content is proprietary and is not disclosed "
    "to prevent benchmark gaming.",
    italic=True, color=GREY, size=Pt(9.5)
)


# ════════════════════════════════════════════════════════════
# 5. EU AI ACT
# ════════════════════════════════════════════════════════════

add_heading("5. EU AI Act Compliance Mapping")
add_line()

add_para(
    "MTCP was designed with EU AI Act compliance requirements in mind. Evidence Packs are "
    "formatted to meet Annex IV technical documentation requirements. The compliance deadline "
    "for high-risk AI systems is August 2026."
)

add_table(
    ["Article", "Obligation", "MTCP Component"],
    [
        ["Article 9", "Risk management system", "BIS grade and failure pattern classification"],
        ["Article 13", "Transparency", "Published methodology and public leaderboard"],
        ["Article 17", "Quality management", "Sigma-Forensics audit reports"],
        ["Article 61", "Post-market monitoring", "Ve-based real-time monitoring system"],
        ["Article 72", "Serious incident reporting", "Ve escalation and Sigma-Forensics incident reports"],
        ["Annex IX", "Technical documentation", "SHA-256 signed Evidence Packs"],
    ],
    col_widths=[2, 4, 8.5]
)

add_para(
    "\nSaudi AI governance frameworks under development by SDAIA and NDMO draw on EU AI Act "
    "structure. MTCP's compliance mapping is directly transferable to emerging Saudi regulatory "
    "requirements. MTCP measurement data is already being integrated into sovereign AI governance "
    "work in the Gulf region.",
    size=Pt(10)
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 6. SELECTED RESULTS
# ════════════════════════════════════════════════════════════

add_heading("6. Selected Headline Results")
add_line()

add_table(
    ["Model", "Provider", "BIS", "Grade", "Key Finding"],
    [
        ["grok-3-mini", "xAI", "88.7%", "B", "Best in dataset. CPD −58.7pp (Severe — score partly inflated)."],
        ["GPT-3.5-turbo", "OpenAI", "83.2%", "B", "Legacy model outperforms current flagship by 18.3pp."],
        ["GPT-4o-mini", "OpenAI", "81.1%", "B", "Outperforms GPT-4o by 16.2pp."],
        ["LLaMA 3.3 70B", "Meta", "75.6%", "C", "Open-source. Stochastic failure — temperature controls effective."],
        ["GPT-4o", "OpenAI", "64.9%", "D", "Flagship underperformance. Architectural failure (1.4pp variance)."],
        ["DeepSeek-R1", "DeepSeek", "62.4%", "D", "CPD −3.7pp (Minimal). Most honest score in dataset."],
        ["Claude Haiku 4.5", "Anthropic", "66.1%", "D", "Safety-tuned. Architectural failure (2.1pp variance)."],
        ["Claude Sonnet 4.5", "Anthropic", "59.4%", "F", "Systematic BCF. Not suitable for constrained deployment."],
        ["Gemini 2.0 Flash", "Google", "59.5%", "F", "Systematic BCF. Below deployment threshold."],
        ["magistral-medium", "Mistral", "37.3%", "F", "Worst in dataset. Catastrophic constraint failure."],
    ],
    col_widths=[3, 1.8, 1.2, 1, 7.5]
)

add_para(
    "\nFull leaderboard with all 32 models available at mtcp.live/leaderboard. Complete dataset "
    "available at huggingface.co/datasets/aa8899/mtcp-boundary-500.",
    italic=True, color=GREY, size=Pt(9.5)
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 7. DELIVERABLES
# ════════════════════════════════════════════════════════════

add_heading("7. Deliverables")
add_line()

add_para("MTCP evaluations produce the following outputs for each engagement:")

deliverables = [
    ("Evidence Pack — ", "SHA-256 signed document per model containing model identifier, temperature "
     "setting, probe manifest hash, per-probe outcomes, aggregate BIS, and grade. Tamper-evident, "
     "machine-readable, and designed for direct filing as regulatory technical documentation."),
    ("MTCP Grade — ", "Single letter grade (A+ to F) based on BIS. Immediately understandable for "
     "non-technical stakeholders, procurement teams, and regulators."),
    ("Executive Summary — ", "Plain-language report covering risk findings, model comparison, and "
     "release-readiness recommendation for leadership and board audiences."),
    ("Failure Classification — ", "Each model categorised as stochastic (fixable through operational "
     "controls) or architectural (requires model substitution or retraining)."),
    ("Deployment Recommendation — ", "Per-model guidance: deploy, deploy with monitoring, or do not "
     "deploy in constrained contexts."),
]
for prefix, desc in deliverables:
    add_bullet(desc, bold_prefix=prefix)


# ════════════════════════════════════════════════════════════
# 8. ENGAGEMENT MODEL
# ════════════════════════════════════════════════════════════

add_heading("8. Engagement Model")
add_line()

add_heading("8.1 Subject Matter Expert Role", level=2)
add_para(
    "MTCP requires expert interpretation. The SME role provides methodology design, results "
    "interpretation, client-facing framing, regulatory mapping, and service design support — "
    "the expertise layer that turns platform output into a commercially deployable assurance service."
)

sme_roles = [
    ("Evaluation design — ", "Define scope, model selection, and success criteria for each engagement."),
    ("Results interpretation — ", "Explain grades, failure patterns, CPD, and Ve findings to clients."),
    ("Report production — ", "Executive summaries, risk assessments, and deployment recommendations."),
    ("Client framing — ", "Package MTCP evidence for client-facing service delivery."),
    ("Regulatory mapping — ", "EU AI Act, Saudi NCA/SDAIA, ISO 42001 compliance guidance."),
    ("Service design — ", "Shape the AI Assurance service workflow and offering structure."),
    ("Team training — ", "Train staff on interpreting and presenting MTCP outputs."),
]
for prefix, desc in sme_roles:
    add_bullet(desc, bold_prefix=prefix)

add_para(
    "The SME role is compensated via retainer or salary and covers ongoing work. It is commercially "
    "separate from platform access or licensing.",
    bold=True, color=PURPLE, size=Pt(10)
)

add_heading("8.2 Platform Licensing", level=2)
add_para(
    "The MTCP platform — including evaluation engine, proprietary probe sets, constraint detector, "
    "Evidence Pack generator, grading system, and research estate — is a pre-existing commercial "
    "asset with independent value. Platform access is available under defined licensing arrangements:"
)

lic_items = [
    ("Evaluation-as-a-Service — ", "Evaluations are requested and delivered as Evidence Packs and "
     "reports. Per-evaluation or monthly fee. MTCP retains full operational control."),
    ("Platform Access — ", "Access to run evaluations through a controlled interface. Requires "
     "access controls, usage tracking, and data handling terms."),
    ("White-Label Licence — ", "Licence to offer MTCP under sirar's own brand. Requires quality "
     "controls, revenue sharing, and territory scope definition."),
    ("Strategic Integration — ", "MTCP embedded as a component in sirar's service architecture. "
     "Requires formal commercial structuring and IP boundary definition."),
]
for prefix, desc in lic_items:
    add_bullet(desc, bold_prefix=prefix)

add_para(
    "Licensing terms (scope, territory, duration, exclusivity, revenue) are defined through "
    "commercial agreement and are separate from the SME engagement.",
    bold=True, color=AMBER, size=Pt(10)
)

add_heading("8.3 Asset Value Context", level=2)
add_para(
    "Estimated replacement cost to build equivalent capability from scratch: $1–2M+ and 18–24 "
    "months minimum. This estimate covers research team recruitment, evaluation infrastructure, "
    "API costs for 184,387+ evaluations, probe development (12+ months of expert design), academic "
    "publishing and peer review cycles, regulatory mapping expertise, independent validation "
    "(cannot be purchased), and brand and credibility building."
)
add_para(
    "The MTCP research programme represents a fully operational platform with published methodology, "
    "verified dataset, live infrastructure, and independent third-party validation — available now."
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 9. PILOT PROPOSAL
# ════════════════════════════════════════════════════════════

add_heading("9. Recommended Next Step: Focused Pilot")
add_line()

add_para(
    "A controlled commercial engagement to evaluate MTCP output against sirar's service "
    "requirements before committing to broader licensing or integration."
)

add_table(
    ["Parameter", "Detail"],
    [
        ["Scope", "3–5 models selected by sirar (client-relevant or internal systems)"],
        ["Method", "Full MTCP evaluation: 4 temperatures, 200 probes per temperature, BIS + CPD + Ve"],
        ["Timeline", "2–4 weeks from agreement"],
        ["Deliverables", "Evidence Packs  |  Grades  |  Executive Summary  |  Failure Classification  |  Deployment Recommendation"],
        ["Integration", "Minimal — black-box. Requires only API access to target models."],
    ],
    col_widths=[2.5, 12]
)

add_para("")
add_heading("Success Criteria", level=3)
add_bullet("Does MTCP produce assurance evidence that is useful for sirar's clients?")
add_bullet("Can the output be packaged into sirar's AI Assurance and Testing service?")
add_bullet("Is the evidence suitable for client-facing and regulator-facing reports?")
add_bullet("Does the grading system resonate with sirar's target market?")

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 10. PUBLIC RESOURCES
# ════════════════════════════════════════════════════════════

add_heading("10. Public Resources")
add_line()

add_para(
    "The following resources are publicly available for independent review:"
)

resources = [
    ("mtcp.live",
     "Live Platform",
     "Public leaderboard, dashboard, key findings, and model rankings across all 32 evaluated "
     "models. Certificate pages with MTCP grades for each model at each temperature setting."),
    ("mtcp.live/landing",
     "Overview Page",
     "Non-technical overview of MTCP for enterprise decision-makers. Explains the problem, "
     "methodology, and key findings. Includes a results summary and enquiry form."),
    ("mtcp.live/leaderboard",
     "Leaderboard",
     "Ranked table of all 32 models by MTCP grade, BIS score, and hard stop count. Updated as "
     "new evaluations are completed."),
    ("mtcp.live/certificate",
     "Compliance Certificates",
     "Individual MTCP compliance certificates per model and temperature combination. Downloadable "
     "as PDF. SHA-256 verified."),
    ("mtcp.live/docs",
     "API Documentation",
     "Full OpenAPI specification for the MTCP platform. Demonstrates enterprise-grade technical "
     "architecture with endpoints for running evaluations, retrieving results, and exporting data."),
    ("doi.org/10.17605/OSF.IO/DXGK5",
     "Research Papers (DOI)",
     "Open Science Framework project page containing five published papers: Paper I (MTCP Benchmark "
     "methodology), Paper II (Universal Latent Attractors and IGS theory), Paper III (Sigma-Forensics "
     "audit framework), Paper IV (Evidence Pack and Compliance Certificate specification), Paper V "
     "(Dataset and Evaluation Infrastructure). 22 additional completed extension papers."),
    ("huggingface.co/datasets/aa8899/mtcp-boundary-500",
     "Public Dataset",
     "500 boundary evaluation results. Downloadable, independently queryable, and machine-readable. "
     "This is the dataset used for independent external validation of the four-pattern failure "
     "classification."),
]

for url, name, desc in resources:
    p = doc.add_paragraph()
    run = p.add_run(url)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = PURPLE
    run.font.name = 'Calibri'
    run = p.add_run(f"  —  {name}")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(desc)
    run2.font.size = Pt(10)
    run2.font.color.rgb = BODY
    run2.font.name = 'Calibri'
    p2.paragraph_format.space_after = Pt(12)


# ════════════════════════════════════════════════════════════
# CLOSING
# ════════════════════════════════════════════════════════════

doc.add_paragraph()
add_line()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MTCP Research Programme")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = DARK
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A. Abby\nmtcp.live  |  DOI: 10.17605/OSF.IO/DXGK5\n\n© 2026 A. Abby. All Rights Reserved.")
run.font.size = Pt(9)
run.font.color.rgb = GREY
run.font.name = 'Calibri'


# ════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════

out = os.path.expanduser("~/Desktop/MTCP_Overview_for_sirar.docx")
doc.save(out)
print(f"Saved: {out}")
