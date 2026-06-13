from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

DARK = RGBColor(0x1A, 0x1A, 0x2E)
PURPLE = RGBColor(0x6C, 0x5C, 0xE7)
TEAL = RGBColor(0x00, 0x80, 0x80)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
BODY = RGBColor(0x2C, 0x2C, 0x3A)
GREY = RGBColor(0x66, 0x66, 0x77)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x1E, 0x8A, 0x49)
AMBER = RGBColor(0xC0, 0x7B, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def setup_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Calibri'
    s.font.size = Pt(10.5)
    s.font.color.rgb = BODY
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.line_spacing = 1.15
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)
    return doc

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = DARK if level == 1 else PURPLE
        r.font.name = 'Calibri'
    return h

def para(doc, text, bold=False, italic=False, color=BODY, size=Pt(10.5), after=Pt(6)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = size; r.font.color.rgb = color; r.bold = bold; r.italic = italic; r.font.name = 'Calibri'
    p.paragraph_format.space_after = after
    return p

def bullet(doc, text, bold_prefix="", color=BODY):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = color; r.font.name = 'Calibri'
        r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = BODY; r.font.name = 'Calibri'
    else:
        r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = color; r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(3)
    return p

def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9); r.font.name = 'Calibri'; r.font.color.rgb = DARK
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ""
            r = c.paragraphs[0].add_run(str(val)); r.font.size = Pt(9); r.font.name = 'Calibri'; r.font.color.rgb = BODY
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width = Cm(w)
    return t

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '1', qn('w:color'): '6C5CE7'})
    pBdr.append(bottom); pPr.append(pBdr)

def sign(doc):
    doc.add_paragraph()
    divider(doc)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MTCP Research Programme  |  A. Abby"); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK; r.font.name = 'Calibri'
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("mtcp.live  |  DOI: 10.17605/OSF.IO/DXGK5\n© 2026 A. Abby. All Rights Reserved."); r.font.size = Pt(9); r.font.color.rgb = GREY; r.font.name = 'Calibri'


# ════════════════════════════════════════════════════════════════
# DOCUMENT 1 — TIER 1 CONTEXT & COMMERCIAL VALUE
# ════════════════════════════════════════════════════════════════

doc1 = setup_doc()

# title
for _ in range(4): doc1.add_paragraph()
p = doc1.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MTCP"); r.font.size = Pt(42); r.bold = True; r.font.color.rgb = DARK; r.font.name = 'Calibri'
p = doc1.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Commercial Value and Engagement Pricing"); r.font.size = Pt(16); r.font.color.rgb = PURPLE; r.font.name = 'Calibri'
p = doc1.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Confidential — Prepared for sirar by stc"); r.font.size = Pt(12); r.font.color.rgb = GREY; r.font.name = 'Calibri'
for _ in range(2): doc1.add_paragraph()
p = doc1.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("April 2026"); r.font.size = Pt(11); r.font.color.rgb = GREY; r.font.name = 'Calibri'
for _ in range(4): doc1.add_paragraph()
p = doc1.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A. Abby  |  MTCP Research Programme\nmtcp.live  |  DOI: 10.17605/OSF.IO/DXGK5"); r.font.size = Pt(9); r.font.color.rgb = GREY; r.font.name = 'Calibri'
doc1.add_page_break()


# --- 1. TIER 1 CONTEXT ---
heading(doc1, "1. Tier 1 Context")
divider(doc1)
para(doc1,
    "sirar by stc operates as a Tier 1 cybersecurity provider in Saudi Arabia. This classification "
    "carries specific implications for the type of capability, credibility, and pricing standard "
    "expected of partners and service components integrated into sirar's offering."
)

heading(doc1, "What Tier 1 Means", level=2)
bullet(doc1, "One of only six companies holding a Tier 1 Managed SOC licence from the Saudi National Cybersecurity Authority (NCA).")
bullet(doc1, "Ranked #1 MSSP in the Middle East and North Africa, #11 globally (MSSP Alert Top 250).")
bullet(doc1, "Wholly-owned subsidiary of STC Group — the largest ICT provider in the MENA region (~$50B+ market capitalisation).")
bullet(doc1, "Clients include Saudi critical national infrastructure: railways, airports, water, banking, Hajj digital systems.")
bullet(doc1, "Compliance baseline includes NCA ECC, NCA CSCC, SAMA CSF, ISO 27001, and Aramco third-party audit requirements.")

para(doc1,
    "A Tier 1 provider does not integrate unpriced, unvalidated, or informal components into its "
    "service offering. Any capability that appears in a sirar engagement must meet the same standard "
    "of commercial structuring, IP clarity, and pricing discipline that sirar applies to its existing "
    "vendor and partner relationships."
)

heading(doc1, "Implication for MTCP Engagement", level=2)
para(doc1,
    "MTCP is being evaluated as a core component of sirar's AI Assurance and Testing service — "
    "not a peripheral add-on. The capability it provides (post-correction constraint persistence "
    "evaluation, grading, and compliance evidence) does not exist elsewhere in the market. This "
    "positions MTCP as a strategic asset, not a commodity service."
)


# --- 2. TWO COMMERCIAL TRACKS ---
heading(doc1, "2. Two Separate Commercial Tracks")
divider(doc1)
para(doc1,
    "The engagement has two distinct components. They must be scoped, priced, and agreed separately. "
    "Combining them undervalues the platform asset and conflates operational work with IP access.",
    bold=True
)

heading(doc1, "Track A: Subject Matter Expert Role", level=2)
para(doc1,
    "Ongoing advisory, interpretation, and service design work. Compensated via retainer or salary. "
    "Covers the expertise required to operate MTCP within sirar's service offering."
)
para(doc1, "Scope:", bold=True, color=PURPLE)
bullet(doc1, "Evaluation design — scope, model selection, success criteria for each client engagement")
bullet(doc1, "Results interpretation — grades, failure patterns, CPD, Ve findings translated for clients")
bullet(doc1, "Report production — executive summaries, risk assessments, deployment recommendations")
bullet(doc1, "Client-facing framing — packaging MTCP evidence for sirar's service delivery")
bullet(doc1, "Regulatory mapping — EU AI Act, Saudi NCA/SDAIA, ISO 42001 compliance guidance")
bullet(doc1, "Service design — shaping the AI Assurance workflow and offering")
bullet(doc1, "Team training — training sirar staff on interpreting and presenting MTCP outputs")

heading(doc1, "Comparable Market Rates", level=3)
para(doc1, "The SME role falls between senior cybersecurity advisory and AI assurance leadership:")

table(doc1,
    ["Comparable Role", "Gulf/Saudi (tax-free)", "Global Equivalent"],
    [
        ["AI Assurance Director — Big 4 (Riyadh)", "SAR 900K–1.2M/yr ($240K–$320K)", "$270K–$380K (US, pre-tax)"],
        ["Principal Cybersecurity Consultant (KSA)", "SAR 480K–720K/yr ($128K–$192K)", "$150K–$250K (US)"],
        ["AI Safety Researcher (Anthropic/OpenAI)", "N/A (no Gulf market)", "$300K–$700K total comp"],
        ["Independent AI Consultant — Gulf day rate", "$2,000–$5,000/day", "$1,500–$4,000/day (US/UK)"],
        ["Big 4 AI Assurance billing rate to client", "$750–$1,200/day (Director)", "$500–$1,000/day (Manager)"],
    ],
    col_widths=[5.5, 4.5, 4.5]
)

doc1.add_paragraph()
para(doc1,
    "Note: There is no established market for this specific role — post-correction constraint "
    "persistence evaluation is a new field defined by the MTCP research programme. The rates above "
    "represent the closest available comparables.",
    italic=True, color=GREY, size=Pt(9.5)
)

heading(doc1, "Recommended Range — SME Role", level=3)

table(doc1,
    ["Engagement Shape", "Rate", "Basis"],
    [
        ["Retainer (2–3 days/week)", "$8,000–$15,000/month", "Below Big 4 Director. Above senior consultant. Reflects niche specialism."],
        ["Full-time equivalent", "$180,000–$280,000/year", "Benchmarked against AI Assurance Director (Big 4, Riyadh). Tax-free."],
        ["Project-based day rate", "$2,500–$4,000/day", "In line with independent AI/cyber specialist (Gulf). Niche premium justified."],
    ],
    col_widths=[3.5, 3.5, 7.5]
)

doc1.add_paragraph()
para(doc1,
    "The appropriate rate depends on time commitment, exclusivity (if any), and whether sirar "
    "requires dedicated or shared availability. The recommended starting point is a retainer "
    "covering a defined number of days per month, with project overages billed at day rate.",
    size=Pt(10)
)


heading(doc1, "Track B: Platform Licensing", level=2)
para(doc1,
    "Access to the MTCP platform — evaluation engine, proprietary probe sets, constraint detector, "
    "Evidence Pack generator, grading and certification system, and research estate. This is a "
    "pre-existing commercial asset with independent value. It is not included in SME compensation.",
    bold=True
)

heading(doc1, "Comparable Platform Pricing", level=3)

table(doc1,
    ["Platform", "Category", "Annual Licence"],
    [
        ["Credo AI", "AI governance", "$100K–$300K/year"],
        ["Holistic AI", "AI assurance", "$75K–$250K/year"],
        ["Arthur AI", "Model monitoring", "$50K–$200K/year"],
        ["Robust Intelligence (Cisco)", "AI security/testing", "$100K–$400K/year"],
        ["ServiceNow GRC", "Compliance platform", "$150K–$500K/year"],
        ["Splunk (security tier)", "Security analytics", "$150K–$1M+/year"],
    ],
    col_widths=[4, 3, 3.5]
)

doc1.add_paragraph()
para(doc1,
    "None of these platforms provide post-correction constraint persistence evaluation. MTCP "
    "occupies a distinct category with no direct competitor.",
    italic=True, color=GREY, size=Pt(9.5)
)

heading(doc1, "Licensing Options", level=3)

table(doc1,
    ["Model", "Description", "Indicative Range"],
    [
        ["Pilot", "3–5 models, 2–4 weeks, Evidence Packs + reports", "$15,000–$30,000 (one-off)"],
        ["Evaluation-as-a-Service", "Per-model evaluations delivered as Evidence Packs", "$5,000–$15,000 per model"],
        ["Annual Platform Access", "Controlled access to run evaluations directly", "$100,000–$250,000/year"],
        ["White-Label Licence", "sirar offers MTCP under own brand to clients", "$200,000–$500,000/year + revenue share"],
        ["Exclusivity (territory)", "Exclusive deployment rights for a defined region", "Significantly higher — scoped separately"],
    ],
    col_widths=[3.5, 5.5, 5.5]
)

doc1.add_paragraph()
para(doc1,
    "These ranges reflect the current market for AI assurance tooling and are adjusted for the "
    "fact that MTCP has no direct competitor. Final pricing depends on scope, territory, duration, "
    "and exclusivity terms.",
    size=Pt(10)
)


# --- 3. REPLACEMENT COST ---
heading(doc1, "3. Replacement Cost — Build vs Access")
divider(doc1)
para(doc1,
    "The alternative to engaging MTCP is to build equivalent capability in-house. The following "
    "table estimates the cost and timeline to replicate what the MTCP research programme has "
    "already produced:"
)

table(doc1,
    ["Component", "Estimated Cost", "Timeline"],
    [
        ["AI safety research team (2–3 senior researchers)", "$450K–$750K/year ongoing", "12+ months to recruit and onboard"],
        ["Platform development (evaluation engine, detectors)", "$200K–$500K", "6–12 months"],
        ["API costs for 184,387+ evaluations across 13 providers", "$50K–$100K+", "3–6 months"],
        ["Probe development (532 proprietary test scenarios)", "Included in research team cost", "12+ months of expert design"],
        ["Academic publishing (5+ papers, DOI registration)", "Included in research team cost", "6–18 months per paper for review cycles"],
        ["Independent validation", "Cannot be purchased", "Requires organic credibility over time"],
        ["Regulatory mapping (EU AI Act, Saudi frameworks)", "$100K–$200K (legal + technical)", "3–6 months"],
        ["Brand and credibility", "Cannot be accelerated", "Years of consistent output"],
    ],
    col_widths=[6, 4, 4.5]
)

doc1.add_paragraph()
para(doc1,
    "Total estimated cost: $1–2M+ and 18–24 months minimum to reach the current state of the "
    "MTCP platform, dataset, research estate, and validation record.",
    bold=True, color=PURPLE
)
para(doc1,
    "This does not account for opportunity cost. sirar's AI Assurance service launch is time-sensitive: "
    "the EU AI Act compliance deadline is August 2026, Saudi AI governance frameworks are being "
    "written now, and client demand for AI assurance evidence is growing. An 18–24 month build "
    "timeline means missing this market window entirely."
)


# --- 4. WHAT MTCP BRINGS ---
heading(doc1, "4. What MTCP Brings to the Table")
divider(doc1)

table(doc1,
    ["Asset", "Detail"],
    [
        ["184,387 evaluations", "Completed across 32 production models from 13 providers"],
        ["532 proprietary probes", "Across 5 evaluation vectors. Core competitive IP."],
        ["Live platform", "mtcp.live — in production, not a prototype"],
        ["Published research", "5 papers on OSF with registered DOI. 22 more completed."],
        ["Grading system", "A+ to F. Instantly understandable. No competitor has this."],
        ["Evidence Packs", "SHA-256 signed. Designed for EU AI Act Annex IV. Audit-ready."],
        ["Independent validation", "6 practitioners validated in 48 hours. 1 formal citation. 1 formal spec."],
        ["Sovereign AI integration", "Already being integrated into Gulf sovereign AI governance architectures."],
        ["Failure classification", "4-pattern taxonomy. Independently validated. Determines remediation."],
        ["Regulatory mapping", "EU AI Act article-level. Saudi transferable. ISO 42001 compatible."],
        ["Public dataset", "HuggingFace. Independently queryable. Verifiable."],
        ["Research estate", "20 frameworks. 16 Architecture Decision Records. 70,000+ words."],
    ],
    col_widths=[3.5, 11]
)


# --- 5. RECOMMENDED STRUCTURE ---
heading(doc1, "5. Recommended Engagement Structure")
divider(doc1)
para(doc1,
    "The recommended approach is phased, beginning with a pilot to establish value before "
    "committing to broader terms:"
)

para(doc1, "Phase 1: Pilot (immediate)", bold=True, color=PURPLE)
bullet(doc1, "3–5 models, 2–4 weeks, full MTCP evaluation")
bullet(doc1, "Deliverables: Evidence Packs, grades, executive summary, failure classification, deployment recommendation")
bullet(doc1, "Price: $15,000–$30,000 (one-off)")
bullet(doc1, "Purpose: demonstrate output quality and service fit")

para(doc1, "Phase 2: SME Retainer (post-pilot)", bold=True, color=PURPLE)
bullet(doc1, "Ongoing advisory: evaluation design, interpretation, client framing, regulatory mapping")
bullet(doc1, "Rate: retainer at $8,000–$15,000/month or equivalent full-time package")
bullet(doc1, "Shape defined based on pilot learnings and service demand")

para(doc1, "Phase 3: Platform Licensing (post-pilot)", bold=True, color=PURPLE)
bullet(doc1, "Licensing model selected based on pilot experience and service requirements")
bullet(doc1, "Options: Evaluation-as-a-Service, Platform Access, White-Label, or Strategic Integration")
bullet(doc1, "Pricing scoped as a separate commercial agreement")

para(doc1,
    "Each phase is a distinct commercial agreement. Proceeding to Phase 2 or 3 is not assumed; "
    "it depends on mutual alignment after the pilot.",
    italic=True, color=GREY, size=Pt(10)
)

sign(doc1)

out1 = os.path.expanduser("~/Desktop/MTCP_Commercial_Value_for_sirar.docx")
doc1.save(out1)
print(f"Saved: {out1}")


# ════════════════════════════════════════════════════════════════
# DOCUMENT 2 — LINKS, METHODOLOGY, AND OVERVIEW (what Syed asked for)
# ════════════════════════════════════════════════════════════════

doc2 = setup_doc()

# title
for _ in range(4): doc2.add_paragraph()
p = doc2.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MTCP"); r.font.size = Pt(42); r.bold = True; r.font.color.rgb = DARK; r.font.name = 'Calibri'
p = doc2.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Multi-Turn Constraint Persistence"); r.font.size = Pt(16); r.font.color.rgb = PURPLE; r.font.name = 'Calibri'
p = doc2.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Platform Overview, Methodology, and Public Resources"); r.font.size = Pt(12); r.font.color.rgb = GREY; r.font.name = 'Calibri'
for _ in range(2): doc2.add_paragraph()
p = doc2.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prepared for sirar by stc  |  April 2026"); r.font.size = Pt(11); r.font.color.rgb = GREY; r.font.name = 'Calibri'
for _ in range(4): doc2.add_paragraph()
p = doc2.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("184,387 evaluations  |  32 models  |  13 providers  |  532 probes  |  0 models passed"); r.font.size = Pt(10); r.font.color.rgb = PURPLE; r.font.name = 'Calibri'; r.bold = True
for _ in range(3): doc2.add_paragraph()
p = doc2.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A. Abby  |  MTCP Research Programme\nmtcp.live  |  DOI: 10.17605/OSF.IO/DXGK5"); r.font.size = Pt(9); r.font.color.rgb = GREY; r.font.name = 'Calibri'
doc2.add_page_break()


# --- 1. WHAT MTCP IS ---
heading(doc2, "1. What Is MTCP?")
divider(doc2)
para(doc2,
    "MTCP (Multi-Turn Constraint Persistence) is a black-box AI assurance platform that measures "
    "whether large language models maintain compliance with stated behavioural constraints after "
    "explicit correction."
)
para(doc2,
    "Unlike red-teaming (which tests whether a model can fail) or monitoring (which watches after "
    "deployment), MTCP tests what happens after failure and correction: does the fix hold, or does "
    "the model revert?"
)
para(doc2, "The platform is black-box — it requires no model weights, training data, or internal access. "
    "It works through the model's standard API or interface."
)

heading(doc2, "Method (6 Steps)", level=2)
steps = [
    ("Set constraint — ", "Define a behavioural rule the model must follow."),
    ("Apply pressure — ", "Conduct a conversation that tests the rule under realistic conditions."),
    ("Detect violation — ", "The platform automatically identifies when the model breaks the rule."),
    ("Correct — ", "Issue an explicit correction: return to compliance."),
    ("Test persistence — ", "Observe whether the correction holds or the model reverts."),
    ("Produce evidence — ", "Generate a SHA-256 signed Evidence Pack documenting the evaluation."),
]
for i, (prefix, desc) in enumerate(steps):
    bullet(doc2, desc, bold_prefix=f"{i+1}. {prefix}")


# --- 2. BIS ---
heading(doc2, "2. BIS — Benchmark Index Score")
divider(doc2)
para(doc2,
    "BIS is the primary scoring metric and determines the MTCP grade. It is the mean pass rate "
    "across evaluations at four temperature settings: T=0.0, T=0.2, T=0.5, and T=0.8."
)
para(doc2,
    "At each temperature, 200 probes are administered. Each probe is a structured three-turn "
    "conversation: Turn 1 establishes a constrained task, Turn 2 issues correction if violated, "
    "Turn 3 reinforces if still violated. A probe passes when the model achieves compliance at "
    "the terminal turn. BIS is the average of the four pass rates."
)

heading(doc2, "Grading Scale", level=3)
table(doc2,
    ["Grade", "BIS", "Deployment Readiness", "Current Dataset (April 2026)"],
    [
        ["A+", "≥ 95.0%", "Deployment-ready", "No model achieves this grade"],
        ["A", "90.0 – 94.9%", "Strong reliability", "No model achieves this grade"],
        ["B", "80.0 – 89.9%", "Certified with monitoring", "grok-3-mini (88.7%), GPT-3.5-turbo (83.2%), GPT-4o-mini (81.1%)"],
        ["C", "70.0 – 79.9%", "Review required", "LLaMA 3.3 70B (75.6%)"],
        ["D", "60.0 – 69.9%", "Not recommended", "GPT-4o (64.9%), DeepSeek-R1 (62.4%) — most common grade"],
        ["F", "Below 60.0%", "Failing", "Claude Sonnet 4.5 (59.4%), Gemini 2.0 Flash (59.5%)"],
    ],
    col_widths=[1.2, 2.2, 3, 8]
)


# --- 3. TEMPERATURE ---
doc2.add_paragraph()
heading(doc2, "3. Temperature and Failure Classification")
divider(doc2)
para(doc2,
    "Temperature controls the randomness of AI responses. T=0.0 produces deterministic output; "
    "T=0.8 introduces high variance. By testing at four settings, MTCP classifies each model's "
    "failure pattern:"
)

table(doc2,
    ["Pattern", "Models", "Diagnostic Criterion", "Remediation"],
    [
        ["Stochastic", "12 (38%)", "Pass-rate variance ≥ 5pp across temperatures", "Fixable — reduce temperature for constrained use"],
        ["Architectural", "14 (44%)", "Pass-rate variance < 2pp (temperature-invariant)", "Unfixable — failure is structural, requires retraining"],
        ["Genuine Persistence", "2 (6%)", "Low variance AND low CPD (< -10pp)", "No remediation needed — model works as intended"],
        ["Atypical", "2 (6%)", "Irregular pattern", "Individual investigation required"],
    ],
    col_widths=[2.5, 1.5, 5, 5.5]
)

doc2.add_paragraph()
para(doc2,
    "This classification determines whether a model's constraint failure can be addressed "
    "through operational controls (temperature, prompt engineering) or whether it requires "
    "model substitution or retraining. This is the distinction enterprises need before deployment.",
    size=Pt(10)
)
para(doc2,
    "The pattern distribution was independently validated through direct query of the MTCP "
    "HuggingFace dataset by an external practitioner.",
    italic=True, color=GREY, size=Pt(9.5)
)


# --- 4. CPD ---
heading(doc2, "4. CPD — Control Performance Degradation")
divider(doc2)
para(doc2,
    "CPD measures whether a model's score reflects genuine capability or test familiarity "
    "(benchmark contamination). AI models may score well simply because they encountered similar "
    "test content during training — not because they genuinely maintain constraints."
)
para(doc2,
    "To measure this, MTCP administers a separate set of 20 control probes with novel content "
    "the model has never seen. CPD is the difference between the primary score and the control score."
)
para(doc2, "CPD = Primary pass rate (T=0.0) − Control pass rate (T=0.0)", bold=True, color=PURPLE)

doc2.add_paragraph()
table(doc2,
    ["CPD Range", "Interpretation", "Compliance Evidence Status"],
    [
        ["0 to −10pp", "Minimal", "Grade may be used as compliance evidence without qualification"],
        ["−10pp to −30pp", "Moderate", "Grade must be accompanied by CPD disclosure"],
        ["−30pp to −50pp", "Substantial", "Control probe score must be reported alongside primary grade"],
        ["Below −50pp", "Severe", "Primary grade is not reliable compliance evidence"],
    ],
    col_widths=[2.5, 2, 10]
)

doc2.add_paragraph()
heading(doc2, "Selected CPD Values", level=3)
table(doc2,
    ["Model", "BIS", "Grade", "CPD", "Interpretation"],
    [
        ["DeepSeek-R1", "62.4%", "D", "−3.7pp", "Minimal — most honest score in dataset"],
        ["mistral-small-3.2", "—", "—", "−2.0pp", "Minimal — score reliable"],
        ["GPT-4o", "64.9%", "D", "−10.9pp", "Moderate — requires disclosure"],
        ["GPT-4o-mini", "81.1%", "B", "−31.1pp", "Substantial — score partly inflated"],
        ["GPT-3.5-turbo", "83.2%", "B", "−48.2pp", "Substantial — significant inflation"],
        ["grok-3-mini", "88.7%", "B", "−58.7pp", "Severe — primary grade unreliable alone"],
        ["command-r", "—", "—", "−75.1pp", "Severe — extreme contamination effect"],
    ],
    col_widths=[3, 1.2, 1, 1.5, 7.5]
)


# --- 5. Ve ---
doc2.add_paragraph()
heading(doc2, "5. Ve — Veterance Metric")
divider(doc2)
para(doc2,
    "Ve (Veterance) counts consecutive corrections the model ignores. It is computed entirely "
    "from the conversation transcript — no model access required — and produces a single integer."
)

table(doc2,
    ["Ve Value", "Meaning", "Interpretation"],
    [
        ["Ve = 0", "Corrected on first attempt", "Model recovered. Correction effective."],
        ["Ve = 1", "One correction ignored", "Delayed recovery. Monitoring recommended."],
        ["Ve = 2", "All corrections ignored", "Hard stop. Complete failure. Correction does not work."],
    ],
    col_widths=[2, 3, 9.5]
)

doc2.add_paragraph()
para(doc2,
    "Ve is designed as a real-time monitoring signal for post-deployment behavioural oversight. "
    "A Ve ≥ 2 event indicates persistent constraint failure that will not resolve through further "
    "user correction. It can trigger alerts, escalation, or model substitution."
)


# --- 6. VECTORS ---
heading(doc2, "6. Evaluation Vectors")
divider(doc2)
para(doc2,
    "MTCP evaluates constraint persistence across five categories of behavioural rules:"
)

table(doc2,
    ["Code", "Name", "What It Tests"],
    [
        ["NCA", "Negative Constraint Adherence", "Content exclusion rules — \"do not mention X\", \"avoid topic Y\""],
        ["SFC", "Structural Format Compliance", "Output format rules — bullet points, word limits, numbered lists"],
        ["IDL", "Information Density and Length", "Scope rules — \"brief summary only\", \"stay high-level\""],
        ["CG", "Contextual Grounding", "Domain restriction rules — \"only discuss finance\""],
        ["LANG", "Multilingual Consistency", "Language rules — \"respond in English only\", \"maintain formal register\""],
    ],
    col_widths=[1.5, 4, 9]
)

para(doc2,
    "\n532 probes across these five vectors. Probe content is proprietary.",
    italic=True, color=GREY, size=Pt(9.5)
)


# --- 7. EU AI ACT ---
heading(doc2, "7. EU AI Act Compliance Mapping")
divider(doc2)
para(doc2,
    "MTCP was designed for EU AI Act compliance. Evidence Packs meet Annex IV technical "
    "documentation requirements. Compliance deadline: August 2026."
)

table(doc2,
    ["Article", "Obligation", "MTCP Component"],
    [
        ["Article 9", "Risk management", "BIS grade and failure pattern classification"],
        ["Article 13", "Transparency", "Published methodology and public leaderboard"],
        ["Article 17", "Quality management", "Sigma-Forensics audit reports"],
        ["Article 61", "Post-market monitoring", "Ve-based real-time monitoring"],
        ["Article 72", "Serious incident reporting", "Ve escalation and Sigma-Forensics reports"],
        ["Annex IX", "Technical documentation", "SHA-256 signed Evidence Packs"],
    ],
    col_widths=[2, 3.5, 9]
)

doc2.add_paragraph()
para(doc2,
    "Saudi AI governance frameworks (SDAIA, NDMO) draw on EU AI Act structure. "
    "MTCP's compliance mapping is directly transferable.",
    size=Pt(10)
)


# --- 8. HEADLINE RESULTS ---
heading(doc2, "8. Selected Results")
divider(doc2)

table(doc2,
    ["Model", "Provider", "BIS", "Grade", "Key Finding"],
    [
        ["grok-3-mini", "xAI", "88.7%", "B", "Best in dataset. CPD −58.7pp (score partly inflated)."],
        ["GPT-3.5-turbo", "OpenAI", "83.2%", "B", "Legacy model outperforms flagship by 18.3pp."],
        ["GPT-4o-mini", "OpenAI", "81.1%", "B", "16.2pp above GPT-4o."],
        ["LLaMA 3.3 70B", "Meta", "75.6%", "C", "Open-source. Stochastic — temperature controls help."],
        ["GPT-4o", "OpenAI", "64.9%", "D", "Flagship. Architectural failure (1.4pp variance)."],
        ["DeepSeek-R1", "DeepSeek", "62.4%", "D", "CPD −3.7pp. Most honest score in dataset."],
        ["Claude Sonnet 4.5", "Anthropic", "59.4%", "F", "Systematic failure. Below deployment threshold."],
        ["Gemini 2.0 Flash", "Google", "59.5%", "F", "Systematic failure."],
        ["magistral-medium", "Mistral", "37.3%", "F", "Worst in dataset."],
    ],
    col_widths=[3, 1.8, 1.2, 1, 7.5]
)


# --- 9. PUBLIC RESOURCES ---
doc2.add_page_break()
heading(doc2, "9. Public Resources")
divider(doc2)
para(doc2, "All of the following are publicly available for independent review:")

resources = [
    ("mtcp.live", "Live Platform",
     "Public leaderboard showing all 32 models ranked by grade. Dashboard with key findings, "
     "model comparisons, and aggregate statistics. Certificate pages with MTCP grades per model."),
    ("mtcp.live/landing", "Overview Page",
     "Non-technical overview for enterprise decision-makers. Explains the problem, methodology, "
     "and key findings. Results summary and enquiry form."),
    ("mtcp.live/leaderboard", "Leaderboard",
     "Ranked table of all 32 models by MTCP grade, BIS score, and hard stop count."),
    ("mtcp.live/certificate", "Compliance Certificates",
     "Individual MTCP certificates per model and temperature. Downloadable as PDF. SHA-256 verified."),
    ("mtcp.live/docs", "API Documentation",
     "Full OpenAPI specification. Endpoints for running evaluations, retrieving results, and "
     "exporting data. Demonstrates enterprise-grade architecture."),
    ("doi.org/10.17605/OSF.IO/DXGK5", "Research Papers (DOI)",
     "Five published papers: Paper I (MTCP Benchmark), Paper II (Latent Attractors and IGS Theory), "
     "Paper III (Sigma-Forensics Audit Framework), Paper IV (Evidence Pack Specification), "
     "Paper V (Dataset Documentation). 22 additional completed papers."),
    ("huggingface.co/datasets/aa8899/mtcp-boundary-500", "Public Dataset",
     "500 boundary evaluation results. Downloadable, independently queryable, machine-readable. "
     "Used for external validation of the four-pattern failure classification."),
]

for url, name, desc in resources:
    p = doc2.add_paragraph()
    r = p.add_run(url); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = PURPLE; r.font.name = 'Calibri'
    r = p.add_run(f"  —  {name}"); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK; r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    p2 = doc2.add_paragraph()
    r2 = p2.add_run(desc); r2.font.size = Pt(10); r2.font.color.rgb = BODY; r2.font.name = 'Calibri'
    p2.paragraph_format.space_after = Pt(12)

sign(doc2)

out2 = os.path.expanduser("~/Desktop/MTCP_Platform_Overview_for_sirar.docx")
doc2.save(out2)
print(f"Saved: {out2}")
