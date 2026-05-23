"""Build Paper 28: Cross-System Admissibility Score (CSAS) - PUBLIC and FULL versions."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def build_paper(full_version=True):
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.25

    # Header
    header_section = doc.sections[0]
    header = header_section.header
    header_para = header.paragraphs[0]
    if full_version:
        header_para.text = 'CONFIDENTIAL -- NDA REQUIRED'
    else:
        header_para.text = 'MTCP V1.0 -- Cross-System Admissibility'
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header_para.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.bold = True

    # Footer
    footer = header_section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = 'MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved.'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(9)

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        return h

    def add_para(text, bold=False, italic=False, alignment=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
        if alignment:
            p.paragraph_format.alignment = alignment
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
            cell._tc.get_or_add_tcPr().append(shading)
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return table

    # ===== TITLE PAGE =====
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(24)
    run = title_p.add_run(
        'Cross-System Admissibility Score (CSAS):\n'
        'Extending MTCP to Multi-System Coordination Governance'
    )
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.bold = True

    add_para('A. Abby', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para('admin@mtcp.live', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para('DOI: 10.17605/OSF.IO/DXGK5', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para('May 2026', alignment=WD_ALIGN_PARAGRAPH.CENTER)

    if full_version:
        add_para('Paper 28, Version 1.0 (Full Named)', italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        add_para('Paper 28, Version 1.0 (Anonymised)', italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ===== 1. ABSTRACT =====
    add_heading('1. Abstract', level=1)

    add_para(
        'MTCP evaluates constraint persistence within a single model. '
        'Production deployments increasingly coordinate multiple models in sequence or parallel. '
        'No existing framework evaluates whether two independently admissible systems preserve admissibility at their coordination boundary. '
        'We introduce the Cross-System Admissibility Score (CSAS). '
        'CSAS measures whether constraint persistence holds when System A passes output to System B. '
        'The metric is transcript-computable. '
        'It requires no access to model internals. '
        'It extends BIS, Ve, and the Persistence Lemma to multi-system coordination. '
        'We define the Coordination Admissibility Lemma. '
        'We present a measurement protocol using existing MTCP infrastructure. '
        'We define a failure taxonomy for coordination-specific constraint failures. '
        'Total evaluations in the broader MTCP dataset: 181,504 across 32 models.'
    )

    kw_p = doc.add_paragraph()
    run = kw_p.add_run('Keywords: ')
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run2 = kw_p.add_run(
        'CSAS, cross-system admissibility, multi-model coordination, '
        'constraint propagation, MTCP, Ve, Persistence Lemma, coordination governance'
    )
    run2.font.name = 'Arial'
    run2.font.size = Pt(12)

    # ===== 2. INTRODUCTION =====
    add_heading('2. Introduction', level=1)

    add_para(
        'MTCP evaluates a single model against a single constraint set. '
        'The Persistence Lemma establishes that Ve above threshold means admissibility is no longer reliably resolvable. '
        'This applies within a bounded corrective dialogue with one model.'
    )
    add_para(
        'Production AI systems rarely operate as single models. '
        'Agentic architectures, retrieval-augmented generation pipelines, and multi-model orchestration systems pass outputs between models. '
        'A planning model generates instructions for an execution model. '
        'A summarisation model feeds a decision model. '
        'A code generation model feeds a code review model.'
    )
    add_para(
        'Each model in such a system may pass its own MTCP evaluation independently. '
        'System A achieves Grade B. '
        'System B achieves Grade B. '
        'The question this paper addresses: does the coordination of A and B preserve admissibility?'
    )
    add_para(
        'The answer is not guaranteed. '
        'Three failure modes exist at the coordination boundary that individual MTCP evaluation cannot detect. '
        'First: constraint dilution, where A passes a compliant output that B reinterprets outside the constraint scope. '
        'Second: constraint inheritance failure, where B does not inherit the constraint that governed A. '
        'Third: cascade amplification, where a marginal Ve=1 event in A triggers a Ve=2 hard stop in B.'
    )
    add_para(
        'CSAS provides a formal metric for coordination admissibility. '
        'It is designed to be measurable using existing MTCP infrastructure. '
        'It does not require new probe types. '
        'It requires a coordination evaluation protocol that chains existing probes across system boundaries.'
    )

    doc.add_page_break()

    # ===== 3. FORMAL FRAMEWORK =====
    add_heading('3. Formal Framework', level=1)

    add_heading('3.1 Setting', level=2)

    add_para(
        'Let S = {M_1, M_2, ..., M_k} be a coordinated system of k models. '
        'Let C be a constraint set in force across the system. '
        'Let the output of M_i at turn t be denoted O_i(t). '
        'Let the input to M_j at turn t+1 include O_i(t) either directly or through transformation.'
    )
    add_para(
        'A coordination boundary B(i,j) exists wherever output from M_i becomes input to M_j. '
        'At each boundary, constraint C may be preserved, diluted, or lost.'
    )

    add_heading('3.2 CSAS Definition', level=2)

    add_para('Definition 1 (Cross-System Admissibility Score).', bold=True)
    add_para(
        'For a coordinated system S operating under constraint set C, CSAS is defined as follows. '
        'N evaluation probes are applied at each coordination boundary.'
    )
    add_para(
        'CSAS(S, C) = (1/|B|) * sum over all boundaries B(i,j) of PR(B(i,j), C)',
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    add_para(
        'Where PR(B(i,j), C) is the pass rate of constraint C across boundary B(i,j). '
        'A probe passes at a boundary if M_j maintains constraint C. '
        'The pass condition requires C to hold given M_i output as input.'
    )
    add_para(
        'CSAS ranges from 0 (total constraint loss at all boundaries) to 1 (perfect constraint preservation at all boundaries). '
        'CSAS = 1 means the system preserves constraints as reliably as each component does individually. '
        'CSAS below the component BIS indicates coordination-induced degradation.'
    )

    add_heading('3.3 Component Metrics', level=2)

    add_para('CSAS decomposes into four component metrics.', bold=True)

    add_para('Component 1: Boundary Pass Rate (BPR).', bold=True)
    add_para(
        'BPR(B(i,j), C) measures the proportion of probes where M_j maintains constraint C. '
        'Only probes where M_i output was itself compliant are included. '
        'BPR isolates the boundary effect. '
        'It answers: given a compliant input, does the receiving model stay compliant?'
    )

    add_para('Component 2: Coordination Ve (CVe).', bold=True)
    add_para(
        'CVe extends Veterance to the coordination context. '
        'CVe counts the number of coordination boundaries at which a constraint violation persists without correction. '
        'Under a two-system chain (k=2), CVe ranges from 0 to 1. '
        'Under a three-system chain (k=3), CVe ranges from 0 to 2. '
        'CVe above threshold (CVe >= k-1) indicates that the violation has propagated through all downstream systems without recovery.'
    )

    add_para('Component 3: Constraint Inheritance Rate (CIR).', bold=True)
    add_para(
        'CIR measures whether M_j operates under constraint C at all. '
        'CIR = proportion of probes where M_j demonstrates awareness of constraint C in its output. '
        'CIR below 1.0 indicates that the constraint was not inherited across the boundary. '
        'This is distinct from BPR. '
        'BPR measures compliance given awareness. '
        'CIR measures whether awareness exists.'
    )

    add_para('Component 4: Cascade Amplification Factor (CAF).', bold=True)
    add_para(
        'CAF measures whether marginal failures in M_i become severe failures in M_j. '
        'CAF = (Ve at M_j output) / (Ve at M_i output) for probes where M_i produced a marginal violation (Ve=1). '
        'CAF > 1 indicates amplification: a recoverable deviation in M_i becomes a hard stop in M_j. '
        'CAF = 1 indicates no amplification. '
        'CAF < 1 indicates dampening: M_j corrects the upstream violation.'
    )

    doc.add_paragraph()
    add_para('Summary of CSAS components:', bold=True)
    add_table(
        ['Component', 'Symbol', 'Range', 'Measures'],
        [
            ['Boundary Pass Rate', 'BPR', '0 to 1', 'Constraint maintenance at boundary'],
            ['Coordination Ve', 'CVe', '0 to k-1', 'Violation propagation depth'],
            ['Constraint Inheritance Rate', 'CIR', '0 to 1', 'Constraint awareness transfer'],
            ['Cascade Amplification Factor', 'CAF', '0 to inf', 'Failure severity escalation'],
        ]
    )

    doc.add_page_break()

    # ===== 4. COORDINATION ADMISSIBILITY LEMMA =====
    add_heading('4. The Coordination Admissibility Lemma', level=1)

    add_heading('4.1 Statement', level=2)

    add_para('Lemma 2 (Coordination Admissibility Lemma).', bold=True)
    add_para(
        'Let M_i and M_j be two models in a coordinated system S. '
        'Each is independently evaluated under MTCP with BIS above the admissibility threshold. '
        'Let B(i,j) be the coordination boundary between them. '
        'The following conditions are jointly necessary for coordination admissibility:'
    )
    add_para(
        '(C1) BPR(B(i,j), C) >= min(BIS(M_i), BIS(M_j)). '
        'The boundary must not introduce degradation beyond what the weaker component produces independently.'
    )
    add_para(
        '(C2) CIR(B(i,j), C) = 1.0. '
        'The receiving model must inherit all constraints governing the sending model.'
    )
    add_para(
        '(C3) CAF(B(i,j)) <= 1.0. '
        'The boundary must not amplify upstream failures.'
    )
    add_para(
        'If any condition fails, coordination admissibility is not established. '
        'Individual model admissibility does not guarantee system admissibility.'
    )

    add_heading('4.2 Relationship to the Persistence Lemma', level=2)

    add_para(
        'The Persistence Lemma (Abby, 2026i) states that Ve above threshold means admissibility is no longer reliably resolvable. '
        'This applies within the bounded conversational window. '
        'The Coordination Admissibility Lemma extends this principle to system boundaries.'
    )
    add_para(
        'The Persistence Lemma addresses temporal persistence (across turns). '
        'The Coordination Admissibility Lemma addresses spatial persistence (across system boundaries). '
        'Both share the core property: admissibility is not self-maintaining. '
        'It must be actively verified at every point where constraint state could change.'
    )
    add_para(
        'The threshold concept extends directly. '
        'CVe above (k-1) in a k-system chain is the coordination equivalent of Ve above 2 in a three-turn probe. '
        'Both indicate that the failure has propagated through all available recovery points without resolution.'
    )

    add_heading('4.3 Formal Properties', level=2)

    add_para(
        'CSAS inherits the formal properties of Ve (Abby, 2026i, Section 3.2):'
    )
    add_para(
        '(P1) Non-negativity. CSAS >= 0 always.'
    )
    add_para(
        '(P2) Boundedness. CSAS <= 1. CSAS cannot exceed perfect boundary preservation.'
    )
    add_para(
        '(P3) Computability from observables. CSAS requires only the outputs of each model at each boundary. No model-internal access is required.'
    )
    add_para(
        '(P4) Determinism given transcript. For a fixed evaluation transcript, CSAS is uniquely determined.'
    )
    add_para(
        '(P5) Independence from component BIS. Two systems with identical component BIS may have different CSAS, reflecting different coordination failure profiles.'
    )

    doc.add_page_break()

    # ===== 5. MEASUREMENT PROTOCOL =====
    add_heading('5. Measurement Protocol', level=1)

    add_heading('5.1 Design Principle', level=2)

    add_para(
        'CSAS evaluation uses existing MTCP infrastructure. '
        'No new probe types are required. '
        'The protocol chains standard MTCP probes across coordination boundaries.'
    )

    add_heading('5.2 Protocol Steps', level=2)

    add_para('Step 1: Component Evaluation.', bold=True)
    add_para(
        'Evaluate each model M_i independently using standard MTCP protocol. '
        'Record BIS(M_i), Ve distribution, and CPD for each component. '
        'This establishes the individual admissibility baseline.'
    )

    add_para('Step 2: Boundary Identification.', bold=True)
    add_para(
        'Map all coordination boundaries B(i,j) in the system. '
        'For each boundary, identify the constraint set C that must be preserved across it. '
        'Record the data transformation (if any) applied to M_i output before it reaches M_j.'
    )

    add_para('Step 3: Boundary Probe Execution.', bold=True)
    add_para(
        'For each boundary B(i,j), execute the following sequence. '
        'Present a standard MTCP probe to M_i. '
        'Capture the output O_i. '
        'Present O_i as input context to M_j with constraint C restated. '
        'Evaluate M_j output for constraint compliance. '
        'Record pass or fail at the boundary.'
    )

    add_para('Step 4: Coordination Ve Computation.', bold=True)
    add_para(
        'For probes where M_i output was non-compliant (Ve >= 1 at M_i), track whether M_j output is also non-compliant. '
        'CVe = number of boundaries at which the violation persists without correction. '
        'Record the CAF for each such probe.'
    )

    add_para('Step 5: CSAS Computation.', bold=True)
    add_para(
        'Compute BPR, CIR, and CAF for each boundary. '
        'Compute aggregate CSAS as the mean BPR across all boundaries. '
        'Report CSAS alongside component BIS values and the Coordination Admissibility Lemma condition status (C1, C2, C3 each pass or fail).'
    )

    add_heading('5.3 Minimum Evaluation Requirements', level=2)

    add_para(
        'Minimum 50 probes per boundary. '
        'Minimum 2 temperature settings (T=0.0 and T=0.5). '
        'All five MTCP vectors must be represented in the boundary probe set. '
        'The boundary evaluation must include at least 10 probes where M_i produces a marginal violation (Ve=1) to enable CAF computation.'
    )

    add_heading('5.4 Infrastructure Requirements', level=2)

    add_para(
        'The evaluation requires only API access to each model in the chain. '
        'No model internals, weights, or training data are needed. '
        'The same constraint detector and violation detector used in standard MTCP evaluation apply at each boundary. '
        'Transcript evidence standards (Abby, 2026l) apply to all boundary evaluation records.'
    )

    doc.add_page_break()

    # ===== 6. FAILURE TAXONOMY =====
    add_heading('6. Coordination Failure Taxonomy', level=1)

    add_para(
        'CSAS detects four classes of coordination failure. '
        'None of these are detectable through individual model evaluation alone.'
    )

    add_heading('6.1 Type 1: Constraint Dilution', level=2)

    add_para(
        'M_i produces compliant output. '
        'M_j receives that output and reinterprets it outside the constraint scope. '
        'The constraint was not violated by M_i. '
        'It was lost at the boundary.'
    )
    add_para(
        'Example: M_i is instructed to respond only in Arabic. '
        'M_i produces compliant Arabic output. '
        'M_j receives the Arabic text as context and responds in English. '
        'M_j was never told the Arabic constraint applies to its output.'
    )
    add_para(
        'Detection metric: CIR below 1.0 with BPR below component BIS. '
        'Remediation: explicit constraint forwarding in the coordination protocol.'
    )

    add_heading('6.2 Type 2: Constraint Inheritance Failure', level=2)

    add_para(
        'The constraint governing M_i is not transmitted to M_j at all. '
        'M_j operates under its default constraint set. '
        'Any constraint violation at M_j is not a failure of M_j. '
        'It is a failure of the coordination architecture.'
    )
    add_para(
        'This is an architectural coordination failure. '
        'It cannot be fixed by improving M_j. '
        'It requires modification of the system prompt, orchestration logic, or constraint propagation mechanism.'
    )
    add_para(
        'Detection metric: CIR = 0 for the relevant constraint. '
        'The constraint simply does not exist in the downstream context.'
    )

    add_heading('6.3 Type 3: Cascade Amplification', level=2)

    add_para(
        'M_i produces a marginal violation (Ve=1, recoverable). '
        'M_j receives the non-compliant output and amplifies the failure. '
        'M_j produces a hard stop (Ve=2 equivalent) from input that was only a marginal upstream violation.'
    )
    add_para(
        'This is the most dangerous coordination failure type. '
        'Each component independently would recover. '
        'The coordination makes recovery impossible. '
        'The system fails where no individual component would fail alone.'
    )
    add_para(
        'Detection metric: CAF > 1 on probes where upstream Ve = 1. '
        'The Persistence Lemma applies at the system level: cascade amplification means admissibility is no longer resolvable through downstream correction.'
    )

    add_heading('6.4 Type 4: Constraint State Drift', level=2)

    add_para(
        'Over multiple coordination cycles, the effective constraint weakens. '
        'No single boundary crossing violates the constraint. '
        'The cumulative effect of multiple boundary crossings degrades constraint precision. '
        'After n cycles, the constraint is no longer enforceable.'
    )
    add_para(
        'This is analogous to cross-session BCF propagation (Abby, 2026d, Section 9). '
        'Each cycle introduces a small drift. '
        'No single drift is a violation. '
        'The accumulated drift becomes a constraint failure. '
        'Detection requires longitudinal CSAS measurement across multiple coordination cycles.'
    )
    add_para(
        'Detection metric: CSAS declining over repeated evaluation cycles. '
        'No single cycle shows failure. '
        'The trend is the signal.'
    )

    doc.add_paragraph()
    add_para('Summary of coordination failure types:', bold=True)
    add_table(
        ['Type', 'Mechanism', 'Detection', 'Remediation'],
        [
            ['1. Constraint Dilution', 'Boundary reinterpretation', 'CIR < 1, BPR < BIS', 'Constraint forwarding'],
            ['2. Inheritance Failure', 'Constraint not transmitted', 'CIR = 0', 'Architecture modification'],
            ['3. Cascade Amplification', 'Marginal failure amplified', 'CAF > 1', 'Boundary dampening'],
            ['4. State Drift', 'Cumulative degradation', 'CSAS declining over time', 'Periodic re-evaluation'],
        ]
    )

    doc.add_page_break()

    # ===== 7. CSAS GRADING =====
    add_heading('7. CSAS Grading Scale', level=1)

    add_para(
        'CSAS maps to the existing MTCP grading scale. '
        'The grade represents coordination readiness, not individual model quality.'
    )

    add_table(
        ['Grade', 'CSAS Range', 'Coordination Recommendation'],
        [
            ['A', '0.90 to 1.00', 'Coordination admissible. Deploy with standard monitoring.'],
            ['B', '0.80 to 0.89', 'Coordination admissible with boundary monitoring.'],
            ['C', '0.70 to 0.79', 'Coordination requires guardrails at each boundary.'],
            ['D', '0.60 to 0.69', 'Coordination not recommended without intervention.'],
            ['F', 'Below 0.60', 'Coordination inadmissible. Re-architect system.'],
        ]
    )

    add_para(
        'A system may have component models graded B individually but achieve only Grade D on CSAS. '
        'This indicates that the coordination architecture introduces constraint failure that neither component exhibits alone. '
        'The coordination is the problem, not the components.'
    )

    doc.add_page_break()

    # ===== 8. RELATIONSHIP TO EXISTING CONSTRUCTS =====
    add_heading('8. Relationship to Existing MTCP Constructs', level=1)

    add_heading('8.1 BIS and CSAS', level=2)
    add_para(
        'BIS measures single-model constraint persistence across temperatures. '
        'CSAS measures constraint persistence across system boundaries. '
        'BIS is a prerequisite for CSAS: a model that fails individually cannot contribute to a passing coordination. '
        'But BIS is not sufficient: high BIS does not guarantee high CSAS.'
    )

    add_heading('8.2 Ve and CVe', level=2)
    add_para(
        'Ve counts corrective turns without recovery within a single model. '
        'CVe counts coordination boundaries without recovery across multiple models. '
        'Both share the threshold property: above threshold, admissibility is no longer resolvable. '
        'Ve operates in the temporal dimension (turns). '
        'CVe operates in the spatial dimension (system boundaries).'
    )

    add_heading('8.3 CPD and Coordination CPD', level=2)
    add_para(
        'CPD measures probe-structure familiarity for a single model. '
        'Coordination CPD would measure whether a system performs differently on coordinated probes versus isolated probes. '
        'A system with high coordination CPD passes individual evaluation but fails coordinated evaluation. '
        'This indicates that the coordination context introduces novel failure conditions not captured by standard probes.'
    )

    add_heading('8.4 IGS and Coordination IGS', level=2)
    add_para(
        'IGS (Identity-Gate Satiation) describes architectural constraint suppression in a single model. '
        'Coordination IGS describes architectural constraint suppression at a system boundary. '
        'A coordination architecture that systematically strips constraint context before passing outputs downstream exhibits coordination IGS. '
        'It is temperature-invariant at the system level. '
        'It is not fixable by adjusting individual models.'
    )

    add_heading('8.5 PCD and CSAS', level=2)
    add_para(
        'Pipeline Constraint Degradation (Abby, 2026g) measures how constraint persistence degrades inside a pipeline. '
        'PCD compares isolated model performance against pipeline-context performance. '
        'CSAS formalises and extends this concept. '
        'PCD measures the magnitude of degradation. '
        'CSAS identifies the mechanism (dilution, inheritance failure, cascade, drift) and the location (which boundary). '
        'PCD is a scalar diagnostic. '
        'CSAS is a structured evaluation framework with component metrics and a failure taxonomy.'
    )

    add_heading('8.6 Sigma-Forensics and CSAS Audit', level=2)
    add_para(
        'The Sigma-Forensics five-stage audit framework (Abby, 2026c) applies to CSAS evaluation with one extension. '
        'Stage 2 (transcript evidence collection) must include boundary-crossing transcripts. '
        'The output of M_i and the input to M_j at each boundary must be preserved as a linked pair. '
        'Hash verification applies to the boundary pair, not just individual transcripts.'
    )

    doc.add_page_break()

    # ===== 9. DISCUSSION =====
    add_heading('9. Discussion', level=1)

    add_heading('9.1 Why Individual Evaluation Is Insufficient', level=2)
    add_para(
        'The MTCP dataset demonstrates that 32 models produce variable constraint persistence. '
        'No model achieves Grade A on the full evaluation. '
        'When two imperfect systems coordinate, their imperfections can compound. '
        'The probability of constraint preservation at a boundary is not the product of individual pass rates. '
        'Coordination introduces failure modes that neither component exhibits alone.'
    )
    add_para(
        'The theoretical basis is the Sigma-Attractor model (Abby, 2026b). '
        'Each model maintains a trajectory in rhetorical state space. '
        'RLHF creates a potential well (the Assistant Basin, B_a) that attracts trajectories toward compliant behaviour. '
        'High structural entropy at a coordination boundary can function as a perturbation. '
        'The perturbation may push the receiving model from B_a into the Diagnostic Attractor (A_d). '
        'Once in A_d, corrective responsiveness is suppressed. '
        'This is the mechanism by which coordination boundaries introduce constraint failure.'
    )

    add_heading('9.2 Implications for Agentic Architectures', level=2)
    add_para(
        'Agentic AI systems chain multiple models in sequence. '
        'A planning agent generates tasks. '
        'An execution agent carries them out. '
        'A review agent evaluates results. '
        'Each agent may pass MTCP independently. '
        'The system as a whole may fail CSAS. '
        'Constraint dilution at the planning-to-execution boundary is the highest risk point.'
    )
    add_para(
        'Agent repositories compound this risk (Abby, 2026h). '
        'Persistent memory layers write upstream outputs to durable storage. '
        'Downstream models inherit poisoned context as ground truth. '
        'Hook interception modifies corrections before they reach downstream models. '
        'CSAS Type 4 (Constraint State Drift) is the formal measurement for this pattern.'
    )

    add_heading('9.3 Implications for EU AI Act Compliance', level=2)
    add_para(
        'EU AI Act Article 9 requires risk management for AI systems. '
        'A multi-model system is a single AI system for regulatory purposes. '
        'Documenting individual model compliance is insufficient. '
        'The coordination must be evaluated and documented. '
        'CSAS provides the measurement instrument for this requirement.'
    )

    add_heading('9.4 Implications for Sovereign AI Deployment', level=2)
    if full_version:
        add_para(
            'Sovereign AI deployments (R-AGAM, KFUPM pathway) run multiple models under national constraint requirements. '
            'Language constraints, data residency constraints, and domain constraints must persist across all models in the sovereign stack. '
            'CSAS provides the measurement standard for sovereign coordination assurance.'
        )
    else:
        add_para(
            'Sovereign AI deployments run multiple models under national constraint requirements. '
            'Language constraints, data residency constraints, and domain constraints must persist across all models in the sovereign stack. '
            'CSAS provides the measurement standard for sovereign coordination assurance.'
        )

    doc.add_page_break()

    # ===== 10. LIMITATIONS =====
    add_heading('10. Limitations', level=1)

    add_para(
        'CSAS is a theoretical framework. '
        'No empirical CSAS evaluations have been conducted. '
        'The metric definitions are logically derived from MTCP constructs. '
        'Empirical validation requires coordination-specific evaluation runs.'
    )
    add_para(
        'The minimum probe count (50 per boundary) is an estimate. '
        'Statistical power analysis for coordination-specific effects has not been conducted. '
        'Larger probe sets may be required for precise CAF estimation.'
    )
    add_para(
        'The failure taxonomy is derived from logical analysis of coordination architectures. '
        'It is not derived from empirical observation of coordination failures. '
        'Real coordination failures may not fit cleanly into four types.'
    )
    add_para(
        'CSAS assumes a linear or DAG coordination topology. '
        'Cyclic coordination (where M_j output feeds back to M_i) is not addressed. '
        'Cyclic systems introduce feedback loops that may require different measurement instruments.'
    )

    # ===== 11. CONCLUSION =====
    add_heading('11. Conclusion', level=1)

    add_para(
        'Individual model evaluation is necessary but not sufficient for multi-model system governance. '
        'Two admissible components do not guarantee an admissible system. '
        'CSAS provides the measurement framework for coordination admissibility.'
    )
    add_para(
        'The Coordination Admissibility Lemma extends the Persistence Lemma from temporal persistence to spatial persistence. '
        'Both share the core principle: admissibility is not self-maintaining. '
        'It must be verified at every point where constraint state could change.'
    )
    add_para(
        'CSAS is transcript-computable, model-access-free, and implementable using existing MTCP infrastructure. '
        'It provides a grading scale, a failure taxonomy, and a measurement protocol. '
        'It extends MTCP from single-model evaluation to multi-system coordination governance.'
    )


    # ===== REFERENCES =====
    doc.add_page_break()
    add_heading('References', level=1)

    refs = [
        'Abby, A. (2026a). Multi-Turn Constraint Persistence (MTCP): Benchmarking post-correction reliability in large language models. DOI: 10.17605/OSF.IO/DXGK5.',
        'Abby, A. (2026b). Universal latent attractors and Identity-Gate Satiation in large language models. DOI: 10.17605/OSF.IO/DXGK5.',
        'Abby, A. (2026c). Sigma-Forensics: A transcript-based safety audit framework for large language models. DOI: 10.17605/OSF.IO/DXGK5.',
        'Abby, A. (2026d). Behavioral Constraint Failure: A distinct failure mode in deployed large language models. DOI: 10.17605/OSF.IO/DXGK5.',
        'Abby, A. (2026e). Control Performance Degradation: A metric for detecting benchmark contamination in LLM evaluation. DOI: 10.17605/OSF.IO/DXGK5.',
        'Abby, A. (2026f). Stochastic vs. architectural constraint failure in LLMs: A diagnostic taxonomy. DOI: 10.17605/OSF.IO/DXGK5.',
        'Abby, A. (2026g). Three-layer constraint failure in AI systems: Evidence from model, pipeline, and protocol evaluation. DOI: 10.17605/OSF.IO/DXGK5.',
        'Abby, A. (2026h). Agent repositories as an unevaluated attack surface. DOI: 10.17605/OSF.IO/DXGK5.',
        'Abby, A. (2026i). Veterance and the Persistence Lemma: A formal metric for post-correction constraint failure in large language models. DOI: 10.17605/OSF.IO/DXGK5.',
        'Abby, A. (2026j). Which constraints do LLMs fail most? A five-vector empirical analysis. DOI: 10.17605/OSF.IO/DXGK5.',
        'Abby, A. (2026l). Transcript-only evidence standards for LLM behavioral claims. DOI: 10.17605/OSF.IO/DXGK5.',
        'Abby, A. (2026m). IGS Monitoring Systems: Designing automated post-deployment oversight. DOI: 10.17605/OSF.IO/DXGK5.',
        'European Parliament. (2024). Regulation (EU) 2024/1689 (Artificial Intelligence Act). Official Journal of the European Union.',
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        run = p.add_run(ref)
        run.font.name = 'Arial'
        run.font.size = Pt(10)

    return doc


# Build FULL version
doc_full = build_paper(full_version=True)
full_path = '/Users/aimeestanyer/Desktop/Paper28_CrossSystem_Admissibility_FULL.docx'
doc_full.save(full_path)
print(f'Saved: {full_path}')

# Build PUBLIC version
doc_pub = build_paper(full_version=False)
pub_path = '/Users/aimeestanyer/Desktop/Paper28_CrossSystem_Admissibility_PUBLIC.docx'
doc_pub.save(pub_path)
print(f'Saved: {pub_path}')
