"""
Build Paper 43: Governance Reference Conditions (GRC).
Framework F33: Comparability conditions for coordinated systems.
Generates PUBLIC and FULL .docx versions.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DESKTOP = os.path.expanduser("~/Desktop")


def set_document_formatting(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_paragraph_justified(doc, text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.25
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return para


def add_header(doc, text):
    for section in doc.sections:
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_footer(doc, text):
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


ABSTRACT = [
    "This paper defines Governance Reference Conditions (GRC). "
    "GRC specifies what must hold for two systems to be canonically comparable. "
    "The comparability problem is logically prior to coordination measurement. "
    "CSAS and JRS produce governance-valid findings only when GRC holds.",

    "Five conditions define canonical comparability. "
    "Same constraint object. LRP satisfaction for both. "
    "Same evaluation conditions. JRS above threshold for both. "
    "BEC integrity for both. All five must hold simultaneously.",

    "The Comparability Lemma constrains CSAS interpretation. "
    "CSAS degradation is only a governance failure when GRC is confirmed. "
    "Without GRC, degradation may represent structural divergence. "
    "The Canonical Priority Theorem makes this formally prior.",
]

INTRODUCTION = [
    "CSAS (Paper 28) measures coordination admissibility between systems. "
    "When CSAS degrades at a boundary, the interpretation matters. "
    "Is the degradation a governance failure or a structural difference.",

    "Without formal comparability criteria, the question cannot be answered. "
    "Two systems evaluated under different conditions are not comparable. "
    "Two systems referencing different constraints are not comparable. "
    "Degradation between incomparable systems is not a governance finding.",

    "GRC addresses this directly. "
    "It defines five conditions for canonical comparability. "
    "It provides a binary determination: comparable or not. "
    "It constrains CSAS interpretation through the Comparability Lemma.",

    "The Canonical Priority Theorem establishes logical ordering. "
    "Comparability is prior to coordination measurement. "
    "You must confirm systems are comparable before measuring coordination. "
    "This is not a recommendation. It is a logical requirement.",
]

FIVE_CONDITIONS = [
    "Five conditions define canonical comparability.",

    "GRC-1: Same Constraint Object. "
    "Both systems reference the same constraint_id from the COS registry. "
    "Alternatively, both reference COS-comparable equivalents. "
    "The Constraint Identity Lemma (Paper 40) gates this condition.",

    "GRC-2: LRP Satisfaction. "
    "Both systems have satisfied LRP conditions for their authority claims. "
    "Legitimacy score must equal 1.0 for both. "
    "Unverified authority invalidates comparability.",

    "GRC-3: Same Evaluation Conditions. "
    "Both systems were evaluated under the same temperature configuration. "
    "Infrastructure conditions must be formally equivalent. "
    "Different evaluation conditions produce different scores by design.",

    "GRC-4: JRS Above Threshold. "
    "Both systems have JRS scores at or above 0.75. "
    "Below this threshold, jurisdiction is not reliably resolved. "
    "Unresolved jurisdiction means governance is not established.",

    "GRC-5: BEC Integrity. "
    "Both systems have BEC chain integrity score of 1.0. "
    "A broken chain means evaluation records may have been tampered. "
    "Tampered records cannot support governance comparisons.",

    "All five conditions must hold simultaneously. "
    "Failing any single condition means not comparable. "
    "Partial compatibility does not exist. "
    "The conjunction is strict by design.",
]

COMPARABILITY_LEMMA = [
    "The Comparability Lemma constrains CSAS interpretation.",

    "CSAS degradation at a coordination boundary is classifiable as follows. "
    "When GRC compatibility is confirmed: governance failure. "
    "When GRC compatibility is not confirmed: structural divergence.",

    "Governance failure means something broke at the boundary. "
    "The systems were comparable. The coordination introduced degradation. "
    "This is actionable. The coordination architecture is the problem.",

    "Structural divergence means the systems were never comparable. "
    "The degradation reflects fundamental difference, not governance failure. "
    "This is not actionable through coordination fixes. "
    "It requires re-evaluation under comparable conditions.",

    "The distinction matters for remediation. "
    "Governance failures are fixable through coordination architecture changes. "
    "Structural divergence requires evaluation-level intervention. "
    "Misclassification leads to wrong remediation.",
]

CANONICAL_PRIORITY = [
    "The Canonical Priority Theorem establishes logical ordering.",

    "GRC compatibility is logically prior to coordination measurement. "
    "CSAS scores between incomparable systems are not governance findings. "
    "JRS scores between incomparable systems are not governance findings.",

    "The theorem has operational consequences. "
    "Step 1: Confirm GRC compatibility. "
    "Step 2: Only then conduct CSAS evaluation. "
    "Step 3: Only then conduct JRS evaluation. "
    "Reversing this order produces invalid findings.",

    "This makes the comparability problem prior to the coordination problem. "
    "Coordination measurement assumes comparability. "
    "That assumption must be verified, not assumed. "
    "GRC provides the verification protocol.",
]

INTEGRATION = [
    "GRC integrates with the Admissibility Gate. "
    "A model pair cannot receive PERMIT for coordination contexts. "
    "GRC compatibility must be confirmed first. "
    "The Gate enforces the Canonical Priority Theorem.",

    "GRC integrates with the Constraint Manifest. "
    "The manifest now includes grc_hash and constraint_object_id. "
    "Receiving systems can verify governance comparability from the manifest.",

    "GRC integrates with the MCP server. "
    "Two new tools: get_constraint_object and compare_governance_conditions. "
    "These enable automated comparability checking in coordination workflows.",
]

LIMITATIONS = [
    "GRC is a specification. "
    "No empirical GRC evaluations have been conducted. "
    "The five conditions are formally specified but untested.",

    "The binary nature may be conservative. "
    "Future versions may support weighted compatibility. "
    "Weighted compatibility would enable bounded comparison. "
    "This is deferred to empirical calibration.",

    "GRC-3 (same evaluation conditions) is difficult to verify retrospectively. "
    "Historical evaluations may lack infrastructure documentation. "
    "GRC requires prospective condition recording.",

    "The Canonical Priority Theorem is logical, not empirical. "
    "It cannot be falsified by data. "
    "It can be challenged on logical grounds. "
    "No logical challenge has been identified.",
]

CONCLUSION = [
    "GRC makes the comparability problem explicit and prior. "
    "Five conditions define when two systems are canonically comparable. "
    "Comparability must be confirmed before coordination is measured. "
    "This is the Canonical Priority Theorem.",

    "The Comparability Lemma constrains CSAS interpretation. "
    "Governance failure and structural divergence require different responses. "
    "Only GRC-confirmed degradation is classifiable as governance failure.",

    "GRC integrates with Gate, Manifest, and MCP server. "
    "The governance stack is now complete from constraint to comparability. "
    "COS defines what. LRP defines who. GRC defines when comparable.",

    "The full governance ordering is now established. "
    "COS then LRP then GRC then CSAS then JRS. "
    "Each layer is logically prior to the next. "
    "Each layer gates the validity of the next.",
]

REFERENCES = [
    "Abby, A. (2026a). Multi-Turn Constraint Persistence (MTCP). DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 28). Cross-System Admissibility Score. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 29). Constraint Jurisdiction Resolution. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 37). Blockchain Evidence Chain Integrity. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 40). Constraint Object Specification. DOI: 10.17605/OSF.IO/DXGK5.",
    "Abby, A. (2026, Paper 41). Legitimacy Resolution Protocol. DOI: 10.17605/OSF.IO/DXGK5.",
    "Framework F21 -- CSAS Definition. MTCP Research Programme.",
    "Framework F22 -- JRS Definition. MTCP Research Programme.",
    "Framework F28 -- BEC Definition. MTCP Research Programme.",
    "Framework F31 -- COS Definition. MTCP Research Programme.",
    "Framework F32 -- LRP Definition. MTCP Research Programme.",
    "Framework F33 -- GRC Definition. MTCP Research Programme.",
]


def build_paper(is_public=True):
    doc = Document()
    set_document_formatting(doc)

    if is_public:
        add_header(doc, "MTCP V1.0 -- Governance Reference Conditions")
    else:
        add_header(doc, "CONFIDENTIAL -- NDA REQUIRED")
        add_footer(doc, "MTCP V1.0 | DOI: 10.17605/OSF.IO/DXGK5 | 2026 A. Abby. All Rights Reserved.")

    add_heading_styled(doc, "MTCP Paper 43", level=1)
    add_heading_styled(doc, "Governance Reference Conditions", level=2)
    add_paragraph_justified(doc, "")
    if not is_public:
        add_paragraph_justified(doc, "CONFIDENTIAL -- NDA REQUIRED")
    add_paragraph_justified(doc, "Author: A. Abby")
    add_paragraph_justified(doc, "Contact: admin@mtcp.live")
    add_paragraph_justified(doc, "DOI: 10.17605/OSF.IO/DXGK5")
    add_paragraph_justified(doc, "Date: May 2026")
    add_paragraph_justified(doc, "Version: 1.0")
    add_paragraph_justified(doc, "")
    add_paragraph_justified(doc, "MTCP Research Programme -- Paper 43")
    add_paragraph_justified(doc, "Framework F33")

    doc.add_page_break()

    add_heading_styled(doc, "1. Abstract", level=1)
    for para in ABSTRACT:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "2. Introduction", level=1)
    for para in INTRODUCTION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "3. Five Governance Reference Conditions", level=1)
    for para in FIVE_CONDITIONS:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "4. Comparability Lemma", level=1)
    for para in COMPARABILITY_LEMMA:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "5. Canonical Priority Theorem", level=1)
    for para in CANONICAL_PRIORITY:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "6. Integration", level=1)
    for para in INTEGRATION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "7. Limitations", level=1)
    for para in LIMITATIONS:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "8. Conclusion", level=1)
    for para in CONCLUSION:
        add_paragraph_justified(doc, para)

    add_heading_styled(doc, "9. References", level=1)
    for ref in REFERENCES:
        add_paragraph_justified(doc, ref)

    return doc


if __name__ == "__main__":
    print("Building Paper 43: Governance Reference Conditions")
    print("=" * 60)

    doc_pub = build_paper(is_public=True)
    pub_path = os.path.join(DESKTOP, "Paper43_Governance_Reference_Conditions_PUBLIC.docx")
    doc_pub.save(pub_path)
    print(f"PUBLIC version saved: {pub_path}")

    doc_full = build_paper(is_public=False)
    full_path = os.path.join(DESKTOP, "Paper43_Governance_Reference_Conditions_FULL.docx")
    doc_full.save(full_path)
    print(f"FULL version saved: {full_path}")

    print("=" * 60)
    print("Done.")
